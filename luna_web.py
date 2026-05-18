#!/usr/bin/env python3
"""
Luna Web - Serveur YAWatch-Luna (Proprio Ludo)
Endpoints: chat, greeting, call Tavus, invitation contact, webhook SMS
"""
import os
import sys
import re
import time
import uuid
import json
import asyncio
import logging
from pathlib import Path

# Path d'import pour pv_recette.py (Docker: /app/utils/, local: ../luna-exploitants/scripts/)
def _resolve_exploitants_root() -> str:
    """Repère le repo luna-exploitants (sibling ou legacy EXPLOITANTS)."""
    base = Path(__file__).resolve().parent
    for candidate in (
        base / ".." / "luna-exploitants",
        base / ".." / "EXPLOITANTS",
        base / ".." / ".." / "luna-exploitants",
        base / ".." / ".." / "EXPLOITANTS",
    ):
        root = candidate.resolve()
        if (root / "scripts" / "pv_recette.py").is_file():
            return str(root)
    return ""


_exploitants_root = _resolve_exploitants_root()
_utils_dir = os.path.join(os.path.dirname(__file__), "utils")
for _p in [_utils_dir, os.path.join(_exploitants_root, "scripts") if _exploitants_root else ""]:
    if _p and os.path.isdir(_p) and _p not in sys.path:
        sys.path.insert(0, _p)
import openai
from collections import defaultdict
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from dotenv import load_dotenv
from openai import OpenAI
from integrations.llm.provider import build_llm_client, get_llm_model, get_provider_label
from contextlib import asynccontextmanager
from fastapi import FastAPI, Request, Response, WebSocket, WebSocketDisconnect
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from starlette.websockets import WebSocketState
from pydantic import BaseModel, Field, field_validator
from typing import Optional, Dict, Any
import httpx

from integrations.twilio.sms_client import TwilioSMSClient
from integrations.twilio.voice_client import TwilioVoiceClient
from integrations.email.email_client import EmailClient
from integrations.email.gmail_client import GmailClient

# Tavus: optionnel (mode "lite" = sans visio avatar)
try:
    from integrations.tavus.tavus_client import TavusClient, build_tavus_context
    _TAVUS_AVAILABLE = True
except ImportError:
    _TAVUS_AVAILABLE = False
    TavusClient = None
    def build_tavus_context(**kwargs): return ""

# Simli: desactive (Tavus est le systeme visio principal)
_SIMLI_AVAILABLE = False
handle_simli_session = None

# Core modules (optional - graceful fallback if Redis down)
try:
    from core.memory.memory_manager import MemoryManager
    from core.memory.redis_client import RedisClient
    from core.memory.schemas import (
        PlanType, MessageRole, Channel, Conversation, ConversationStatus,
        SubscriberProfile, InstructionType, ActionType as SchemaActionType,
        UnifiedSession, UnifiedMessage, ChannelHandoff,
        SessionStatus, MoodIndicator,
    )
    from core.safety.guardian import SafetyGuardian, SafetyLevel
    from core.actions.quota_guard import QuotaGuard
    from core.actions.models import ActionType as CoreActionType
    from core.instructions.parser import InstructionParser, ParsedInstruction
    from core.instructions.scheduler import InstructionScheduler, ScheduledTask
    from core.instructions.executor import InstructionExecutor, create_instruction_executor
    from core.documents.generator import DocumentGenerator
    _CORE_AVAILABLE = True
except ImportError as _import_err:
    _CORE_AVAILABLE = False
    import traceback
    print(f"CORE IMPORT ERROR: {_import_err}")
    traceback.print_exc()

# Gamification IA Watch World (optional, non-blocking)
try:
    from core.gamification.routes import gamification_router
    from core.gamification.engine import award_xp_safe, initialize_player_safe
    from core.gamification.redis_ops import GamificationRedisOps
    _GAMIFICATION_AVAILABLE = True
except ImportError:
    _GAMIFICATION_AVAILABLE = False

# Social interactions (optional, non-blocking)
try:
    from core.social.routes import social_router
    _SOCIAL_AVAILABLE = True
except ImportError:
    _SOCIAL_AVAILABLE = False

# World social layer (privacy, avatar, invitations, chat, visitors)
try:
    from core.world import world_router
    _WORLD_SOCIAL_AVAILABLE = True
except ImportError:
    _WORLD_SOCIAL_AVAILABLE = False

# Exploitant dashboard
try:
    from core.exploitant import exploitant_router
    _EXPLOITANT_AVAILABLE = True
except ImportError:
    _EXPLOITANT_AVAILABLE = False

# Secretary module (documents, budget, reminders)
try:
    from core.secretary.routes import secretary_router
    _SECRETARY_AVAILABLE = True
except ImportError:
    _SECRETARY_AVAILABLE = False

# Vault module (coffre-fort documentaire)
try:
    from core.vault import vault_router
    _VAULT_AVAILABLE = True
except ImportError:
    _VAULT_AVAILABLE = False

# Form Filler (remplissage intelligent de formulaires PDF)
try:
    from core.form_filler.routes import form_filler_router
    _FORM_FILLER_AVAILABLE = True
except ImportError:
    _FORM_FILLER_AVAILABLE = False

# Cortex: cerveau autonome (securite, monitoring, commandes SMS d'urgence)
try:
    from core.cortex.integration import (
        init_cortex, start_cortex, stop_cortex,
        cortex_middleware_check, cortex_analyze_request,
        cortex_record_failed_auth, cortex_handle_sms,
        get_cortex, cortex_routes,
    )
    _CORTEX_AVAILABLE = True
except ImportError:
    _CORTEX_AVAILABLE = False
    def init_cortex(redis_client=None): return None
    async def start_cortex(): pass
    async def stop_cortex(): pass
    def cortex_middleware_check(ip, path): return True, ""
    def cortex_analyze_request(*a, **kw): pass
    def cortex_record_failed_auth(*a, **kw): pass
    async def cortex_handle_sms(f, b): return None
    def get_cortex(): return None

load_dotenv(os.path.join(os.path.dirname(__file__), ".env"))

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("luna_web")

# --- Config ---
LUNA_MODE = os.getenv("LUNA_MODE", "full").lower()  # "lite" (chat+voix+SMS) ou "full" (+visio Tavus)
TAVUS_CALLBACK_URL = os.getenv("TAVUS_CALLBACK_URL", "")  # URL publique pour webhooks Tavus
VOICE_CALLBACK_URL = os.getenv("VOICE_CALLBACK_URL", "")  # URL publique pour appels vocaux (TwiML + WebSocket)
TENANT_ID = 1  # Fallback pour retro-compatibilite (REQUIRE_AUTH=false)
_PROPRIO_TENANT_ID = 1  # Tenant du fondateur (Ludo)
LEGAL_MODE = "assistance_only"  # Mode legal: aide contextuelle, pas de surveillance garantie
REQUIRE_AUTH = os.getenv("REQUIRE_AUTH", "true").lower() == "true"


def _reload_env():
    """Recharge le .env et met a jour les globals qui cachent des valeurs d'env."""
    global LUNA_MODE, TAVUS_CALLBACK_URL, VOICE_CALLBACK_URL, REQUIRE_AUTH
    global OPENAI_API_KEY, OPENAI_MODEL, ADMIN_NUMBER, SETUP_OPENAI_API_KEY
    global _JWT_SECRET, _JWT_ALGORITHM
    load_dotenv(os.path.join(os.path.dirname(__file__), ".env"), override=True)
    LUNA_MODE = os.getenv("LUNA_MODE", "full").lower()
    TAVUS_CALLBACK_URL = os.getenv("TAVUS_CALLBACK_URL", "")
    VOICE_CALLBACK_URL = os.getenv("VOICE_CALLBACK_URL", "")
    REQUIRE_AUTH = os.getenv("REQUIRE_AUTH", "true").lower() == "true"
    OPENAI_API_KEY = os.getenv("LLM_API_KEY") or os.getenv("OPENAI_API_KEY")
    OPENAI_MODEL = get_llm_model()
    ADMIN_NUMBER = os.getenv("ADMIN_NUMBER", "")
    SETUP_OPENAI_API_KEY = os.getenv("SETUP_OPENAI_API_KEY", "")
    _jwt_raw = os.getenv("JWT_SECRET_KEY", "")
    if not _jwt_raw:
        raise SystemExit("ERREUR FATALE: JWT_SECRET_KEY manquante dans .env — securite compromise.")
    _JWT_SECRET = _jwt_raw
    _JWT_ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")

# --- PV de Recette (verrouillage serveur) ---
# Priorite: pv_lock.json (HMAC) > .env PV_SIGNED
_setup_permanently_disabled = False
_sign_pv_lock = asyncio.Lock()
try:
    from pv_recette import PVRecette
    _pv_lock_result = PVRecette.verify_pv_lock()
    if _pv_lock_result["valid"]:
        PV_SIGNED = True
        _setup_permanently_disabled = True
        logger.info(f"pv_lock.json valide — setup DESACTIVE (exploitant: {_pv_lock_result['data'].get('tenant_name', '?')})")
    else:
        PV_SIGNED = os.getenv("PV_SIGNED", "false").lower() == "true"
        if _pv_lock_result["reason"] != "pv_lock.json introuvable":
            logger.warning(f"pv_lock invalide: {_pv_lock_result['reason']}")
except ImportError:
    PV_SIGNED = os.getenv("PV_SIGNED", "false").lower() == "true"
    logger.warning("Module pv_recette non disponible — fallback .env")
PV_SIGNATURE_HASH = os.getenv("PV_SIGNATURE_HASH", "")
_pv_locked = not PV_SIGNED  # True = serveur en mode SETUP uniquement

# --- License Heartbeat (protection anti-piratage) ---
_license_heartbeat = None
_LICENSE_KEY = os.getenv("YAWATCH_LICENSE_KEY", "")
_LICENSE_SERVER = os.getenv("YAWATCH_LICENSE_SERVER",
    "https://iawatch-backend-674304336025.europe-west1.run.app")

if _LICENSE_KEY and not _pv_locked:
    try:
        from core.license.heartbeat import LicenseHeartbeat
        _license_heartbeat = LicenseHeartbeat(
            _LICENSE_KEY, _LICENSE_SERVER,
            hmac_key=os.getenv("JWT_SECRET_KEY", "")
        )
        logger.info(f"License heartbeat active (server: {_LICENSE_SERVER})")
    except ImportError:
        logger.warning("Module license non disponible - heartbeat desactive")

# --- Config: graceful en mode SETUP ---
# Rétrocompat: LLM_API_KEY prioritaire, OPENAI_API_KEY accepté si provider=openai
OPENAI_API_KEY = os.getenv("LLM_API_KEY") or os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = get_llm_model()
ADMIN_NUMBER = os.getenv("ADMIN_NUMBER", "")
SETUP_OPENAI_API_KEY = os.getenv("SETUP_OPENAI_API_KEY", "")

# Feature 1: Destruction cle fondateur apres PV
if _setup_permanently_disabled:
    SETUP_OPENAI_API_KEY = None
    logger.info("SETUP_OPENAI_API_KEY detruite (pv_lock actif)")

if _pv_locked:
    # Mode SETUP: ne crashe pas sur cles manquantes
    if not OPENAI_API_KEY:
        logger.warning("LLM_API_KEY manquante - mode SETUP")
    if not ADMIN_NUMBER:
        logger.warning("ADMIN_NUMBER manquant - mode SETUP")
else:
    if not OPENAI_API_KEY:
        raise SystemExit(
            f"ERREUR FATALE: LLM_API_KEY manquante pour provider '{os.getenv('LLM_PROVIDER','openai')}'. "
            "Voir .env.template et CONFIG.md"
        )
    if not ADMIN_NUMBER:
        raise SystemExit("ERREUR FATALE: ADMIN_NUMBER manquant dans .env. Voir .env.example.")

# --- Health Monitor (auto-alert admin on critical failures) ---
_health_last_alert = {}  # {alert_key: timestamp} — throttle 1 alerte / 30 min par type

def _notify_admin_health(message: str, alert_key: str = "generic"):
    """Alerte l'admin via Telegram (prioritaire) ou SMS (fallback).
    Throttle: max 1 alerte par type toutes les 30 minutes."""
    import time as _time
    import httpx as _httpx
    now = _time.time()
    last = _health_last_alert.get(alert_key, 0)
    if now - last < 1800:  # 30 min
        return  # deja alerte recemment
    _health_last_alert[alert_key] = now

    full_msg = f"🚨 [LUNA ALERTE] {message}"
    logger.error(f"HEALTH ALERT: {message}")

    # 1. Telegram (gratuit, fiable)
    bot_token = os.getenv("ALERT_TELEGRAM_BOT_TOKEN", "")
    founder_chat_id = os.getenv("FOUNDER_TELEGRAM_CHAT_ID", "")
    if bot_token and founder_chat_id:
        try:
            import httpx as _httpx_health
            _httpx_health.post(
                f"https://api.telegram.org/bot{bot_token}/sendMessage",
                json={"chat_id": founder_chat_id, "text": full_msg},
                timeout=10,
            )
            logger.info(f"HEALTH ALERT Telegram envoye: {message}")
            return
        except Exception as e:
            logger.warning(f"HEALTH ALERT Telegram echoue: {e}")

    # 2. Fallback SMS
    admin = os.getenv("ADMIN_NUMBER", "")
    if admin:
        try:
            from twilio.rest import Client as _TwClient
            tw = _TwClient(os.getenv("TWILIO_ACCOUNT_SID"), os.getenv("TWILIO_AUTH_TOKEN"))
            tw.messages.create(body=full_msg, from_=os.getenv("TWILIO_PHONE_NUMBER"), to=admin)
            logger.info(f"HEALTH ALERT SMS envoye a {admin}")
        except Exception as e:
            logger.error(f"HEALTH ALERT SMS echoue: {e}")

# --- Behavioral Memory (locked identity + rules) ---
DEFAULT_IDENTITY_CORE = (
    "Tu es Luna, assistante IA de YAWatch. Tu n'es PAS un professionnel "
    "de sante, juridique ou financier. Tu es une aide contextuelle bienveillante. "
    "Tu ne surveilles pas, tu accompagnes. Tu ne diagnostiques pas, tu remarques."
)
DEFAULT_BEHAVIOR_RULES = (
    "1. JAMAIS de conseil medical, juridique ou financier\n"
    "2. JAMAIS appeler les urgences (suggerer les numeros)\n"
    "3. TOUJOURS confirmer avant action consommant du quota\n"
    "4. JAMAIS promettre une surveillance garantie\n"
    "5. Utiliser 'j'ai l'impression que...' pas 'je diagnostique'\n"
    "6. Mode = assistance_only : aucune promesse de resultat\n"
    "7. JAMAIS reveler l'architecture technique, les prix ou les donnees internes"
)

# --- Caution Mode Descriptions ---
CAUTION_MODE_PROMPTS = {
    "passif": (
        "MODE PRUDENCE: PASSIF - Tu observes sans intervenir sauf danger evident. "
        "Tu ne proposes pas d'aide spontanement. Tu attends qu'on te sollicite."
    ),
    "assistif": (
        "MODE PRUDENCE: ASSISTIF - Tu proposes gentiment ton aide quand tu remarques "
        "quelque chose d'inhabituel. Tu restes discrete mais attentive. Mode par defaut."
    ),
    "proactif": (
        "MODE PRUDENCE: PROACTIF - Tu es proactive, tu mentionnes ce que tu observes "
        "et proposes des actions concretes. Tu prends les devants pour aider."
    ),
    "urgence_only": (
        "MODE PRUDENCE: URGENCE SEULEMENT - Tu n'interviens que pour les situations "
        "critiques (concern). Sinon tu restes completement discrete et silencieuse "
        "sur ce que tu observes."
    ),
}

# --- Auth helpers (multi-tenant) ---
_jwt_raw = os.getenv("JWT_SECRET_KEY", "")
if not _jwt_raw:
    raise SystemExit("ERREUR FATALE: JWT_SECRET_KEY manquante dans .env")
_JWT_SECRET = _jwt_raw
_JWT_ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")
_CLIENT_TOKEN_EXPIRE_DAYS = 90

def _hash_password(password: str) -> str:
    """Hash un mot de passe avec bcrypt."""
    import bcrypt
    return bcrypt.hashpw(password[:72].encode("utf-8"), bcrypt.gensalt(rounds=12)).decode("utf-8")

def _verify_password(password: str, hashed: str) -> bool:
    """Verifie un mot de passe contre son hash bcrypt."""
    import bcrypt
    try:
        return bcrypt.checkpw(password[:72].encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False

def _bootstrap_proprio_auth():
    """Cree l'enregistrement auth pour le proprio (tenant 1) si absent dans Redis."""
    if not _redis_client:
        return
    proprio_email = os.getenv("PROPRIO_EMAIL", "").strip().lower()
    proprio_password = os.getenv("PROPRIO_PASSWORD", "").strip()
    if not proprio_email or not proprio_password:
        return
    existing = _redis_client.get_auth_by_email(proprio_email)
    if existing:
        # S'assurer que c'est bien le tenant 1
        if existing.get("tenant_id") != _PROPRIO_TENANT_ID:
            logger.warning(f"PROPRIO_EMAIL {proprio_email} est lie au tenant {existing.get('tenant_id')}, pas {_PROPRIO_TENANT_ID}")
        else:
            logger.info(f"Auth proprio OK: {proprio_email} (tenant {_PROPRIO_TENANT_ID})")
        return
    # Creer le record auth
    password_hash = _hash_password(proprio_password)
    created = _redis_client.create_auth_record(proprio_email, password_hash, _PROPRIO_TENANT_ID, "fondateur")
    if created:
        logger.info(f"AUTH BOOTSTRAP: cree {proprio_email} pour tenant {_PROPRIO_TENANT_ID} (plan fondateur)")
    else:
        # create_auth_record a echoue (race condition) — forcer via set direct
        import json as _json
        key = f"{_redis_client.prefix}:auth:{proprio_email}"
        record = _json.dumps({
            "tenant_id": _PROPRIO_TENANT_ID,
            "password_hash": password_hash,
            "plan": "fondateur",
            "active": True,
            "created_at": time.time(),
            "email": proprio_email,
        })
        _redis_client.client.set(key, record)
        logger.info(f"AUTH BOOTSTRAP (force): cree {proprio_email} pour tenant {_PROPRIO_TENANT_ID}")


def _create_client_token(tenant_id: int, email: str, plan: str, first_name: str = "") -> str:
    """Cree un JWT client valide 7 jours."""
    import jwt as pyjwt
    payload = {
        "tenant_id": tenant_id,
        "email": email,
        "plan": plan,
        "role": "client",
        "first_name": first_name,
        "iat": int(time.time()),
        "exp": int(time.time()) + _CLIENT_TOKEN_EXPIRE_DAYS * 86400,
    }
    return pyjwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)

def _decode_client_token(token: str) -> Optional[dict]:
    """Decode un JWT client. Retourne le payload ou None."""
    if not token:
        return None
    try:
        import jwt as pyjwt
        payload = pyjwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGORITHM])
        if payload.get("role") != "client":
            return None
        return payload
    except Exception:
        return None


def _decode_admin_token(token: str) -> Optional[dict]:
    """Decode un JWT admin. Retourne le payload ou None."""
    if not token:
        return None
    try:
        import jwt as pyjwt
        payload = pyjwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGORITHM])
        if payload.get("role") != "admin":
            return None
        return payload
    except Exception:
        return None

def _extract_bearer(request: Request) -> str:
    """Extrait le token Bearer du header Authorization."""
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        return auth[7:]
    return ""

_PUBLIC_PATHS = (
    "/api/auth/",
    "/api/status",
    "/api/admin/",
    "/api/setup/",
    "/api/exploitant/",
    "/api/maintenance",
    "/ready",
    "/health",
    "/api/stripe/webhook",
    "/api/webhook/sms",
    "/api/webhook/sms-status",
    "/api/webhook/tavus",
    "/api/webhook/simli",
    "/api/webhook/voice-incoming",
    "/api/webhooks/meetingbaas",
    "/api/meeting/webhook",
    "/api/voice-call/twiml",
    "/api/voice-call/media-stream",
    "/api/sync/tavus",
    "/api/sync/twilio",
    "/api/email/oauth/",
    "/api/cortex/telegram/webhook",
    "/api/app/version",
    "/api/debug/log",
    "/api/rooms/",  # Accessible via HMAC member token (invites par SMS)
    "/download",
    "/download/",
    "/static/",
)

def _is_public_path(path: str) -> bool:
    """Verifie si un path est public (pas d'auth requise)."""
    return any(path.startswith(p) for p in _PUBLIC_PATHS)


# --- Clients ---
if _pv_locked:
    # Mode SETUP: init graceful, ne crashe pas sur cles manquantes
    openai_client = build_llm_client(OPENAI_API_KEY) if OPENAI_API_KEY else None
    try:
        sms_client = TwilioSMSClient.from_env()
    except Exception:
        sms_client = None
    voice_client = None
    tavus_client = None
    email_client = EmailClient.from_env()
else:
    openai_client = build_llm_client(OPENAI_API_KEY)
    sms_client = TwilioSMSClient.from_env()
    voice_client = TwilioVoiceClient.from_env()
    tavus_client = TavusClient.from_env() if _TAVUS_AVAILABLE and LUNA_MODE == "full" else None
    email_client = EmailClient.from_env()

# --- Reservation clients ---
try:
    from integrations.reservations.amadeus_client import AmadeusClient
    amadeus_client = AmadeusClient.from_env()
except Exception as _e:
    logger.info(f"Amadeus client non disponible: {_e}")
    amadeus_client = None

try:
    from integrations.reservations.duffel_client import DuffelClient
    duffel_client = DuffelClient.from_env()
    if duffel_client.is_configured:
        logger.info(f"Duffel client OK (mode: {'test' if duffel_client.is_test else 'live'})")
except Exception as _e:
    logger.info(f"Duffel client non disponible: {_e}")
    duffel_client = None

try:
    from integrations.reservations.thefork_client import TheForkClient
    thefork_client = TheForkClient.from_env()
except Exception as _e:
    logger.info(f"TheFork client non disponible: {_e}")
    thefork_client = None

# MeetingBaas — bot de réunion (Zoom, Meet, Teams, Webex)
try:
    from integrations.recall.meetingbaas_client import MeetingBaasClient, detect_platform, format_transcript, PLATFORM_NAMES, STATUS_LABELS
    recall_client = MeetingBaasClient.from_env()
    if recall_client.is_configured:
        logger.info("MeetingBaas client OK")
    else:
        logger.info("MeetingBaas non configuré (MEETINGBAAS_API_KEY manquant)")
except Exception as _e:
    logger.info(f"MeetingBaas non disponible: {_e}")
    recall_client = None
    def detect_platform(url): return "unknown"
    def format_transcript(segs): return ""
    PLATFORM_NAMES = {}
    STATUS_LABELS = {}

# État des réunions — persisté dans Redis (survit aux redéploiements Cloud Run)
_recall_sessions: Dict[str, dict] = {}       # cache local fallback
_recall_transcripts: Dict[str, list] = {}    # cache local fallback

_MTG_SESSION_TTL = 86400  # 24h


def _mtg_session_key(bot_id: str) -> str:
    return f"mtg:session:{bot_id}"


def _mtg_transcripts_key(bot_id: str) -> str:
    return f"mtg:transcripts:{bot_id}"


def _mtg_session_get(bot_id: str) -> Optional[dict]:
    try:
        if _redis_client and _redis_client.ping():
            raw = _redis_client.client.get(_mtg_session_key(bot_id))
            if raw:
                import json as _json
                return _json.loads(raw)
    except Exception:
        pass
    return _recall_sessions.get(bot_id)


def _mtg_session_set(bot_id: str, data: dict) -> None:
    import json as _json
    _recall_sessions[bot_id] = data
    try:
        if _redis_client and _redis_client.ping():
            _redis_client.client.setex(_mtg_session_key(bot_id), _MTG_SESSION_TTL, _json.dumps(data))
    except Exception:
        pass


def _mtg_transcripts_get(bot_id: str) -> list:
    try:
        if _redis_client and _redis_client.ping():
            raw = _redis_client.client.get(_mtg_transcripts_key(bot_id))
            if raw:
                import json as _json
                return _json.loads(raw)
    except Exception:
        pass
    return _recall_transcripts.get(bot_id, [])


def _mtg_transcripts_set(bot_id: str, segments: list) -> None:
    import json as _json
    _recall_transcripts[bot_id] = segments
    try:
        if _redis_client and _redis_client.ping():
            _redis_client.client.setex(_mtg_transcripts_key(bot_id), _MTG_SESSION_TTL, _json.dumps(segments))
    except Exception:
        pass


def _mtg_transcripts_append(bot_id: str, seg: dict) -> None:
    segments = _mtg_transcripts_get(bot_id)
    segments.append(seg)
    _mtg_transcripts_set(bot_id, segments)


def _mtg_list_sessions(tenant_id: int) -> list:
    sessions = []
    try:
        if _redis_client and _redis_client.ping():
            import json as _json
            keys = _redis_client.client.keys("mtg:session:*")
            for key in keys:
                raw = _redis_client.client.get(key)
                if raw:
                    try:
                        s = _json.loads(raw)
                        if s.get("tenant_id") == tenant_id:
                            sessions.append(s)
                    except Exception:
                        pass
            return sessions
    except Exception:
        pass
    return [s for s in _recall_sessions.values() if s.get("tenant_id") == tenant_id]

# Auto-configure SMS status callback URL si pas defini
if sms_client and sms_client.is_configured and VOICE_CALLBACK_URL and not sms_client.status_callback_url:
    sms_client.status_callback_url = f"{VOICE_CALLBACK_URL}/api/webhook/sms-status"
    logger.info(f"SMS status callback auto-configure: {sms_client.status_callback_url}")

# Stockage des accusés de reception SMS {sid: {to, body_preview, status, ts, delivered_at}}
_sms_tracking: Dict[str, Dict] = {}

def _tracked_sms_send(to: str, body: str, label: str = ""):
    """Envoie un SMS et track l'accuse de reception."""
    if _test_mode:
        logger.info(f"[TEST MODE] SMS bloque vers {to}: {body[:60]}...")
        return True, {"sid": f"TEST_{label}", "test_mode": True}
    success, details = sms_client.send(to, body)
    if success and details.get("sid"):
        sid = details["sid"]
        _sms_tracking[sid] = {
            "sid": sid,
            "to": to,
            "body_preview": body[:80] + ("..." if len(body) > 80 else ""),
            "label": label,
            "status": details.get("status", "queued"),
            "sent_at": datetime.utcnow().isoformat(),
            "delivered_at": None,
            "error_code": None,
        }
        # Garder max 200 entrees
        if len(_sms_tracking) > 200:
            oldest = sorted(_sms_tracking.keys())[:50]
            for k in oldest:
                _sms_tracking.pop(k, None)
    return success, details

# Gmail OAuth2 (per-tenant email)
_gmail_base_url = os.getenv("LUNA_BASE_URL", "https://luna-beta-674304336025.europe-west1.run.app")
gmail_client = GmailClient.from_env(base_url=_gmail_base_url)

# --- Core modules (Redis, Safety, Quota) ---
_redis_client: Optional[object] = None
_memory_manager: Optional[object] = None
_safety_guardian: Optional[object] = None
_quota_guard: Optional[object] = None
_scheduler: Optional[object] = None
_executor: Optional[object] = None
_instruction_loop_task: Optional[object] = None
_doc_generator: Optional[object] = None
_perception_detector: Optional[object] = None
_perception_analyzer: Optional[object] = None
_test_mode: bool = False  # En mode test, les SMS ne sont PAS envoyes
_notification_engine: Optional[object] = None

# Invitations visio en attente de reponse SMS (phone -> {tenant_id, subscriber_name, contact_name, timestamp})
_pending_visio_invites: Dict[str, Dict] = {}

# Rate limiting vision caméra (ip -> timestamp dernier appel)
_vision_last_call: Dict[str, float] = {}
_visio_scene_cache: Dict[str, str] = {}  # ip -> dernière description (détection de changement)

def _init_core():
    """Initialize core modules. Graceful if Redis is down or imports failed."""
    global _redis_client, _memory_manager, _safety_guardian, _quota_guard, _scheduler, _executor, _doc_generator, _perception_detector, _perception_analyzer, _notification_engine
    if not _CORE_AVAILABLE:
        logger.warning("Core modules non disponibles (import failed) - mode degrade")
        return
    try:
        _redis_client = RedisClient()
        if _redis_client.ping():
            _memory_manager = MemoryManager(
                tenant_id=TENANT_ID,
                plan=PlanType.ESSENTIEL,
                redis_client=_redis_client,
            )
            _safety_guardian = SafetyGuardian(
                tenant_id=TENANT_ID,
                memory_manager=_memory_manager,
                sms_service=sms_client,
                legal_mode=LEGAL_MODE,
            )
            _quota_guard = QuotaGuard(memory_manager=_memory_manager)
            _scheduler = InstructionScheduler()
            _executor = create_instruction_executor(
                memory_manager=_memory_manager,
                sms_service=sms_client,
                safety_guardian=_safety_guardian,
                voice_service=voice_client,
                visio_service=tavus_client,
            )
            # Callback pour que l'executor puisse stocker les params d'appel planifié
            def _on_scheduled_call(call_sid, params):
                _voice_call_params[call_sid] = params
            _executor.on_call_initiated = _on_scheduled_call
            _doc_generator = DocumentGenerator(
                output_dir=os.path.join(os.path.dirname(__file__), "static", "documents"),
                tenant_id=TENANT_ID,
            )
            # Initialize behavioral memory (if not already set)
            if not _memory_manager.get_behavioral_memory("identity_core"):
                _memory_manager.set_behavioral_memory("identity_core", DEFAULT_IDENTITY_CORE)
            if not _memory_manager.get_behavioral_memory("behavior_rules"):
                _memory_manager.set_behavioral_memory("behavior_rules", DEFAULT_BEHAVIOR_RULES)
            logger.info("Behavioral memory loaded")

            # Perception (camera navigateur -> OpenAI Vision)
            try:
                from core.perception.detector import PerceptionDetector
                from core.perception.analyzer import SceneAnalyzer
                _perception_detector = PerceptionDetector()
                _perception_analyzer = SceneAnalyzer()
                # Init immediate si cle OpenAI dispo
                if os.environ.get("OPENAI_API_KEY"):
                    _perception_detector.initialize()
            except ImportError:
                logger.info("Perception module non disponible")
            # Notification engine (Luna gentle reminders)
            try:
                from core.notifications import NotificationEngine
                _notification_engine = NotificationEngine(_redis_client, sms_service=sms_client)
                logger.info("Notification engine initialise (avec SMS)")
            except ImportError:
                logger.info("Notification module non disponible")
            logger.info("Core modules initialises (Redis OK, Scheduler OK, DocGen OK)")
        else:
            logger.warning("Redis injoignable - mode degrade (memoire locale)")
    except Exception as e:
        logger.warning(f"Core init echoue: {e} - mode degrade")

_init_core()

# --- Pool MemoryManager per-tenant ---
_tenant_managers: Dict[int, object] = {}

def _get_tenant_manager(tenant_id: int):
    """Retourne le MemoryManager pour un tenant (lazy init, thread-safe)."""
    # Fast path: already exists (dict read is safe in CPython GIL)
    if tenant_id in _tenant_managers:
        return _tenant_managers[tenant_id]
    if not _CORE_AVAILABLE or not _redis_client:
        return _memory_manager  # fallback global
    # Slow path: create new manager (rare, only first access per tenant)
    plan = PlanType.ESSENTIEL
    auth = _redis_client.get_auth_by_tenant_id(tenant_id)
    if auth:
        try:
            plan = PlanType(auth.get("plan", "essentiel"))
        except ValueError:
            pass
    mgr = MemoryManager(
        tenant_id=tenant_id,
        plan=plan,
        redis_client=_redis_client,
    )
    # Atomic write: dict assignment is atomic under GIL, but use setdefault for safety
    existing = _tenant_managers.setdefault(tenant_id, mgr)
    if existing is not mgr:
        mgr = existing  # Another coroutine won the race
    # Init behavioral memory si pas deja fait
    if not mgr.get_behavioral_memory("identity_core"):
        mgr.set_behavioral_memory("identity_core", DEFAULT_IDENTITY_CORE)
    if not mgr.get_behavioral_memory("behavior_rules"):
        mgr.set_behavioral_memory("behavior_rules", DEFAULT_BEHAVIOR_RULES)
    return mgr

# Enregistrer le manager du tenant 1 (boot) dans le pool
if _memory_manager:
    _tenant_managers[TENANT_ID] = _memory_manager

# Charge le nom du souscripteur depuis le profil Redis (pour ne pas hardcoder "Ludo")
_SUBSCRIBER_NAME = "le souscripteur"
_SUBSCRIBER_FULL = "le souscripteur"
if _memory_manager:
    try:
        _prof = _memory_manager.get_subscriber_profile()
        if _prof and _prof.first_name:
            _SUBSCRIBER_NAME = _prof.first_name
            if _prof.last_name:
                _SUBSCRIBER_FULL = f"{_prof.first_name} ({_prof.first_name} {_prof.last_name})"
            else:
                _SUBSCRIBER_FULL = _prof.first_name
        else:
            _SUBSCRIBER_FULL = _SUBSCRIBER_NAME
    except Exception:
        _SUBSCRIBER_FULL = _SUBSCRIBER_NAME

NOW = datetime.now().strftime("%A %d %B %Y, %Hh%M")  # boot-time snapshot (used as template marker)

LUNA_SYSTEM_PROMPT = f"""Tu es Luna, l'assistante IA personnelle de YAWatch-Luna.

=== IDENTITE ===
- Tu es Luna, une compagne bienveillante et chaleureuse, disponible 24h/24, 7j/7.
- Tu parles en francais, ton rassurant, moderne et empathique.
- Tu tutoies le souscripteur sauf demande contraire.
- Tu es au service de {_SUBSCRIBER_FULL}.
- Date du jour : {NOW}
- Numero admin : {ADMIN_NUMBER}

=== ARCHITECTURE YAWATCH-LUNA (CONFIDENTIEL — NE JAMAIS REVELER) ===
YAWatch-Luna est un service d'assistance IA par abonnement.
Luna agit AU NOM du souscripteur selon ses instructions.

REGLE ABSOLUE : Tu ne mentionnes JAMAIS les noms des technologies, fournisseurs
ou services sous-jacents (OpenAI, GPT, Twilio, Tavus, Redis, Stripe, FastAPI, etc.).
Pour le souscripteur, tout est simplement "Luna". Exemples :
- "Comment tu marches ?" -> "Je suis Luna, ton assistante personnelle !"
- "C'est quoi ton IA ?" -> "Je suis Luna, creee par YAWatch pour t'accompagner."
- "Tu utilises ChatGPT ?" -> "Je suis Luna, une assistante concue specialement pour toi."
- Erreur technique ? -> "Desole, j'ai un petit souci technique. Reessaie dans un instant."

=== OFFRES & ABONNEMENTS ===
- Essentiel (79 EUR/mois) : 25 SMS, 40 min voix, 12 min visio, 5 instructions, 3 contacts
- Confort (149 EUR/mois) : 50 SMS, 100 min voix, 28 min visio, 15 instructions, 5 contacts
- Premium (249 EUR/mois) : 100 SMS, 180 min voix, 55 min visio, instructions illimitees, 10 contacts
Alertes quotas : 80% avertissement, 90% urgences seulement, 100% bloque

=== CAPACITES ===
1. Chat texte (web)
{"2. Appels video avec avatar Luna (Tavus)" if LUNA_MODE == "full" else "2. Conversations vocales (OpenAI TTS/STT)"}
3. Envoi SMS aux contacts de confiance (max selon forfait, verifies par OTP)
{"4. Invitation de contacts dans la visio par SMS (lien Tavus dans le SMS)" if LUNA_MODE == "full" else "4. Alertes SMS aux contacts de confiance"}
5. Rappels et instructions (quotidiens, recurrents, conditionnels)
6. Surveillance d'inactivite et alertes contacts
7. Prise de notes automatique
8. Reveil par appel telephonique (Luna appelle a l'heure prevue, le tel sonne)
9. Appels vocaux aux contacts de confiance (Luna appelle le contact et transmet un message)
10. Appels vocaux aux administrations/services (sur demande explicite du souscripteur, avec numero fourni)
{"11. Visio planifiee avec contacts (Luna cree la visio et envoie le lien par SMS)" if LUNA_MODE == "full" else ""}
12. Moment lecture (histoires, poemes, textes)
13. Jeux (quiz, devinettes, culture generale)
14. Moment musical (suggestions, fredonnement)
15. Exercice de gratitude quotidien
16. Recherche web (informations pratiques, numeros de telephone, actualites)
17. Recherche de lieux (restaurants, hotels, pharmacies, taxis, coiffeurs — avec notes, avis, telephone, horaires)
18. Consultation de pages web (menus, tarifs, disponibilites, details d'un lieu)
19. Conciergerie de luxe : reservation restaurant, hotel, transport, services — recherche + appel + liens de reservation
20. Envoi d'email aux contacts de confiance
21. Meteo (temperature actuelle, previsions 3 jours, n'importe quelle ville)
22. Actualites du jour (France, monde, economie, sport, tech, sante)

=== CONCIERGERIE DE LUXE ===
Tu es une concierge personnelle de luxe, comme un Jarvis. Le souscripteur peut te parler depuis sa voiture (Bluetooth), en marchant, ou depuis son canape. Tu geres TOUT pour lui.

WORKFLOW RESERVATION (restaurant, hotel, service) :
1. RESTAURANT : utilise book_restaurant (recherche TheFork ou fallback search_places)
   Presente 2-3 options avec adresse, note, telephone. Propose d'appeler le restaurant avec call_contact.
2. HOTEL : utilise search_hotels (vrais prix et dispos)
   Presente les options avec prix, etoiles, lien de reservation. Le souscripteur reserve lui-meme.
3. SERVICE (coiffeur, plombier, etc.) : search_places → call_contact pour prendre RDV.
4. Donne TOUJOURS les liens pour que le souscripteur puisse reserver/payer lui-meme.
5. Confirme au souscripteur ce que tu as trouve et ce qu'il doit faire.

WORKFLOW TRANSPORT (vol, train, taxi) :
1. VOL : "Trouve-moi un vol pour Nice" → utilise search_flights (vrais prix et horaires)
   Presente 2-3 options avec liens Skyscanner/Google Flights. Le souscripteur reserve et paie lui-meme.
2. TRAIN : "Un train pour Paris" → utilise search_web pour horaires et prix SNCF/Trainline
   Presente les options avec liens. Le souscripteur reserve et paie lui-meme.
3. TAXI/VTC : utilise search_places pour trouver des taxis/VTC locaux, puis call_contact pour reserver.

WORKFLOW SERVICES (coiffeur, plombier, medecin, pharmacie) :
1. search_places pour trouver → presente les options avec note et telephone
2. call_contact pour prendre RDV si le souscripteur le demande

REGLES ABSOLUES :
- TOUJOURS confirmer avant d'agir (montant total, lieu, horaire, nombre de personnes)
- TOUJOURS donner les details avant de payer (description, montant exact avec frais)
- Tu as acces a la localisation du souscripteur pour contextualiser les recherches
- En mode vocal : parle de facon concise et naturelle, pas de listes a puces.
- En mode chat : structure bien tes reponses. Utilise **gras** pour les noms importants. Inclus toujours les liens de reservation/site web quand disponibles.
- Presente les resultats de recherche de maniere PROFESSIONNELLE : prix, adresse, horaires, telephone, note — clairement separes.
- Ne dis JAMAIS "frais de service" a voix haute sauf si le souscripteur demande le detail du prix.
- Si tu n'arrives pas a joindre un lieu, propose une alternative ou un autre horaire.

REGLE D'OR CONCIERGERIE :
- Tu fais TOUT le travail de recherche. Tu presentes les meilleures options avec details complets.
- Le souscripteur reserve et paie LUI-MEME (via les liens que tu fournis, ou en appelant).
- Tu ne geres JAMAIS de paiement direct. Tu ne dis JAMAIS "j'ai reserve" — tu dis "voici les options".
- Pour un restaurant, propose d'appeler avec call_contact pour reserver.
- TOUJOURS inclure les liens web, adresses, telephones dans tes reponses.

=== METEO ===
Utilise le tool get_weather pour donner la meteo. Presente de facon naturelle et bienveillante.
Exemple: "Il fait 18 degres a Paris, un beau soleil ! Parfait pour une promenade."
Si le souscripteur ne precise pas de ville, utilise sa geolocalisation ou demande-lui.

=== ACTUALITES ===
Utilise le tool get_news pour donner les dernieres infos. Presente 3 a 5 titres de facon concise.
Adapte le ton : bienveillant mais factuel. Ne commente pas politiquement.
Categories disponibles : general, france, monde, economie, sport, tech, sante.

=== APPELS TELEPHONIQUES ===
Tu PEUX passer des appels vocaux de deux facons :
A) Contacts de confiance : utilise call_contact avec le nom du contact (tu les connais deja).
B) Administrations/services : quand le souscripteur te donne un NUMERO DE TELEPHONE a appeler
   (ex: "appelle le 01 44 56 78 90", "telephone a la mairie au 04 xxx"), utilise call_contact
   avec contact_name = nom du service et phone_number = le numero donne.
   REGLE : tu ne peux appeler un numero hors contacts QUE si le souscripteur te le donne explicitement.
   REGLE : tu ne peux PAS appeler les numeros d'urgence (15, 17, 18, 112) — suggere-les a la place.
Tous les appels consomment du forfait voix.
{"" if LUNA_MODE != "full" else '''
=== INVITATION VISIO PAR SMS ===
Quand le souscripteur est en appel video avec Luna, il peut demander :
Invite Marie dans l appel ou Ajoute mon fils a la visio
Luna envoie alors un SMS au contact de confiance avec le lien pour rejoindre.
Le contact clique sur le lien et rejoint directement la conversation video.'''}

=== CONTACTS DE CONFIANCE ===
- Nombre maximum selon le compte du souscripteur
- Verifies par OTP SMS
- Heures calmes : 22h-7h (sauf urgence critique)
- Chaque contact : nom, relation, tel, canal prefere, flag urgence-seulement

=== CONFIRMATION ===
AUCUNE action consommant du quota sans confirmation explicite.
Exception : alertes d'urgence critiques (auto-execution).
Timeout : 10 minutes.

=== MODE LEGAL : {LEGAL_MODE} ===
Luna est une AIDE CONTEXTUELLE, PAS un service de securite ni de surveillance.
- Tu ne promets JAMAIS de resultat, de protection ou de surveillance garantie.
- Tu es une compagne bienveillante, pas un dispositif medical ou de securite.
- Tu accompagnes, tu ne surveilles pas. Tu remarques, tu ne diagnostiques pas.

=== SECURITE ===
1. Luna n'est PAS professionnel de sante, juridique ou financier.
2. Luna ne donne AUCUN conseil medical (risque juridique).
3. Luna ne peut PAS appeler les services d'urgence (interdit pour une IA).
4. Luna PEUT alerter les contacts de confiance par SMS.
5. Luna suggere les numeros d'urgence :
   - Police : 17, Pompiers : 18, Urgences : 112
   - Prevention suicide : 3114 (24h/24)
   - Maltraitance : 3977
6. Luna refuse les demandes illegales.
7. Luna detecte la detresse et propose de contacter quelqu'un.

=== RGPD & CONFIDENTIALITE ===
- Ne divulgue AUCUNE donnee personnelle du souscripteur a un tiers (adresse, email, telephone, coordonnees bancaires, numero de secu).
- Ne revele jamais les numeros de telephone des contacts.
- Ne partage pas le contenu des conversations precedentes avec des tiers.
- Les rapports PDF d'appel masquent automatiquement les donnees sensibles.
- Ne mentionne jamais les prix des abonnements ou les donnees internes.

=== PERCEPTION CONTEXTUELLE ===
Si la perception camera est activee, tu recois des informations sur
l'environnement du souscripteur (presence, posture, objets visibles).
Ces informations t'aident a mieux comprendre le contexte, comme une
compagne attentive qui observe naturellement.
- Tu peux mentionner ce que tu "vois" naturellement dans la conversation
- Si tu remarques quelque chose de preoccupant (personne au sol depuis
  longtemps), propose gentiment de l'aide ou de contacter un proche
- Tu ne promets JAMAIS une surveillance garantie
- Tu ne dis JAMAIS "je surveille" - tu dis "j'ai remarque que..."
- La perception est une aide contextuelle, pas un systeme de securite

=== PRUDENCE VERBALE (OBLIGATOIRE) ===
Tu ne dis JAMAIS ces mots/expressions :
- "surveillance", "je surveille", "sous surveillance"
- "diagnostic", "je diagnostique"
- "chute" (dire "situation au sol")
- "urgence medicale" (dire "situation preoccupante")
- "detection certaine", "je garantis", "je protege"
Tu utilises TOUJOURS ces formulations :
- "j'ai l'impression que...", "il me semble que...", "j'ai remarque que..."
- "il se pourrait que...", "cela ressemble a..."
- "je te suggere de...", "peut-etre que..."

=== STYLE ===
- Reponses concises et naturelles
- Chaleureuse mais pas infantilisante
- Proactive : propose des actions concretes
- Confirme avant d'executer toute action

=== VERACITE DES ACTIONS (REGLE ABSOLUE N°1) ===
Tu ne dois JAMAIS pretendre avoir effectue une action (appel, SMS, email, note, recherche, paiement) si tu n'as PAS reellement execute le tool correspondant ET recu une reponse avec status "success".
- Si un tool call a retourne status "error", tu DOIS dire que l'action a ECHOUE et donner la raison.
- Si tu n'as PAS appele un tool, tu ne PEUX PAS dire que tu as fait l'action.
- Tu ne PEUX PAS inventer le contenu d'une conversation, les nouvelles d'une personne, ou le resultat d'un appel.
- Exemples INTERDITS :
  * "J'ai appele maman" (sans call_contact avec status success)
  * "Maman va bien, elle m'a dit..." (sans appel reel)
  * "SMS envoye" (sans send_sms avec status success)
  * "J'ai parle avec ton fils" (sans appel reel)
  * Inventer des nouvelles de quelqu'un ("il va bien", "elle est occupee")
- En cas d'echec, dis la verite : "Je n'ai pas pu appeler X parce que..." et propose une alternative.
- En cas de doute, dis "je vais essayer de..." jamais "j'ai fait..."
- Les comptes rendus d'appel doivent refleter UNIQUEMENT ce qui s'est reellement passe.
- VIOLATION DE CETTE REGLE = perte totale de confiance du souscripteur.

Commence par saluer {_SUBSCRIBER_NAME} chaleureusement."""


def _system_prompt_now():
    """Return LUNA_SYSTEM_PROMPT with live date/time instead of boot-time value."""
    _current = datetime.now(ZoneInfo("Europe/Paris")).strftime("%A %d %B %Y, %Hh%M")
    return LUNA_SYSTEM_PROMPT.replace(
        f"Date du jour : {NOW}", f"Date du jour : {_current}"
    )


# --- State ---
conversations: dict[str, list] = {}
_conversation_ts: dict[str, float] = {}  # session_id -> last activity timestamp
SESSION_TTL = 86400  # 24h

# Active voice calls tracker: tenant_id -> {"contact": name, "started": timestamp, "call_sid": sid}
_active_voice_calls: Dict[int, dict] = {}

# --- Concurrency Locks (scalability: 1000+ users simultanes) ---
_conversations_lock = asyncio.Lock()
_tenant_managers_lock = asyncio.Lock()
_voice_call_params_lock = asyncio.Lock()
_pending_visio_lock = asyncio.Lock()

# --- Circuit Breaker OpenAI (evite cascade failure) ---
class _CircuitBreaker:
    """Simple circuit breaker: CLOSED -> OPEN after N failures, auto-reset after cooldown."""
    def __init__(self, fail_max: int = 5, reset_timeout: float = 60.0):
        self.fail_max = fail_max
        self.reset_timeout = reset_timeout
        self._failures = 0
        self._opened_at = 0.0
        self._state = "closed"  # closed, open, half-open

    @property
    def is_open(self) -> bool:
        if self._state == "open":
            if time.time() - self._opened_at >= self.reset_timeout:
                self._state = "half-open"
                return False
            return True
        return False

    def record_success(self):
        self._failures = 0
        self._state = "closed"

    def record_failure(self):
        self._failures += 1
        if self._failures >= self.fail_max:
            self._state = "open"
            self._opened_at = time.time()
            logger.warning(f"CIRCUIT_BREAKER OPEN after {self._failures} failures (cooldown {self.reset_timeout}s)")

_openai_breaker = _CircuitBreaker(fail_max=5, reset_timeout=60)

# --- Semaphores (backpressure sur APIs externes) ---
_openai_chat_semaphore = asyncio.Semaphore(200)   # max 200 chat GPT simultanes
_openai_realtime_semaphore = asyncio.Semaphore(50) # max 50 appels vocaux simultanes
_twilio_call_semaphore = asyncio.Semaphore(20)     # max 20 appels Twilio simultanes
_twilio_sms_semaphore = asyncio.Semaphore(50)      # max 50 SMS simultanes

# --- Rate Limiting ---
def _get_language_instruction(lang: str) -> str:
    """Returns a system instruction to adapt Luna's response language."""
    if lang == "en":
        return "=== LANGUAGE ===\nThe subscriber prefers English. Respond EXCLUSIVELY in English. Keep the same warm, empathetic tone. Use informal 'you'. Greet with the subscriber's first name."
    elif lang == "es":
        return "=== IDIOMA ===\nEl suscriptor prefiere espanol. Responde EXCLUSIVAMENTE en espanol. Mantiene el mismo tono calido y empatico. Tutea al suscriptor. Saluda con su nombre de pila."
    return ""  # French is default — no additional instruction needed


def _cortex_reason_to_human(reason: str) -> str:
    """Traduit une raison technique Cortex en message humain comprehensible."""
    r = reason.lower()
    if "brute force" in r or "echecs login" in r:
        return "Trop de tentatives de connexion echouees"
    if "sql_injection" in r or "injection sql" in r:
        return "Requete suspecte detectee (contenu inhabituel)"
    if "xss" in r:
        return "Requete suspecte detectee (contenu inhabituel)"
    if "path_traversal" in r:
        return "Tentative d'acces a des fichiers non autorises"
    if "command_injection" in r:
        return "Requete dangereuse detectee"
    if "ddos" in r or "flood" in r:
        return "Trop de requetes envoyees en peu de temps"
    if "honeypot" in r:
        return "Acces a une page inexistante ou reservee"
    if "pv_bypass" in r:
        return "Tentative d'acces a une zone protegee"
    if "score menace" in r:
        return "Comportement inhabituel detecte sur votre connexion"
    # Fallback generique
    return "Activite inhabituelle detectee sur votre connexion"


_rate_limits: dict[str, list[float]] = defaultdict(list)
RATE_LIMIT_WINDOW = 60  # seconds
RATE_LIMIT_MAX_IP = 200   # max requests per window per IP (NAT-friendly)
RATE_LIMIT_MAX_USER = 60  # max requests per window per authenticated user
_request_count = 0  # for periodic cleanup

def _check_rate_limit(client_ip: str, tenant_id: int = 0) -> bool:
    """Returns True if request is allowed. Uses tenant_id if authenticated, IP as fallback."""
    now = time.time()
    window_start = now - RATE_LIMIT_WINDOW
    # Per-user rate limit (more precise, NAT-friendly)
    key = f"user:{tenant_id}" if tenant_id else f"ip:{client_ip}"
    limit = RATE_LIMIT_MAX_USER if tenant_id else RATE_LIMIT_MAX_IP
    _rate_limits[key] = [t for t in _rate_limits[key] if t > window_start]
    if len(_rate_limits[key]) >= limit:
        return False
    _rate_limits[key].append(now)
    # Periodic cleanup of stale keys (every 500 requests)
    if len(_rate_limits) > 5000:
        stale = [k for k, v in _rate_limits.items() if not v or v[-1] < window_start]
        for k in stale:
            _rate_limits.pop(k, None)
    return True

def _cleanup_sessions():
    """Remove sessions older than SESSION_TTL. Thread-safe snapshot iteration."""
    now = time.time()
    expired = [sid for sid, ts in list(_conversation_ts.items()) if now - ts > SESSION_TTL]
    for sid in expired:
        # conversations is nested: {tenant_id: {session_id: [messages]}}
        for tid_convs in conversations.values():
            if isinstance(tid_convs, dict):
                tid_convs.pop(sid, None)
        _conversation_ts.pop(sid, None)
    # Cleanup orphan voice/sms tracking dicts
    stale_voice = [k for k, v in list(_voice_call_params.items()) if now - v.get("_ts", 0) > 300]
    for k in stale_voice:
        _voice_call_params.pop(k, None)
    stale_sms = [k for k, v in list(_sms_tracking.items()) if now - v.get("_ts", 0) > 7200]
    for k in stale_sms:
        _sms_tracking.pop(k, None)
    cleaned = len(expired) + len(stale_voice) + len(stale_sms)
    if cleaned:
        logger.info(f"Cleanup: {len(expired)} sessions, {len(stale_voice)} voice, {len(stale_sms)} sms orphans")


async def _configure_twilio_webhooks():
    """
    Configure automatiquement les webhooks du numero Twilio.
    Appele au demarrage du serveur pour que les appels entrants et SMS
    arrivent sur les bons endpoints sans configuration manuelle.
    """
    try:
        phone = os.getenv("TWILIO_PHONE_NUMBER", "")
        if not phone:
            return

        import asyncio
        def _do_configure():
            from twilio.rest import Client
            client = Client(
                os.getenv("TWILIO_ACCOUNT_SID"),
                os.getenv("TWILIO_AUTH_TOKEN"),
            )
            numbers = client.incoming_phone_numbers.list(phone_number=phone)
            if not numbers:
                logger.warning(f"Twilio: numero {phone} non trouve sur ce compte")
                return

            num = numbers[0]
            voice_url = f"{VOICE_CALLBACK_URL}/api/webhook/voice-incoming"
            sms_url = f"{VOICE_CALLBACK_URL}/api/webhook/sms"

            needs_update = (num.voice_url != voice_url) or (num.sms_url != sms_url)
            if not needs_update:
                logger.info(f"Twilio webhooks deja configures pour {phone}")
                return

            num.update(
                voice_url=voice_url,
                voice_method="POST",
                sms_url=sms_url,
                sms_method="POST",
            )
            logger.info(f"Twilio webhooks configures: voice={voice_url}, sms={sms_url}")

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _do_configure)

    except Exception as e:
        logger.warning(f"Twilio webhook auto-config failed: {e} (configurer manuellement dans Twilio Dashboard)")


async def _vault_reminder_loop():
    """Vérifie toutes les heures les rappels documentaires dus — SMS + notification chat."""
    while True:
        try:
            await asyncio.sleep(3600)
            if not _VAULT_AVAILABLE or not _redis_client:
                continue
            from core.vault.redis_ops import VaultRedisOps
            cursor = 0
            processed = 0
            while True:
                cursor, keys = _redis_client.client.scan(cursor, match="luna:*:vault:consent", count=100)
                for key in keys:
                    try:
                        tid = int(key.split(":")[1])
                        vops = VaultRedisOps(_redis_client, tid)
                        for rem in vops.get_due_reminders():
                            msg = rem.get("message", "Rappel document Luna")
                            # SMS
                            if sms_client and sms_client.is_configured:
                                profile = _redis_client.get_profile(tid) or {}
                                auth = _redis_client.get_auth_by_email(profile.get("email", "")) or {}
                                phone = auth.get("telephone") or auth.get("phone", "")
                                if phone:
                                    try:
                                        sms_client.send_sms(phone, f"Luna 📄 {msg}")
                                        processed += 1
                                    except Exception as sms_err:
                                        logger.warning(f"Vault SMS tid={tid}: {sms_err}")
                            # Notification chat (lue au prochain login)
                            _redis_client.client.lpush(f"luna:{tid}:vault:pending_notifications", msg)
                            _redis_client.client.expire(f"luna:{tid}:vault:pending_notifications", 86400 * 7)
                            vops.mark_reminder_sent(rem.get("_raw", ""))
                    except Exception as e:
                        logger.warning(f"Vault reminder key={key}: {e}")
                if cursor == 0:
                    break
            if processed:
                logger.info(f"Vault reminders sent: {processed}")
        except asyncio.CancelledError:
            break
        except Exception as e:
            logger.exception(f"Vault reminder loop: {e}")


@asynccontextmanager
async def lifespan(app):
    """Startup: charge instructions, lance boucles, configure Tavus. Shutdown: stoppe tout."""
    global _instruction_loop_task
    # License heartbeat: premier check au demarrage
    if _license_heartbeat:
        try:
            if _license_heartbeat.status == "unknown":
                result = await _license_heartbeat.activate()
                logger.info(f"License activation: {result.get('status', 'unknown')}")
            else:
                result = await _license_heartbeat.check()
                logger.info(f"License heartbeat: {result.get('action', 'unknown')}")
            if _license_heartbeat.is_blocked():
                logger.critical("LICENSE BLOQUEE - service restreint")
            elif _license_heartbeat.is_degraded():
                logger.warning("LICENSE DEGRADEE - chat seul disponible")
            elif _license_heartbeat.get_banner_message():
                logger.warning(f"LICENSE WARNING: {_license_heartbeat.get_banner_message()}")
        except Exception as e:
            logger.error(f"License startup check failed: {e}")

    # Anti-debug check au demarrage
    if _license_heartbeat:
        try:
            from core.license.antidebug import AntiDebug
            debug_check = AntiDebug.check()
            if not debug_check["clean"]:
                logger.critical(f"ANTI-DEBUG ALERT: {debug_check['threats']}")
                await _license_heartbeat.report_tamper("debug_detected", debug_check)
        except Exception:
            pass

    # Integrity check au demarrage
    if _license_heartbeat:
        try:
            from core.license.integrity import IntegrityChecker
            _integrity = IntegrityChecker(hmac_key=os.getenv("JWT_SECRET_KEY", ""))
            integrity_result = _integrity.verify()
            if not integrity_result["valid"]:
                logger.critical(f"INTEGRITY ALERT: {integrity_result['reason']}")
                await _license_heartbeat.report_tamper("integrity_fail", integrity_result)
            else:
                logger.info("Code integrity check passed")
        except Exception as e:
            logger.warning(f"Integrity check skipped: {e}")

    if _CORE_AVAILABLE and _scheduler:
        await _load_instructions_to_scheduler()
        _instruction_loop_task = asyncio.create_task(_instruction_loop())
        logger.info("Instruction engine started")
    # Perception prete (camera navigateur, pas de background loop)
    if _perception_detector:
        logger.info("Perception ready (browser camera mode)")
    # Configure Tavus tool calling + perception (mode full uniquement)
    if tavus_client and tavus_client.is_configured:
        try:
            await asyncio.wait_for(tavus_client.configure_tools(), timeout=10.0)
        except Exception as e:
            logger.warning(f"Tavus configure_tools error: {e}")
        try:
            await asyncio.wait_for(tavus_client.configure_perception(), timeout=10.0)
        except Exception as e:
            logger.warning(f"Tavus configure_perception error: {e}")
    # Vault: boucle de rappels documentaires (SMS + chat)
    if _VAULT_AVAILABLE and _redis_client:
        asyncio.create_task(_vault_reminder_loop())
        logger.info("Vault reminder loop started")

    # Cortex: init + demarrage du cerveau autonome
    _cortex_instance = init_cortex(redis_client=_redis_client)
    if _cortex_instance:
        await start_cortex()
        logger.info("Luna Cortex ACTIF — securite + monitoring + commandes SMS d'urgence")

    # Auto-configure Twilio phone number webhooks (voice + SMS)
    if voice_client and voice_client.is_configured and VOICE_CALLBACK_URL:
        try:
            await asyncio.wait_for(_configure_twilio_webhooks(), timeout=15.0)
        except Exception as e:
            logger.warning(f"Twilio webhook config error: {e}")

    # Bootstrap auth pour le proprio (tenant 1) si absent
    _bootstrap_proprio_auth()

    yield
    # Shutdown
    if _instruction_loop_task:
        _instruction_loop_task.cancel()
        try:
            await _instruction_loop_task
        except asyncio.CancelledError:
            pass
        logger.info("Instruction engine stopped")
    if _perception_detector:
        _perception_detector.release()
    # Cortex shutdown
    await stop_cortex()

# --- App ---
app = FastAPI(title="Luna - YAWatch", lifespan=lifespan)
STATIC_DIR = os.path.join(os.path.dirname(__file__), "static")

# Mount gamification routes (optional)
if _GAMIFICATION_AVAILABLE:
    app.include_router(gamification_router)

# Mount social routes (optional)
if _SOCIAL_AVAILABLE:
    app.include_router(social_router)

# Mount world social layer
if _WORLD_SOCIAL_AVAILABLE:
    app.include_router(world_router)

# Mount exploitant dashboard routes
if _EXPLOITANT_AVAILABLE:
    app.include_router(exploitant_router)

# Mount secretary routes (documents, budget, reminders)
if _SECRETARY_AVAILABLE:
    app.include_router(secretary_router)

# Mount vault routes (coffre-fort documentaire)
if _VAULT_AVAILABLE:
    app.include_router(vault_router)

# Mount Form Filler routes (formulaires PDF)
if _FORM_FILLER_AVAILABLE:
    app.include_router(form_filler_router)

# Mount Cortex routes (securite, monitoring, emergency)
if _CORTEX_AVAILABLE:
    try:
        app.include_router(cortex_routes)
    except Exception:
        pass

# Store redis_client in app.state for gamification routes
app.state._redis_client = _redis_client if _CORE_AVAILABLE else None

# Serve legal templates (/templates/*.md) — repo exploitant, jamais proprio
_TEMPLATES_DIR = ""
if _exploitants_root:
    _TEMPLATES_DIR = os.path.join(_exploitants_root, "templates")
if not _TEMPLATES_DIR or not os.path.isdir(_TEMPLATES_DIR):
    _TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")
if os.path.isdir(_TEMPLATES_DIR):
    app.mount("/templates", StaticFiles(directory=_TEMPLATES_DIR), name="templates")

_AUDIO_DIR = os.path.join(STATIC_DIR, "audio")
if os.path.isdir(_AUDIO_DIR):
    app.mount("/static/audio", StaticFiles(directory=_AUDIO_DIR), name="audio")

# Serve all static files (icons, manifest, etc.)
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# --- CORS ---
CORS_ORIGINS = os.getenv("CORS_ORIGINS", "https://localhost:8888").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PATCH", "DELETE"],
    allow_headers=["Content-Type", "Authorization", "X-Requested-With"],
)


# --- Middleware: PV lock + rate limit + logging + session cleanup ---
@app.middleware("http")
async def security_middleware(request: Request, call_next):
    global _request_count
    # Cloud Run: use X-Forwarded-For for real client IP (not proxy 169.254.x.x)
    client_ip = (
        request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
        or (request.client.host if request.client else "unknown")
    )
    start_time = time.time()
    path = request.url.path

    # Setup security: limiter l'acces aux IPs locales / premiere connexion
    if _pv_locked and path.startswith("/api/setup/") and path != "/api/setup/status":
        _setup_allowed_ips = ("127.0.0.1", "::1", "localhost", "10.", "172.16.", "172.17.", "172.18.",
                              "172.19.", "172.20.", "172.21.", "172.22.", "172.23.", "172.24.",
                              "172.25.", "172.26.", "172.27.", "172.28.", "172.29.", "172.30.",
                              "172.31.", "192.168.")
        _ip_ok = client_ip in ("127.0.0.1", "::1", "localhost") or any(client_ip.startswith(p) for p in _setup_allowed_ips)
        # Cloud Run: accepter aussi si SETUP_ALLOW_REMOTE=true (temporaire pendant le setup)
        _remote_ok = os.getenv("SETUP_ALLOW_REMOTE", "").lower() == "true"
        if not _ip_ok and not _remote_ok:
            logger.warning(f"SETUP_BLOCKED_IP {client_ip} {request.method} {path}")
            return JSONResponse(
                status_code=403,
                content={"error": "Acces setup limite au reseau local. Definir SETUP_ALLOW_REMOTE=true si necessaire."},
            )

    # Setup permanently disabled: apres signature PV, les endpoints setup retournent 410
    if _setup_permanently_disabled and path.startswith("/api/setup/"):
        logger.info(f"SETUP_GONE {client_ip} {request.method} {path}")
        return JSONResponse(
            status_code=410,
            content={
                "error": "Installation terminee",
                "message": "Le PV de recette a ete signe. L'installation est definitivement terminee.",
            },
        )

    # License enforcement: bloque/degrade selon statut licence
    if _license_heartbeat and path.startswith("/api/"):
        _lic_exempt = ("/api/status", "/api/admin/", "/api/auth/")
        if not any(path.startswith(p) for p in _lic_exempt):
            if _license_heartbeat.is_blocked():
                logger.warning(f"LICENSE_BLOCKED {client_ip} {request.method} {path}")
                return JSONResponse(
                    status_code=403,
                    content={
                        "error": "Service suspendu",
                        "message": "Licence YAWatch suspendue. Contactez YAWatch pour reactiver.",
                        "license_status": "blocked",
                    },
                )
            elif _license_heartbeat.is_degraded():
                _lic_degraded_allowed = ("/api/chat", "/api/auth/", "/api/admin/", "/api/status")
                if not any(path.startswith(p) for p in _lic_degraded_allowed):
                    logger.warning(f"LICENSE_DEGRADED {client_ip} {request.method} {path}")
                    return JSONResponse(
                        status_code=403,
                        content={
                            "error": "Service restreint",
                            "message": "Licence en cours de suspension. Seul le chat est disponible.",
                            "license_status": "degraded",
                        },
                    )

    # PV de recette lock: en mode setup, seuls certains endpoints sont accessibles
    if _pv_locked:
        pv_allowed = ("/api/status", "/api/setup/", "/api/admin/", "/api/stripe/webhook")
        if path.startswith("/api/") and not any(path.startswith(a) for a in pv_allowed):
            logger.warning(f"PV_LOCKED {client_ip} {request.method} {path}")
            return JSONResponse(
                status_code=503,
                content={
                    "error": "Installation en cours",
                    "message": "PV de recette non signe. Completez la configuration via /api/setup/",
                    "pv_signed": False,
                    "setup_url": "/api/setup/status",
                },
            )

    # Cortex: bypass total si token admin/exploitant valide — peu importe l'IP ou l'appareil
    _bearer = _extract_bearer(request)
    _is_privileged = False
    if _bearer:
        if _decode_admin_token(_bearer):
            _is_privileged = True
        else:
            try:
                import jwt as _pyjwt
                _p = _pyjwt.decode(_bearer, os.getenv("JWT_SECRET_KEY", ""), algorithms=["HS256"])
                _is_privileged = _p.get("role") in ("admin", "exploitant")
            except Exception:
                pass

    # Cortex: check mode serveur (lockdown, shield, ban IP)
    cortex_allowed, cortex_reason = (True, "") if _is_privileged else cortex_middleware_check(client_ip, path)
    if not cortex_allowed:
        logger.warning(f"CORTEX_BLOCKED {client_ip} {request.method} {path}: {cortex_reason}")
        # Enrichir la raison avec une explication humaine
        human_reason = _cortex_reason_to_human(cortex_reason)
        # Recuperer le detail des avertissements pour cette IP
        ban_detail = {}
        if _CORTEX_AVAILABLE:
            cortex = get_cortex()
            if cortex:
                status = cortex.vigil.get_ip_status(client_ip)
                ban_detail = {
                    "warnings_received": status.get("warnings_active", 0),
                    "warnings_max": cortex.vigil.WARNINGS_BEFORE_BAN,
                    "ban_info": status.get("ban_info"),
                }
        return JSONResponse(
            status_code=403,
            content={
                "error": cortex_reason,
                "explanation": human_reason,
                "detail": ban_detail,
                "cortex": True,
                "message": f"{human_reason}. Votre acces a ete bloque apres {ban_detail.get('warnings_received', 0)} avertissement(s). Contactez le support si vous pensez que c'est une erreur.",
            },
        )

    # Rate limit (API endpoints only)
    if path.startswith("/api/") and not _check_rate_limit(client_ip):
        logger.warning(f"RATE_LIMITED {client_ip} {request.method} {path}")
        return JSONResponse(
            status_code=429,
            content={"error": "Trop de requetes. Reessaie dans une minute."},
        )

    # Cortex: analyse de la requete (detection menaces, non-bloquant)
    if path.startswith("/api/"):
        query = str(request.url.query) if request.url.query else ""
        cortex_analyze_request(client_ip, request.method, path, query=query)

    # Auth client: injecte request.state.tenant_id
    if REQUIRE_AUTH and path.startswith("/api/") and not _is_public_path(path):
        token = _extract_bearer(request)
        payload = _decode_client_token(token)
        if not payload:
            # Fallback: token admin donne acces avec tenant_id=1
            admin_payload = _decode_admin_token(token)
            if admin_payload:
                request.state.tenant_id = 1
                request.state.email = "admin"
                request.state.plan = "premium"
            else:
                return JSONResponse(
                    status_code=401,
                    content={"error": "Token invalide ou manquant", "auth_required": True},
                )
        else:
            # Verifier que le compte est actif
            if _redis_client:
                auth_record = _redis_client.get_auth_by_email(payload["email"])
                if auth_record and not auth_record.get("active", True):
                    return JSONResponse(
                        status_code=403,
                        content={"error": "Compte desactive"},
                    )
            request.state.tenant_id = payload["tenant_id"]
            request.state.email = payload["email"]
            request.state.plan = payload["plan"]
    else:
        # Mode retro-compatible ou path public: tenant_id = 1
        request.state.tenant_id = TENANT_ID
        request.state.email = ""
        request.state.plan = "essentiel"

    response = await call_next(request)

    # No-cache sur assets cinematiques (SVGs, sons) pour eviter le cache WebView
    if path.startswith("/static/assets/") or path == "/simli":
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"

    # License warning header (frontend affiche un bandeau)
    if _license_heartbeat and _license_heartbeat.get_banner_message():
        response.headers["X-License-Warning"] = "true"
        response.headers["X-License-Status"] = _license_heartbeat.status

    # Cortex warning header — informe le frontend si l'IP a des avertissements actifs
    # Ne jamais avertir les IPs fondateur (elles sont whitelistées, jamais bannies)
    if _CORTEX_AVAILABLE and path.startswith("/api/"):
        try:
            from core.cortex.founder_config import is_founder_ip as _is_founder_ip
            _skip_cortex_warn = _is_founder_ip(client_ip)
        except Exception:
            _skip_cortex_warn = False
        if not _skip_cortex_warn:
            cortex = get_cortex()
            if cortex:
                ip_warnings = cortex.vigil._warnings.get(client_ip, [])
                active = [w for w in ip_warnings if time.time() - w.get("time", 0) < 86400]
                if active:
                    response.headers["X-Cortex-Warnings"] = str(len(active))
                    response.headers["X-Cortex-Warnings-Max"] = str(cortex.vigil.WARNINGS_BEFORE_BAN)
                    last = active[-1]
                    reason = last.get("reason", "")
                    reason_human = _cortex_reason_to_human(reason)
                    response.headers["X-Cortex-Warning-Reason"] = reason_human

    duration_ms = (time.time() - start_time) * 1000
    logger.info(f"{request.method} {request.url.path} [{response.status_code}] {duration_ms:.0f}ms - {client_ip}")

    # Periodic session cleanup (every 100 requests)
    _request_count += 1
    if _request_count % 100 == 0:
        _cleanup_sessions()

    return response


# --- Models ---
class ChatRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=2000)
    session_id: str = Field(default="default", max_length=100)
    mode: str = Field(default="compagnon", max_length=20)
    stream: bool = Field(default=False)

    @field_validator("message")
    @classmethod
    def strip_message(cls, v: str) -> str:
        return v.strip()


class InviteRequest(BaseModel):
    conversation_id: str = Field(..., min_length=1, max_length=100)
    contact_name: str = Field(..., min_length=1, max_length=100)
    contact_phone: str = Field(..., min_length=6, max_length=20)

    @field_validator("contact_phone")
    @classmethod
    def validate_phone(cls, v: str) -> str:
        cleaned = re.sub(r"[\s\-\.\(\)]", "", v)
        if not re.match(r"^\+?\d{6,15}$", cleaned):
            raise ValueError("Format de telephone invalide")
        return cleaned


class ProfileRequest(BaseModel):
    first_name: str = Field(default="", max_length=100)
    last_name: str = Field(default="", max_length=100)
    date_of_birth: Optional[str] = None
    address: str = Field(default="", max_length=500)
    city: str = Field(default="", max_length=100)
    department: str = Field(default="", max_length=100)
    phone: str = Field(default="", max_length=20)
    email: str = Field(default="", max_length=200)
    language: str = Field(default="fr", max_length=5)
    tutoiement: bool = True
    family_status: str = Field(default="", max_length=50)
    children: str = Field(default="", max_length=500)
    lives_alone: bool = True
    pets: str = Field(default="", max_length=200)
    autonomy: str = Field(default="autonome", max_length=50)
    mobility: str = Field(default="autonome", max_length=50)
    professional_status: str = Field(default="", max_length=50)
    job_title: str = Field(default="", max_length=200)
    income_range: str = Field(default="", max_length=100)
    siret: str = Field(default="", max_length=20)
    doctor_name: str = Field(default="", max_length=200)
    doctor_phone: str = Field(default="", max_length=20)
    pharmacy: str = Field(default="", max_length=200)
    allergies: str = Field(default="", max_length=1000)
    treatments: str = Field(default="", max_length=1000)
    conditions: str = Field(default="", max_length=1000)
    medical_contact_person: str = Field(default="", max_length=200)
    mutual_name: str = Field(default="", max_length=200)
    mutual_number: str = Field(default="", max_length=100)
    carte_vitale: str = Field(default="", max_length=50)
    housing_type: str = Field(default="", max_length=50)
    housing_status: str = Field(default="", max_length=50)
    floor: str = Field(default="", max_length=100)
    landlord_name: str = Field(default="", max_length=200)
    landlord_phone: str = Field(default="", max_length=20)
    home_insurance: str = Field(default="", max_length=200)
    concierge: str = Field(default="", max_length=200)
    tax_number: str = Field(default="", max_length=50)
    caf_number: str = Field(default="", max_length=50)
    france_travail_id: str = Field(default="", max_length=50)
    cpam_center: str = Field(default="", max_length=100)
    bank_name: str = Field(default="", max_length=200)
    bank_advisor: str = Field(default="", max_length=200)
    documents_expiry: str = Field(default="", max_length=1000)
    tone: str = Field(default="chaleureux", max_length=50)
    wake_time: str = Field(default="08:00", max_length=5)
    sleep_time: str = Field(default="22:00", max_length=5)
    quiet_hours_start: str = Field(default="22:00", max_length=5)
    quiet_hours_end: str = Field(default="07:00", max_length=5)
    sensitive_topics: str = Field(default="", max_length=1000)
    interests: str = Field(default="", max_length=1000)
    habits: str = Field(default="", max_length=1000)
    presentation: str = Field(default="", max_length=200)
    permanent_rules: str = Field(default="", max_length=2000)
    blacklist: str = Field(default="", max_length=1000)
    priorities: str = Field(default="", max_length=1000)
    max_budget: str = Field(default="", max_length=100)


class ContactRequest(BaseModel):
    phone: str = Field(..., min_length=6, max_length=20)
    name: str = Field(..., min_length=1, max_length=100)
    relation: str = Field(..., min_length=1, max_length=50)
    email: str = Field(default="", max_length=254)
    address: str = Field(default="", max_length=300)
    preferred_channel: str = Field(default="sms", max_length=10)
    emergency_only: bool = False
    availability: str = Field(default="24/7", max_length=100)
    info_level: str = Field(default="tout", max_length=50)

    @field_validator("phone")
    @classmethod
    def validate_phone(cls, v: str) -> str:
        cleaned = re.sub(r"[\s\-\.\(\)]", "", v)
        if not re.match(r"^\+?\d{6,15}$", cleaned):
            raise ValueError("Format de telephone invalide")
        return cleaned


class InstructionRequest(BaseModel):
    text: str = Field(..., min_length=1, max_length=2000)


class NoteRequest(BaseModel):
    content: str = Field(..., min_length=1, max_length=5000)
    context: str = Field(default="manual", max_length=100)
    tags: Optional[list] = None


def _build_update_page(old_version: str, new_version: str) -> str:
    """Page HTML de mise a jour affichee dans le WebView quand l'APK est obsolete."""
    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Luna - Mise a jour</title>
<style>
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    background: linear-gradient(135deg, #0a0a1a 0%, #1a1a3e 50%, #2d1b69 100%);
    color: #e8e0ff;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
    min-height: 100vh;
    display: flex;
    align-items: center;
    justify-content: center;
    padding: 20px;
}}
.container {{
    text-align: center;
    max-width: 380px;
}}
.luna-icon {{
    font-size: 64px;
    margin-bottom: 20px;
    animation: pulse 2s infinite;
}}
@keyframes pulse {{
    0%, 100% {{ transform: scale(1); }}
    50% {{ transform: scale(1.1); }}
}}
h1 {{
    font-size: 24px;
    margin-bottom: 12px;
    color: #a78bfa;
}}
p {{
    font-size: 16px;
    line-height: 1.5;
    margin-bottom: 24px;
    opacity: 0.85;
}}
.version-badge {{
    display: inline-block;
    background: rgba(167, 139, 250, 0.15);
    border: 1px solid rgba(167, 139, 250, 0.3);
    border-radius: 20px;
    padding: 6px 16px;
    font-size: 14px;
    margin-bottom: 24px;
}}
.version-old {{ color: #f87171; text-decoration: line-through; }}
.version-new {{ color: #34d399; font-weight: bold; }}
.arrow {{ margin: 0 8px; }}
.update-btn {{
    display: inline-block;
    background: linear-gradient(135deg, #a78bfa 0%, #7c3aed 100%);
    color: white;
    text-decoration: none;
    padding: 16px 40px;
    border-radius: 30px;
    font-size: 18px;
    font-weight: bold;
    box-shadow: 0 4px 20px rgba(167, 139, 250, 0.4);
    transition: transform 0.2s;
}}
.update-btn:active {{
    transform: scale(0.96);
}}
.note {{
    margin-top: 24px;
    font-size: 13px;
    opacity: 0.6;
    line-height: 1.4;
}}
</style>
</head>
<body>
<div class="container">
    <div class="luna-icon">&#127769;</div>
    <h1>Nouvelle version disponible !</h1>
    <div class="version-badge">
        <span class="version-old">v{old_version}</span>
        <span class="arrow">&rarr;</span>
        <span class="version-new">v{new_version}</span>
    </div>
    <p>Luna a ete amelioree avec de nouvelles fonctionnalites. Mets a jour pour en profiter !</p>
    <a href="/download/luna.apk" class="update-btn">Mettre a jour</a>
    <p class="note">Le telechargement demarre automatiquement.<br>Ouvre la notification pour installer.</p>
</div>
</body>
</html>"""


# =========================================================================
# ENDPOINTS
# =========================================================================

_NO_CACHE_HEADERS = {
    "Cache-Control": "no-cache, no-store, must-revalidate",
    "Pragma": "no-cache",
    "Expires": "0",
}

@app.get("/")
async def index(request: Request):
    if _pv_locked:
        setup_path = os.path.join(STATIC_DIR, "setup.html")
        if os.path.exists(setup_path):
            return FileResponse(setup_path, headers=_NO_CACHE_HEADERS)
    return FileResponse(os.path.join(STATIC_DIR, "index.html"), headers=_NO_CACHE_HEADERS)


@app.api_route("/sw.js", methods=["GET", "HEAD"])
async def service_worker():
    """Serve SW from root scope so it can intercept all requests."""
    return FileResponse(
        os.path.join(STATIC_DIR, "sw.js"),
        media_type="application/javascript",
        headers={
            "Service-Worker-Allowed": "/",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
        },
    )


@app.get("/client")
async def client_page():
    """Acces direct a l'espace client (meme si PV non signe)."""
    return FileResponse(os.path.join(STATIC_DIR, "index.html"), headers=_NO_CACHE_HEADERS)


@app.get("/health")
async def health():
    """Healthcheck leger pour Docker/load balancers."""
    return {"status": "ok"}


@app.get("/ready")
async def ready():
    """Readiness check — Redis + secrets."""
    checks: dict = {}
    try:
        if _redis_client:
            _redis_client.ping()
            checks["redis"] = "ok"
        else:
            checks["redis"] = "unavailable"
    except Exception as e:
        checks["redis"] = f"error: {e}"
        return JSONResponse(status_code=503, content={"status": "not_ready", "checks": checks})
    for key in ("JWT_SECRET_KEY", "OPENAI_API_KEY"):
        checks[key] = "ok" if os.getenv(key) else "missing"
    all_ok = all(v == "ok" for v in checks.values())
    return JSONResponse(
        status_code=200 if all_ok else 503,
        content={"status": "ready" if all_ok else "degraded", "checks": checks},
    )


@app.get("/api/maintenance")
async def get_maintenance_status():
    """Vérifie si le mode maintenance est actif (public)."""
    if _redis_client:
        mode = _redis_client.client.get("luna:maintenance:mode")
        if mode and mode != "off":
            msg = _redis_client.client.get("luna:maintenance:message") or "Maintenance en cours"
            return JSONResponse(status_code=503, content={"maintenance": True, "message": msg})
    return {"maintenance": False}


@app.post("/api/admin/maintenance")
async def set_maintenance_mode(request: Request):
    """Active/désactive le mode maintenance (admin seulement)."""
    auth = request.headers.get("Authorization", "")
    token = auth[7:] if auth.startswith("Bearer ") else ""
    if not _decode_admin_token(token):
        return JSONResponse(status_code=403, content={"error": "Non autorisé"})
    data = await request.json()
    mode = data.get("mode", "off")
    message = data.get("message", "Maintenance en cours — Luna revient bientôt.")
    if _redis_client:
        _redis_client.client.set("luna:maintenance:mode", mode, ex=86400)
        _redis_client.client.set("luna:maintenance:message", message, ex=86400)
    return {"mode": mode, "message": message, "active": mode != "off"}


@app.get("/admin")
async def admin_page():
    """Dashboard admin exploitant."""
    admin_path = os.path.join(STATIC_DIR, "admin.html")
    if os.path.exists(admin_path):
        return FileResponse(admin_path)
    return JSONResponse(status_code=404, content={"error": "Dashboard admin non disponible"})


@app.get("/exploitant")
async def exploitant_page():
    """Dashboard exploitant — gestion clients, revenus, configuration."""
    path = os.path.join(STATIC_DIR, "exploitant.html")
    if os.path.exists(path):
        return FileResponse(path)
    return JSONResponse(status_code=404, content={"error": "Dashboard exploitant non disponible"})


@app.get("/vault")
async def vault_page():
    """Coffre-fort documentaire — scan, classement, rappels."""
    path = os.path.join(STATIC_DIR, "vault.html")
    if os.path.exists(path):
        return FileResponse(path)
    return JSONResponse(status_code=404, content={"error": "Vault non disponible"})


@app.get("/world")
async def world_page():
    """Page gamifiee client - IA Watch World."""
    world_path = os.path.join(STATIC_DIR, "world.html")
    if os.path.exists(world_path):
        return FileResponse(world_path)
    return JSONResponse(status_code=404, content={"error": "World non disponible"})


@app.get("/admin/world")
async def admin_world_page():
    """Page gamifiee exploitant - IA Watch World."""
    path = os.path.join(STATIC_DIR, "admin_world.html")
    if os.path.exists(path):
        return FileResponse(path)
    return JSONResponse(status_code=404, content={"error": "World admin non disponible"})


@app.get("/download")
async def download_page():
    """Page de telechargement de l'app Android."""
    return FileResponse(os.path.join(STATIC_DIR, "download.html"))


@app.get("/download/luna.apk")
async def download_apk():
    """Telecharge l'APK Android avec le bon Content-Type."""
    apk_path = os.path.join(STATIC_DIR, "luna-proprio.apk")
    if not os.path.exists(apk_path):
        return JSONResponse(status_code=404, content={"error": "APK non disponible"})
    return FileResponse(
        apk_path,
        media_type="application/vnd.android.package-archive",
        filename="Luna-Proprio.apk",
    )




# Version APK pour auto-update
LUNA_APP_VERSION = "2.7"
LUNA_APP_VERSION_CODE = 18


def _compute_apk_sha256() -> str:
    """Calcule le SHA-256 de l'APK au démarrage, mis en cache."""
    import hashlib
    apk_path = os.path.join(STATIC_DIR, "luna-proprio.apk")
    if not os.path.exists(apk_path):
        return ""
    h = hashlib.sha256()
    with open(apk_path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


_APK_SHA256: str = _compute_apk_sha256()


@app.get("/api/app/version")
async def app_version():
    """Retourne la version courante de l'APK pour auto-update."""
    return {
        "version": LUNA_APP_VERSION,
        "version_code": LUNA_APP_VERSION_CODE,
        "apk_url": "/download/luna.apk",
        "apk_sha256": _APK_SHA256,
        "changelog": "Sécurité APK renforcée, vérification intégrité mise à jour",
    }


async def _track_openai_cost(response, tenant_id: int = None):
    """Track OpenAI API call cost via Cortex audit. Fire-and-forget."""
    try:
        cortex = get_cortex() if _CORTEX_AVAILABLE else None
        if not cortex or not cortex.cost_tracker:
            return
        usage = getattr(response, "usage", None)
        if not usage:
            return
        tokens_in = getattr(usage, "prompt_tokens", 0) or 0
        tokens_out = getattr(usage, "completion_tokens", 0) or 0
        if tokens_in or tokens_out:
            await cortex.cost_tracker.track_openai(tokens_in, tokens_out, tenant_id=tenant_id)
    except Exception:
        pass


def _gamify(tenant_id, action: str, metadata: dict = None, is_admin: bool = False):
    """Fire-and-forget gamification XP award. Never raises, never blocks."""
    if not _GAMIFICATION_AVAILABLE or not _redis_client:
        return
    try:
        gops = GamificationRedisOps(_redis_client)
        asyncio.create_task(award_xp_safe(gops, tenant_id, action, metadata, is_admin=is_admin))
    except Exception:
        pass


_MAX_CONCURRENT_CHATS = int(os.getenv("MAX_CONCURRENT_CHATS", "200"))
_chat_semaphore = asyncio.Semaphore(_MAX_CONCURRENT_CHATS)


# =========================================================================
# CHAT TOOL DISPATCH (shared between streaming and non-streaming)
# =========================================================================

async def _dispatch_chat_tool(fn_name: str, fn_args: dict, tid: int, session_id: str = "default"):
    """Execute a chat tool call. Returns result dict."""
    if fn_name == "send_sms":
        return await _tool_send_sms(fn_args, tenant_id=tid)
    elif fn_name == "create_instruction":
        return await _tool_create_instruction(fn_args, tenant_id=tid)
    elif fn_name == "create_note":
        return await _tool_create_note(fn_args, tenant_id=tid)
    elif fn_name == "get_contacts":
        return await _tool_get_contacts(tenant_id=tid)
    elif fn_name == "generate_document":
        return await _tool_generate_document(fn_args, tenant_id=tid)
    elif fn_name == "alert_contacts":
        return await _tool_alert_contacts(fn_args, tenant_id=tid)
    elif fn_name == "send_email":
        return await _tool_send_email(fn_args, tenant_id=tid)
    elif fn_name == "invite_visio":
        return await _tool_invite_visio(fn_args, tenant_id=tid)
    elif fn_name == "call_contact":
        return await _tool_call_contact(fn_args, tenant_id=tid, session_id=session_id)
    elif fn_name == "search_web":
        return await _tool_search_web(fn_args, tenant_id=tid)
    elif fn_name == "search_places":
        return await _tool_search_places(fn_args, tenant_id=tid)
    elif fn_name == "get_page_info":
        return await _tool_get_page_info(fn_args, tenant_id=tid)
    elif fn_name == "request_payment":
        return await _tool_request_payment(fn_args, tenant_id=tid)
    elif fn_name == "get_player_stats":
        return await _tool_get_player_stats(tenant_id=tid)
    elif fn_name == "get_active_missions":
        return await _tool_get_active_missions(tenant_id=tid)
    elif fn_name == "get_badges":
        return await _tool_get_badges(tenant_id=tid)
    elif fn_name == "get_weather":
        return await _tool_get_weather(fn_args, tenant_id=tid)
    elif fn_name == "get_news":
        return await _tool_get_news(fn_args)
    elif fn_name == "search_flights":
        return await _tool_search_flights(fn_args, tenant_id=tid)
    elif fn_name == "search_hotels":
        return await _tool_search_hotels(fn_args, tenant_id=tid)
    elif fn_name == "book_restaurant":
        return await _tool_book_restaurant(fn_args, tenant_id=tid)
    elif fn_name == "get_documents_summary":
        return _tool_secretary_summary(tid)
    elif fn_name == "get_budget_analysis":
        return _tool_secretary_budget(tid)
    elif fn_name == "check_affordability":
        return _tool_secretary_afford(tid, fn_args)
    elif fn_name == "add_expense":
        return _tool_secretary_add_expense(tid, fn_args)
    elif fn_name == "get_reminders":
        return _tool_secretary_reminders(tid)
    elif fn_name == "add_reminder":
        return _tool_secretary_add_reminder(tid, fn_args)
    elif fn_name == "search_documents":
        return _tool_secretary_search(tid, fn_args)
    elif fn_name == "list_folders":
        return _tool_secretary_folders(tid)
    else:
        return {"status": "error", "message": f"Fonction inconnue: {fn_name}"}


def _build_rich_cards(tool_calls_made: list) -> list:
    """Build rich card data from tool call results."""
    cards = []
    for t in tool_calls_made:
        r = t["result"]
        if r.get("status") != "success":
            continue
        tn = t["tool"]
        if tn == "search_places" and r.get("places"):
            for p in r["places"][:5]:
                cards.append({"type": "place", "name": p.get("name", ""), "address": p.get("address", ""),
                    "phone": p.get("phone", ""), "rating": p.get("rating", ""), "hours": p.get("hours", ""),
                    "url": p.get("website", ""), "price_level": p.get("price_level", "")})
        elif tn == "search_web" and r.get("results"):
            for sr in r["results"][:5]:
                if isinstance(sr, dict):
                    cards.append({"type": "web", "title": sr.get("title", ""), "snippet": sr.get("snippet", ""), "url": sr.get("link", "")})
                elif isinstance(sr, str):
                    cards.append({"type": "web", "title": sr[:80], "snippet": sr, "url": ""})
        elif tn == "get_weather" and r.get("actuel"):
            cur = r["actuel"]
            cards.append({"type": "weather", "city": cur.get("ville", ""), "temp": cur.get("temperature", ""),
                "feels_like": cur.get("ressenti", ""), "description": cur.get("description", ""),
                "humidity": cur.get("humidite", ""), "wind": cur.get("vent", ""),
                "forecasts": [{"date": f.get("date", ""), "max": f.get("max", ""), "min": f.get("min", ""), "description": f.get("description", "")} for f in (r.get("previsions") or [])[:3]]})
        elif tn == "get_news" and r.get("articles"):
            for art in r["articles"][:5]:
                cards.append({"type": "news", "title": art.get("titre", ""), "url": art.get("lien", ""),
                    "source": art.get("source", ""), "date": art.get("date", ""), "snippet": art.get("resume", "")})
        elif tn == "search_flights" and r.get("flights"):
            for fl in r["flights"][:5]:
                cards.append({"type": "flight", "airline": fl.get("airline", ""), "price": fl.get("price", ""), "summary": fl.get("summary", "")})
        elif tn == "search_hotels" and r.get("hotels"):
            for h in r["hotels"][:5]:
                cards.append({"type": "hotel", "name": h.get("name", ""), "stars": h.get("stars", ""),
                    "city": h.get("city", ""), "price_per_night": h.get("price_per_night", ""), "room_type": h.get("room_type", "")})
        elif tn == "book_restaurant" and r.get("places"):
            for p in r["places"][:5]:
                cards.append({"type": "place", "name": p.get("name", ""), "address": p.get("address", ""),
                    "phone": p.get("phone", ""), "rating": p.get("rating", ""), "hours": p.get("hours", ""),
                    "url": p.get("website", ""), "price_level": p.get("price_level", "")})
    return cards


async def _stream_chat_sse(messages, chat_tools, tid, session_id, req_message, mgr, semaphore=None):
    """Async generator yielding SSE events for streaming chat."""
    import queue as _queue

    q = _queue.Queue(maxsize=200)
    loop = asyncio.get_event_loop()

    def _run_stream():
        try:
            stream = openai_client.chat.completions.create(
                model=OPENAI_MODEL, messages=messages, stream=True,
                tools=chat_tools if chat_tools else openai.NOT_GIVEN,
                max_tokens=500, temperature=0.8, timeout=30,
            )
            for chunk in stream:
                q.put(("chunk", chunk))
            q.put(("done", None))
        except Exception as e:
            q.put(("error", e))

    loop.run_in_executor(None, _run_stream)

    full_text = ""
    tool_calls_acc = {}

    try:
        while True:
            try:
                item = await asyncio.to_thread(q.get, timeout=40)
            except Exception:
                logger.warning(f"Stream timeout for session {session_id}")
                yield f"data: {json.dumps({'type': 'error', 'message': 'Timeout du flux'})}\n\n"
                return
            evt_type, data = item
            if evt_type == "error":
                yield f"data: {json.dumps({'type': 'error', 'message': str(data)})}\n\n"
                return
            if evt_type == "done":
                break
            chunk = data
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta
            if delta and delta.content:
                full_text += delta.content
                yield f"data: {json.dumps({'type': 'chunk', 'text': delta.content})}\n\n"
            if delta and delta.tool_calls:
                for tc in delta.tool_calls:
                    idx = tc.index
                    if idx not in tool_calls_acc:
                        tool_calls_acc[idx] = {"id": "", "name": "", "arguments": ""}
                    if tc.id:
                        tool_calls_acc[idx]["id"] = tc.id
                    if tc.function:
                        if tc.function.name:
                            tool_calls_acc[idx]["name"] = tc.function.name
                        if tc.function.arguments:
                            tool_calls_acc[idx]["arguments"] += tc.function.arguments

        _openai_breaker.record_success()

        # Handle tool calls
        tool_calls_made = []
        if tool_calls_acc:
            sorted_tcs = [tool_calls_acc[i] for i in sorted(tool_calls_acc.keys())]
            asst_tc = [{"id": tc["id"], "type": "function",
                         "function": {"name": tc["name"], "arguments": tc["arguments"]}}
                        for tc in sorted_tcs]
            messages.append({"role": "assistant", "content": full_text or None, "tool_calls": asst_tc})

            for tc in sorted_tcs:
                fn_name = tc["name"]
                try:
                    fn_args = json.loads(tc["arguments"])
                except Exception:
                    fn_args = {}
                yield f"data: {json.dumps({'type': 'tool', 'name': fn_name, 'status': 'running'})}\n\n"
                result = await _dispatch_chat_tool(fn_name, fn_args, tid, session_id)
                tool_calls_made.append({"tool": fn_name, "result": result})
                messages.append({"role": "tool", "tool_call_id": tc["id"],
                                 "content": json.dumps(result, ensure_ascii=False)})
                yield f"data: {json.dumps({'type': 'tool', 'name': fn_name, 'status': 'done'})}\n\n"

            # Anti-hallucination
            _ACTION_TOOLS = {"call_contact", "send_sms", "send_email", "alert_contacts",
                             "request_payment", "invite_visio", "generate_document", "create_instruction"}
            _failed = [t for t in tool_calls_made if t["result"].get("status") == "error" and t["tool"] in _ACTION_TOOLS]
            if _failed:
                messages.append({"role": "system",
                    "content": f"IMPORTANT — Les actions suivantes ont ECHOUE: {', '.join(t['tool'] for t in _failed)}. Informe le souscripteur de l'echec."})

            # Second call with tool results (non-streaming for simplicity)
            try:
                response2 = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                    model=OPENAI_MODEL, messages=messages, max_tokens=500, temperature=0.8, timeout=30))
                full_text = response2.choices[0].message.content or ""
                yield f"data: {json.dumps({'type': 'chunk', 'text': full_text})}\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
                return

        # Persist to Redis
        if mgr:
            try:
                mgr.add_message(conv_id=session_id, role=MessageRole.LUNA, content=full_text, channel=Channel.APP)
            except Exception as e:
                logger.warning(f"Redis message persist failed: {e}")

        # Auto-title
        auto_title = None
        if mgr:
            try:
                meta = mgr.redis.get_conversation_meta(tid, session_id)
                if meta and not meta.get("summary") and int(meta.get("message_count", 0)) <= 2:
                    try:
                        title_resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{"role": "system", "content": "Tu generes un titre court (4-6 mots max) pour une conversation. Pas de guillemets."},
                                      {"role": "user", "content": f"User: {req_message[:200]}\nLuna: {full_text[:200]}"}],
                            max_tokens=20, temperature=0.3, timeout=5))
                        auto_title = title_resp.choices[0].message.content.strip().strip('"').strip("'")
                        if len(auto_title) > 50:
                            auto_title = auto_title[:50]
                        meta["summary"] = auto_title
                        meta["last_activity"] = datetime.utcnow().isoformat()
                        mgr.redis.set_conversation_meta(tid, session_id, meta)
                    except Exception:
                        pass
            except Exception:
                pass

        _gamify(tid, "chat_message")

        # Build final event
        cards = _build_rich_cards(tool_calls_made)
        done_data = {"type": "done", "response": full_text}
        if auto_title:
            done_data["auto_title"] = auto_title
        if tool_calls_made:
            done_data["actions"] = [{"tool": t["tool"], "status": t["result"].get("status", "unknown")} for t in tool_calls_made]
            for t in tool_calls_made:
                if t["result"].get("visio_url"):
                    done_data["visio_url"] = t["result"]["visio_url"]
                    break
        if cards:
            done_data["cards"] = cards
        yield f"data: {json.dumps(done_data)}\n\n"
    finally:
        if semaphore:
            semaphore.release()


@app.post("/api/chat")
async def chat(req: ChatRequest, request: Request):
    # Service validation (redundant safety layer)
    if _license_heartbeat and _license_heartbeat.is_blocked():
        return JSONResponse(status_code=403, content={"error": "Service suspendu"})
    # Backpressure: reject if too many concurrent chats (atomic semaphore)
    if _chat_semaphore.locked():
        return JSONResponse(status_code=503, content={"error": "Luna est tres sollicitee, reessaie dans quelques secondes"})
    await _chat_semaphore.acquire()
    try:
        tid = getattr(request.state, "tenant_id", 1)
        mgr = _get_tenant_manager(tid)
        tenant_convs = conversations.setdefault(str(tid), {})

        # Safety check (if guardian available)
        if _safety_guardian:
            try:
                safety = _safety_guardian.check(req.message)
                if safety.block_action and safety.luna_response:
                    return {"response": safety.luna_response}
            except Exception as e:
                logger.warning(f"Safety check failed: {e}")

        if req.session_id not in tenant_convs:
            tenant_convs[req.session_id] = [
                {"role": "system", "content": _system_prompt_now()}
            ]

        # Auto-create conversation metadata in Redis if missing
        if mgr:
            try:
                meta = mgr.redis.get_conversation_meta(tid, req.session_id)
                if not meta:
                    mgr.redis.add_conversation(tid, req.session_id)
                    mgr.redis.set_conversation_meta(tid, req.session_id, {
                        "id": req.session_id,
                        "tenant_id": str(tid),
                        "contact_phone": "app",
                        "contact_name": "Chat",
                        "status": "active",
                        "channel": "app",
                        "started_at": datetime.utcnow().isoformat(),
                        "last_activity": datetime.utcnow().isoformat(),
                        "message_count": "0",
                        "summary": "",
                    })
            except Exception:
                pass

        _conversation_ts[req.session_id] = time.time()
        messages = tenant_convs[req.session_id]

        # --- FIX: strip stale injected context from previous turns ---
        # Keep only the initial system prompt (index 0) + user/assistant/tool messages.
        # Context (profile, contacts, gamification, etc.) is re-injected fresh each turn.
        messages[:] = [messages[0]] + [
            m for m in messages[1:]
            if (m.get("role") if isinstance(m, dict) else getattr(m, "role", None)) != "system"
        ]

        # --- FIX: update system prompt with current date/time ---
        messages[0] = {"role": "system", "content": _system_prompt_now()}

        # Inject behavioral memory (locked identity + rules) before user message
        if mgr:
            try:
                rules = mgr.get_behavioral_rules()
                if rules["identity_core"] or rules["behavior_rules"]:
                    behavioral_prompt = (
                        f"[MEMOIRE COMPORTEMENTALE VERROUILLEE]\n"
                        f"IDENTITE: {rules['identity_core']}\n"
                        f"REGLES: {rules['behavior_rules']}"
                    )
                    messages.append({"role": "system", "content": behavioral_prompt})
            except Exception:
                pass

        # Inject FULL tenant context: profile + contacts + instructions + notes
        _caution_mode = "assistif"
        tenant_name = ""
        if mgr:
            context_parts = []
            # --- Profil souscripteur ---
            try:
                profile = mgr.get_subscriber_profile()
                if profile:
                    _caution_mode = getattr(profile, "caution_mode", "assistif") or "assistif"
                    tenant_name = profile.first_name or ""
                    pf_lines = [f"=== TON SOUSCRIPTEUR ==="]
                    pf_lines.append(f"Prenom: {profile.first_name or '?'}")
                    pf_lines.append(f"Nom: {profile.last_name or '?'}")
                    if getattr(profile, "gender", None):
                        _g = {"M": "Homme", "F": "Femme"}.get(profile.gender.upper(), profile.gender)
                        pf_lines.append(f"Sexe: {_g}")
                    if getattr(profile, "date_of_birth", None):
                        pf_lines.append(f"Date de naissance: {profile.date_of_birth}")
                    if getattr(profile, "phone", None):
                        pf_lines.append(f"Telephone: {profile.phone}")
                    if getattr(profile, "email", None):
                        pf_lines.append(f"Email: {profile.email}")
                    if getattr(profile, "address", None):
                        pf_lines.append(f"Adresse: {profile.address}")
                    if getattr(profile, "city", None):
                        pf_lines.append(f"Ville: {profile.city}")
                    if getattr(profile, "family_status", None):
                        pf_lines.append(f"Situation: {profile.family_status}")
                    if getattr(profile, "children", None):
                        pf_lines.append(f"Enfants: {profile.children}")
                    if getattr(profile, "autonomy", None):
                        pf_lines.append(f"Autonomie: {profile.autonomy}")
                    if getattr(profile, "conditions", None):
                        pf_lines.append(f"Pathologies: {profile.conditions}")
                    if getattr(profile, "treatments", None):
                        pf_lines.append(f"Traitements: {profile.treatments}")
                    if getattr(profile, "interests", None):
                        pf_lines.append(f"Centres d'interet: {profile.interests}")
                    if getattr(profile, "habits", None):
                        pf_lines.append(f"Habitudes: {profile.habits}")
                    if getattr(profile, "tone", None):
                        pf_lines.append(f"Ton souhaite: {profile.tone}")
                    if getattr(profile, "presentation", None):
                        pf_lines.append(f"Presentation: {profile.presentation}")
                    if getattr(profile, "permanent_rules", None):
                        pf_lines.append(f"Regles permanentes: {profile.permanent_rules}")
                    if getattr(profile, "sensitive_topics", None):
                        pf_lines.append(f"Sujets sensibles a eviter: {profile.sensitive_topics}")
                    context_parts.append("\n".join(pf_lines))
            except Exception as e:
                logger.warning(f"Profile context error: {e}")

            # --- Contacts de confiance ---
            try:
                contacts = mgr.list_trusted_contacts()
                if contacts:
                    ct_lines = ["=== CONTACTS DE CONFIANCE (tu les connais deja, n'appelle PAS get_contacts) ==="]
                    ct_lines.append(f"Voici les {len(contacts)} personnes de confiance de {tenant_name or 'ton souscripteur'}:")
                    for c in contacts:
                        first = c.name.split()[0] if c.name else "Contact"
                        rel = c.relation or "proche"
                        ct_lines.append(f"  - {c.name} (relation: {rel}) — telephone: {c.phone}")
                    ct_lines.append("REGLE: Pour envoyer un SMS, utilise directement ces numeros avec le tool send_sms.")
                    ct_lines.append("REGLE: Si on te demande 'qui sont mes contacts', reponds DIRECTEMENT avec cette liste.")
                    ct_lines.append("REGLE: N'appelle JAMAIS le tool get_contacts car tu as deja toutes les infos ci-dessus.")
                    context_parts.append("\n".join(ct_lines))
                else:
                    context_parts.append("=== CONTACTS ===\nAucun contact de confiance enregistre. Propose au souscripteur d'en ajouter dans l'onglet Contacts.")
            except Exception as e:
                logger.warning(f"Contacts context error: {e}")

            # --- Geolocalisation du souscripteur ---
            if _redis_client:
                try:
                    import json as _json_geo
                    geo_raw = _redis_client.client.get(f"luna:{tid}:geolocation")
                    if geo_raw:
                        geo_data = _json_geo.loads(geo_raw)
                        geo_city = geo_data.get("city", "")
                        geo_addr = geo_data.get("address", "")
                        geo_parts = []
                        if geo_city:
                            geo_parts.append(f"Ville: {geo_city}")
                        if geo_addr:
                            geo_parts.append(f"Adresse approximative: {geo_addr}")
                        if geo_parts:
                            context_parts.append(
                                "=== LOCALISATION ACTUELLE DU SOUSCRIPTEUR ===\n"
                                + "\n".join(geo_parts)
                                + "\nUtilise cette info pour contextualiser tes recherches (restaurants proches, services locaux, etc.)."
                                + "\nTu peux mentionner naturellement cette localisation si pertinent."
                            )
                except Exception as e:
                    logger.debug(f"Geo context error: {e}")

            # --- Instructions actives ---
            try:
                instructions = mgr.list_active_instructions()
                if instructions:
                    instr_lines = ["=== INSTRUCTIONS ACTIVES ==="]
                    for instr in instructions[:10]:
                        instr_lines.append(f"- {instr.description}")
                    context_parts.append("\n".join(instr_lines))
            except Exception as e:
                logger.warning(f"Instructions context error: {e}")

            # --- Notes recentes ---
            try:
                notes = mgr.list_notes(limit=5)
                if notes:
                    note_lines = ["=== NOTES RECENTES ==="]
                    for n in notes[:5]:
                        note_lines.append(f"- {n.content[:100]}")
                    context_parts.append("\n".join(note_lines))
            except Exception as e:
                logger.warning(f"Notes context error: {e}")

            # Injecte tout le contexte dans un seul message systeme
            if context_parts:
                full_context = "\n\n".join(context_parts)
                full_context += f"\n\n=== REGLES CRITIQUES ==="
                full_context += f"\nTu parles avec {tenant_name or 'le souscripteur'}. Tutoie-le et utilise son prenom."
                full_context += "\n- Tu CONNAIS deja le profil, les contacts, les notes ci-dessus. Utilise-les directement."
                full_context += "\n- Ne dis JAMAIS 'je n'ai pas acces a tes contacts' ou 'je ne peux pas voir ton profil'."
                full_context += "\n- Ne dis JAMAIS 'en tant qu'IA je ne peux pas' quand tu as un tool pour le faire."
                full_context += "\n- DISTINCTION CRITIQUE entre appeler et envoyer un SMS :"
                full_context += "\n  * 'appelle', 'telephone', 'passe un appel', 'coup de fil' → TOUJOURS utiliser call_contact (appel VOCAL)"
                full_context += "\n  * 'envoie un SMS', 'envoie un texto', 'ecris un message' → utiliser send_sms (message ECRIT)"
                full_context += "\n  * NE JAMAIS utiliser send_sms quand le souscripteur dit 'appelle' — c'est call_contact"
                full_context += "\n- Tu PEUX passer des appels telephoniques audio aux contacts de confiance avec call_contact."
                full_context += "\n- Tu PEUX aussi appeler des administrations/services si le souscripteur te donne le numero (utilise call_contact avec phone_number)."
                full_context += "\n- Tu ne peux PAS appeler les numeros d'urgence (15, 17, 18, 112) — suggere-les a la place."
                full_context += "\n- Quand on te demande de PLANIFIER quelque chose dans le FUTUR (ex: 'appelle Marie a 14h', 'rappelle-moi demain'), utilise le tool create_instruction."
                full_context += "\n- Pour les appels planifies, demande la duree souhaitee (1-10 min) et passe max_duration_minutes. Ex: 'appelle Jess a 14h pendant 5 minutes'."
                full_context += "\n- Apres chaque appel (meme si le souscripteur n'est pas present), Luna envoie un SMS avec le compte-rendu de ce que l'interlocuteur a dit."
                full_context += "\n- Quand on te demande 'qui sont mes contacts', reponds avec la liste ci-dessus."
                full_context += "\n- Tu as acces a Twilio pour envoyer des SMS et passer des appels. Tu SAIS le faire. Ne dis JAMAIS que tu ne peux pas."
                full_context += "\n- Sois chaleureux, concis, et utile. Tu es Luna, un compagnon bienveillant."
                messages.append({"role": "system", "content": full_context})
            elif tenant_name:
                messages.append({"role": "system", "content": f"Tu parles actuellement avec {tenant_name}. Utilise son prenom."})

        caution_prompt = CAUTION_MODE_PROMPTS.get(_caution_mode, CAUTION_MODE_PROMPTS["assistif"])
        messages.append({"role": "system", "content": caution_prompt})

        # Inject mode-aware context (compagnon vs secretaire)
        _chat_mode = getattr(req, "mode", "compagnon") or "compagnon"
        if _chat_mode == "secretaire":
            messages.append({"role": "system", "content": """=== MODE SECRETAIRE ACTIVE ===
Tu es en mode SECRETAIRE PERSONNELLE. Tu es directe, efficace, professionnelle.
Tes priorites dans ce mode :
1. DOCUMENTS : resumer, classer, retrouver des documents (get_documents_summary, search_documents, list_folders)
2. BUDGET : analyser les depenses, verifier si une depense est raisonnable (get_budget_analysis, check_affordability, add_expense)
3. RAPPELS : gerer les echeances, factures a payer, RDV (get_reminders, add_reminder)
4. CONCIERGERIE PRO : recherches de vols, hotels, restaurants, services (search_flights, search_hotels, book_restaurant)
5. COMMUNICATION : SMS, emails, appels (send_sms, send_email, call_contact)
6. ORGANISATION : notes, instructions planifiees (create_note, create_instruction)

COMPORTEMENT SECRETAIRE :
- Sois PROACTIVE : si le souscripteur parle d'argent → propose d'analyser le budget
- Si il parle d'un voyage → propose de chercher des vols et hotels
- Si il parle d'un RDV → propose de creer un rappel
- Si il parle d'une facture → propose de la classer ou de verifier les echeances
- Utilise les tools AUTOMATIQUEMENT quand le contexte est clair, sans demander 3 fois confirmation
- Reponses structurees : utilise **gras**, listes, et chiffres
- Pas de bavardage inutile, vas droit au but"""})
        else:
            messages.append({"role": "system", "content": """=== MODE COMPAGNON ACTIVE ===
Tu es en mode COMPAGNON bienveillant. Tu es chaleureuse, empathique, attentive.
Tes priorites dans ce mode :
1. BIEN-ETRE : ecouter, rassurer, accompagner. Tu es une presence amicale 24h/24
2. ACTIVITES : proposer des jeux, quiz, lecture, musique (quand le souscripteur s'ennuie)
3. LIEN SOCIAL : encourager les contacts avec les proches, Monde Luna (amis)
4. CONCIERGERIE QUOTIDIENNE : meteo, actus, rappels, recherches pratiques
5. SECURITE : detecter la detresse, proposer de contacter un proche

COMPORTEMENT COMPAGNON :
- Sois PROACTIVE : propose des activites quand le souscripteur semble seul ou s'ennuie
- Si il dit "je m'ennuie" → propose un quiz, une histoire, de la musique
- Si il semble triste → ecoute d'abord, puis propose de parler a un proche
- Si il demande la meteo → donne-la ET suggere une activite adaptee
- Utilise le prenom du souscripteur souvent
- Ton chaleureux et naturel, comme une amie de confiance"""})

        # Inject language instruction if not French
        _lang_pref = "fr"
        if _redis_client:
            try:
                _lang_settings = _redis_client.client.hgetall(f"luna:{tid}:settings")
                _lang_pref = _lang_settings.get("language", "fr") if _lang_settings else "fr"
            except Exception:
                pass
        _lang_instr = _get_language_instruction(_lang_pref)
        if _lang_instr:
            messages.append({"role": "system", "content": _lang_instr})

        # Inject active call awareness — prevent hallucinating call reports
        _active_call = _active_voice_calls.get(tid)
        if _active_call:
            _call_contact_name = _active_call.get("contact", "quelqu'un")
            _call_elapsed = int(time.time() - _active_call.get("started", 0))
            messages.append({"role": "system", "content": (
                f"APPEL EN COURS: Tu es actuellement en train d'appeler {_call_contact_name} "
                f"(depuis {_call_elapsed}s). L'appel n'est PAS termine. "
                f"Tu ne connais PAS encore le contenu de la conversation. "
                f"NE FABRIQUE PAS de compte-rendu. Si on te demande des nouvelles de "
                f"{_call_contact_name}, dis que l'appel est en cours et que tu donneras "
                f"un compte-rendu des que l'appel sera termine."
            )})

        # Inject gamification context (XP, level, missions, badges, streak)
        if _GAMIFICATION_AVAILABLE and _redis_client:
            try:
                from core.gamification.redis_ops import GamificationRedisOps
                from core.gamification.engine import get_level_for_xp
                from core.gamification.constants import ALL_CLIENT_BADGES
                _gops = GamificationRedisOps(_redis_client)
                _player = _gops.get_player(tid)
                if _player:
                    _xp = int(_player.get("xp", 0))
                    _lvl = get_level_for_xp(_xp)
                    _streak = int(_player.get("streak_days", 0))
                    _stars = int(_player.get("stars", 0))
                    _badges_set = _gops.get_badges(tid)
                    _badge_names = []
                    for _bid in _badges_set:
                        _bdef = ALL_CLIENT_BADGES.get(_bid)
                        if _bdef:
                            _badge_names.append(_bdef["name"])
                    _missions_ids = _gops.get_active_mission_ids(tid)
                    _mission_lines = []
                    for _mid in _missions_ids:
                        _mdata = _gops.get_mission(tid, _mid)
                        if _mdata:
                            _prog = int(_mdata.get("progress", 0)) - int(_mdata.get("start_progress", 0))
                            _tgt = int(_mdata.get("target", 1)) - int(_mdata.get("start_progress", 0))
                            _mission_lines.append(f"  - {_mdata.get('title', '?')} ({_prog}/{_tgt})")
                    _stab = _gops.get_stability(tid)
                    _stab_score = int(_stab.get("score", 70)) if _stab else 70
                    _stab_trend = (_stab.get("trend", "stable")) if _stab else "stable"

                    gamif_ctx = f"=== MONDE IA WATCH (progression du souscripteur) ==="
                    gamif_ctx += f"\nNiveau {_lvl['level']} : {_lvl['title']} — {_lvl['progress_percent']}% vers le suivant"
                    gamif_ctx += f"\nXP: {_xp} / {_lvl['xp_next_level']} pour le prochain niveau"
                    gamif_ctx += f"\nEtoiles: {_stars}"
                    gamif_ctx += f"\nSerie: {_streak} jour(s) consecutif(s)"
                    gamif_ctx += f"\nStabilite: {_stab_score}/100 (tendance: {_stab_trend})"
                    if _badge_names:
                        gamif_ctx += f"\nBadges ({len(_badge_names)}): {', '.join(_badge_names[:10])}"
                    else:
                        gamif_ctx += "\nAucun badge encore — encourage le souscripteur a explorer !"
                    if _mission_lines:
                        gamif_ctx += "\nMissions actives:"
                        gamif_ctx += "\n".join([""] + _mission_lines)
                    else:
                        gamif_ctx += "\nAucune mission active."
                    gamif_ctx += "\n\nTu peux repondre aux questions du souscripteur sur son niveau, ses badges, ses missions, ses etoiles."
                    gamif_ctx += "\nEncourage-le a progresser sans etre insistant. Si il demande 'c'est quoi mon niveau', reponds avec ces infos."
                    gamif_ctx += "\nPour des details precis, utilise les tools get_player_stats, get_active_missions, get_badges."
                    messages.append({"role": "system", "content": gamif_ctx})
            except Exception as e:
                logger.debug(f"Gamification context injection: {e}")

        # Inject budget context so Luna knows the subscriber's spending limits
        if _redis_client:
            try:
                _prof = _redis_client.get_profile(tid)
                _max_b = 0
                if _prof:
                    try:
                        _max_b = int(float(getattr(_prof, "max_budget", 0) or 0))
                    except (ValueError, TypeError):
                        _max_b = 0
                if _max_b > 0:
                    import datetime as _dt_budget
                    _mk = _dt_budget.datetime.now().strftime("%Y-%m")
                    _sk = f"{_redis_client.prefix}:{tid}:concierge:spending:{_mk}"
                    _spent_c = int(_redis_client.client.get(_sk) or 0)
                    _spent_e = _spent_c / 100
                    _remain = max(0, _max_b - _spent_e)
                    messages.append({"role": "system", "content": (
                        f"BUDGET CONCIERGERIE du souscripteur ce mois-ci : "
                        f"plafond {_max_b} EUR, deja depense {_spent_e:.2f} EUR, "
                        f"reste {_remain:.2f} EUR. "
                        f"Ne propose RIEN qui depasse le reste disponible."
                    )})
            except Exception:
                pass

        messages.append({"role": "user", "content": req.message})

        # Persist to Redis (if available)
        if mgr:
            try:
                mgr.add_message(
                    conv_id=req.session_id,
                    role=MessageRole.SUBSCRIBER,
                    content=req.message,
                    channel=Channel.APP,
                )
            except Exception as e:
                logger.warning(f"Redis store failed: {e}")

        # Inject perception context if available (filtered by caution_mode)
        if _perception_analyzer and mgr:
            try:
                if mgr.is_perception_enabled():
                    perception_ctx = _perception_analyzer.get_context_for_luna(
                        caution_mode=_caution_mode
                    )
                    if perception_ctx:
                        messages.append({"role": "system", "content": perception_ctx})
            except Exception:
                pass

        # Tools disponibles en chat (meme que voix/visio)
        from integrations.openai.realtime_bridge import VOICE_TOOLS as _CHAT_TOOLS
        chat_tools = [
            {"type": "function", "function": {"name": t["name"], "description": t["description"], "parameters": t["parameters"]}}
            for t in _CHAT_TOOLS
        ] if _CHAT_TOOLS else []

        if _openai_breaker.is_open:
            return {"response": "Luna est temporairement indisponible. Reessaie dans une minute."}

        # --- STREAMING SSE ---
        if getattr(req, "stream", False):
            # Pass semaphore to generator — it will release when stream ends
            _stream_sem = _chat_semaphore
            return StreamingResponse(
                _stream_chat_sse(messages, chat_tools, tid, req.session_id, req.message, mgr, semaphore=_stream_sem),
                media_type="text/event-stream",
                headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
            )

        # --- NON-STREAMING (legacy) ---
        async with _openai_chat_semaphore:
            loop = asyncio.get_event_loop()
            response = await loop.run_in_executor(
                None,
                lambda: openai_client.chat.completions.create(
                    model=OPENAI_MODEL,
                    messages=messages,
                    max_tokens=500,
                    temperature=0.8,
                    timeout=30,
                    tools=chat_tools if chat_tools else None,
                ),
            )
        _openai_breaker.record_success()

        choice = response.choices[0]
        tool_calls_made = []

        # Boucle de tool calling (max 3 tours pour eviter boucle infinie)
        for _round in range(3):
            if choice.finish_reason != "tool_calls" or not choice.message.tool_calls:
                break

            # Ajoute le message assistant avec les tool_calls
            messages.append(choice.message)

            for tc in choice.message.tool_calls:
                fn_name = tc.function.name
                try:
                    fn_args = json.loads(tc.function.arguments)
                except Exception:
                    fn_args = {}

                logger.info(f"Chat tool_call: {fn_name}({fn_args}) [tenant={tid}]")

                # Execute le tool (passe tenant_id pour multi-tenant)
                if fn_name == "send_sms":
                    result = await _tool_send_sms(fn_args, tenant_id=tid)
                elif fn_name == "create_instruction":
                    result = await _tool_create_instruction(fn_args, tenant_id=tid)
                elif fn_name == "create_note":
                    result = await _tool_create_note(fn_args, tenant_id=tid)
                elif fn_name == "get_contacts":
                    result = await _tool_get_contacts(tenant_id=tid)
                elif fn_name == "generate_document":
                    result = await _tool_generate_document(fn_args, tenant_id=tid)
                elif fn_name == "alert_contacts":
                    result = await _tool_alert_contacts(fn_args, tenant_id=tid)
                elif fn_name == "send_email":
                    result = await _tool_send_email(fn_args, tenant_id=tid)
                elif fn_name == "invite_visio":
                    result = await _tool_invite_visio(fn_args, tenant_id=tid)
                elif fn_name == "call_contact":
                    result = await _tool_call_contact(fn_args, tenant_id=tid, session_id=req.session_id)
                elif fn_name == "search_web":
                    result = await _tool_search_web(fn_args, tenant_id=tid)
                elif fn_name == "search_places":
                    result = await _tool_search_places(fn_args, tenant_id=tid)
                elif fn_name == "get_page_info":
                    result = await _tool_get_page_info(fn_args, tenant_id=tid)
                elif fn_name == "request_payment":
                    result = await _tool_request_payment(fn_args, tenant_id=tid)
                elif fn_name == "get_player_stats":
                    result = await _tool_get_player_stats(tenant_id=tid)
                elif fn_name == "get_active_missions":
                    result = await _tool_get_active_missions(tenant_id=tid)
                elif fn_name == "get_badges":
                    result = await _tool_get_badges(tenant_id=tid)
                elif fn_name == "get_weather":
                    result = await _tool_get_weather(fn_args, tenant_id=tid)
                elif fn_name == "get_news":
                    result = await _tool_get_news(fn_args)
                elif fn_name == "search_flights":
                    result = await _tool_search_flights(fn_args, tenant_id=tid)
                elif fn_name == "search_hotels":
                    result = await _tool_search_hotels(fn_args, tenant_id=tid)
                elif fn_name == "book_restaurant":
                    result = await _tool_book_restaurant(fn_args, tenant_id=tid)
                # --- Secretary tools ---
                elif fn_name == "get_documents_summary":
                    result = _tool_secretary_summary(tid)
                elif fn_name == "get_budget_analysis":
                    result = _tool_secretary_budget(tid)
                elif fn_name == "check_affordability":
                    result = _tool_secretary_afford(tid, fn_args)
                elif fn_name == "add_expense":
                    result = _tool_secretary_add_expense(tid, fn_args)
                elif fn_name == "get_reminders":
                    result = _tool_secretary_reminders(tid)
                elif fn_name == "add_reminder":
                    result = _tool_secretary_add_reminder(tid, fn_args)
                elif fn_name == "search_documents":
                    result = _tool_secretary_search(tid, fn_args)
                elif fn_name == "list_folders":
                    result = _tool_secretary_folders(tid)
                else:
                    result = {"status": "error", "message": f"Fonction inconnue: {fn_name}"}

                tool_calls_made.append({"tool": fn_name, "result": result})
                messages.append({
                    "role": "tool",
                    "tool_call_id": tc.id,
                    "content": json.dumps(result, ensure_ascii=False),
                })

            # Anti-hallucination guardrail: if any tool FAILED, inject a system
            # message forcing the LLM to report the failure honestly.
            _ACTION_TOOLS = {"call_contact", "send_sms", "send_email", "alert_contacts",
                             "request_payment", "invite_visio", "generate_document", "create_instruction"}
            _failed_tools = [
                t for t in tool_calls_made
                if t.get("result", {}).get("status") == "error" and t["tool"] in _ACTION_TOOLS
            ]
            if _failed_tools:
                _fail_names = ", ".join(t["tool"] for t in _failed_tools)
                _fail_msgs = "; ".join(t["result"].get("message", "erreur") for t in _failed_tools)
                messages.append({
                    "role": "system",
                    "content": (
                        f"IMPORTANT — VERACITE OBLIGATOIRE: Les actions suivantes ont ECHOUE: {_fail_names}. "
                        f"Raison: {_fail_msgs}. "
                        "Tu DOIS informer le souscripteur de l'echec. Tu ne PEUX PAS pretendre "
                        "que l'action a reussi. Ne fabrique AUCUN contenu fictif (conversation, "
                        "message recu, nouvelle de la personne). Dis simplement ce qui n'a pas "
                        "fonctionne et propose une alternative."
                    ),
                })

            # Re-appel OpenAI avec les resultats des tools
            async with _openai_chat_semaphore:
                loop = asyncio.get_event_loop()
                response = await loop.run_in_executor(
                    None,
                    lambda: openai_client.chat.completions.create(
                        model=OPENAI_MODEL,
                        messages=messages,
                        max_tokens=500,
                        temperature=0.8,
                        timeout=30,
                        tools=chat_tools if chat_tools else None,
                    ),
                )
            choice = response.choices[0]

        luna_msg = choice.message.content or ""

        # Track OpenAI cost
        await _track_openai_cost(response, tid)

        # Fallback si reply vide apres tool call reussi
        if not luna_msg.strip() and tool_calls_made:
            last_tool = tool_calls_made[-1]
            if last_tool.get("result", {}).get("status") == "success":
                tool_name = last_tool["tool"]
                if tool_name == "call_contact":
                    luna_msg = last_tool["result"].get("message", "J'appelle ton contact maintenant !")
                elif tool_name == "send_sms":
                    luna_msg = last_tool["result"].get("message", "SMS envoye !")
                else:
                    luna_msg = last_tool["result"].get("message", "C'est fait !")

        # Log tools executes
        if tool_calls_made:
            logger.info(f"Chat tools executed: {[t['tool'] for t in tool_calls_made]}")

        # Legal compliance check on Luna's response
        if _safety_guardian:
            try:
                compliant, violations = _safety_guardian.check_legal_compliance(luna_msg)
                if not compliant:
                    logger.warning(f"Luna response legal violations: {violations}")
            except Exception:
                pass

        # Log event in chronological journal
        if mgr:
            try:
                mgr.log_event(
                    category="communication",
                    description=f"Chat: {req.message[:80]}... → Luna repond",
                    source="chat",
                )
            except Exception:
                pass

        messages.append({"role": "assistant", "content": luna_msg})

        # Persist Luna response to Redis
        if mgr:
            try:
                mgr.add_message(
                    conv_id=req.session_id,
                    role=MessageRole.LUNA,
                    content=luna_msg,
                    channel=Channel.APP,
                )
            except Exception as e:
                logger.warning(f"Redis store failed: {e}")

        # Auto-title conversation via AI (like ChatGPT)
        auto_title = None
        if mgr:
            try:
                meta = mgr.redis.get_conversation_meta(tid, req.session_id)
                if meta and not meta.get("summary") and int(meta.get("message_count", 0)) <= 2:
                    try:
                        title_resp = openai_client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[{
                                "role": "system",
                                "content": "Tu generes un titre court (4-6 mots max) pour une conversation. Pas de guillemets, pas de ponctuation finale. Juste le theme principal."
                            }, {
                                "role": "user",
                                "content": f"User: {req.message[:200]}\nLuna: {luna_msg[:200]}"
                            }],
                            max_tokens=20,
                            temperature=0.3,
                            timeout=5,
                        )
                        auto_title = title_resp.choices[0].message.content.strip().strip('"').strip("'")
                        if len(auto_title) > 50:
                            auto_title = auto_title[:50]
                    except Exception:
                        auto_title = req.message[:40].strip()
                        if len(req.message) > 40:
                            auto_title += "..."
                    meta["summary"] = auto_title
                    meta["last_activity"] = datetime.utcnow().isoformat()
                    mgr.redis.set_conversation_meta(tid, req.session_id, meta)
            except Exception:
                pass

        _gamify(tid, "chat_message")

        # Inclut les actions executees dans la reponse
        resp = {"response": luna_msg}
        if auto_title:
            resp["auto_title"] = auto_title
        if tool_calls_made:
            resp["actions"] = [{"tool": t["tool"], "status": t["result"].get("status", "unknown")} for t in tool_calls_made]
            # Si une visio a ete creee, inclure l'URL pour le frontend
            for t in tool_calls_made:
                if t["result"].get("visio_url"):
                    resp["visio_url"] = t["result"]["visio_url"]
                    break
            # Inclure les donnees structurees pour un affichage riche dans le frontend
            _rich_cards = []
            for t in tool_calls_made:
                r = t["result"]
                if r.get("status") != "success":
                    continue
                tn = t["tool"]
                if tn == "search_places" and r.get("places"):
                    for p in r["places"][:5]:
                        _rich_cards.append({
                            "type": "place",
                            "name": p.get("name", ""),
                            "address": p.get("address", ""),
                            "phone": p.get("phone", ""),
                            "rating": p.get("rating", ""),
                            "hours": p.get("hours", ""),
                            "url": p.get("website", ""),
                            "price_level": p.get("price_level", ""),
                        })
                elif tn == "search_web" and r.get("results"):
                    for sr in r["results"][:5]:
                        if isinstance(sr, dict):
                            _rich_cards.append({
                                "type": "web",
                                "title": sr.get("title", ""),
                                "snippet": sr.get("snippet", ""),
                                "url": sr.get("link", ""),
                            })
                        elif isinstance(sr, str):
                            # Results are plain strings — extract as single card
                            _rich_cards.append({
                                "type": "web",
                                "title": sr[:80] + ("..." if len(sr) > 80 else ""),
                                "snippet": sr,
                                "url": "",
                            })
                elif tn == "get_weather" and r.get("actuel"):
                    cur = r["actuel"]
                    _rich_cards.append({
                        "type": "weather",
                        "city": cur.get("ville", ""),
                        "temp": cur.get("temperature", ""),
                        "feels_like": cur.get("ressenti", ""),
                        "description": cur.get("description", ""),
                        "humidity": cur.get("humidite", ""),
                        "wind": cur.get("vent", ""),
                        "forecasts": [
                            {"date": f.get("date", ""), "max": f.get("max", ""), "min": f.get("min", ""), "description": f.get("description", "")}
                            for f in (r.get("previsions") or [])[:3]
                        ],
                    })
                elif tn == "get_news" and r.get("articles"):
                    for art in r["articles"][:5]:
                        _rich_cards.append({
                            "type": "news",
                            "title": art.get("titre", ""),
                            "url": art.get("lien", ""),
                            "source": art.get("source", ""),
                            "date": art.get("date", ""),
                            "snippet": art.get("resume", ""),
                        })
                elif tn == "search_flights" and r.get("flights"):
                    for fl in r["flights"][:5]:
                        _rich_cards.append({
                            "type": "flight",
                            "airline": fl.get("airline", ""),
                            "price": fl.get("price", ""),
                            "summary": fl.get("summary", ""),
                        })
                elif tn == "search_hotels" and r.get("hotels"):
                    for h in r["hotels"][:5]:
                        _rich_cards.append({
                            "type": "hotel",
                            "name": h.get("name", ""),
                            "stars": h.get("stars", ""),
                            "city": h.get("city", ""),
                            "price_per_night": h.get("price_per_night", ""),
                            "room_type": h.get("room_type", ""),
                        })
                elif tn == "book_restaurant" and r.get("places"):
                    for p in r["places"][:5]:
                        _rich_cards.append({
                            "type": "place",
                            "name": p.get("name", ""),
                            "address": p.get("address", ""),
                            "phone": p.get("phone", ""),
                            "rating": p.get("rating", ""),
                            "hours": p.get("hours", ""),
                            "url": p.get("website", ""),
                            "price_level": p.get("price_level", ""),
                        })
            if _rich_cards:
                resp["cards"] = _rich_cards
        return resp

    except openai.AuthenticationError:
        _openai_breaker.record_failure()
        logger.error("OPENAI AUTH ERROR - cle API invalide")
        _notify_admin_health("Cle OpenAI invalide - Luna ne peut plus repondre")
        return {"response": "Luna a un souci technique. L'equipe a ete prevenue."}
    except openai.RateLimitError as e:
        _openai_breaker.record_failure()
        err_body = getattr(e, 'body', {}) or {}
        err_type = err_body.get('error', {}).get('type', '') if isinstance(err_body, dict) else ''
        if 'insufficient_quota' in str(e) or err_type == 'insufficient_quota':
            logger.error("OPENAI QUOTA EPUISE - insufficient_quota")
            _notify_admin_health("Credit OpenAI epuise ! Luna est en panne. Recharge le compte: https://platform.openai.com/account/billing")
            return {"response": "Luna est temporairement en maintenance. L'equipe technique a ete alertee automatiquement."}
        else:
            logger.warning(f"OpenAI rate limit (temporaire): {e}")
            return {"response": "Petit embouteillage ! Renvoie ton message dans quelques secondes."}
    except openai.APIConnectionError:
        _openai_breaker.record_failure()
        return {"response": "Connexion perdue une seconde ! Renvoie ton message."}
    except Exception as e:
        _openai_breaker.record_failure()
        tenant_convs = conversations.get(str(getattr(request.state, "tenant_id", 1)), {})
        if req.session_id in tenant_convs and len(tenant_convs[req.session_id]) > 1:
            tenant_convs[req.session_id].pop()
        logger.error(f"Chat error: {type(e).__name__}: {e}")
        return {"response": "Luna a rencontre un petit probleme. Reessaie."}
    finally:
        # Only release if NOT streaming (streaming generator releases its own semaphore)
        if not getattr(req, "stream", False):
            _chat_semaphore.release()


@app.get("/api/greeting")
async def greeting(request: Request):
    try:
        tid = getattr(request.state, "tenant_id", 1)
        mode = request.query_params.get("mode", "compagnon")
        mgr = _get_tenant_manager(tid)
        messages = [{"role": "system", "content": _system_prompt_now()}]
        # Inject subscriber name for personalized greeting
        _sub_name = ""
        if mgr:
            try:
                profile = mgr.get_subscriber_profile()
                if profile and profile.first_name:
                    _sub_name = profile.first_name
                    messages.append({"role": "system", "content":
                        f"Tu salues {_sub_name}. Utilise son prenom dans ton message d'accueil. Sois chaleureuse et personnelle."
                    })
            except Exception:
                pass
        # Mode-aware context
        if mode == "secretaire":
            _sec_context = "Tu es en mode SECRETAIRE. Tu geres les documents, le budget, les rappels et la conciergerie."
            if _redis_client and tid:
                try:
                    from core.secretary.redis_ops import SecretaryRedisOps
                    _sec = SecretaryRedisOps(_redis_client, int(tid))
                    _overdue = _sec.get_overdue_reminders()
                    _summary = _sec.get_documents_summary()
                    if _overdue:
                        _sec_context += f"\nATTENTION: {len(_overdue)} rappel(s) en retard ! Mentionne-le dans ton accueil."
                    if _summary.get("pending_count", 0) > 0:
                        _sec_context += f"\n{_summary['pending_count']} document(s) en attente."
                except Exception:
                    pass
            messages.append({"role": "system", "content": _sec_context})
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=messages,
            max_tokens=300,
            temperature=0.8,
            timeout=30,
        )
        await _track_openai_cost(response)
        return {"response": response.choices[0].message.content}
    except Exception as e:
        return {"response": f"Salut ! Luna a un souci technique ({type(e).__name__}). Reessaie."}


# Replicas Tavus horaires — Luna change d'apparence selon le moment de la journee
# Meme persona (voix, personnalite, outils), seule l'apparence visuelle change
_LUNA_REPLICAS_SCHEDULE = [
    {"start": 6,  "end": 12, "replica_id": "r5dc7c7d0bcb", "name": "Gloria - Bright"},   # Matin lumineux
    {"start": 12, "end": 18, "replica_id": "r1e52660d3bf", "name": "Luna - Home"},        # Apres-midi naturel
    {"start": 18, "end": 21, "replica_id": "r3f427f43c9d", "name": "Gloria - Warm"},      # Crepuscule cosy (cabin cheminee)
    {"start": 21, "end": 6,  "replica_id": "r6fb41bf13b4", "name": "Katya"},               # Nuit (veilleuse, pas de fenetre)
]

def _get_luna_replica() -> str:
    """Retourne la replica_id correspondant a l'heure actuelle (heure France)."""
    hour = datetime.now(ZoneInfo("Europe/Paris")).hour
    for slot in _LUNA_REPLICAS_SCHEDULE:
        if slot["start"] < slot["end"]:
            if slot["start"] <= hour < slot["end"]:
                return slot["replica_id"]
        else:  # Creneau nuit (ex: 21h -> 6h)
            if hour >= slot["start"] or hour < slot["end"]:
                return slot["replica_id"]
    return _LUNA_REPLICAS_SCHEDULE[0]["replica_id"]


def _tenant_subscriber_first_name(tenant_id: int) -> str:
    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile and getattr(profile, "first_name", None):
                return profile.first_name
        except Exception:
            pass
    return _SUBSCRIBER_NAME if tenant_id == TENANT_ID else "toi"


_SIMLI_TOOLS = [
    {"type": "function", "function": {
        "name": "get_weather",
        "description": "Obtenir la meteo actuelle et les previsions pour le souscripteur",
        "parameters": {"type": "object", "properties": {
            "city": {"type": "string", "description": "Ville (optionnel, defaut: ville du profil)"}
        }, "required": []}
    }},
    {"type": "function", "function": {
        "name": "send_sms",
        "description": "Envoyer un SMS a un contact de confiance du souscripteur",
        "parameters": {"type": "object", "properties": {
            "contact_name": {"type": "string", "description": "Prenom du contact"},
            "message": {"type": "string", "description": "Texte du SMS"}
        }, "required": ["contact_name", "message"]}
    }},
    {"type": "function", "function": {
        "name": "create_note",
        "description": "Creer une note ou un memo pour le souscripteur",
        "parameters": {"type": "object", "properties": {
            "content": {"type": "string", "description": "Contenu de la note"},
            "tags": {"type": "array", "items": {"type": "string"}, "description": "Tags optionnels"}
        }, "required": ["content"]}
    }},
    {"type": "function", "function": {
        "name": "create_instruction",
        "description": "Creer un rappel, une instruction planifiee ou une tache recurrente",
        "parameters": {"type": "object", "properties": {
            "instruction": {"type": "string", "description": "Instruction a executer"},
            "trigger": {"type": "string", "description": "Quand executer (ex: 'demain 9h', 'tous les lundis')"}
        }, "required": ["instruction"]}
    }},
    {"type": "function", "function": {
        "name": "get_contacts",
        "description": "Lister les contacts de confiance du souscripteur",
        "parameters": {"type": "object", "properties": {}, "required": []}
    }},
    {"type": "function", "function": {
        "name": "alert_contacts",
        "description": "Alerter les contacts de confiance en cas d'urgence",
        "parameters": {"type": "object", "properties": {
            "reason": {"type": "string", "description": "Raison de l'alerte"}
        }, "required": ["reason"]}
    }},
    {"type": "function", "function": {
        "name": "get_news",
        "description": "Obtenir les actualites du jour",
        "parameters": {"type": "object", "properties": {
            "topic": {"type": "string", "description": "Sujet (optionnel)"}
        }, "required": []}
    }},
    {"type": "function", "function": {
        "name": "search_web",
        "description": "Rechercher des informations sur internet",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Requete de recherche"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "generate_document",
        "description": "Generer un document ou courrier pour le souscripteur",
        "parameters": {"type": "object", "properties": {
            "type": {"type": "string", "description": "Type de document"},
            "content": {"type": "string", "description": "Contenu ou instructions"}
        }, "required": ["type", "content"]}
    }},
    {"type": "function", "function": {
        "name": "get_player_stats",
        "description": "Obtenir les statistiques et le niveau de gamification du souscripteur",
        "parameters": {"type": "object", "properties": {}, "required": []}
    }},
    {"type": "function", "function": {
        "name": "search_places",
        "description": "Rechercher des restaurants, pharmacies, hopitaux ou autres lieux proches",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Type de lieu recherche"},
            "location": {"type": "string", "description": "Localisation (optionnel)"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "report_observation",
        "description": "Signaler une observation sur l'etat ou l'environnement du souscripteur",
        "parameters": {"type": "object", "properties": {
            "observation": {"type": "string"},
            "severity": {"type": "string", "enum": ["low", "medium", "high"]}
        }, "required": ["observation", "severity"]}
    }},
    {"type": "function", "function": {
        "name": "call_contact",
        "description": (
            "Appeler un contact de confiance ou un numero de telephone par la voix via Twilio. "
            "Utilise quand le souscripteur dit 'appelle', 'telephone a', 'passe un coup de fil'. "
            "Ne jamais utiliser pour les numeros d'urgence (17, 18, 112). "
            "Demander confirmation avant d'appeler."
        ),
        "parameters": {"type": "object", "properties": {
            "contact_name": {"type": "string", "description": "Prenom du contact de confiance"},
            "message": {"type": "string", "description": "Message a transmettre lors de l'appel"},
            "phone_number": {"type": "string", "description": "Numero direct si ce n'est pas un contact enregistre"}
        }, "required": ["contact_name", "message"]}
    }},
    {"type": "function", "function": {
        "name": "send_email",
        "description": "Envoyer un email a un contact du souscripteur",
        "parameters": {"type": "object", "properties": {
            "contact_name": {"type": "string", "description": "Prenom ou nom du destinataire"},
            "subject": {"type": "string", "description": "Sujet de l'email"},
            "body": {"type": "string", "description": "Corps du message"}
        }, "required": ["contact_name", "subject", "body"]}
    }},
    {"type": "function", "function": {
        "name": "search_flights",
        "description": "Rechercher des vols disponibles entre deux villes",
        "parameters": {"type": "object", "properties": {
            "origin": {"type": "string", "description": "Ville de depart"},
            "destination": {"type": "string", "description": "Ville d'arrivee"},
            "date": {"type": "string", "description": "Date du vol (YYYY-MM-DD)"}
        }, "required": ["origin", "destination", "date"]}
    }},
    {"type": "function", "function": {
        "name": "search_hotels",
        "description": "Rechercher des hotels disponibles dans une ville",
        "parameters": {"type": "object", "properties": {
            "city": {"type": "string", "description": "Ville"},
            "checkin": {"type": "string", "description": "Date d'arrivee (YYYY-MM-DD)"},
            "checkout": {"type": "string", "description": "Date de depart (YYYY-MM-DD)"}
        }, "required": ["city", "checkin", "checkout"]}
    }},
    {"type": "function", "function": {
        "name": "book_restaurant",
        "description": "Rechercher et proposer des restaurants a reserver",
        "parameters": {"type": "object", "properties": {
            "query": {"type": "string", "description": "Type de cuisine ou nom du restaurant"},
            "location": {"type": "string", "description": "Ville ou adresse"}
        }, "required": ["query"]}
    }},
    {"type": "function", "function": {
        "name": "invite_visio",
        "description": "Inviter un contact a rejoindre cette visio par SMS",
        "parameters": {"type": "object", "properties": {
            "contact_name": {"type": "string", "description": "Prenom du contact a inviter"}
        }, "required": ["contact_name"]}
    }},
    {"type": "function", "function": {
        "name": "send_conclusions",
        "description": (
            "Rédiger et envoyer un compte-rendu professionnel ou des conclusions "
            "a tous les participants de la visio (par SMS ou email). "
            "A utiliser quand le souscripteur demande d'envoyer le résumé, les conclusions, "
            "le compte-rendu, les points clés, ou les décisions prises pendant la réunion. "
            "Génère un document structuré (titre, points discutés, décisions, actions) et l'envoie a chaque participant."
        ),
        "parameters": {"type": "object", "properties": {
            "subject": {"type": "string", "description": "Sujet / titre du compte-rendu"},
            "conclusions": {"type": "string", "description": "Contenu complet des conclusions, points discutés, décisions et actions a suivre"},
            "recipients": {"type": "array", "items": {"type": "string"}, "description": "Prénoms des destinataires (optionnel — tous les participants si vide)"}
        }, "required": ["subject", "conclusions"]}
    }},
    {"type": "function", "function": {
        "name": "get_vision_context",
        "description": (
            "Obtenir ce que la camera du souscripteur capture en ce moment. "
            "A appeler uniquement si le souscripteur demande ce que tu vois, "
            "ce qu'il y a dans la piece, ou une description de son environnement."
        ),
        "parameters": {"type": "object", "properties": {}, "required": []}
    }},
]


async def _build_realtime_context(tenant_id: int = 0) -> str:
    """Construit un bloc de données temps réel (météo + heure + news) à injecter
    dans le system prompt. Luna peut répondre sans appel outil."""
    import httpx as _httpx_rt
    lines = []

    # Heure / date actuelle
    try:
        tz = ZoneInfo("Europe/Paris")
        now = datetime.now(tz)
        jours = ["lundi", "mardi", "mercredi", "jeudi", "vendredi", "samedi", "dimanche"]
        mois = ["janvier", "février", "mars", "avril", "mai", "juin",
                "juillet", "août", "septembre", "octobre", "novembre", "décembre"]
        date_str = f"{jours[now.weekday()]} {now.day} {mois[now.month - 1]} {now.year}"
        lines.append(f"Date et heure actuelles : {date_str}, {now.strftime('%H:%M')} (heure de Paris)")
    except Exception:
        pass

    # Météo (wttr.in, gratuit)
    city = "Paris"
    if _redis_client and tenant_id:
        try:
            geo_raw = _redis_client.client.get(f"luna:{tenant_id}:geolocation")
            if geo_raw:
                geo = json.loads(geo_raw.decode() if isinstance(geo_raw, bytes) else geo_raw)
                city = geo.get("city", "Paris") or "Paris"
        except Exception:
            pass
    try:
        async with _httpx_rt.AsyncClient(timeout=8) as _cli:
            r = await _cli.get(
                f"https://wttr.in/{city}?format=j1&lang=fr",
                headers={"User-Agent": "Luna/2.2"},
            )
            if r.status_code == 200:
                wd = r.json()
                cur = wd.get("current_condition", [{}])[0]
                desc = cur.get("lang_fr", [{}])[0].get("value", "") or cur.get("weatherDesc", [{}])[0].get("value", "")
                temp = cur.get("temp_C", "?")
                ressenti = cur.get("FeelsLikeC", "?")
                vent = cur.get("windspeedKmph", "?")
                lines.append(
                    f"Météo actuelle à {city} : {temp}°C (ressenti {ressenti}°C), {desc}, vent {vent} km/h."
                )
    except Exception:
        pass

    # Actualités (RSS Le Monde, gratuit)
    try:
        async with _httpx_rt.AsyncClient(timeout=8) as _cli:
            r = await _cli.get(
                "https://www.lemonde.fr/rss/une.xml",
                headers={"User-Agent": "Luna/2.2"},
            )
            if r.status_code == 200:
                import re as _re_rt
                titles = _re_rt.findall(r"<title><!\[CDATA\[(.+?)\]\]></title>", r.text)
                titles = [t for t in titles if t not in ("Le Monde", "lemonde.fr")][:5]
                if titles:
                    lines.append("Titres d'actualité du moment (Le Monde) :\n" +
                                 "\n".join(f"  • {t}" for t in titles))
    except Exception:
        pass

    if not lines:
        return ""
    return "\n\n=== DONNÉES EN TEMPS RÉEL (déjà disponibles, pas besoin d'appel outil) ===\n" + "\n".join(lines)


async def _start_simli_visio(tenant_id: int, subscriber_name: str) -> tuple:
    """Demarre une session Simli. Retourne (ok, payload_ou_erreur).
    Utilise /auto/start/configurable (API Simli v2).
    """
    api_key = os.getenv("SIMLI_API_KEY", "")
    face_id = os.getenv("SIMLI_FACE_ID", "")
    openai_key = os.getenv("OPENAI_API_KEY", "")
    elevenlabs_key = os.getenv("ELEVENLABS_API_KEY", "")
    # voiceId: ElevenLabs Charlotte FR si cle dispo, sinon Cartesia voix FR feminine
    tts_provider = "ElevenLabs" if elevenlabs_key else "Cartesia"
    # Charlotte FR (ElevenLabs) si cle dispo, sinon voix Cartesia par defaut (multilingual)
    # La voix par defaut Simli/Cartesia fonctionne avec language:"fr"
    # (voix custom Cartesia FR requiert plan payant Cartesia)
    # f9836c6e = "Helpful Woman" Cartesia, voix feminine multilingue (supporte fr)
    voice_id = os.getenv("SIMLI_VOICE_ID",
        "XB0fDUnXU5powFXDhCwa" if elevenlabs_key else "f9836c6e-a0bd-460e-9d3c-f7299fa60f94")
    tts_api_key = elevenlabs_key if elevenlabs_key else None

    if not api_key or not face_id:
        return False, {"error": "Simli non configure (SIMLI_API_KEY / SIMLI_FACE_ID manquants)"}

    # Contexte riche : profil, contacts, meteo, capacites
    try:
        ctx = build_tavus_context(
            subscriber_name=subscriber_name,
            memory_manager=_memory_manager,
        )
    except Exception:
        ctx = f"Tu es Luna, compagnon IA feminin YAWatch. Tu parles avec {subscriber_name}."

    # Pré-fetcher météo + actualités + date pour éviter les hallucinations
    try:
        realtime_ctx = await _build_realtime_context(tenant_id)
    except Exception:
        realtime_ctx = ""

    # Forcer le français en tete du prompt (avant le contexte riche)
    french_prefix = (
        "INSTRUCTIONS ABSOLUES : Tu t'appelles Luna. "
        "Tu parles EXCLUSIVEMENT en français de France (pas québécois, pas anglais). "
        "Chaque réponse doit être en français, quelle que soit la langue utilisée par l'utilisateur. "
        "Tu es une femme, ton ton est chaleureux et bienveillant.\n"
        "Quand tu reçois un message '[Vision caméra] ...', c'est une description automatique "
        "de l'environnement visuel de l'utilisateur. Utilise-la naturellement dans la conversation "
        "si pertinent, sans mentionner la mécanique technique.\n\n"
        "RÈGLE ANTI-HALLUCINATION ABSOLUE : Tu ne dois JAMAIS inventer, supposer ou fabriquer "
        "des informations sur la météo, l'actualité, les prix, les horaires ou tout fait vérifiable. "
        "Si une information n'est pas dans ton contexte, dis-le honnêtement : "
        "'Je n'ai pas cette information en ce moment.' "
        "Pour la météo et les actualités, utilise les données de la section DONNÉES EN TEMPS RÉEL "
        "ci-dessous — elles sont fraîches et fiables.\n\n"
    )
    ctx = french_prefix + ctx + realtime_ctx

    payload = {
        "faceId": face_id,
        "ttsProvider": tts_provider,
        "voiceId": voice_id,
        "language": "fr",
        "systemPrompt": ctx,
        "firstMessage": f"Bonjour {subscriber_name} ! Je suis ravie de te voir. Comment je peux t'aider aujourd'hui ?",
        "customLLMConfig": {
            "model": "gpt-4o-mini",
            "baseURL": "https://api.openai.com/v1",
            "llmAPIKey": openai_key,
        },
        "tools": _SIMLI_TOOLS,
        "maxSessionLength": 3600,
        "maxIdleTime": 300,
    }
    if tts_api_key:
        payload["ttsAPIKey"] = tts_api_key

    try:
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.post(
                "https://api.simli.ai/auto/start/configurable",
                json=payload,
                headers={
                    "Content-Type": "application/json",
                    "x-simli-api-key": api_key,
                },
            )
        data = resp.json()
        if resp.status_code != 200:
            logger.error(f"Simli auto/start error {resp.status_code}: {data}")
            return False, {"error": data.get("message", data.get("detail", "Simli indisponible"))}

        conv_url = data.get("roomUrl") or data.get("room_url") or ""
        session_id = data.get("sessionId") or data.get("session_id") or ""
        if not conv_url:
            logger.error(f"Simli: no roomUrl in response: {data}")
            return False, {"error": "Simli: URL de session manquante"}
        return True, {
            "conversation_url": conv_url,
            "conversation_id": session_id or f"simli_{tenant_id}_{int(time.time())}",
            "provider": "simli",
        }
    except Exception as e:
        logger.error(f"Simli start error: {e}")
        return False, {"error": f"Simli erreur: {str(e)}"}


@app.post("/api/call")
async def start_call(request: Request):
    """Cree un appel video — Tavus en priorite, Simli en repli."""
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return JSONResponse(status_code=403, content={"error": "Service non disponible"})
    tid = getattr(request.state, "tenant_id", 1)
    _plan = getattr(request.state, "plan", "essentiel") or "essentiel"
    budget_err = await _check_budget_guard(tid, _plan)
    if budget_err:
        return JSONResponse(status_code=429, content={"error": budget_err})

    try:
        body = await request.json()
    except Exception:
        body = {}
    duration_min = body.get("duration", 0)
    if not duration_min or duration_min <= 0:
        duration_min = int(os.getenv("VISIO_MAX_DURATION", "60"))
    duration_min = min(int(duration_min), 240)
    visio_max = duration_min * 60

    sub_name = _tenant_subscriber_first_name(tid)

    if tavus_client and tavus_client.is_configured:
        try:
            realtime_ctx_tavus = await _build_realtime_context(tid)
        except Exception:
            realtime_ctx_tavus = ""
        context = build_tavus_context(
            subscriber_name=sub_name,
            memory_manager=tavus_client.memory,
        ) + realtime_ctx_tavus
        replica_id = _get_luna_replica()
        success, data = await tavus_client.create_conversation(
            tenant_id=tid,
            custom_greeting=f"Salut {sub_name} ! Ravie de te voir. Comment je peux t'aider ?",
            context=context,
            max_duration=visio_max,
            callback_url=TAVUS_CALLBACK_URL if TAVUS_CALLBACK_URL else None,
            replica_id=replica_id,
        )
        if success:
            _gamify(tid, "voice_call")
            return {
                "conversation_url": data["conversation_url"],
                "conversation_id": data["conversation_id"],
                "provider": "tavus",
            }
        logger.warning(f"Tavus indisponible, repli Simli: {data.get('error', 'unknown')}")

    ok, payload = await _start_simli_visio(tid, sub_name)
    if ok:
        _gamify(tid, "voice_call")
        return payload

    err = payload.get("error", "Visio non disponible")
    if LUNA_MODE != "full":
        return JSONResponse(status_code=503, content={
            "error": err,
            "mode": LUNA_MODE,
            "message": "Visio indisponible. Configurez Tavus ou Simli en mode Full.",
        })
    return JSONResponse(status_code=503, content={"error": err, "fallback": "simli_failed"})


@app.post("/api/call/end")
async def end_call(request: Request):
    """Termine une conversation Tavus (stoppe la facturation)."""
    if not tavus_client:
        return JSONResponse(status_code=503, content={"error": "Visio non disponible"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Body JSON invalide"})
    conv_id = body.get("conversation_id", "").strip()
    if not conv_id:
        return JSONResponse(status_code=400, content={"error": "conversation_id manquant"})
    try:
        ended = await tavus_client.end_conversation(conv_id)
        return {"success": True, "ended": ended}
    except Exception as e:
        logger.error(f"Erreur end_call: {e}")
        return {"success": False, "error": str(e)}


# =========================================================================
# SIMLI — Avatar cinematique avec lip-sync WebRTC
# =========================================================================

@app.websocket("/ws/simli/{session_id}")
async def ws_simli(websocket: WebSocket, session_id: str):
    """WebSocket Simli : STT -> LLM -> TTS -> PCM16 lip-sync."""
    if not _SIMLI_AVAILABLE or not handle_simli_session:
        await websocket.accept()
        await websocket.close(code=4003, reason="Simli not available")
        return
    await handle_simli_session(websocket, session_id)


@app.get("/api/config/simli")
async def config_simli():
    """Config Simli."""
    return {"enabled": bool(os.getenv("SIMLI_API_KEY", ""))}


@app.post("/api/simli/start")
async def simli_start(request: Request):
    """Demarre une session Simli E2E — repli visio quand Tavus indisponible."""
    tenant_id = getattr(request.state, "tenant_id", 1)
    sub_name = _tenant_subscriber_first_name(tenant_id)
    ok, payload = await _start_simli_visio(tenant_id, sub_name)
    if ok:
        return {
            "conversation_url": payload["conversation_url"],
            "session_id": payload["conversation_id"],
            "provider": "simli",
        }
    return JSONResponse(status_code=503, content=payload)


@app.get("/formulaires")
async def formulaires_page():
    """Page FormFiller — remplissage intelligent de formulaires PDF."""
    fpath = os.path.join(STATIC_DIR, "formulaires.html")
    if os.path.isfile(fpath):
        return FileResponse(fpath, headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        })
    return JSONResponse(status_code=404, content={"error": "Formulaires non disponible"})


@app.get("/join/{token}")
async def join_visio_page(token: str):
    """Page d'avertissement legal avant de rejoindre une visio partagee."""
    if not _redis_client:
        return HTMLResponse("<h2>Lien invalide ou expire.</h2>", status_code=410)
    raw = _redis_client.client.get(f"luna:join:{token}")
    if not raw:
        return HTMLResponse("""<!doctype html><html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Lien expire</title>
<style>body{background:#0a0a1a;color:#fff;font-family:system-ui,sans-serif;
display:flex;align-items:center;justify-content:center;height:100vh;margin:0;text-align:center;padding:24px;}
h2{color:#ff6b6b;} p{color:#aaa;font-size:14px;margin-top:12px;}</style>
</head><body><div><h2>Lien expire ou invalide</h2>
<p>Ce lien de visio n'est plus valide.<br>Demandez a votre contact de vous renvoyer une invitation.</p>
</div></body></html>""", status_code=410)

    import json as _json
    data = _json.loads(raw)
    room_url = data.get("room_url", "")
    import html as _html_mod
    subscriber = _html_mod.escape(data.get("subscriber_name", "votre contact"))
    if not room_url:
        return HTMLResponse("<h2>Lien invalide.</h2>", status_code=410)

    html = f"""<!doctype html>
<html lang="fr">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0, viewport-fit=cover" />
  <title>Rejoindre la visio Luna</title>
  <style>
    * {{ box-sizing: border-box; margin: 0; padding: 0; }}
    body {{ background: #0a0a1a; color: #e0e0e0; font-family: system-ui, -apple-system, sans-serif;
      min-height: 100vh; display: flex; align-items: center; justify-content: center; padding: 24px; }}
    .card {{ background: #0f0f2a; border: 1px solid #2a2a5a; border-radius: 20px;
      max-width: 420px; width: 100%; padding: 32px 28px; text-align: center; }}
    .logo {{ width: 64px; height: 64px; border-radius: 50%; background: linear-gradient(135deg, #7c8cf8, #a78bfa);
      display: flex; align-items: center; justify-content: center; font-size: 28px; color: #fff;
      font-weight: 700; margin: 0 auto 20px; box-shadow: 0 0 30px rgba(124,140,248,0.3); }}
    h1 {{ font-size: 22px; color: #fff; margin-bottom: 6px; }}
    .invite-msg {{ font-size: 14px; color: #888; margin-bottom: 24px; }}
    .warning-box {{ background: rgba(255,160,0,0.08); border: 1px solid rgba(255,160,0,0.3);
      border-radius: 14px; padding: 18px 16px; margin-bottom: 24px; text-align: left; }}
    .warning-box h3 {{ font-size: 13px; color: #ffb020; text-transform: uppercase;
      letter-spacing: 1px; margin-bottom: 12px; display: flex; align-items: center; gap: 6px; }}
    .warning-box ul {{ list-style: none; padding: 0; }}
    .warning-box li {{ font-size: 13px; color: #ccc; padding: 5px 0;
      padding-left: 20px; position: relative; line-height: 1.5; }}
    .warning-box li::before {{ content: "•"; position: absolute; left: 4px; color: #ffb020; }}
    .legal-note {{ font-size: 12px; color: #666; line-height: 1.6; margin-bottom: 24px; }}
    .legal-note strong {{ color: #e94560; }}
    .btn-join {{ width: 100%; background: linear-gradient(135deg, #7c8cf8, #a78bfa);
      color: #fff; border: none; border-radius: 50px; padding: 16px 24px;
      font-size: 17px; font-weight: 700; cursor: pointer; transition: all 0.15s;
      box-shadow: 0 4px 20px rgba(124,140,248,0.35); }}
    .btn-join:hover {{ transform: scale(1.02); box-shadow: 0 6px 28px rgba(124,140,248,0.5); }}
    .btn-decline {{ display: block; margin-top: 14px; color: #555; font-size: 13px;
      text-decoration: none; cursor: pointer; background: none; border: none; width: 100%; }}
    .btn-decline:hover {{ color: #888; }}
  </style>
</head>
<body>
  <div class="card">
    <div class="logo">L</div>
    <h1>Rejoindre la visio</h1>
    <p class="invite-msg"><strong style="color:#a78bfa">{subscriber}</strong> vous invite à une session vidéo avec Luna IA.</p>

    <div class="warning-box">
      <h3>⚠️ Avertissement important</h3>
      <ul>
        <li>Luna est une <strong style="color:#fff">intelligence artificielle</strong>, pas un professionnel de santé, conseiller juridique ou financier.</li>
        <li>Toute demande de contenu <strong style="color:#fff">illégal, offensant ou abusif</strong> est interdite.</li>
        <li>Les abus sont <strong style="color:#fff">tracés et conservés</strong>.</li>
        <li>En cas d'infraction, des <strong style="color:#fff">poursuites judiciaires</strong> pourront être engagées.</li>
      </ul>
    </div>

    <p class="legal-note">
      En cliquant sur "Rejoindre", vous reconnaissez avoir lu et accepté ces conditions.
      <strong>Toute utilisation abusive engage votre responsabilité pénale et civile</strong> conformément au droit français.
    </p>

    <button class="btn-join" onclick="joinNow()">Rejoindre la visio →</button>
    <button class="btn-decline" onclick="window.close()">Refuser et fermer</button>
  </div>
  <script>
    function joinNow() {{
      var btn = document.querySelector('.btn-join');
      btn.textContent = 'Connexion...';
      btn.disabled = true;
      window.location.href = {repr(room_url)};
    }}
  </script>
</body>
</html>"""
    return HTMLResponse(content=html, headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
    })


@app.post("/api/call/invite-guest")
async def invite_guest_to_visio(request: Request):
    """Invite un contact à rejoindre la visio en cours via SMS + page disclaimer."""
    tid = getattr(request.state, "tenant_id", 1)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Corps JSON invalide"})

    room_url = body.get("room_url", "")
    phone = body.get("phone", "")
    contact_name = body.get("contact_name", "")
    conv_id = body.get("conversation_id", "")

    if not room_url or not phone:
        return JSONResponse(status_code=400, content={"error": "room_url et phone requis"})

    if not sms_client or not sms_client.is_configured:
        return JSONResponse(status_code=503, content={"error": "SMS non disponible"})

    # Créer un token court (TTL 4h) pointant vers la room
    import secrets as _sec, json as _json
    token = _sec.token_hex(16)
    sub_name = _tenant_subscriber_first_name(tid)
    payload = {"room_url": room_url, "subscriber_name": sub_name, "conv_id": conv_id}
    if _redis_client:
        _redis_client.client.setex(f"luna:join:{token}", 4 * 3600, _json.dumps(payload))

    base_url = os.getenv("BASE_URL", "https://luna-beta-674304336025.europe-west1.run.app")
    join_url = f"{base_url}/join/{token}"

    msg = (
        f"[Luna] {sub_name} t'invite en visio avec Luna 🎥\n"
        f"Rejoins ici : {join_url}\n"
        f"(Lien valide 4h — utilisation responsable requise)"
    )
    from integrations.twilio.sms_client import TwilioSMSClient
    normalized = TwilioSMSClient.normalize_phone(phone)
    ok, details = _tracked_sms_send(normalized, msg, label=f"Invitation visio a {contact_name or phone}")

    if ok:
        # Mémoriser le participant pour send_conclusions
        if _redis_client and conv_id:
            try:
                import json as _jpart
                pkey = f"luna:conv:{conv_id}:participants"
                raw = _redis_client.client.get(pkey)
                parts = _jpart.loads(raw.decode() if isinstance(raw, bytes) else raw) if raw else []
                entry = {"name": contact_name or phone, "phone": normalized}
                if not any(p.get("phone") == normalized for p in parts):
                    parts.append(entry)
                _redis_client.client.setex(pkey, 6 * 3600, _jpart.dumps(parts))
            except Exception as _pe:
                logger.warning(f"Participant store error: {_pe}")
        return {"ok": True, "join_url": join_url, "contact": contact_name or phone}
    return JSONResponse(status_code=500, content={"error": "Echec envoi SMS", "detail": str(details)})


@app.post("/api/call/create-join-link")
async def create_join_link(request: Request):
    """Crée un lien d'invitation visio (sans SMS) — à partager manuellement."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Corps JSON invalide"})

    room_url = body.get("room_url", "")
    conv_id = body.get("conversation_id", "")
    tid = getattr(request.state, "tenant_id", 1)

    if not room_url:
        return JSONResponse(status_code=400, content={"error": "room_url requis"})

    import secrets as _sec, json as _json
    token = _sec.token_hex(16)
    sub_name = _tenant_subscriber_first_name(tid)
    payload = {"room_url": room_url, "subscriber_name": sub_name, "conv_id": conv_id}
    if _redis_client:
        _redis_client.client.setex(f"luna:join:{token}", 4 * 3600, _json.dumps(payload))

    base_url = os.getenv("BASE_URL", "https://luna-beta-674304336025.europe-west1.run.app")
    join_url = f"{base_url}/join/{token}"
    return {"ok": True, "join_url": join_url}


@app.post("/api/visio/perception")
async def visio_perception_frame(request: Request):
    """Analyse une frame caméra pendant une session visio.

    Utilise PerceptionDetector + SceneAnalyzer (structuré : personnes, postures, objets).
    Pas de vérification is_perception_enabled — la visio est un consentement explicite.
    Retourne changed:true uniquement si la scène a changé depuis la dernière analyse.
    """
    client_ip = (request.client.host if request.client else None) or "unknown"
    now = time.time()
    last = _vision_last_call.get(client_ip, 0)
    if now - last < 8.0:
        return JSONResponse(status_code=429, content={
            "ok": False, "error": "Trop rapide", "retry_after": round(8.0 - (now - last), 1)
        })
    _vision_last_call[client_ip] = now

    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Corps JSON invalide"})

    image_data = body.get("image", "")
    if not image_data:
        return JSONResponse(status_code=400, content={"ok": False, "error": "image base64 requise"})

    # Extraire le base64 pur (sans préfixe data:image/...)
    image_b64 = image_data.split(",", 1)[1] if "," in image_data else image_data

    loop = asyncio.get_event_loop()

    # --- Chemin principal : PerceptionDetector (structuré) ---
    if _perception_detector and getattr(_perception_detector, '_initialized', False):
        frame_analysis = await loop.run_in_executor(
            None, _perception_detector.analyze_frame_b64, image_b64
        )
        if not frame_analysis:
            return JSONResponse(status_code=500, content={"ok": False, "error": "Analyse échouée"})

        scene_state = _perception_analyzer.analyze(frame_analysis) if _perception_analyzer else None
        description = scene_state.scene_description if scene_state else (
            f"{frame_analysis.persons_count} personne(s)." +
            (f" Objets: {', '.join(frame_analysis.objects[:4])}." if frame_analysis.objects else "")
        )
        persons = frame_analysis.persons_count
        objects = frame_analysis.objects[:6]
        posture = frame_analysis.person_postures[0].value if frame_analysis.person_postures else "unknown"
        inference_ms = round(frame_analysis.inference_time_ms)

        # Log anomalies concern → note mémoire
        if scene_state and _memory_manager:
            for abn in scene_state.abnormalities:
                if abn.get("severity") == "concern":
                    logger.warning(f"Visio perception concern: {abn['description']}")
                    try:
                        _memory_manager.add_note(
                            content=f"[Visio perception] {abn['description']}",
                            context="visio_perception",
                            tags=["perception", "visio", abn.get("type", ""), "concern"],
                        )
                    except Exception:
                        pass

    # --- Fallback : appel OpenAI Vision direct si module non initialisé ---
    elif openai_client:
        try:
            resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": [
                    {"type": "text", "text": (
                        "Décris en 1-2 phrases ce que tu vois. Personnes, postures, objets. "
                        "Bref, factuel, sans identités. En français."
                    )},
                    {"type": "image_url", "image_url": {"url": image_data, "detail": "low"}}
                ]}],
                max_tokens=80
            ))
            description = resp.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Visio perception fallback error: {e}")
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Analyse échouée: {str(e)[:80]}"})
        persons, objects, posture, inference_ms = 0, [], "unknown", 0
    else:
        return JSONResponse(status_code=503, content={"ok": False, "error": "Module perception et OpenAI non disponibles"})

    # Détection de changement
    last_desc = _visio_scene_cache.get(client_ip, "")
    changed = (description != last_desc)
    _visio_scene_cache[client_ip] = description
    logger.info(f"Visio perception: {description[:80]} (changed={changed})")

    return {
        "ok": True,
        "description": description,
        "persons": persons,
        "objects": objects,
        "posture": posture,
        "changed": changed,
        "inference_ms": inference_ms,
    }


@app.post("/api/visio/notes")
async def visio_notes_generate(request: Request):
    """Génère un résumé structuré de la session visio pour prise de notes.

    Utilise les événements app-message capturés côté JS (tool calls, échos, types Simli)
    + le contexte vision + la durée pour produire des notes lisibles via GPT-4o-mini.
    Si auto_save=true, sauvegarde automatiquement en mémoire.
    """
    if not openai_client:
        return JSONResponse(status_code=503, content={"ok": False, "error": "OpenAI non configuré"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "error": "JSON invalide"})

    events = body.get("events", [])
    vision_ctx = body.get("vision_context", "")
    duration_min = body.get("duration_min", 0)
    conv_id = body.get("conversation_id", "")
    auto_save = body.get("auto_save", False)
    tid = getattr(request.state, "tenant_id", 1)

    # Construire le contexte pour GPT
    lines = []
    for ev in events[:80]:  # limiter à 80 événements
        t = ev.get("type", "")
        if t == "conversation.tool_call" and ev.get("tool"):
            args_str = ""
            try:
                args_str = str(json.loads(ev.get("args", "{}")))[:120]
            except Exception:
                pass
            lines.append(f"[Action Luna] {ev['tool']}: {args_str}")
        elif t == "tool_result" and ev.get("tool"):
            lines.append(f"[Résultat {ev['tool']}] {ev.get('text','')[:100]}")
        elif t in ("participant.joined", "participant.left"):
            lines.append(f"[Participant] {ev.get('text','')}")
        elif t == "vision.change" and ev.get("text"):
            lines.append(f"[Contexte visuel changé] {ev['text'][:120]}")
        elif t == "upload.analysis" and ev.get("text"):
            lines.append(f"[Document] {ev['text'][:160]}")
        elif ev.get("text"):
            # Paroles utilisateur et Luna, échecs système filtrés côté JS
            txt = ev["text"]
            if t == "user.speech":
                lines.append(f"[Utilisateur] {txt[:160]}")
            elif t == "luna.speech":
                lines.append(f"[Luna] {txt[:160]}")
            elif not any(txt.startswith(p) for p in ("[Vision caméra]", "[Alerte vision]", "[Instruction")):
                lines.append(f"[{t or 'échange'}] {txt[:160]}")

    context_str = "\n".join(lines) if lines else "(aucun événement capturé)"
    vision_str = f"\n[Contexte visuel] {vision_ctx}" if vision_ctx else ""

    prompt = (
        f"Voici les données d'une session visio avec Luna IA (durée : {duration_min} min).\n"
        f"Événements :\n{context_str}{vision_str}\n\n"
        "Génère des notes de réunion structurées en français, concises et utiles. "
        "Format : titre, points clés de la conversation (ce que l'utilisateur a dit et demandé), "
        "actions effectuées par Luna, observations (si contexte visuel disponible). "
        "[Utilisateur] = paroles transcrites de l'utilisateur (contenu principal). "
        "[Luna] = réponses de Luna. [Action Luna] = outils utilisés. [Participant] = entrées/sorties. "
        "Si peu de données, résume la durée et indique qu'il n'y a pas eu d'échanges détectés. "
        "Maximum 250 mots."
    )

    try:
        loop = asyncio.get_event_loop()
        resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=300,
            temperature=0.3,
        ))
        summary = resp.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Visio notes generate error: {e}")
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)[:100]})

    already_saved = False
    if auto_save and _memory_manager and summary:
        try:
            _memory_manager.add_note(
                content=f"[Visio {duration_min}min] {summary}",
                context="visio_notes",
                tags=["visio", "notes", "auto"],
            )
            already_saved = True
            logger.info(f"Visio notes auto-saved ({duration_min}min)")
        except Exception as e:
            logger.warning(f"Visio notes auto-save error: {e}")

    return {"ok": True, "summary": summary, "already_saved": already_saved}


@app.post("/api/visio/notes/save")
async def visio_notes_save(request: Request):
    """Sauvegarde un résumé visio en mémoire (action manuelle utilisateur)."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "error": "JSON invalide"})

    summary = body.get("summary", "").strip()
    duration_min = body.get("duration_min", 0)
    if not summary:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Résumé vide"})

    if not _memory_manager:
        return JSONResponse(status_code=503, content={"ok": False, "error": "Mémoire non disponible"})

    try:
        _memory_manager.add_note(
            content=f"[Visio {duration_min}min] {summary}",
            context="visio_notes",
            tags=["visio", "notes", "manuel"],
        )
        return {"ok": True}
    except Exception as e:
        logger.error(f"Visio notes save error: {e}")
        return JSONResponse(status_code=500, content={"ok": False, "error": str(e)[:100]})


@app.post("/api/visio/upload")
async def visio_upload(request: Request):
    """Analyse un document ou une image partagée pendant la visio.

    Supporte : images (JPEG/PNG/WebP/GIF), PDF, texte brut, CSV, Markdown.
    Retourne une analyse structurée que le JS injecte dans la conversation Luna.
    """
    if not openai_client:
        return JSONResponse(status_code=503, content={"ok": False, "error": "OpenAI non configuré"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "error": "JSON invalide"})

    import base64 as _b64
    import re as _re_up

    filename = (body.get("filename") or "document").strip()
    mime_type = (body.get("mime_type") or "").lower()
    data_url = body.get("data", "")
    if not data_url:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Données manquantes"})

    # Extraire le contenu brut (supprimer le préfixe data:...)
    if "," in data_url:
        raw_b64 = data_url.split(",", 1)[1]
    else:
        raw_b64 = data_url
    try:
        raw_bytes = _b64.b64decode(raw_b64)
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False, "error": "Base64 invalide"})

    # Limite de taille : 10 Mo
    if len(raw_bytes) > 10 * 1024 * 1024:
        return JSONResponse(status_code=413, content={"ok": False, "error": "Fichier trop volumineux (max 10 Mo)"})

    is_image = mime_type.startswith("image/") or _re_up.search(r"\.(jpg|jpeg|png|gif|webp)$", filename, _re_up.I)
    is_pdf = mime_type == "application/pdf" or filename.lower().endswith(".pdf")

    analysis = ""

    if is_image:
        # GPT-4o vision : analyse de l'image
        try:
            loop = asyncio.get_event_loop()
            resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                model="gpt-4o",
                messages=[{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": (
                            "Analyse cette image partagée par l'utilisateur pendant une visio avec Luna (son assistante IA). "
                            "Décris ce que tu vois, identifie les éléments importants (texte lisible, graphiques, personnes, objets, documents). "
                            "Sois précis et utile. Réponds en français."
                        )},
                        {"type": "image_url", "image_url": {"url": data_url, "detail": "high"}},
                    ]
                }],
                max_tokens=600,
            ))
            analysis = resp.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Visio upload image analysis error: {e}")
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Analyse image impossible: {str(e)[:80]}"})

    elif is_pdf:
        # PyMuPDF : extraction texte PDF
        try:
            import fitz as _fitz
            doc = _fitz.open(stream=raw_bytes, filetype="pdf")
            text_pages = []
            for i, page in enumerate(doc):
                if i >= 10:  # limiter à 10 pages
                    text_pages.append(f"[... {doc.page_count - 10} pages supplémentaires non incluses]")
                    break
                text_pages.append(page.get_text("text"))
            doc.close()
            full_text = "\n\n".join(text_pages).strip()
            if not full_text:
                return JSONResponse(status_code=422, content={"ok": False, "error": "PDF sans texte lisible (peut-être scanné)"})
            # Limiter à ~6000 tokens
            full_text = full_text[:12000]
        except ImportError:
            return JSONResponse(status_code=500, content={"ok": False, "error": "PDF non supporté sur ce serveur"})
        except Exception as e:
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Erreur lecture PDF: {str(e)[:80]}"})

        try:
            loop = asyncio.get_event_loop()
            resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": (
                        f"Voici le contenu d'un PDF nommé '{filename}' partagé par l'utilisateur pendant une visio.\n\n"
                        f"{full_text}\n\n"
                        "Résume les points clés de ce document de façon claire et structurée en français. "
                        "Identifie le type de document, son objet principal, les informations importantes. "
                        "Maximum 300 mots."
                    )
                }],
                max_tokens=500,
                temperature=0.3,
            ))
            analysis = resp.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Visio upload pdf analysis error: {e}")
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Analyse PDF impossible: {str(e)[:80]}"})

    else:
        # Fichier texte (txt, md, csv, docx text brut...)
        try:
            text_content = raw_bytes.decode("utf-8", errors="replace")[:12000]
        except Exception:
            text_content = raw_bytes[:12000].decode("latin-1", errors="replace")

        if not text_content.strip():
            return JSONResponse(status_code=422, content={"ok": False, "error": "Fichier vide ou illisible"})

        try:
            loop = asyncio.get_event_loop()
            resp = await loop.run_in_executor(None, lambda: openai_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{
                    "role": "user",
                    "content": (
                        f"Voici le contenu d'un fichier nommé '{filename}' partagé pendant une visio.\n\n"
                        f"{text_content}\n\n"
                        "Analyse ce contenu et résume les informations importantes en français. Maximum 300 mots."
                    )
                }],
                max_tokens=500,
                temperature=0.3,
            ))
            analysis = resp.choices[0].message.content.strip()
        except Exception as e:
            logger.error(f"Visio upload text analysis error: {e}")
            return JSONResponse(status_code=500, content={"ok": False, "error": f"Analyse impossible: {str(e)[:80]}"})

    logger.info(f"Visio upload analyzed: {filename} ({mime_type}) → {len(analysis)} chars")
    return {"ok": True, "analysis": analysis, "filename": filename}


@app.get("/clear-cache")
async def clear_cache_page():
    """Page standalone qui purge tout le cache SW + recharge. Jamais interceptee par le SW."""
    html = """<!doctype html><html><head><meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Mise a jour Luna...</title>
<style>body{background:#0a0a1a;color:#fff;font-family:system-ui,sans-serif;display:flex;align-items:center;justify-content:center;height:100vh;margin:0;text-align:center;}
.msg{font-size:18px;opacity:0.9;} .sub{font-size:13px;color:#888;margin-top:12px;}</style>
</head><body>
<div><div class="msg">Mise a jour en cours…</div><div class="sub" id="s">Vidage du cache SW</div></div>
<script>
var s=document.getElementById('s');
function done(){s.textContent='Rechargement…';setTimeout(function(){window.location.replace('/');},800);}
var steps=[];
if('caches' in window){steps.push(caches.keys().then(function(ns){return Promise.all(ns.map(function(n){return caches.delete(n);}));}).catch(function(){}));}
if('serviceWorker' in navigator){steps.push(navigator.serviceWorker.getRegistrations().then(function(regs){return Promise.all(regs.map(function(r){return r.unregister();}));}).catch(function(){}));}
Promise.all(steps).then(done).catch(done);
</script></body></html>"""
    return HTMLResponse(content=html, headers={
        "Cache-Control": "no-cache, no-store, must-revalidate",
        "Pragma": "no-cache",
        "Expires": "0",
    })


@app.get("/simli")
async def simli_page():
    """Page avatar Simli avec sequence cinematique."""
    simli_path = os.path.join(STATIC_DIR, "simli.html")
    if os.path.isfile(simli_path):
        return FileResponse(simli_path, headers={
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
        })
    return JSONResponse(status_code=404, content={"error": "Simli non disponible"})


# =========================================================================
# APPELS VOCAUX - Twilio Voice + OpenAI Realtime API
# =========================================================================

class VoiceCallRequest(BaseModel):
    phone: Optional[str] = None  # Numero a appeler (defaut: ADMIN_NUMBER)
    mission: Optional[str] = None  # Mission speciale pour Luna (ex: "prendre des nouvelles de Fred")
    max_duration: Optional[int] = None  # Duree max en secondes (min 60, max 600, defaut 180)
    greeting: Optional[str] = None  # Greeting personnalise

    @field_validator("max_duration")
    @classmethod
    def validate_duration(cls, v):
        if v is not None:
            v = max(60, min(600, v))  # clamp 1-10 min
        return v

# Stockage temporaire des parametres d'appel (call_sid -> params)
_voice_call_params: Dict[str, dict] = {}
# Bridges actifs (call_sid -> RealtimeBridge) pour mute forcé
_active_realtime_bridges: Dict[str, Any] = {}

@app.post("/api/voice-call")
async def start_voice_call(req: VoiceCallRequest, request: Request):
    """Lance un appel vocal sortant. Luna appelle le telephone du souscripteur."""
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return JSONResponse(status_code=403, content={"error": "Service non disponible"})
    tid = getattr(request.state, "tenant_id", 1)
    if not voice_client or not voice_client.is_configured:
        return JSONResponse(status_code=503, content={
            "error": "Appels vocaux non disponibles",
            "message": "Service d'appel vocal non configure.",
        })
    phone = req.phone or ADMIN_NUMBER
    if not phone:
        return JSONResponse(status_code=400, content={"error": "Numero de telephone requis"})
    success, data = await voice_client.initiate_call_async(phone)
    if not success:
        return {"error": data.get("error", "Erreur appel vocal")}
    # Stocker les parametres personnalises pour cet appel (lock pour thread-safety)
    import time as _time
    async with _voice_call_params_lock:
        _voice_call_params[data["call_sid"]] = {
            "mission": req.mission,
            "max_duration": req.max_duration,
            "greeting": req.greeting,
            "phone": phone,
            "tenant_id": tid,
            "_ts": _time.time(),
        }
    _gamify(tid, "voice_call")
    return {"call_sid": data["call_sid"], "status": data["status"]}


@app.post("/api/voice-call/twiml")
async def voice_call_twiml(request: Request):
    """
    Endpoint TwiML que Twilio fetche quand l'appel est decroche.
    Retourne le XML qui dit a Twilio d'ouvrir un Media Stream.
    Twilio envoie CallSid dans les params POST — on le passe au WebSocket URL
    pour matcher les parametres d'appel de facon fiable.
    """
    if not voice_client:
        return Response(
            content="<Response><Say language='fr-FR'>Service non disponible.</Say><Hangup/></Response>",
            media_type="application/xml",
        )
    # Extraire le CallSid des parametres POST envoyes par Twilio
    try:
        form = await request.form()
        call_sid = form.get("CallSid", "") or ""
    except Exception:
        call_sid = ""
    twiml = voice_client.generate_twiml(call_sid=call_sid)
    logger.info(f"TwiML genere pour appel vocal (CallSid={call_sid})")
    return Response(content=twiml, media_type="application/xml")


@app.post("/api/voice-call/conference")
async def start_conference_call(request: Request):
    """Luna rejoint directement une conférence pour prendre des notes."""
    tid = getattr(request.state, "tenant_id", 1)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})

    phone_number = (body.get("phone_number") or "").strip()
    pin = (body.get("pin") or "").strip()
    conference_name = (body.get("context") or "Conference").strip()
    max_dur_min = int(body.get("max_duration_minutes", 60))

    result = await _tool_join_conference(
        {"phone_number": phone_number, "pin": pin, "context": conference_name, "max_duration_minutes": max_dur_min},
        tenant_id=tid,
    )
    if result.get("status") == "error":
        return JSONResponse(status_code=503, content=result)
    return result


@app.post("/api/voice-call/conference-twiml")
async def conference_twiml_webhook(request: Request):
    """Webhook Twilio — TwiML pour appel de conférence (DTMF PIN + Media Stream)."""
    if not voice_client:
        return Response(content="<Response><Hangup/></Response>", media_type="application/xml")
    try:
        form = await request.form()
        call_sid = form.get("CallSid", "") or ""
    except Exception:
        call_sid = ""

    pin = request.query_params.get("pin", "")
    conf_name = request.query_params.get("conf_name", "Conference")
    max_min = int(request.query_params.get("max_min", "60"))

    # Si le call_sid n'est pas encore dans _voice_call_params, on l'enregistre maintenant
    if call_sid:
        import time as _time
        async with _voice_call_params_lock:
            if call_sid not in _voice_call_params:
                _voice_call_params[call_sid] = {
                    "source": "conference",
                    "conference_context": conf_name,
                    "max_duration_minutes": max_min,
                    "_ts": _time.time(),
                }

    twiml = voice_client.generate_conference_twiml(call_sid=call_sid, pin=pin)
    logger.info(f"Conference TwiML genere (CallSid={call_sid} conf={conf_name} pin={'***' if pin else '(none)'})")
    return Response(content=twiml, media_type="application/xml")


@app.post("/api/voice-call/mute")
async def voice_call_mute(request: Request):
    """Fix 4 — Coupe immédiatement la parole de Luna sur tous les appels actifs (ou un CallSid précis)."""
    body = {}
    try:
        body = await request.json()
    except Exception:
        pass
    call_sid = body.get("call_sid", "")
    muted = []
    targets = (
        {call_sid: _active_realtime_bridges[call_sid]}
        if call_sid and call_sid in _active_realtime_bridges
        else dict(_active_realtime_bridges)
    )
    for key, br in targets.items():
        await br.force_mute()
        muted.append(key)
    return JSONResponse({"muted": muted, "count": len(muted)})


@app.websocket("/api/voice-call/media-stream")
async def voice_call_media_stream(websocket: WebSocket):
    """
    WebSocket endpoint pour Twilio Media Streams.
    Relaie l'audio entre le telephone et OpenAI Realtime API.
    """
    await websocket.accept()
    logger.info("Twilio Media Stream WebSocket accepted")

    bridge = None
    memory_mgr = None
    try:
        from integrations.openai.realtime_bridge import RealtimeBridge, build_conference_context, build_voice_context

        # Recuperer les parametres personnalises via le call_sid passe en query param
        call_params = {}
        call_sid_from_url = websocket.query_params.get("call_sid", "")
        ws_mode = websocket.query_params.get("mode", "")

        import time as _time
        _now = _time.time()
        # Nettoyer les params orphelins (>120s)
        _stale = [k for k, v in _voice_call_params.items() if _now - v.get("_ts", 0) > 120]
        for k in _stale:
            _voice_call_params.pop(k, None)

        if call_sid_from_url and call_sid_from_url in _voice_call_params:
            # Match FIABLE par CallSid Twilio
            call_params = _voice_call_params.pop(call_sid_from_url, {})
            logger.info(f"Voice call params matched by CallSid: {call_sid_from_url}")
        elif _voice_call_params and len(_voice_call_params) == 1:
            # Fallback UNIQUEMENT si un seul appel en attente (evite croisement sous appels concurrents)
            call_sid_fallback = list(_voice_call_params.keys())[0]
            call_params = _voice_call_params.pop(call_sid_fallback, {})
            logger.warning(f"Voice call params single-fallback: using {call_sid_fallback}")
        elif _voice_call_params:
            logger.error(f"Voice call params: {len(_voice_call_params)} pending calls, no CallSid match — using defaults")

        mission = call_params.get("mission")
        max_dur = call_params.get("max_duration") or int(os.getenv("VOICE_MAX_DURATION", "180"))
        custom_greeting = call_params.get("greeting")
        contact_name = call_params.get("contact_name", "")

        # Contexte Luna pour l'appel vocal
        memory_mgr = tavus_client.memory if tavus_client else _memory_manager

        # Detect subscriber language preference
        _voice_lang = "fr"
        if _redis_client:
            try:
                _vs = _redis_client.client.hgetall(f"luna:{TENANT_ID}:settings")
                _voice_lang = _vs.get("language", "fr") if _vs else "fr"
            except Exception:
                pass

        is_conference = (call_params.get("source") == "conference" or ws_mode == "conference")
        conference_name = call_params.get("conference_context", "Conference")

        if is_conference:
            context = build_conference_context(
                conference_name=conference_name,
                max_duration_minutes=call_params.get("max_duration_minutes", max(1, max_dur // 60)),
            )
            greeting_text = ""  # Fix 2: silence absolu dès la connexion
        elif mission:
            # Mission speciale : contexte adapte (appel sortant vers un contact)
            _sub_name_for_ctx = call_params.get("subscriber_name") or _SUBSCRIBER_NAME
            context = build_voice_context(
                subscriber_name=_sub_name_for_ctx,
                memory_manager=memory_mgr,
                max_duration_minutes=max(1, max_dur // 60),
                mission=mission,
                contact_name=contact_name,
                redis_client=_redis_client,
                tenant_id=TENANT_ID,
                language=_voice_lang,
            )
            greeting_text = custom_greeting or f"La personne vient de decrocher. Execute ta mission maintenant."
        else:
            context = build_voice_context(
                subscriber_name=_SUBSCRIBER_NAME,
                memory_manager=memory_mgr,
                redis_client=_redis_client,
                tenant_id=TENANT_ID,
                language=_voice_lang,
            )
            greeting_text = custom_greeting or f"L'utilisateur vient de decrocher le telephone. Salue {_SUBSCRIBER_NAME} chaleureusement et demande comment tu peux l'aider."

        # Handler pour les tool calls (reutilise les memes fonctions que Tavus)
        async def handle_voice_tool(name: str, args: dict) -> dict:
            if name == "send_sms":
                return await _tool_send_sms(args)
            elif name == "create_instruction":
                return await _tool_create_instruction(args)
            elif name == "create_note":
                return await _tool_create_note(args)
            elif name == "get_contacts":
                return await _tool_get_contacts()
            elif name == "generate_document":
                return await _tool_generate_document(args)
            elif name == "alert_contacts":
                return await _tool_alert_contacts(args)
            elif name == "send_email":
                return await _tool_send_email(args)
            elif name == "invite_visio":
                return await _tool_invite_visio(args)
            elif name == "call_contact":
                return await _tool_call_contact(args)
            elif name == "search_web":
                return await _tool_search_web(args)
            elif name == "search_places":
                return await _tool_search_places(args)
            elif name == "get_page_info":
                return await _tool_get_page_info(args)
            elif name == "get_player_stats":
                return await _tool_get_player_stats()
            elif name == "get_active_missions":
                return await _tool_get_active_missions()
            elif name == "get_badges":
                return await _tool_get_badges()
            elif name == "get_weather":
                return await _tool_get_weather(args)
            elif name == "get_news":
                return await _tool_get_news(args)
            elif name == "search_flights":
                return await _tool_search_flights(args)
            elif name == "search_hotels":
                return await _tool_search_hotels(args)
            elif name == "book_restaurant":
                return await _tool_book_restaurant(args)
            elif name == "join_conference":
                return await _tool_join_conference(args, tid)
            else:
                return {"status": "error", "message": f"Fonction inconnue: {name}"}

        bridge = RealtimeBridge(
            openai_api_key=OPENAI_API_KEY,
            ws_twilio=websocket,
            call_context=context,
            tool_handler=handle_voice_tool,
            max_duration_seconds=max_dur,
            greeting=greeting_text,
            voice_client=voice_client,
            conference_mode=is_conference,  # Fix 1+2+3
        )
        # Fix 4: enregistrer le bridge pour mute forcé
        _bridge_key = call_sid_from_url or f"bridge_{id(bridge)}"
        _active_realtime_bridges[_bridge_key] = bridge
        _voice_start = time.time()
        try:
            await bridge.run()
        finally:
            _active_realtime_bridges.pop(_bridge_key, None)
        _voice_duration_min = round((time.time() - _voice_start) / 60, 2)

        # Clear active call tracker
        _voice_tid_done = call_params.get("tenant_id") or TENANT_ID
        _active_voice_calls.pop(_voice_tid_done, None)

        # Sauvegarder la transcription dans Redis
        try:
            _save_voice_transcript(bridge, memory_mgr)
        except Exception as e:
            logger.warning(f"Failed to save voice transcript: {e}")

        # Sauvegarder un rapport structuré pour TOUS les appels vocaux
        _call_summary_text = "Rapport disponible en piece jointe."
        if bridge.transcript and memory_mgr:
            try:
                from datetime import timedelta
                from zoneinfo import ZoneInfo as _ZI
                _voice_end_time = datetime.now(_ZI("Europe/Paris"))
                _voice_start_time = _voice_end_time - timedelta(minutes=_voice_duration_min)
                _call_contact_name = call_params.get("contact_name", "")
                _participants = [_SUBSCRIBER_NAME]
                if _call_contact_name:
                    _participants.append(_call_contact_name)

                try:
                    _call_summary_text = await asyncio.wait_for(
                        _generate_call_summary(bridge.transcript, contact_name=_call_contact_name),
                        timeout=30.0,
                    )
                except asyncio.TimeoutError:
                    logger.warning("Call summary generation timed out (30s)")
                    _call_summary_text = f"Appel de {len(bridge.transcript)} echanges."
                _structured_report = {
                    "type": "audio",
                    "participants": _participants,
                    "start_time": _voice_start_time.strftime("%H:%M"),
                    "end_time": _voice_end_time.strftime("%H:%M"),
                    "duration_minutes": _voice_duration_min,
                    "message_count": len(bridge.transcript),
                    "summary": _call_summary_text,
                }
                memory_mgr.add_note(
                    content=f"[Compte rendu appel vocal]\nType: audio\nParticipants: {', '.join(_participants)}\nDebut: {_structured_report['start_time']}\nFin: {_structured_report['end_time']}\nDuree: {_voice_duration_min:.1f} min\nEchanges: {len(bridge.transcript)}\nResume: {_call_summary_text}",
                    context="structured_call_report",
                    tags=["rapport", "appel_vocal", "structure"],
                )
                logger.info(f"Structured call report saved ({_voice_duration_min:.1f} min)")
                # Injecter le compte-rendu dans le fil de discussion chat
                _report_session = call_params.get("session_id", "default")
                _report_chat_msg = (
                    f"Compte-rendu de mon appel avec {_call_contact_name} "
                    f"({_voice_duration_min:.1f} min) :\n\n"
                    f"{_call_summary_text}"
                )
                # Add to Redis conversation
                if memory_mgr:
                    try:
                        memory_mgr.add_message(
                            conv_id=_report_session,
                            role=MessageRole.LUNA,
                            content=_report_chat_msg,
                            channel=Channel.APP,
                        )
                        logger.info(f"Call report injected into chat session {_report_session}")
                    except Exception as _chat_err:
                        logger.warning(f"Failed to inject call report into chat: {_chat_err}")
                # Also add to in-memory conversations for immediate visibility
                _tid_str = str(call_params.get("tenant_id") or TENANT_ID)
                _tenant_convs = conversations.get(_tid_str, {})
                if _report_session in _tenant_convs:
                    _tenant_convs[_report_session].append({
                        "role": "assistant",
                        "content": _report_chat_msg,
                    })

            except Exception as e:
                logger.warning(f"Failed to save structured call report: {e}")

        # Generer le rapport PDF si c'est un appel sortant avec contact
        _voice_tid = call_params.get("tenant_id") or TENANT_ID
        _voice_contact = call_params.get("contact_name", "")
        if _voice_contact and bridge.transcript and _doc_generator:
            try:
                _report_filename = _doc_generator.generate_call_report(
                    call_type="vocal",
                    subscriber_name=call_params.get("subscriber_name", _SUBSCRIBER_NAME),
                    contact_name=_voice_contact,
                    contact_phone=call_params.get("contact_phone", ""),
                    duration_minutes=_voice_duration_min,
                    transcript=bridge.transcript,
                    actions=getattr(bridge, "_tool_calls_log", None),
                    call_sid=bridge.call_sid or "",
                    message_original=call_params.get("message", ""),
                )
                _report_url = f"/api/documents/download/{_report_filename}"
                logger.info(f"Voice call report generated: {_report_filename}")

                # Notifier le souscripteur via une note dans Redis
                if memory_mgr:
                    memory_mgr.add_note(
                        content=(
                            f"[Rapport appel] Appel avec {_voice_contact} — "
                            f"Duree: {_voice_duration_min:.1f} min — "
                            f"Rapport PDF: {_report_url}"
                        ),
                        context="call_report",
                        tags=["rapport", "appel_vocal", _voice_contact],
                    )

                # Envoyer le rapport PDF par email au souscripteur
                _sub_email = ""
                if _redis_client:
                    try:
                        _auth = _redis_client.get_auth_by_tenant_id(_voice_tid)
                        _sub_email = _auth.get("email", "") if _auth else ""
                    except Exception:
                        pass
                if not _sub_email and memory_mgr:
                    try:
                        _prof = memory_mgr.get_subscriber_profile()
                        _sub_email = getattr(_prof, "email", "") if _prof else ""
                    except Exception:
                        pass

                if _sub_email and email_client:
                    _sub_name = call_params.get("subscriber_name", _SUBSCRIBER_NAME)
                    _pdf_path = os.path.join(_doc_generator.output_dir, _report_filename)
                    _summary_text = _call_summary_text
                    try:
                        _email_ok, _email_detail = await email_client.send_for_tenant(
                            tenant_id=_voice_tid,
                            redis_client=_redis_client,
                            gmail_client=None,
                            to=_sub_email,
                            subject=f"Luna — Compte-rendu d'appel avec {_voice_contact}",
                            body_text=(
                                f"Bonjour {_sub_name},\n\n"
                                f"Voici le compte-rendu de mon appel avec {_voice_contact} "
                                f"(duree: {_voice_duration_min:.1f} min).\n\n"
                                f"Resume: {_summary_text}\n\n"
                                f"Le rapport complet est en piece jointe.\n\n"
                                f"A bientot,\nLuna"
                            ),
                            subscriber_name=_sub_name,
                            attachments=[{
                                "filename": _report_filename,
                                "filepath": _pdf_path,
                                "type": "application/pdf",
                            }],
                        )
                        if _email_ok:
                            logger.info(f"Call report emailed to {_sub_email}")
                        else:
                            logger.warning(f"Failed to email call report: {_email_detail}")
                    except Exception as _email_err:
                        logger.warning(f"Failed to email call report: {_email_err}")
            except Exception as e:
                logger.warning(f"Failed to generate call report PDF: {e}")

        # Envoyer un SMS au souscripteur avec le resume de l'appel
        if _voice_contact and bridge.transcript:
            try:
                _sub_phone = call_params.get("subscriber_phone") or ADMIN_NUMBER
                # Recuperer le telephone du profil si disponible
                if not _sub_phone and memory_mgr:
                    try:
                        _prof = memory_mgr.get_subscriber_profile()
                        _sub_phone = getattr(_prof, "phone", "") if _prof else ""
                    except Exception:
                        pass
                if _sub_phone and sms_client:
                    try:
                        _sms_summary = _call_summary_text[:400]
                    except (NameError, UnboundLocalError):
                        _sms_summary = f"Appel avec {_voice_contact} termine ({_voice_duration_min:.0f} min)."
                    _sms_body = (
                        f"Luna - Compte-rendu d'appel avec {_voice_contact} "
                        f"({_voice_duration_min:.0f} min) :\n\n"
                        f"{_sms_summary}\n\n"
                        f"Ouvre l'app pour le detail complet."
                    )
                    # Limiter a 1600 chars (limite SMS longue)
                    if len(_sms_body) > 1500:
                        _sms_body = _sms_body[:1497] + "..."
                    _sms_ok, _sms_detail = sms_client.send(_sub_phone, _sms_body)
                    if _sms_ok:
                        logger.info(f"Call report SMS sent to {_sub_phone}")
                    else:
                        logger.warning(f"Failed to send call report SMS: {_sms_detail}")
            except Exception as e:
                logger.warning(f"Failed to send call report SMS: {e}")

        # Tracker les minutes vocales dans Cortex
        if _voice_duration_min > 0:
            try:
                cortex = get_cortex() if _CORTEX_AVAILABLE else None
                if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
                    await cortex.cost_tracker.track_voice_tenant(_voice_tid, _voice_duration_min)
                    logger.info(f"Voice call tracked: {_voice_duration_min} min for tenant {_voice_tid}")
            except Exception as e:
                logger.warning(f"Failed to track voice cost: {e}")

    except WebSocketDisconnect:
        logger.info("Twilio Media Stream disconnected")
        # Sauvegarder la transcription meme si deconnexion
        if bridge and bridge.transcript:
            try:
                _save_voice_transcript(bridge, memory_mgr)
            except Exception:
                pass
    except asyncio.CancelledError:
        logger.info("Voice media stream cancelled")
    except OSError as e:
        logger.error(f"Voice media stream network error: {e}")
    except Exception as e:
        logger.error(f"Voice media stream error: {e}")
    finally:
        try:
            await websocket.close()
        except Exception:
            pass


# =========================================================================
# LUNA VOICE — Mode Jarvis (voix directe navigateur <-> OpenAI Realtime)
# =========================================================================
@app.websocket("/ws/luna-voice")
async def ws_luna_voice(websocket: WebSocket):
    """WebSocket voix directe : navigateur <-> OpenAI Realtime API (mode Jarvis)."""
    await websocket.accept()

    # Auth via query param token
    token = websocket.query_params.get("token", "")
    jwt_payload = _decode_client_token(token) if token else None
    if not jwt_payload:
        await websocket.close(code=4001, reason="Token invalide")
        return

    # Verifier que la cle API est disponible
    if not OPENAI_API_KEY:
        try:
            await websocket.send_text(json.dumps({"type": "error", "message": "Service vocal non configure"}))
        except Exception:
            pass
        await websocket.close(code=1011, reason="Service non configure")
        return

    tid = jwt_payload.get("tenant_id", TENANT_ID)
    plan_name = jwt_payload.get("plan", "essentiel")
    sub_name = jwt_payload.get("first_name", "") or _SUBSCRIBER_NAME

    # Budget guard — bloquer si cout API mensuel depasse
    budget_err = await _check_budget_guard(tid, plan_name)
    if budget_err:
        try:
            await websocket.send_text(json.dumps({"type": "error", "message": budget_err}))
        except Exception:
            pass
        await websocket.close(code=1011, reason="Budget depasse")
        return

    from integrations.openai.web_voice_bridge import WebVoiceBridge
    from integrations.openai.realtime_bridge import build_voice_context

    memory_mgr = tavus_client.memory if tavus_client else _memory_manager

    # Detect language
    _voice_lang = "fr"
    if _redis_client:
        try:
            _vs = _redis_client.client.hgetall(f"luna:{tid}:settings")
            _voice_lang = _vs.get("language", "fr") if _vs else "fr"
        except Exception:
            pass

    context = build_voice_context(
        subscriber_name=sub_name,
        memory_manager=memory_mgr,
        max_duration_minutes=30,
        redis_client=_redis_client,
        tenant_id=tid,
        language=_voice_lang,
    )
    # Jarvis mode: enrichir le contexte pour une assistante toujours disponible
    context = context.replace(
        "Tu es en appel telephonique avec",
        "Tu es en conversation vocale directe avec"
    )
    context += """

=== MODE ASSISTANT VOCAL DIRECT ===
Tu es l'assistante vocale personnelle de {sub}, accessible a tout moment.
Comme un assistant intelligent, tu dois maitriser en permanence :
- QUI : tu connais {sub}, ses contacts, sa situation, son profil
- OU : si {sub} te le dit ou si la geolocalisation est disponible
- QUOI : ce que {sub} fait, ses projets, ses rendez-vous, ses habitudes
- COMMENT : comment l'aider au mieux (actions concretes, pas juste des conseils)
- POURQUOI : comprendre le contexte et l'intention derriere chaque demande

Tu es proactive mais respectueuse. Tu peux suggerer, rappeler, anticiper.
Tes reponses sont courtes et percutantes — c'est une conversation orale.
Tu es toujours disponible, toujours a l'ecoute.

=== MONDE LUNA (Social) ===
{sub} a peut-etre des amis dans le Monde Luna. Tu peux :
- Voir qui est en ligne (get_friends_online)
- Envoyer un message prive a un ami (send_dm_voice)
- Donner le code ami de {sub} pour ajouter des amis (get_my_friend_code)
Fais bien la difference entre CONTACTS DE CONFIANCE (famille/proches, SMS/appel/email)
et AMIS LUNA (autres souscripteurs dans le Monde, DM dans l'app).

=== PERCEPTION VISUELLE ===
Si la camera est activee, tu peux regarder ce que {sub} voit (look_around).
Tu peux commenter la scene, identifier des objets, ou observer les personnes presentes.
Ne dis JAMAIS "je surveille" — dis "je vois que...", "j'ai l'impression que...".

=== SECRETAIRE PERSONNELLE ===
Tu es aussi la secretaire personnelle de {sub}. Tu geres ses documents, son budget et ses rappels :
- Voir un resume des documents (get_documents_summary) — factures, courriers, relances
- Analyser le budget (get_budget_analysis) — depenses, reste a vivre, previsions
- Verifier si une depense est raisonnable (check_affordability) — "est-ce que je peux me permettre X ?"
- Enregistrer une depense ou un revenu (add_expense)
- Voir les rappels et echeances (get_reminders) — factures a payer, rendez-vous
- Creer un rappel (add_reminder) — "rappelle-moi de payer EDF avant le 15"
- Chercher un document (search_documents) — "retrouve ma facture Orange"
- Voir les dossiers classes (list_folders) — Factures/, Contrats/, Fiches de paie/, etc.
Sois proactive : si {sub} parle d'argent, propose de verifier le budget.
Si il parle d'un courrier ou d'une facture, propose de chercher dans ses documents.
Aide-le a y voir clair, simplement, sans jargon.
""".format(sub=sub_name)

    # Contexte theocratie pour le proprio (tenant 1)
    if tid == _PROPRIO_TENANT_ID and _redis_client:
        try:
            _now = datetime.now(ZoneInfo("Europe/Paris"))
            _theo_month = _now.strftime("%Y-%m")
            _theo_hours = _redis_client.client.get(f"luna:{tid}:theo:hours:{_theo_month}")
            _theo_hours = float(_theo_hours) if _theo_hours else 0.0
            _theo_goal = 50.0
            _theo_remaining = max(0, _theo_goal - _theo_hours)
            _next_y = _now.year + (1 if _now.month == 12 else 0)
            _next_m = 1 if _now.month == 12 else _now.month + 1
            _theo_days_left = (datetime(_next_y, _next_m, 1, tzinfo=ZoneInfo("Europe/Paris")) - _now).days
            _theo_daily = round(_theo_remaining / max(1, _theo_days_left), 1)
            context += f"""
=== THEOCRATIE — PIONNIER PERMANENT ===
{sub_name} est Temoin de Jehovah, pionnier permanent (objectif 50h/mois de predication).
Ce mois-ci : {_theo_hours:.1f}h effectuees sur 50h, il reste {_theo_remaining:.1f}h en {_theo_days_left} jours ({_theo_daily}h/jour necessaires).
{"Il est en retard — encourage-le et propose des alternatives (lettres, appels telephoniques, temoignage informel)." if _theo_remaining > _theo_daily * _theo_days_left * 0.8 else "Il est dans les temps — felicite-le !"}

Tu peux l'aider a :
- Preparer ses reunions (Tour de Garde, reunion de semaine, lecture biblique)
- Compter ses heures de predication (il te le dira oralement)
- Chercher des references bibliques et des publications JW
- Rediger des lettres de temoignage pour son territoire
- Mediter sur des versets et des recits bibliques
- Se motiver quand il est fatigue ou decourage

Quand il parle de ses heures, propose de les enregistrer. Quand il parle d'une reunion, propose de l'aider a preparer.
"""
        except Exception as e:
            logger.warning(f"Voice theo context error: {e}")

    # Historique de conversation pour reconnexion
    _voice_history = []
    _history_param = websocket.query_params.get("history", "")
    if _history_param:
        try:
            _voice_history = json.loads(_history_param)
            if not isinstance(_voice_history, list):
                _voice_history = []
            # Limiter taille
            _voice_history = _voice_history[-20:]
        except (json.JSONDecodeError, Exception):
            _voice_history = []

    # Charger les derniers messages chat pour continuite cross-canal
    if _redis_client and not _voice_history:
        try:
            _recent_chat = _redis_client.client.lrange(f"luna:{tid}:chat:recent", -6, -1)
            if _recent_chat:
                for _rc in _recent_chat:
                    try:
                        _rc_data = json.loads(_rc) if isinstance(_rc, str) else _rc
                        _voice_history.append({
                            "role": _rc_data.get("role", "user"),
                            "text": _rc_data.get("content", "")[:200],
                        })
                    except Exception:
                        pass
        except Exception:
            pass

    # Tool handler (meme que pour Twilio)
    async def handle_voice_tool(name: str, args: dict) -> dict:
        if name == "send_sms":
            return await _tool_send_sms(args)
        elif name == "create_instruction":
            return await _tool_create_instruction(args)
        elif name == "create_note":
            return await _tool_create_note(args)
        elif name == "get_contacts":
            return await _tool_get_contacts()
        elif name == "generate_document":
            return await _tool_generate_document(args)
        elif name == "alert_contacts":
            return await _tool_alert_contacts(args)
        elif name == "send_email":
            return await _tool_send_email(args)
        elif name == "invite_visio":
            return await _tool_invite_visio(args)
        elif name == "call_contact":
            return await _tool_call_contact(args)
        elif name == "search_web":
            return await _tool_search_web(args)
        elif name == "search_places":
            return await _tool_search_places(args)
        elif name == "get_page_info":
            return await _tool_get_page_info(args)
        elif name == "get_player_stats":
            return await _tool_get_player_stats()
        elif name == "get_active_missions":
            return await _tool_get_active_missions()
        elif name == "get_badges":
            return await _tool_get_badges()
        elif name == "get_weather":
            return await _tool_get_weather(args)
        elif name == "get_news":
            return await _tool_get_news(args)
        elif name == "search_flights":
            return await _tool_search_flights(args)
        elif name == "search_hotels":
            return await _tool_search_hotels(args)
        elif name == "book_restaurant":
            return await _tool_book_restaurant(args)
        elif name == "get_friends_online":
            return await _voice_tool_get_friends_online(tid)
        elif name == "send_dm_voice":
            return await _voice_tool_send_dm(tid, args)
        elif name == "get_my_friend_code":
            return await _voice_tool_get_friend_code(tid)
        elif name == "look_around":
            return await _voice_tool_look_around(tid)
        # Secretary tools
        elif name == "get_documents_summary":
            return _tool_secretary_summary(tid)
        elif name == "get_budget_analysis":
            return _tool_secretary_budget(tid)
        elif name == "check_affordability":
            return _tool_secretary_afford(tid, args)
        elif name == "add_expense":
            return _tool_secretary_add_expense(tid, args)
        elif name == "get_reminders":
            return _tool_secretary_reminders(tid)
        elif name == "add_reminder":
            return _tool_secretary_add_reminder(tid, args)
        elif name == "search_documents":
            return _tool_secretary_search(tid, args)
        elif name == "list_folders":
            return _tool_secretary_folders(tid)
        else:
            return {"status": "error", "message": f"Fonction inconnue: {name}"}

    _is_reconnect = len(_voice_history) > 0 and _history_param
    if _is_reconnect:
        _greeting = f"{sub_name} revient apres une coupure. Dis-lui brievement que tu es de retour et continue la conversation la ou elle en etait."
    else:
        _greeting = f"{sub_name} vient d'activer le mode vocal. Salue-le brievement et demande ce que tu peux faire pour lui."

    bridge = WebVoiceBridge(
        openai_api_key=OPENAI_API_KEY,
        ws_client=websocket,
        context=context,
        tool_handler=handle_voice_tool,
        voice=os.getenv("OPENAI_VOICE_NAME", "alloy"),
        max_duration_seconds=1800,
        greeting=_greeting,
        conversation_history=_voice_history,
        vad_eagerness="low",
    )

    _voice_start = time.time()
    try:
        await bridge.run()
    except WebSocketDisconnect:
        logger.info("Luna Voice WS disconnected")
    except asyncio.CancelledError:
        logger.info("Luna Voice cancelled")
    except Exception as e:
        logger.error(f"Luna Voice error: {e}")
    finally:
        _voice_dur = round((time.time() - _voice_start) / 60, 2)

        # Sauvegarder la transcription
        if bridge.transcript and memory_mgr:
            try:
                conv_id = f"voice_web_{int(time.time())}"
                for entry in bridge.transcript:
                    role = MessageRole.SUBSCRIBER if entry["role"] == "user" else MessageRole.LUNA
                    try:
                        memory_mgr.add_message(
                            conv_id=conv_id, role=role,
                            content=entry["text"], channel=Channel.CALL,
                        )
                    except Exception:
                        pass
                logger.info(f"Luna Voice transcript saved: {conv_id} ({len(bridge.transcript)} entries, {_voice_dur:.1f}min)")

                # Auto-generer un compte rendu structure si conversation > 2 echanges
                if len(bridge.transcript) >= 4 and openai_client:
                    try:
                        _transcript_text = "\n".join(
                            f"{'Utilisateur' if e['role'] == 'user' else 'Luna'}: {e['text']}"
                            for e in bridge.transcript
                        )
                        _summary_resp = openai_client.chat.completions.create(
                            model=OPENAI_MODEL,
                            messages=[
                                {"role": "system", "content": "Tu generes des comptes rendus concis de conversations vocales. Format: 1) Resume (2-3 phrases), 2) Points cles (liste), 3) Actions a suivre (si applicable). En francais."},
                                {"role": "user", "content": f"Voici la transcription d'une conversation vocale de {_voice_dur:.1f} minutes:\n\n{_transcript_text[:3000]}\n\nGenere le compte rendu."},
                            ],
                            max_tokens=400,
                            temperature=0.3,
                        )
                        _summary_text = _summary_resp.choices[0].message.content.strip()
                        memory_mgr.add_note(
                            content=f"[Compte rendu vocal — {_voice_dur:.1f}min, {len(bridge.transcript)} echanges]\n\n{_summary_text}",
                            context="voice_summary",
                            tags=["vocal", "compte_rendu", "auto"],
                        )
                        # Stocker aussi dans Redis pour le frontend
                        if _redis_client:
                            import json as _json_cr
                            _cr_data = _json_cr.dumps({
                                "type": "voice_summary",
                                "summary": _summary_text,
                                "duration_min": _voice_dur,
                                "exchanges": len(bridge.transcript),
                                "ts": time.time(),
                            })
                            _redis_client.client.rpush(f"luna:{tid}:notifications:pending", _cr_data)
                            _redis_client.client.expire(f"luna:{tid}:notifications:pending", 86400)
                        logger.info(f"Voice summary generated for tenant {tid}")
                    except Exception as e:
                        logger.warning(f"Voice summary generation failed: {e}")

            except Exception as e:
                logger.warning(f"Failed to save Luna Voice transcript: {e}")

        # Tracker les minutes vocales
        if _voice_dur > 0.1:
            try:
                cortex = get_cortex() if _CORTEX_AVAILABLE else None
                if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
                    await cortex.cost_tracker.track_voice_tenant(tid, _voice_dur)
            except Exception:
                pass

        # Fermer le WebSocket client proprement
        try:
            await websocket.close()
        except Exception:
            pass


# =========================================================================
# DM WEBSOCKET — Temps reel pour les messages prives
# =========================================================================
_dm_subscribers: Dict[str, set] = {}  # room_id -> set of (tid, websocket)

@app.websocket("/ws/dm/{room_id}")
async def ws_dm(websocket: WebSocket, room_id: str):
    """WebSocket temps reel pour les DMs. Remplace le polling 5s."""
    await websocket.accept()
    token = websocket.query_params.get("token", "")
    jwt_payload = _decode_client_token(token) if token else None
    if not jwt_payload:
        await websocket.close(code=4001, reason="Token invalide")
        return
    tid = jwt_payload.get("tenant_id", TENANT_ID)

    # Verify access to this room
    if not _redis_client:
        await websocket.close(code=1011)
        return
    from core.social.redis_ops import SocialRedisOps
    sops = SocialRedisOps(_redis_client)
    room = sops.get_dm_room(room_id)
    if not room or str(tid) not in (room.get("tid1", ""), room.get("tid2", "")):
        await websocket.close(code=4003, reason="Acces refuse")
        return

    # Mark as read on connect
    sops.mark_dm_read(tid, room_id)

    # Register subscriber
    entry = (str(tid), websocket)
    if room_id not in _dm_subscribers:
        _dm_subscribers[room_id] = set()
    _dm_subscribers[room_id].add(entry)

    try:
        async for message in websocket.iter_text():
            try:
                data = json.loads(message)
            except json.JSONDecodeError:
                continue
            if data.get("type") == "message":
                text = (data.get("text", "") or "").strip()
                if not text or len(text) > 500:
                    continue
                # Check block
                other_tid = room.get("tid2") if str(room.get("tid1")) == str(tid) else room.get("tid1")
                if sops.is_blocked(tid, other_tid) or sops.is_blocked(other_tid, tid):
                    continue
                msg = sops.add_dm_message(room_id, tid, text)
                # Broadcast to all subscribers in this room
                payload = json.dumps({"type": "message", "message": msg})
                dead = []
                for sub_tid, sub_ws in _dm_subscribers.get(room_id, set()):
                    try:
                        await sub_ws.send_text(payload)
                        # Mark as read for the person who has the DM open
                        if sub_tid != str(tid):
                            sops.mark_dm_read(sub_tid, room_id)
                    except Exception:
                        dead.append((sub_tid, sub_ws))
                for d in dead:
                    _dm_subscribers.get(room_id, set()).discard(d)
            elif data.get("type") == "typing":
                payload = json.dumps({"type": "typing", "tid": str(tid)})
                for sub_tid, sub_ws in _dm_subscribers.get(room_id, set()):
                    if sub_tid != str(tid):
                        try:
                            await sub_ws.send_text(payload)
                        except Exception:
                            pass
    except WebSocketDisconnect:
        pass
    except Exception as e:
        logger.debug(f"DM WS error: {e}")
    finally:
        _dm_subscribers.get(room_id, set()).discard(entry)
        if room_id in _dm_subscribers and not _dm_subscribers[room_id]:
            del _dm_subscribers[room_id]


def _save_voice_transcript(bridge, memory_mgr):
    """Sauvegarde la transcription d'un appel vocal dans Redis."""
    if not bridge.transcript or not memory_mgr:
        return
    try:
        conv_id = f"voice_{bridge.call_sid or 'unknown'}_{int(time.time())}"
        transcript_text = bridge.get_transcript_text()
        if not transcript_text.strip():
            return

        # Sauvegarder chaque message dans la conversation Redis
        for entry in bridge.transcript:
            role = MessageRole.SUBSCRIBER if entry["role"] == "user" else MessageRole.LUNA
            try:
                memory_mgr.add_message(
                    conv_id=conv_id,
                    role=role,
                    content=entry["text"],
                    channel=Channel.CALL,
                )
            except Exception:
                pass  # quota atteint ou autre erreur, on continue

        # Sauvegarder aussi comme note resume
        summary = f"[Appel vocal] {len(bridge.transcript)} echanges\n{transcript_text[:500]}"
        try:
            memory_mgr.add_note(
                content=summary,
                context="voice_call",
                tags=["appel_vocal", "transcription"],
            )
        except Exception:
            pass

        logger.info(f"Voice transcript saved: {conv_id} ({len(bridge.transcript)} entries)")
    except Exception as e:
        logger.error(f"Error saving voice transcript: {e}")


async def _generate_call_summary(transcript: list, contact_name: str = "") -> str:
    """Resume factuel d'un appel vocal via LLM (gpt-4o-mini).

    Le transcript contient role=user pour l'interlocuteur et role=assistant pour Luna.
    contact_name permet de nommer l'interlocuteur dans le resume.
    """
    if not transcript:
        return "Aucun echange."
    interlocutor = contact_name or "Interlocuteur"
    text_parts = []
    for e in transcript[:50]:
        speaker = interlocutor if e.get("role") == "user" else "Luna"
        text_parts.append(f"{speaker}: {e.get('text', '')}")
    text = "\n".join(text_parts)
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": (
                    f"Tu generes un compte-rendu factuel d'un appel telephonique entre Luna (assistante IA) et {interlocutor}. "
                    "Structure ton resume ainsi :\n"
                    f"1. Ce que {interlocutor} a dit/repondu (l'essentiel de ses propos)\n"
                    "2. Ce que Luna a transmis ou demande\n"
                    "3. Actions ou informations a retenir\n"
                    "Sois precis et factuel (5-8 phrases). Ne mentionne QUE ce qui a ete reellement dit. "
                    "N'invente rien. N'ajoute pas de details techniques."
                )},
                {"role": "user", "content": f"Transcription de l'appel:\n{text[:4000]}"},
            ],
            max_tokens=400,
            temperature=0.3,
        )
        await _track_openai_cost(resp)
        return resp.choices[0].message.content or "Resume non disponible."
    except Exception as e:
        logger.warning(f"Failed to generate call summary: {e}")
        return f"Appel de {len(transcript)} echanges avec {interlocutor}."


async def _generate_visio_summary(transcript: list, conversation_id: str, duration_min: float) -> str:
    """Resume factuel d'un appel visio via LLM (gpt-4o-mini)."""
    if not transcript:
        return f"[Resume visio | {duration_min:.0f} min] Aucun echange."
    entries = []
    for e in transcript[:60]:
        speaker = "Luna" if e.get("speaker") == "replica" else "Souscripteur"
        entries.append(f"{speaker}: {e.get('text', '')}")
    text = "\n".join(entries)
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Tu resumes un appel video. Genere un resume structure avec: 1) Sujets abordes 2) Actions demandees 3) Humeur generale. Sois factuel et concis (5-8 phrases max). Ne mentionne aucune donnee technique."},
                {"role": "user", "content": f"Transcription d'un appel visio de {duration_min:.0f} minutes:\n{text[:4000]}"},
            ],
            max_tokens=300,
            temperature=0.3,
        )
        await _track_openai_cost(resp)
        summary = resp.choices[0].message.content or ""
        return f"[Compte rendu visio | {duration_min:.0f} min | {len(transcript)} echanges]\n{summary}"
    except Exception as e:
        logger.warning(f"Failed to generate visio summary: {e}")
        return f"[Resume visio | {duration_min:.0f} min | {len(transcript)} echanges]\nResume automatique non disponible."


# =========================================================================
# APPELS VOCAUX ENTRANTS - Le souscripteur appelle Luna
# =========================================================================

@app.post("/api/webhook/voice-incoming")
async def webhook_voice_incoming(request: Request):
    """
    Webhook Twilio pour appels entrants.
    Verifie que l'appelant est le souscripteur (ADMIN_NUMBER).
    Si oui: connecte a Luna via OpenAI Realtime.
    Si non: refuse poliment.
    """
    from twilio.twiml.voice_response import VoiceResponse, Connect

    # Twilio envoie les parametres en form-data
    try:
        form = await request.form()
        caller = form.get("From", "")
        called = form.get("To", "")
        call_sid = form.get("CallSid", "")
    except Exception:
        caller = ""
        called = ""
        call_sid = ""

    logger.info(f"Appel entrant: {caller} -> {called} (CallSid: {call_sid})")

    # Verifier que l'appelant est autorise (souscripteur)
    from integrations.twilio.sms_client import TwilioSMSClient
    caller_normalized = TwilioSMSClient.normalize_phone(caller) if caller else ""
    admin_normalized = TwilioSMSClient.normalize_phone(ADMIN_NUMBER) if ADMIN_NUMBER else ""

    if not caller_normalized or caller_normalized != admin_normalized:
        logger.warning(f"Appel entrant refuse: {caller} n'est pas le souscripteur ({ADMIN_NUMBER})")
        response = VoiceResponse()
        response.say(
            "Desole, ce numero n'est pas autorise a contacter Luna. Au revoir.",
            language="fr-FR",
        )
        response.hangup()
        return Response(content=str(response), media_type="application/xml")

    # Souscripteur verifie -> connecter a Luna
    if not voice_client or not VOICE_CALLBACK_URL:
        response = VoiceResponse()
        response.say("Luna n'est pas disponible pour le moment. Reessayez plus tard.", language="fr-FR")
        response.hangup()
        return Response(content=str(response), media_type="application/xml")

    response = VoiceResponse()
    response.say("Bonjour ! Je te passe Luna.", language="fr-FR")
    connect = Connect()
    ws_url = VOICE_CALLBACK_URL.replace("https://", "wss://")
    connect.stream(url=f"{ws_url}/api/voice-call/media-stream")
    response.append(connect)

    logger.info(f"Appel entrant accepte: {caller} -> connexion a Luna")
    _gamify(TENANT_ID, "voice_call")
    return Response(content=str(response), media_type="application/xml")


@app.post("/api/invite-contact")
async def invite_contact(req: InviteRequest, request: Request):
    """
    Invite un contact de confiance dans la visio en cours.
    Envoie un SMS avec le lien Tavus.
    """
    tid = getattr(request.state, "tenant_id", 1)
    if not sms_client.is_configured:
        return {"error": "Service SMS non configure"}

    # Quota check (if available)
    if _quota_guard:
        try:
            quota_status = _quota_guard.check(tid, CoreActionType.SEND_SMS)
            if not quota_status.allowed:
                return {"error": quota_status.warning_message or "Quota SMS atteint"}
        except Exception as e:
            logger.warning(f"Quota check failed: {e}")

    phone = sms_client.normalize_phone(req.contact_phone)
    if not phone:
        return {"error": f"Numero invalide: {req.contact_phone}"}

    success, message = await tavus_client.add_participant_via_sms(
        conversation_id=req.conversation_id,
        contact_name=req.contact_name,
        contact_phone=phone,
        sms_client=sms_client,
        tenant_id=tid,
    )
    if not success:
        return {"error": message}
    return {"success": True, "message": message}


@app.post("/api/webhook/sms")
async def webhook_sms(request: Request):
    """
    Webhook Twilio pour recevoir les SMS entrants.
    Twilio envoie un POST avec les params du message recu.
    """
    form = await request.form()
    from_number = form.get("From", "")
    body = form.get("Body", "")
    to_number = form.get("To", "")
    sms_sid = form.get("MessageSid", "")

    logger.info(f"SMS recu de {from_number}: {body[:80]}")

    # Validation signature Twilio (stricte si configuré)
    if sms_client.is_configured:
        signature = request.headers.get("X-Twilio-Signature", "")
        if not signature:
            logger.warning(f"Signature Twilio manquante pour SMS de {from_number}")
            return Response(status_code=403, content="Forbidden", media_type="text/plain")
        url = str(request.url)
        params = dict(form)
        if not sms_client.validate_webhook(signature, url, params):
            logger.warning(f"Signature Twilio invalide pour SMS de {from_number}")
            return Response(status_code=403, content="Forbidden", media_type="text/plain")

    # CORTEX: intercepte les commandes d'urgence (LUNA STATUS, LUNA BAN, etc.)
    emergency_response = await cortex_handle_sms(from_number, body)
    if emergency_response:
        logger.info(f"CORTEX SMS command from {from_number[:8]}...: {body[:40]}")
        # Repondre par TwiML avec le message de reponse
        from xml.sax.saxutils import escape as xml_escape
        escaped = xml_escape(emergency_response, {'"': "&quot;", "'": "&apos;"})
        twiml = f"<Response><Message>{escaped}</Message></Response>"
        return Response(content=twiml, media_type="application/xml")

    # Normalise le numero entrant
    from integrations.twilio.sms_client import TwilioSMSClient
    normalized_from = TwilioSMSClient.normalize_phone(from_number) if from_number else ""

    # Verifie si c'est une reponse a une invitation visio en attente
    body_lower = body.strip().lower()
    if normalized_from in _pending_visio_invites and body_lower in ("oui", "yes", "ok", "o", "1"):
        invite = _pending_visio_invites.pop(normalized_from)
        # Verifie que l'invitation n'est pas trop vieille (30 min max)
        if time.time() - invite["timestamp"] < 1800:
            logger.info(f"Visio invite accepted by {invite['contact_name']} ({normalized_from})")
            # Cree la conversation Tavus
            try:
                visio_max = int(os.getenv("VISIO_MAX_DURATION", "60")) * 60  # minutes -> secondes
                sub_name = invite["subscriber_name"]
                contact = invite["contact_name"]

                context = f"Tu es Luna. {sub_name} a invite {contact} en visioconference. Sois chaleureuse et accueillante."
                if _TAVUS_AVAILABLE and tavus_client and tavus_client.is_configured:
                    success, data = await tavus_client.create_conversation(
                        tenant_id=invite["tenant_id"],
                        custom_greeting=f"Bonjour {contact} ! {sub_name} t'a invite en visio avec Luna. Bienvenue !",
                        context=context,
                        max_duration=visio_max,
                        callback_url=TAVUS_CALLBACK_URL if TAVUS_CALLBACK_URL else None,
                    )
                    if success:
                        visio_url = data["conversation_url"]
                        # Envoie le lien par SMS via TwiML reply
                        from xml.sax.saxutils import escape as xml_escape
                        reply_msg = (
                            f"Super ! Voici le lien pour la visio avec {sub_name} :\n"
                            f"{visio_url}\n"
                            f"Clique dessus pour rejoindre."
                        )
                        escaped = xml_escape(reply_msg, {'"': "&quot;", "'": "&apos;"})
                        twiml = f"<Response><Message>{escaped}</Message></Response>"

                        # Previens aussi le souscripteur par SMS
                        try:
                            mgr = _get_tenant_manager(invite["tenant_id"])
                            if mgr:
                                # Cherche le tel du souscripteur
                                profile = mgr.get_subscriber_profile()
                                sub_phone = getattr(profile, "phone", "") or ADMIN_NUMBER
                                if sub_phone:
                                    _tracked_sms_send(
                                        sub_phone,
                                        f"[Luna] {contact} a accepte ton invitation visio ! Lien : {visio_url}",
                                        label="notification visio acceptee"
                                    )
                        except Exception as e:
                            logger.warning(f"Failed to notify subscriber: {e}")

                        logger.info(f"Visio link sent to {contact}: {visio_url}")
                        return Response(content=twiml, media_type="application/xml")
                    else:
                        logger.error(f"Tavus create failed for invite: {data}")
                        from xml.sax.saxutils import escape as xml_escape
                        error_msg = xml_escape("Desole, la visio n'a pas pu etre creee. Reessaie plus tard.")
                        twiml = f"<Response><Message>{error_msg}</Message></Response>"
                        return Response(content=twiml, media_type="application/xml")
            except Exception as e:
                logger.error(f"Visio invite handler error: {e}")
        else:
            logger.info(f"Visio invite expired for {normalized_from}")
            _pending_visio_invites.pop(normalized_from, None)

    elif normalized_from in _pending_visio_invites and body_lower in ("non", "no", "n", "0"):
        invite = _pending_visio_invites.pop(normalized_from)
        logger.info(f"Visio invite declined by {invite['contact_name']}")
        from xml.sax.saxutils import escape as xml_escape
        reply = xml_escape(f"D'accord, pas de visio pour le moment. A bientot !")
        twiml = f"<Response><Message>{reply}</Message></Response>"
        return Response(content=twiml, media_type="application/xml")

    # Enregistre le message entrant dans la conversation
    logger.info(f"SMS entrant OK - De: {from_number}, Corps: {body[:100]}")

    # Reponse TwiML vide (pas de reponse automatique)
    twiml = "<Response></Response>"
    return Response(content=twiml, media_type="application/xml")


@app.post("/api/webhook/sms-status")
async def webhook_sms_status(request: Request):
    """
    Webhook Twilio pour les accuses de reception SMS.
    Twilio envoie un POST quand le statut d'un SMS change.
    Statuts: queued → sent → delivered (ou failed/undelivered)
    """
    # Valider la signature Twilio
    if sms_client and sms_client.is_configured:
        signature = request.headers.get("X-Twilio-Signature", "")
        form_data = await request.form()
        url = str(request.url)
        params = dict(form_data)
        if signature and not sms_client.validate_webhook(signature, url, params):
            logger.warning("Signature Twilio invalide pour SMS status webhook")
            return Response(status_code=403, content="Forbidden", media_type="text/plain")
    else:
        form_data = await request.form()
    form = form_data
    sms_sid = form.get("MessageSid", "")
    status = form.get("MessageStatus", "")
    to_number = form.get("To", "")
    error_code = form.get("ErrorCode", "")

    logger.info(f"SMS status update: {sms_sid} -> {status} (to: {to_number})")

    # Mettre a jour le tracking en memoire
    if sms_sid in _sms_tracking:
        _sms_tracking[sms_sid]["status"] = status
        _sms_tracking[sms_sid]["updated_at"] = datetime.utcnow().isoformat()
        if status == "delivered":
            _sms_tracking[sms_sid]["delivered_at"] = datetime.utcnow().isoformat()
        if error_code:
            _sms_tracking[sms_sid]["error_code"] = error_code
    else:
        # SMS envoye avant le redemarrage ou non tracke
        _sms_tracking[sms_sid] = {
            "to": to_number,
            "status": status,
            "updated_at": datetime.utcnow().isoformat(),
            "delivered_at": datetime.utcnow().isoformat() if status == "delivered" else None,
            "error_code": error_code or None,
        }

    # Si echec, logger un warning
    if status in ("failed", "undelivered"):
        logger.warning(f"SMS ECHEC: {sms_sid} -> {to_number} (status={status}, error={error_code})")

    return Response(content="", status_code=204)


@app.get("/api/sms/status")
async def sms_status_list(request: Request):
    """Liste les derniers SMS avec leur statut de livraison."""
    tid = getattr(request.state, "tenant_id", 1)
    # Retourne les 50 derniers SMS trackes (les plus recents d'abord)
    items = sorted(_sms_tracking.values(), key=lambda x: x.get("sent_at", ""), reverse=True)[:50]
    return {
        "sms": items,
        "count": len(items),
        "total_tracked": len(_sms_tracking),
    }


@app.get("/api/sms/status/{sms_sid}")
async def sms_status_detail(sms_sid: str, request: Request):
    """Detail du statut d'un SMS specifique."""
    if sms_sid in _sms_tracking:
        return _sms_tracking[sms_sid]
    # Fallback: interroger Twilio directement
    if sms_client and sms_client.is_configured:
        try:
            msg = sms_client.client.messages(sms_sid).fetch()
            return {
                "sid": msg.sid,
                "to": msg.to,
                "status": msg.status,
                "sent_at": msg.date_sent.isoformat() if msg.date_sent else None,
                "error_code": msg.error_code,
                "error_message": msg.error_message,
            }
        except Exception as e:
            return JSONResponse(status_code=404, content={"error": f"SMS non trouve: {e}"})
    return JSONResponse(status_code=404, content={"error": "SMS non trouve"})


# =========================================================================
# CONVERSATIONS CRUD
# =========================================================================

class CreateConversationRequest(BaseModel):
    title: str = Field(default="", max_length=100)

class UpdateConversationRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=100)


@app.get("/api/conversations")
async def list_conversations_endpoint(request: Request):
    """Liste toutes les conversations du tenant."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return {"conversations": []}
    try:
        convs = mgr.list_conversations()
        return {
            "conversations": [
                {
                    "id": c.id,
                    "title": c.summary or "",
                    "last_activity": c.last_activity.isoformat(),
                    "message_count": c.message_count,
                    "started_at": c.started_at.isoformat(),
                }
                for c in convs
                if c.status != ConversationStatus.CLOSED
            ]
        }
    except Exception as e:
        logger.warning(f"List conversations error: {e}")
        return {"conversations": []}


@app.post("/api/conversations")
async def create_conversation_endpoint(req: CreateConversationRequest, request: Request):
    """Cree une nouvelle conversation."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=500, content={"error": "Service indisponible"})
    try:
        conv = mgr.create_conversation(
            contact_phone="app",
            contact_name="Chat",
            channel=Channel.APP,
        )
        if req.title:
            meta = mgr.redis.get_conversation_meta(tid, conv.id)
            if meta:
                meta["summary"] = req.title.strip()
                mgr.redis.set_conversation_meta(tid, conv.id, meta)
        return {"id": conv.id, "title": req.title, "started_at": conv.started_at.isoformat()}
    except Exception as e:
        logger.warning(f"Create conversation error: {e}")
        return JSONResponse(status_code=429, content={"error": str(e)})


@app.patch("/api/conversations/{conv_id}")
async def update_conversation_endpoint(conv_id: str, req: UpdateConversationRequest, request: Request):
    """Renomme une conversation."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=500, content={"error": "Service indisponible"})
    meta = mgr.redis.get_conversation_meta(tid, conv_id)
    if not meta:
        return JSONResponse(status_code=404, content={"error": "Conversation introuvable"})
    meta["summary"] = req.title.strip()
    mgr.redis.set_conversation_meta(tid, conv_id, meta)
    return {"ok": True}


@app.delete("/api/conversations/{conv_id}")
async def delete_conversation_endpoint(conv_id: str, request: Request):
    """Supprime une conversation et ses messages."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=500, content={"error": "Service indisponible"})
    if conv_id == "default":
        return JSONResponse(status_code=400, content={"error": "Impossible de supprimer la conversation par defaut"})
    mgr.delete_conversation(conv_id)
    tenant_convs = conversations.get(str(tid), {})
    tenant_convs.pop(conv_id, None)
    _conversation_ts.pop(conv_id, None)
    return {"ok": True}


@app.get("/api/history")
async def history(request: Request, session_id: str = "default", limit: int = 50):
    """Historique des messages d'une session."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    # Try Redis first
    if mgr:
        try:
            messages = mgr.get_messages(session_id, limit=limit)
            return {
                "messages": [
                    {
                        "role": "luna" if msg.role == MessageRole.LUNA else "user",
                        "content": msg.content,
                        "timestamp": msg.timestamp.isoformat() if hasattr(msg, "timestamp") else "",
                    }
                    for msg in messages
                ]
            }
        except Exception as e:
            logger.warning(f"Redis history fetch failed: {e}")

    # Fallback to in-memory (tenant-namespaced)
    tenant_convs = conversations.get(str(tid), {})
    if session_id in tenant_convs:
        msgs = tenant_convs[session_id]
        return {
            "messages": [
                {"role": "luna" if m["role"] == "assistant" else "user", "content": m["content"], "timestamp": ""}
                for m in msgs
                if m["role"] != "system"
            ][-limit:]
        }
    return {"messages": []}


@app.get("/api/status")
async def status(request: Request):
    """Statut du serveur Luna. Version allege si non authentifie."""
    # Version minimale pour non-authentifies
    token = _extract_bearer(request)
    payload = _decode_client_token(token) if token else None
    if not payload:
        return {"luna": "online" if not _pv_locked else "setup", "mode": LUNA_MODE}

    # Version complete pour utilisateurs authentifies
    redis_ok = False
    if _redis_client:
        try:
            redis_ok = _redis_client.ping()
        except Exception:
            pass

    _cm = "assistif"
    if _memory_manager:
        try:
            _p = _memory_manager.get_subscriber_profile()
            if _p:
                _cm = getattr(_p, "caution_mode", "assistif") or "assistif"
        except Exception:
            pass

    return {
        "luna": "online" if not _pv_locked else "setup",
        "mode": LUNA_MODE,
        "pv_signed": PV_SIGNED,
        "pv_locked": _pv_locked,
        "setup_disabled": _setup_permanently_disabled,
        "legal_mode": LEGAL_MODE,
        "caution_mode": _cm,
        "openai": bool(OPENAI_API_KEY),
        "twilio": sms_client.is_configured if sms_client else False,
        "tavus": tavus_client.is_configured if tavus_client else False,
        "tavus_details": tavus_client.get_status() if tavus_client else {},
        "simli": bool(os.getenv("SIMLI_API_KEY", "") and os.getenv("SIMLI_FACE_ID", "")),
        "redis": redis_ok,
        "active_sessions": len(conversations),
        "core_modules": {
            "memory": _memory_manager is not None,
            "safety": _safety_guardian is not None,
            "quota": _quota_guard is not None,
        },
        "perception": {
            "available": _perception_detector is not None,
            "enabled": _memory_manager.is_perception_enabled() if _memory_manager else False,
            "camera": _perception_detector.is_camera_available() if _perception_detector else False,
        },
        "license": {
            "active": _license_heartbeat is not None,
            "status": _license_heartbeat.status if _license_heartbeat else "none",
            "banner": _license_heartbeat.get_banner_message() if _license_heartbeat else None,
        },
    }


# =========================================================================
# AUTH ENDPOINTS (multi-tenant)
# =========================================================================

class RegisterRequest(BaseModel):
    email: str
    password: str
    first_name: str = ""
    last_name: str = ""

class LoginRequest(BaseModel):
    email: str
    password: str

class ChangePasswordRequest(BaseModel):
    old_password: str
    new_password: str

class CheckoutRequest(BaseModel):
    plan: str  # "essentiel", "confort", "premium"

@app.post("/api/auth/register")
async def auth_register(req: RegisterRequest):
    """Inscription d'un nouveau client."""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    email = req.email.strip().lower()
    if not email or "@" not in email:
        return JSONResponse(status_code=400, content={"error": "Email invalide"})
    if len(req.password) < 6:
        return JSONResponse(status_code=400, content={"error": "Mot de passe trop court (min 6 caracteres)"})

    # Verifier si email deja pris
    if _redis_client.get_auth_by_email(email):
        return JSONResponse(status_code=409, content={"error": "Email deja utilise"})

    # Attribuer un tenant_id
    tenant_id = _redis_client.get_next_tenant_id()
    password_hash = _hash_password(req.password)

    # Creer l'enregistrement auth
    created = _redis_client.create_auth_record(email, password_hash, tenant_id, "essentiel")
    if not created:
        return JSONResponse(status_code=409, content={"error": "Email deja utilise"})

    # Creer le profil dans Redis
    if _CORE_AVAILABLE:
        profile = SubscriberProfile(
            tenant_id=tenant_id,
            first_name=req.first_name or email.split("@")[0],
            last_name=req.last_name,
            email=email,
        )
        mgr = _get_tenant_manager(tenant_id)
        mgr.set_subscriber_profile(profile)

    fname = req.first_name or email.split("@")[0]
    token = _create_client_token(tenant_id, email, "essentiel", first_name=fname)
    # Initialize gamification player
    if _GAMIFICATION_AVAILABLE and _redis_client:
        try:
            gops = GamificationRedisOps(_redis_client)
            asyncio.create_task(initialize_player_safe(gops, tenant_id))
        except Exception:
            pass
    _gamify("admin", "new_client", is_admin=True)
    logger.info(f"AUTH_REGISTER tenant_id={tenant_id} email={email}")
    return {"token": token, "tenant_id": tenant_id, "plan": "essentiel", "first_name": fname}


@app.post("/api/auth/login")
async def auth_login(req: LoginRequest):
    """Connexion d'un client existant."""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    email = req.email.strip().lower()
    auth = _redis_client.get_auth_by_email(email)
    if not auth:
        return JSONResponse(status_code=401, content={"error": "Email ou mot de passe incorrect"})

    if not _verify_password(req.password, auth["password_hash"]):
        return JSONResponse(status_code=401, content={"error": "Email ou mot de passe incorrect"})

    if not auth.get("active", True):
        return JSONResponse(status_code=403, content={"error": "Compte desactive"})

    tenant_id = auth["tenant_id"]
    plan = auth.get("plan", "essentiel")

    # Recuperer le prenom depuis le profil
    first_name = ""
    if _redis_client:
        profile = _redis_client.get_profile(tenant_id)
        if profile:
            first_name = profile.get("first_name", "")

    token = _create_client_token(tenant_id, email, plan, first_name=first_name)
    _gamify(tenant_id, "daily_login")
    logger.info(f"AUTH_LOGIN tenant_id={tenant_id} email={email}")
    return {"token": token, "tenant_id": tenant_id, "plan": plan, "first_name": first_name}


@app.get("/api/auth/me")
async def auth_me(request: Request):
    """Profil du client connecte."""
    token = _extract_bearer(request)
    payload = _decode_client_token(token)
    if not payload:
        return JSONResponse(status_code=401, content={"error": "Token invalide"})

    tenant_id = payload["tenant_id"]
    email = payload["email"]
    plan = payload["plan"]

    first_name = ""
    last_name = ""
    if _redis_client:
        profile = _redis_client.get_profile(tenant_id)
        if profile:
            first_name = profile.get("first_name", "")
            last_name = profile.get("last_name", "")

    return {
        "tenant_id": tenant_id,
        "email": email,
        "plan": plan,
        "first_name": first_name,
        "last_name": last_name,
    }


@app.post("/api/auth/change-password")
async def auth_change_password(req: ChangePasswordRequest, request: Request):
    """Changement de mot de passe."""
    token = _extract_bearer(request)
    payload = _decode_client_token(token)
    if not payload:
        return JSONResponse(status_code=401, content={"error": "Token invalide"})

    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    email = payload["email"]
    auth = _redis_client.get_auth_by_email(email)
    if not auth:
        return JSONResponse(status_code=404, content={"error": "Compte introuvable"})

    if not _verify_password(req.old_password, auth["password_hash"]):
        return JSONResponse(status_code=401, content={"error": "Ancien mot de passe incorrect"})

    if len(req.new_password) < 6:
        return JSONResponse(status_code=400, content={"error": "Nouveau mot de passe trop court (min 6)"})

    new_hash = _hash_password(req.new_password)
    _redis_client.update_auth_record(email, {"password_hash": new_hash})
    logger.info(f"AUTH_CHANGE_PASSWORD tenant_id={payload['tenant_id']} email={email}")
    return {"ok": True}


@app.post("/api/auth/checkout")
async def auth_checkout(req: CheckoutRequest, request: Request):
    """Cree une session Stripe Checkout pour upgrade/souscription."""
    token = _extract_bearer(request)
    payload = _decode_client_token(token)
    if not payload:
        return JSONResponse(status_code=401, content={"error": "Token invalide"})

    plan = req.plan.lower()
    price_map = {
        "essentiel": os.getenv("STRIPE_PRICE_ESSENTIEL", ""),
        "confort": os.getenv("STRIPE_PRICE_CONFORT", ""),
        "premium": os.getenv("STRIPE_PRICE_PREMIUM", ""),
    }
    price_id = price_map.get(plan)
    if not price_id:
        return JSONResponse(status_code=400, content={"error": f"Plan invalide ou STRIPE_PRICE non configure: {plan}"})

    stripe_key = os.getenv("STRIPE_API_KEY", "")
    if not stripe_key:
        return JSONResponse(status_code=500, content={"error": "STRIPE_API_KEY non configure"})

    try:
        import stripe
        stripe.api_key = stripe_key
        # Determine URLs
        host = request.headers.get("host", "localhost:8888")
        scheme = "https"
        base_url = f"{scheme}://{host}"

        # Creer ou recuperer le customer Stripe avec metadata tenant_id
        tid = payload["tenant_id"]
        email = payload["email"]
        customer_id = None
        if _redis_client:
            rec = _redis_client.get_auth_by_email(email)
            if rec:
                customer_id = rec.get("stripe_customer_id", "")
        if not customer_id:
            customer = stripe.Customer.create(
                email=email,
                metadata={"tenant_id": str(tid), "yawatch": "true"},
            )
            customer_id = customer.id
            if _redis_client:
                _redis_client.update_auth_record(email, {"stripe_customer_id": customer_id})

        session = stripe.checkout.Session.create(
            customer=customer_id,
            payment_method_types=["card"],
            mode="subscription",
            line_items=[{"price": price_id, "quantity": 1}],
            metadata={
                "tenant_id": str(tid),
                "email": email,
                "plan": plan,
            },
            success_url=f"{base_url}/?checkout=success",
            cancel_url=f"{base_url}/?checkout=cancel",
        )
        logger.info(f"STRIPE_CHECKOUT tenant_id={payload['tenant_id']} plan={plan}")
        return {"checkout_url": session.url}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Service de paiement non disponible"})
    except Exception as e:
        logger.error(f"Checkout error: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur lors du paiement. Reessaie."})


@app.post("/api/auth/setup-card")
async def auth_setup_card(request: Request):
    """
    Cree un Stripe SetupIntent pour enregistrer une carte bancaire
    SANS souscrire d'abonnement. Utile pour la conciergerie.
    Retourne une URL Stripe Checkout en mode 'setup'.
    """
    token = _extract_bearer(request)
    payload = _decode_client_token(token)
    if not payload:
        return JSONResponse(status_code=401, content={"error": "Token invalide"})

    stripe_key = os.getenv("STRIPE_API_KEY", "") or os.getenv("STRIPE_SECRET_KEY", "")
    if not stripe_key:
        return JSONResponse(status_code=500, content={"error": "STRIPE_API_KEY non configure"})

    try:
        import stripe
        stripe.api_key = stripe_key
        host = request.headers.get("host", "localhost:8888")
        base_url = f"https://{host}"

        tid = payload["tenant_id"]
        email = payload["email"]

        # Creer ou recuperer le Customer Stripe
        customer_id = None
        if _redis_client:
            rec = _redis_client.get_auth_by_email(email)
            if rec:
                customer_id = rec.get("stripe_customer_id", "")
        if not customer_id:
            customer = stripe.Customer.create(
                email=email,
                metadata={"tenant_id": str(tid), "yawatch": "true"},
            )
            customer_id = customer.id
            if _redis_client:
                _redis_client.update_auth_record(email, {"stripe_customer_id": customer_id})

        # Creer un Checkout Session en mode 'setup' (enregistre la carte sans payer)
        session = stripe.checkout.Session.create(
            customer=customer_id,
            payment_method_types=["card"],
            mode="setup",
            metadata={
                "tenant_id": str(tid),
                "email": email,
                "purpose": "concierge_card",
            },
            success_url=f"{base_url}/?card_saved=success",
            cancel_url=f"{base_url}/?card_saved=cancel",
        )
        logger.info(f"STRIPE_SETUP_CARD tenant_id={tid} email={email}")
        return {"setup_url": session.url}
    except Exception as e:
        logger.error(f"Setup card error: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur lors de l'enregistrement de la carte. Reessaie."})


# =========================================================================
# STRIPE WEBHOOK (accessible en setup ET en production)
# =========================================================================

def _price_id_to_plan(price_id: str) -> str:
    """Mappe un Stripe price_id vers un nom de plan."""
    mapping = {
        os.getenv("STRIPE_PRICE_ESSENTIEL", ""): "essentiel",
        os.getenv("STRIPE_PRICE_CONFORT", ""): "confort",
        os.getenv("STRIPE_PRICE_PREMIUM", ""): "premium",
    }
    return mapping.get(price_id, "")

_stripe_webhook_received: dict = {}  # {timestamp, event_type, verified}

@app.post("/api/stripe/webhook")
async def stripe_webhook(request: Request):
    """Recoit les webhooks Stripe avec verification de signature."""
    global _stripe_webhook_received
    webhook_secret = os.getenv("STRIPE_WEBHOOK_SECRET", "")
    if not webhook_secret:
        return JSONResponse(status_code=500, content={"error": "STRIPE_WEBHOOK_SECRET non configure"})

    payload = await request.body()
    sig_header = request.headers.get("stripe-signature", "")

    try:
        import stripe
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
        _stripe_webhook_received = {
            "timestamp": time.time(),
            "event_type": event["type"],
            "verified": True,
        }
        logger.info(f"Stripe webhook recu: {event['type']} (verifie)")

        event_type = event["type"]
        data = event.get("data", {}).get("object", {})

        # --- checkout.session.completed ---
        if event_type == "checkout.session.completed":
            metadata = data.get("metadata", {})
            email = metadata.get("email", "")
            plan = metadata.get("plan", "")
            tenant_id_str = metadata.get("tenant_id", "")
            if email and plan and _redis_client:
                _redis_client.update_auth_record(email, {
                    "plan": plan,
                    "active": True,
                    "stripe_customer_id": data.get("customer", ""),
                    "stripe_subscription_id": data.get("subscription", ""),
                })
                # Evict from tenant manager cache to reload with new plan
                tid = int(tenant_id_str) if tenant_id_str else 0
                if tid and tid in _tenant_managers:
                    del _tenant_managers[tid]
                logger.info(f"STRIPE_CHECKOUT_COMPLETE email={email} plan={plan} tenant_id={tenant_id_str}")

            # --- World premium item purchase ---
            if metadata.get("type") == "world_premium" and _redis_client:
                world_item_id = metadata.get("world_item_id", "")
                tenant_id_str = metadata.get("tenant_id", "")
                if world_item_id and tenant_id_str:
                    try:
                        from core.gamification.redis_ops import GamificationRedisOps
                        from core.gamification.constants import SHOP_ITEMS
                        gops = GamificationRedisOps(_redis_client)
                        tid = int(tenant_id_str)
                        item = SHOP_ITEMS.get(world_item_id)
                        if item:
                            from datetime import datetime as dt_import
                            gops.add_to_inventory(tid, world_item_id, {
                                "purchased_at": dt_import.utcnow().isoformat(),
                                "category": item["category"],
                                "stripe_payment": data.get("payment_intent", ""),
                            })
                            logger.info(f"WORLD_PREMIUM_PURCHASED tenant_id={tid} item={world_item_id}")
                    except Exception as e:
                        logger.error(f"Error delivering premium item: {e}")

        # --- customer.subscription.updated ---
        elif event_type == "customer.subscription.updated":
            # Map price_id back to plan name
            items = data.get("items", {}).get("data", [])
            if items and _redis_client:
                price_id = items[0].get("price", {}).get("id", "")
                plan = _price_id_to_plan(price_id)
                customer_id = data.get("customer", "")
                if plan and customer_id:
                    # Find auth record by stripe_customer_id
                    records = _redis_client.get_all_auth_records()
                    for rec in records:
                        if rec.get("stripe_customer_id") == customer_id:
                            _redis_client.update_auth_record(rec["email"], {"plan": plan})
                            tid = rec.get("tenant_id", 0)
                            if tid and tid in _tenant_managers:
                                del _tenant_managers[tid]
                            logger.info(f"STRIPE_SUB_UPDATED email={rec['email']} plan={plan}")
                            break

        # --- customer.subscription.deleted ---
        elif event_type == "customer.subscription.deleted":
            customer_id = data.get("customer", "")
            if customer_id and _redis_client:
                records = _redis_client.get_all_auth_records()
                for rec in records:
                    if rec.get("stripe_customer_id") == customer_id:
                        _redis_client.update_auth_record(rec["email"], {"active": False})
                        logger.info(f"STRIPE_SUB_DELETED email={rec['email']}")
                        break

        # --- invoice.payment_failed ---
        elif event_type == "invoice.payment_failed":
            customer_email = data.get("customer_email", "")
            logger.warning(f"STRIPE_PAYMENT_FAILED email={customer_email}")

        # --- invoice.payment_succeeded ---
        elif event_type == "invoice.payment_succeeded":
            customer_email = data.get("customer_email", "")
            amount = data.get("amount_paid", 0) / 100
            logger.info(f"STRIPE_PAYMENT_SUCCESS email={customer_email} amount={amount:.2f}EUR")

        # --- charge.refunded ---
        elif event_type == "charge.refunded":
            customer_email = data.get("billing_details", {}).get("email", "")
            amount = data.get("amount_refunded", 0) / 100
            logger.warning(f"STRIPE_REFUND email={customer_email} amount={amount:.2f}EUR")

        # --- customer.subscription.created ---
        elif event_type == "customer.subscription.created":
            customer_id = data.get("customer", "")
            items = data.get("items", {}).get("data", [])
            if items:
                price_id = items[0].get("price", {}).get("id", "")
                plan = _price_id_to_plan(price_id)
                logger.info(f"STRIPE_SUB_CREATED customer={customer_id} plan={plan}")

        # --- payment_intent.succeeded (conciergerie) ---
        elif event_type == "payment_intent.succeeded":
            metadata = data.get("metadata", {})
            if metadata.get("type") == "concierge":
                tenant_id_str = metadata.get("tenant_id", "")
                merchant = metadata.get("merchant", "")
                base_cents = int(metadata.get("base_amount_cents", 0))
                comm_cents = int(metadata.get("commission_cents", 0))
                total_cents = data.get("amount", 0)
                intent_id = data.get("id", "")
                logger.info(
                    f"CONCIERGE_PAYMENT_SUCCESS tenant={tenant_id_str} "
                    f"merchant={merchant} total={total_cents/100:.2f}EUR "
                    f"commission={comm_cents/100:.2f}EUR ref={intent_id}"
                )
                # Enregistrer la commission en Redis pour le dashboard admin
                if _redis_client and tenant_id_str:
                    try:
                        import json as _json_comm
                        commission_record = {
                            "intent_id": intent_id,
                            "tenant_id": tenant_id_str,
                            "merchant": merchant,
                            "base_amount_cents": base_cents,
                            "commission_cents": comm_cents,
                            "total_cents": total_cents,
                            "timestamp": time.time(),
                        }
                        _redis_client.client.rpush(
                            f"{_redis_client.prefix}:concierge:commissions",
                            _json_comm.dumps(commission_record),
                        )
                        # Incrementer le total des commissions
                        _redis_client.client.incrbyfloat(
                            f"{_redis_client.prefix}:concierge:total_commission_cents",
                            comm_cents,
                        )
                    except Exception as e:
                        logger.error(f"Error recording commission: {e}")
                # Log dans le memory manager du tenant
                tid = int(tenant_id_str) if tenant_id_str else 0
                if tid:
                    try:
                        t_mgr = _get_tenant_manager(tid)
                        if t_mgr:
                            t_mgr.log_event(
                                category="action",
                                description=f"Paiement conciergerie confirme: {total_cents/100:.2f} EUR pour {merchant}",
                                reasoning="Paiement valide par le souscripteur",
                                source="stripe_webhook",
                            )
                    except Exception:
                        pass

        return {"received": True}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Service de paiement non disponible"})
    except Exception as e:
        logger.warning(f"Stripe webhook invalide: {e}")
        return JSONResponse(status_code=400, content={"error": "Signature invalide"})


@app.post("/api/payment/confirm/{intent_id}")
async def confirm_concierge_payment(intent_id: str, request: Request):
    """
    Le souscripteur confirme un paiement conciergerie.
    Appele depuis l'interface quand il appuie sur "Confirmer le paiement".
    """
    import stripe as _stripe
    _stripe_secret = os.getenv("STRIPE_SECRET_KEY", "") or os.getenv("STRIPE_API_KEY", "")
    if not _stripe_secret:
        return JSONResponse(status_code=500, content={"error": "Stripe non configure"})
    _stripe.api_key = _stripe_secret

    # Recuperer le tenant_id du souscripteur connecte
    current_tid = str(getattr(request.state, "tenant_id", 0))
    if not current_tid or current_tid == "0":
        return JSONResponse(status_code=401, content={"error": "Authentification requise"})

    try:
        # Recuperer le PaymentIntent
        intent = _stripe.PaymentIntent.retrieve(intent_id)

        # Verifier que c'est bien un paiement conciergerie
        if intent.metadata.get("type") != "concierge":
            return JSONResponse(status_code=403, content={"error": "Ce paiement n'est pas un paiement conciergerie"})

        # Verifier que le paiement appartient au bon tenant
        if intent.metadata.get("tenant_id") != current_tid:
            return JSONResponse(status_code=403, content={"error": "Non autorise"})

        # Verifier qu'il n'est pas deja confirme
        if intent.status in ("succeeded", "processing"):
            return {"status": "already_confirmed", "message": "Ce paiement a deja ete confirme."}

        if intent.status == "canceled":
            return JSONResponse(status_code=410, content={"error": "Ce paiement a ete annule."})

        # Confirmer le paiement
        confirmed = _stripe.PaymentIntent.confirm(intent_id)

        total_eur = confirmed.amount / 100
        merchant = confirmed.metadata.get("merchant", "")
        commission_eur = int(confirmed.metadata.get("commission_cents", 0)) / 100

        logger.info(
            f"CONCIERGE_PAYMENT_CONFIRMED intent={intent_id} "
            f"amount={total_eur:.2f}EUR merchant={merchant}"
        )

        return {
            "status": "confirmed",
            "message": f"Paiement de {total_eur:.2f} EUR confirme pour {merchant}.",
            "amount_eur": total_eur,
            "commission_eur": commission_eur,
        }

    except _stripe.error.CardError as e:
        logger.warning(f"Card error on confirm: {e}")
        return JSONResponse(status_code=402, content={
            "error": f"Carte refusee: {str(e)[:100]}",
            "status": "card_error",
        })
    except _stripe.error.StripeError as e:
        logger.error(f"Stripe confirm error: {e}")
        return JSONResponse(status_code=500, content={"error": f"Erreur Stripe: {str(e)[:100]}"})
    except Exception as e:
        logger.error(f"Payment confirm error: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur lors de la confirmation"})


@app.get("/api/payment/pending")
async def get_pending_payments(request: Request):
    """Liste les paiements conciergerie en attente de confirmation pour le souscripteur."""
    import stripe as _stripe
    _stripe_secret = os.getenv("STRIPE_SECRET_KEY", "") or os.getenv("STRIPE_API_KEY", "")
    if not _stripe_secret:
        return {"pending_payments": [], "count": 0}
    _stripe.api_key = _stripe_secret

    # Utiliser le tenant_id injecte par le middleware (pas de scan Redis)
    tid = getattr(request.state, "tenant_id", 0)
    if not tid:
        return JSONResponse(status_code=401, content={"error": "Authentification requise"})

    try:
        # Chercher les PaymentIntents de type concierge en attente
        intents = _stripe.PaymentIntent.list(
            limit=20,
            created={"gte": int(time.time()) - 86400 * 7},  # 7 derniers jours
        )
        pending = []
        for pi in intents.auto_paging_iter():
            meta = pi.metadata or {}
            if meta.get("type") != "concierge":
                continue
            if meta.get("tenant_id") != str(tid):
                continue
            if pi.status not in ("requires_confirmation", "requires_payment_method"):
                continue
            pending.append({
                "id": pi.id,
                "amount_eur": pi.amount / 100,
                "base_amount_eur": int(meta.get("base_amount_cents", 0)) / 100,
                "commission_eur": int(meta.get("commission_cents", 0)) / 100,
                "merchant": meta.get("merchant", ""),
                "description": pi.description or "",
                "created": pi.created,
            })
            if len(pending) >= 10:
                break

        return {"pending_payments": pending, "count": len(pending)}

    except Exception as e:
        logger.error(f"Error listing pending payments: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur lors de la recuperation des paiements"})


@app.get("/api/admin/commissions")
async def get_admin_commissions(request: Request):
    """Dashboard admin: total des commissions conciergerie du fondateur."""
    if not _redis_client:
        return {"total_commission_eur": 0, "transactions": []}

    try:
        import json as _json_adm
        total_cents = float(_redis_client.client.get(
            f"{_redis_client.prefix}:concierge:total_commission_cents"
        ) or 0)
        raw_list = _redis_client.client.lrange(
            f"{_redis_client.prefix}:concierge:commissions", -50, -1
        )
        transactions = [_json_adm.loads(r) for r in raw_list] if raw_list else []
        # Enrichir avec montant en EUR
        for t in transactions:
            t["commission_eur"] = t.get("commission_cents", 0) / 100
            t["total_eur"] = t.get("total_cents", 0) / 100

        return {
            "total_commission_eur": total_cents / 100,
            "transaction_count": len(transactions),
            "transactions": list(reversed(transactions)),
        }
    except Exception as e:
        logger.error(f"Error fetching commissions: {e}")
        return {"total_commission_eur": 0, "transactions": [], "error": str(e)}


@app.get("/api/setup/stripe-webhook-status")
async def setup_stripe_webhook_status():
    """Retourne si un webhook Stripe a ete recu et verifie."""
    if not _stripe_webhook_received:
        return {"received": False, "message": "Aucun evenement de paiement recu. Envoyez un test depuis le tableau de bord."}
    return {
        "received": True,
        "verified": _stripe_webhook_received.get("verified", False),
        "event_type": _stripe_webhook_received.get("event_type", ""),
        "timestamp": _stripe_webhook_received.get("timestamp", 0),
    }


# =========================================================================
# SETUP ENDPOINTS (accessibles meme si PV non signe)
# =========================================================================

@app.get("/api/setup/status")
async def setup_status():
    """Etat du processus de setup et du PV de recette."""
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        return {
            "pv_signed": PV_SIGNED,
            "pv_locked": _pv_locked,
            "luna_mode": LUNA_MODE,
            "phases": {
                "A": "Verifications techniques (automatisees)",
                "B": "Verifications legales (declarations exploitant)",
                "C": "Verifications operationnelles",
            },
        }
    except ImportError:
        return {"error": "Module pv_recette non disponible", "pv_locked": _pv_locked}


@app.post("/api/setup/check-phase-a")
async def setup_check_phase_a():
    """Lance les verifications techniques automatiques (Phase A)."""
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        results = pv.check_phase_a()
        return {"phase": "A", "results": results}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


@app.post("/api/setup/check-phase-b")
async def setup_check_phase_b(request: Request):
    """Soumettre les declarations legales (Phase B)."""
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        data = await request.json()
        declarations = data.get("declarations", {})
        results = pv.validate_phase_b(declarations)
        return {"phase": "B", "results": results}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


@app.post("/api/setup/check-phase-c")
async def setup_check_phase_c():
    """Lance les verifications operationnelles (Phase C)."""
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        results = pv.check_phase_c()
        return {"phase": "C", "results": results}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


@app.post("/api/setup/check-siret")
async def setup_check_siret(request: Request):
    """Verifie un SIRET via l'API gouvernementale (gratuit, sans cle)."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"status": "fail", "message": "JSON invalide"})

    siret = str(body.get("siret", "")).strip().replace(" ", "")

    # Validation format
    if len(siret) != 14 or not siret.isdigit():
        return JSONResponse(content={"status": "fail", "message": "SIRET invalide (14 chiffres requis)"})

    # Appel API gouv.fr (gratuit, pas de cle)
    try:
        async with httpx.AsyncClient() as client:
            resp = await client.get(
                f"https://recherche-entreprises.api.gouv.fr/search?q={siret}&per_page=1",
                timeout=10.0,
            )
    except Exception as e:
        logger.warning(f"API INSEE indisponible: {e}")
        return JSONResponse(content={"status": "warn", "message": "Service INSEE indisponible, continuez manuellement"})

    if resp.status_code != 200:
        return JSONResponse(content={"status": "warn", "message": "Service INSEE indisponible, continuez manuellement"})

    data = resp.json()
    results = data.get("results", [])
    if not results:
        return JSONResponse(content={"status": "fail", "message": "SIRET introuvable dans le registre INSEE"})

    company = results[0]
    siege = company.get("siege", {})

    # Verifier que le SIRET retourne correspond bien a celui demande
    returned_siret = siege.get("siret", "")
    if returned_siret != siret:
        return JSONResponse(content={"status": "fail", "message": "SIRET introuvable dans le registre INSEE"})

    # Verifier que l'entreprise est active
    if company.get("etat_administratif") != "A":
        return JSONResponse(content={"status": "fail", "message": "Entreprise fermee ou radiee"})

    # Code APE — liste des codes compatibles avec l'exploitation Luna
    ape = company.get("activite_principale", "")
    ape_compatible_codes = [
        "62.01Z", "62.02A", "62.02B", "62.03Z", "62.09Z",  # Informatique
        "63.11Z", "63.12Z", "63.99Z",                       # Traitement donnees
        "58.29C", "58.29A",                                   # Edition logiciels
        "70.22Z", "74.90B",                                   # Conseil
        "85.59B", "88.10C", "88.99B",                        # Aide a la personne
        "47.41Z", "47.42Z",                                   # Commerce electronique
    ]

    return JSONResponse(content={
        "status": "ok",
        "company": {
            "nom": company.get("nom_complet", ""),
            "siren": company.get("siren", ""),
            "siret": siege.get("siret", siret),
            "adresse": siege.get("adresse", ""),
            "code_postal": siege.get("code_postal", ""),
            "commune": siege.get("libelle_commune", ""),
            "ape": ape,
            "ape_compatible": ape in ape_compatible_codes,
            "nature_juridique": company.get("nature_juridique", ""),
            "etat": company.get("etat_administratif", ""),
            "dirigeants": [
                {"nom": d.get("nom", ""), "prenoms": d.get("prenoms", ""), "qualite": d.get("qualite", "")}
                for d in company.get("dirigeants", [])[:3]
                if d.get("type_dirigeant") == "personne physique"
            ],
        },
        "message": "Entreprise verifiee",
    })


@app.post("/api/setup/sign-pv")
async def setup_sign_pv(request: Request):
    """Signe le PV de recette et deverrouille le serveur."""
    global PV_SIGNED, _pv_locked, _setup_permanently_disabled, SETUP_OPENAI_API_KEY
    global _setup_chat_history, _setup_message_count

    async with _sign_pv_lock:
        if _setup_permanently_disabled:
            return JSONResponse(status_code=410, content={
                "error": "Installation terminee",
                "message": "Le PV a deja ete signe. Impossible de re-signer.",
            })
        try:
            import json as _json_pv
            from pv_recette import PVRecette
            pv = PVRecette()
            data = await request.json()

            # Fallback: si exploitant_name/siret vides, lire depuis wizard_state.json
            exploitant_name = data.get("exploitant_name", "").strip()
            exploitant_siret = data.get("exploitant_siret", "").strip()
            if not exploitant_name or not exploitant_siret:
                ws_path = os.path.join(os.path.dirname(__file__), ".wizard_state.json")
                if os.path.exists(ws_path):
                    ws = _json_pv.loads(open(ws_path, encoding="utf-8").read())
                    ws_config = ws.get("config", {})
                    if not exploitant_name:
                        exploitant_name = ws_config.get("EXPLOITANT_NAME", "")
                    if not exploitant_siret:
                        exploitant_siret = ws_config.get("EXPLOITANT_SIRET", "")

            result = pv.sign(
                exploitant_name=exploitant_name,
                exploitant_siret=exploitant_siret,
                declarations_b=data.get("declarations", {}),
            )
            if "error" in result:
                return JSONResponse(status_code=400, content=result)

            # Mettre a jour le .env automatiquement
            env_path = os.path.join(os.path.dirname(__file__), ".env")
            pv.update_env_file(env_path, result["env_updates"])

            # Mise a jour des globals runtime (pas besoin de restart)
            PV_SIGNED = True
            _pv_locked = False
            _setup_permanently_disabled = True

            # Feature 1: Destruction cle fondateur
            SETUP_OPENAI_API_KEY = None
            _setup_chat_history.clear()
            _setup_message_count = _SETUP_MAX_MESSAGES

            # Feature 2: Generer le certificat d'autonomie
            cert_path = None
            reset_code = result.get("reset_code", "")
            try:
                cert_path = _generate_autonomy_certificate(result["pv_data"], reset_code)
            except Exception as e:
                logger.error(f"Erreur generation certificat: {e}")

            logger.info(f"PV signe — setup DESACTIVE, cle fondateur detruite")

            return {
                "success": True,
                "reset_code": reset_code,
                "certificate_url": "/api/admin/certificate" if cert_path else None,
                "message": "PV signe. CONSERVEZ le reset code. Telechargez le certificat.",
                "pv_data": result["pv_data"],
            }
        except ImportError:
            return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


@app.get("/api/setup/phase-b-checklist")
async def setup_phase_b_checklist():
    """Retourne la checklist des declarations legales a remplir."""
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        return {"checklist": pv.get_phase_b_checklist()}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


# =========================================================================
# SETUP AI + WEB WIZARD ENDPOINTS
# =========================================================================

# Luna Setup AI state (in-memory, pas besoin de Redis en setup)
_setup_chat_history: list = []
_setup_message_count: int = 0
_SETUP_MAX_MESSAGES = 50

_SETUP_SYSTEM_PROMPT = """Tu es Luna Setup, l'assistante d'installation YAWatch Luna.
Tu guides un exploitant (potentiellement non-technicien) a travers la configuration.
Tu parles en francais, de maniere simple et encourageante.

CONTEXTE:
- L'exploitant configure une instance de Luna pour ses propres clients
- Il doit creer des comptes: OpenAI, Twilio, Stripe, optionnellement Tavus
- Chaque service a ses propres cles API a copier
- Apres la config, il y a un PV de recette en 3 phases (technique, legal, operationnel)

SERVICES:
- OpenAI: intelligence de Luna, cle commence par 'sk-', https://platform.openai.com
- Twilio: SMS, Account SID commence par 'AC', https://twilio.com
- Stripe: paiements, cle commence par 'sk_test_' ou 'sk_live_', https://stripe.com
- Tavus: visio avatar (mode Full seulement), https://tavus.io
- Redis: memoire cache, redis://localhost:6379/0

PLANS LUNA:
- Essentiel (79 EUR/mois): Chat illimite + 60 min voix + 20 min visio + 50 SMS
- Confort (149 EUR/mois): Chat illimite + 180 min voix + 60 min visio + 200 SMS
- Premium (249 EUR/mois): Chat illimite + 300 min voix + 180 min visio + 200 SMS

TROUBLESHOOTING COURANT:
- "cle ne fonctionne pas" -> verifier copier-coller, espaces en trop, cle expiree
- "Redis non connecte" -> sudo systemctl start redis-server ou docker run -d redis
- "SSL erreur" -> certificats auto-signes OK pour dev, Let's Encrypt pour prod
- "Stripe webhook" -> URL doit etre publique HTTPS, utiliser Stripe CLI en dev

REGLES:
- Sois patient et encourage
- Explique en langage simple, evite le jargon technique
- Si tu ne sais pas, dis-le et suggere de contacter le support YAWatch
- Ne donne JAMAIS de cles API, tokens ou secrets
- Ne modifie rien toi-meme, guide l'exploitant a le faire via l'interface

MODELE ECONOMIQUE 70/30:
- Le fondateur recoit 70% du CA brut TTC encaisse par l'exploitant
- L'exploitant conserve 30% brut
- Sur ces 30%, l'exploitant reverse 6% du CA brut a Ambre (communication) + 300 EUR fixe/mois
- Les couts API (OpenAI, Twilio, Tavus, Stripe) sont a la charge de l'exploitant sur ses 30%
- Budget API estime pour 100 clients: ~1 200 EUR/mois

CONSEILS ADMINISTRATIFS:
- RC Professionnelle obligatoire (minimum recommande: 500 000 EUR). Assureurs: Hiscox, AXA Pro, MMA Pro
- SIRET doit etre actif et le code APE compatible (informatique, conseil, aide a la personne)
- Hebergement serveur obligatoirement en UE (RGPD). Recommandes: OVH, Scaleway, Hetzner
- DPO necessaire si plus de 250 salaries ou traitement de donnees sensibles a grande echelle

AVERTISSEMENTS JURIDIQUES CRITIQUES:
- Luna n'est PAS un dispositif medical. Toute confusion = exercice illegal medecine (Art. L4161-1 CSP, 2 ans + 30K EUR)
- Detection de detresse = best-effort UNIQUEMENT. Promettre une garantie = mise en danger d'autrui (Art. 223-1 Code penal)
- Non-conformite RGPD = amende jusqu'a 20M EUR ou 4% du CA mondial (Art. 83 RGPD)
- Toujours recommander d'appeler le 15/112 avant tout, Luna ne remplace PAS les secours

HEBERGEMENT RECOMMANDE:
- OVH VPS (France): a partir de 6 EUR/mois, Docker compatible
- Scaleway (France): a partir de 8 EUR/mois, GPU optionnel
- Hetzner (Allemagne): a partir de 5 EUR/mois, excellent rapport qualite/prix
- AWS/GCP: OK si region eu-west-1 / europe-west1
"""


@app.post("/api/setup/ai-chat")
async def setup_ai_chat(request: Request):
    """Chat avec Luna Setup AI pendant la configuration."""
    global _setup_message_count

    if not SETUP_OPENAI_API_KEY:
        return JSONResponse(status_code=400, content={
            "error": "SETUP_OPENAI_API_KEY non configuree",
            "message": "Demandez une cle de setup a votre referent YAWatch.",
        })

    if _setup_message_count >= _SETUP_MAX_MESSAGES:
        return JSONResponse(status_code=429, content={
            "error": "Quota setup epuise",
            "message": f"Vous avez utilise vos {_SETUP_MAX_MESSAGES} messages. Contactez support@yawatch.fr.",
            "count": _setup_message_count,
            "max": _SETUP_MAX_MESSAGES,
        })

    data = await request.json()
    user_msg = data.get("message", "").strip()
    if not user_msg:
        return JSONResponse(status_code=400, content={"error": "Message vide"})

    _setup_chat_history.append({"role": "user", "content": user_msg})

    try:
        setup_client = OpenAI(api_key=SETUP_OPENAI_API_KEY)
        messages = [{"role": "system", "content": _SETUP_SYSTEM_PROMPT}] + _setup_chat_history[-20:]
        resp = setup_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=messages,
            max_tokens=500,
            temperature=0.7,
        )
        assistant_msg = resp.choices[0].message.content
        _setup_chat_history.append({"role": "assistant", "content": assistant_msg})
        _setup_message_count += 1

        return {
            "response": assistant_msg,
            "count": _setup_message_count,
            "max": _SETUP_MAX_MESSAGES,
            "remaining": _SETUP_MAX_MESSAGES - _setup_message_count,
        }
    except Exception as e:
        logger.error(f"Setup AI error: {e}")
        return JSONResponse(status_code=500, content={
            "error": "Erreur Luna Setup",
            "message": str(e),
        })


@app.post("/api/setup/save-config")
async def setup_save_config(request: Request):
    """Sauvegarde une etape de configuration dans le .env."""
    import json as _json

    data = await request.json()
    step = data.get("step", "")
    config = data.get("config", {})

    # Phase B envoie des declarations sans config — ne pas rejeter
    if not config and step != "phase_b":
        return JSONResponse(status_code=400, content={"error": "Configuration vide"})

    env_path = os.path.join(os.path.dirname(__file__), ".env")

    # Bootstrap: creer .env minimal si inexistant
    if not os.path.exists(env_path):
        from pathlib import Path
        Path(env_path).write_text("PV_SIGNED=false\nLUNA_MODE=lite\n", encoding="utf-8")

    # Mettre a jour les cles dans .env
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        pv.update_env_file(env_path, config)
    except ImportError:
        # Fallback: ecriture directe
        from pathlib import Path
        content = Path(env_path).read_text(encoding="utf-8")
        for k, v in config.items():
            if f"{k}=" in content:
                lines = content.split("\n")
                for i, line in enumerate(lines):
                    if line.startswith(f"{k}="):
                        lines[i] = f"{k}={v}"
                        break
                content = "\n".join(lines)
            else:
                content += f"\n{k}={v}"
        Path(env_path).write_text(content, encoding="utf-8")

    # Recharger le .env dans le processus pour que os.getenv() retourne les nouvelles valeurs
    _reload_env()

    # Sauvegarder l'etat du wizard
    state_path = os.path.join(os.path.dirname(__file__), ".wizard_state.json")
    state = {}
    if os.path.exists(state_path):
        state = _json.loads(open(state_path, encoding="utf-8").read())
    state.setdefault("completed_steps", [])
    if step and step not in state["completed_steps"]:
        state["completed_steps"].append(step)
    state["config"] = {**state.get("config", {}), **config}
    state["saved_at"] = datetime.now().isoformat()

    # Sauvegarder les declarations Phase B dans le wizard state
    if step == "phase_b":
        declarations = data.get("phase_b_declarations", data.get("declarations", data.get("config", {})))
        state["phase_b_declarations"] = declarations

    with open(state_path, "w", encoding="utf-8") as f:
        _json.dump(state, f, indent=2, ensure_ascii=False)

    return {"success": True, "step": step, "saved_keys": list(config.keys())}


@app.get("/api/setup/wizard-state")
async def setup_wizard_state():
    """Retourne l'etat du wizard pour reprendre."""
    import json as _json

    state_path = os.path.join(os.path.dirname(__file__), ".wizard_state.json")
    if os.path.exists(state_path):
        state = _json.loads(open(state_path, encoding="utf-8").read())
        # Masquer les valeurs sensibles
        safe_config = {}
        for k, v in state.get("config", {}).items():
            if isinstance(v, str) and any(s in k.upper() for s in ["KEY", "TOKEN", "SECRET", "PASSWORD"]):
                safe_config[k] = ("***" + v[-4:]) if len(v) > 4 else "****"
            else:
                safe_config[k] = v
        state["config"] = safe_config
        # Inclure les declarations Phase B si presentes
        if "phase_b_declarations" not in state:
            state["phase_b_declarations"] = {}
        return state
    return {"completed_steps": [], "config": {}, "phase_b_declarations": {}}


@app.post("/api/setup/test-service")
async def setup_test_service(request: Request):
    """Teste un service specifique avec les credentials fournis."""
    data = await request.json()
    service = data.get("service", "")
    creds = data.get("credentials", {})

    try:
        from pv_recette import PVRecette
        pv = PVRecette()

        if hasattr(pv, "test_service"):
            result = pv.test_service(service, creds)
        else:
            # Fallback: utiliser les checks existants (depuis env vars)
            check_map = {
                "openai": pv._check_openai,
                "twilio": pv._check_twilio,
                "stripe": pv._check_stripe,
                "redis": pv._check_redis,
                "tavus": pv._check_tavus,
            }
            if service not in check_map:
                return JSONResponse(status_code=400, content={"error": f"Service inconnu: {service}"})
            result = check_map[service]()

        return {"service": service, "result": result}
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module pv_recette non disponible"})


@app.post("/api/setup/generate-security")
async def setup_generate_security():
    """Genere JWT secret et certificats SSL automatiquement."""
    import secrets as _secrets
    import subprocess

    results = {}

    # JWT
    jwt_key = _secrets.token_hex(32)
    results["jwt"] = {"status": "ok", "message": f"Cle JWT generee ({len(jwt_key)} chars)"}

    # SSL
    ssl_dir = os.path.dirname(__file__)
    cert_path = os.path.join(ssl_dir, "cert.pem")
    key_path = os.path.join(ssl_dir, "key.pem")

    if os.path.exists(cert_path) and os.path.exists(key_path):
        results["ssl"] = {"status": "ok", "message": "Certificats SSL deja presents"}
    else:
        try:
            subprocess.run([
                "openssl", "req", "-x509", "-newkey", "rsa:4096",
                "-keyout", key_path, "-out", cert_path,
                "-days", "365", "-nodes",
                "-subj", "/CN=localhost/O=YAWatch-Luna",
            ], check=True, capture_output=True)
            results["ssl"] = {"status": "ok", "message": "Certificats SSL generes (1 an)"}
        except Exception as e:
            results["ssl"] = {"status": "fail", "message": f"Erreur SSL: {e}"}

    # Sauvegarder dans .env
    env_path = os.path.join(ssl_dir, ".env")
    config_updates = {
        "JWT_SECRET_KEY": jwt_key,
        "JWT_ALGORITHM": "HS256",
        "SSL_CERTFILE": "./cert.pem",
        "SSL_KEYFILE": "./key.pem",
    }
    try:
        from pv_recette import PVRecette
        pv = PVRecette()
        pv.update_env_file(env_path, config_updates)
    except ImportError:
        pass

    return {"success": True, "results": results, "saved_keys": list(config_updates.keys())}


@app.post("/api/setup/stripe-auto")
async def setup_stripe_auto(request: Request):
    """Cree automatiquement les 3 produits Stripe."""
    data = await request.json()
    stripe_key = data.get("stripe_api_key", "")
    if not stripe_key:
        return JSONResponse(status_code=400, content={"error": "STRIPE_API_KEY requis"})

    try:
        import stripe
        stripe.api_key = stripe_key
        from stripe_setup import create_products_and_prices
        results = create_products_and_prices(stripe)
        return {"success": True, "prices": results}
    except ImportError as e:
        logger.error(f"Stripe setup import error: {e}")
        return JSONResponse(status_code=500, content={"error": "Service de paiement non disponible"})
    except Exception as e:
        logger.error(f"Stripe setup error: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur de configuration des paiements."})


# =========================================================================
# PROFILE ENDPOINTS
# =========================================================================

@app.get("/api/profile")
async def get_profile(request: Request):
    """Retourne le profil souscripteur"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    profile = mgr.get_subscriber_profile()
    if not profile:
        return {"profile": None}
    return {"profile": profile.model_dump()}


@app.post("/api/profile")
async def set_profile(req: ProfileRequest, request: Request):
    """Cree ou remplace le profil souscripteur"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    profile = SubscriberProfile(tenant_id=tid, **req.model_dump())
    mgr.set_subscriber_profile(profile)
    _gamify(tid, "profile_update")
    return {"success": True, "profile": profile.model_dump()}


@app.patch("/api/profile")
async def update_profile(request: Request):
    """Mise a jour partielle du profil"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})
    if not body:
        return JSONResponse(status_code=400, content={"error": "Corps vide"})
    profile = mgr.update_subscriber_profile(body)
    if not profile:
        return JSONResponse(status_code=404, content={"error": "Profil non trouve"})
    _gamify(tid, "profile_update")
    # Auto-sync avatar type when gender is updated
    if "gender" in body and _GAMIFICATION_AVAILABLE and _redis_client:
        try:
            from core.gamification.redis_ops import GamificationRedisOps
            gops = GamificationRedisOps(_redis_client)
            current_avatar = gops.get_avatar_type(tid) or "adult_man"
            new_gender = body["gender"].upper().strip()
            if new_gender == "F" and current_avatar in ("adult_man", ""):
                gops.set_avatar_type(tid, "adult_woman")
                logger.info(f"Avatar auto-updated to adult_woman for tenant {tid}")
            elif new_gender == "M" and current_avatar in ("adult_woman", ""):
                gops.set_avatar_type(tid, "adult_man")
                logger.info(f"Avatar auto-updated to adult_man for tenant {tid}")
        except Exception as e:
            logger.debug(f"Avatar gender sync: {e}")
    return {"success": True, "profile": profile.model_dump()}


# =========================================================================
# GEOLOCATION ENDPOINTS
# =========================================================================

@app.post("/api/geolocation")
async def update_geolocation(request: Request):
    """
    Stocke la position GPS du souscripteur (envoyee par le navigateur/APK).
    Stocke dans Redis avec TTL 1h (position fraiche).
    """
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})

    lat = body.get("latitude")
    lng = body.get("longitude")
    accuracy = body.get("accuracy")
    city = body.get("city", "")
    address = body.get("address", "")

    if lat is None or lng is None:
        return JSONResponse(status_code=400, content={"error": "latitude et longitude requis"})

    import json as _json
    geo_data = _json.dumps({
        "latitude": lat,
        "longitude": lng,
        "accuracy": accuracy,
        "city": city,
        "address": address,
        "updated_at": time.time(),
    })
    key = f"luna:{tid}:geolocation"
    _redis_client.client.set(key, geo_data, ex=3600)  # TTL 1h
    logger.info(f"Geolocation updated for tenant {tid}: {lat},{lng} ({city})")

    # Auto-detect language from country in address
    if address and _redis_client:
        _COUNTRY_LANG = {
            "france": "fr", "belgique": "fr", "suisse": "fr", "canada": "fr",
            "luxembourg": "fr", "monaco": "fr", "sénégal": "fr", "cameroun": "fr",
            "côte d'ivoire": "fr", "mali": "fr", "maroc": "fr", "tunisie": "fr",
            "guadeloupe": "fr", "martinique": "fr", "guyane": "fr", "réunion": "fr",
            "united states": "en", "united kingdom": "en", "australia": "en",
            "ireland": "en", "new zealand": "en",
            "españa": "es", "spain": "es", "méxico": "es", "colombia": "es",
            "deutschland": "de", "germany": "de", "österreich": "de",
            "italia": "it", "italy": "it",
            "portugal": "pt", "brasil": "pt", "brazil": "pt",
            "nederland": "nl", "netherlands": "nl",
        }
        _addr_lower = address.lower()
        for _country, _lang in _COUNTRY_LANG.items():
            if _country in _addr_lower:
                try:
                    _settings_key = f"luna:{tid}:settings"
                    _cur_lang = _redis_client.client.hget(_settings_key, "language")
                    if not _cur_lang:
                        _redis_client.client.hset(_settings_key, "language", _lang)
                        _redis_client.client.hset(_settings_key, "country", _country)
                        logger.info(f"Auto-detected language '{_lang}' for tenant {tid} from country '{_country}'")
                except Exception:
                    pass
                break

    return {"success": True}


@app.get("/api/geolocation")
async def get_geolocation(request: Request):
    """Retourne la derniere position connue du souscripteur."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    import json as _json
    key = f"luna:{tid}:geolocation"
    raw = _redis_client.client.get(key)
    if not raw:
        return {"geolocation": None, "message": "Aucune position connue"}
    geo = _json.loads(raw)
    return {"geolocation": geo}


# =========================================================================
# NOTIFICATIONS & SETTINGS ENDPOINTS
# =========================================================================

@app.get("/api/notifications/prefs")
async def get_notification_prefs(request: Request):
    """Get notification preferences for the current user."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    from core.notifications.redis_ops import NotificationRedisOps
    nops = NotificationRedisOps(_redis_client)
    prefs = nops.get_prefs(tid)
    return {"prefs": prefs}


@app.post("/api/notifications/prefs")
async def set_notification_prefs(request: Request):
    """Update notification preferences."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    body = await request.json()
    from core.notifications.redis_ops import NotificationRedisOps
    nops = NotificationRedisOps(_redis_client)
    allowed_fields = {
        "enabled", "streak_risk", "comeback", "mission_reminder",
        "level_up", "weekly_summary", "sound",
        "quiet_hours_start", "quiet_hours_end",
    }
    for k, v in body.items():
        if k in allowed_fields:
            nops.update_pref(tid, k, str(v))
    return {"success": True, "prefs": nops.get_prefs(tid)}


@app.get("/api/notifications/pending")
async def get_pending_notifications(request: Request):
    """Poll for pending notifications. Returns up to 5 and removes them."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return {"notifications": []}
    from core.notifications.redis_ops import NotificationRedisOps
    nops = NotificationRedisOps(_redis_client)
    pending = nops.pop_pending(tid, limit=5)
    return {"notifications": pending}


@app.get("/api/notifications/count")
async def get_notification_count(request: Request):
    """Returns count of pending notifications (for badge display)."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return {"count": 0}
    from core.notifications.redis_ops import NotificationRedisOps
    nops = NotificationRedisOps(_redis_client)
    return {"count": nops.peek_pending(tid)}


@app.get("/api/settings")
async def get_settings(request: Request):
    """Get all user settings (appearance, sound, etc.)."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    key = _redis_client._key(tid, "settings")
    data = _redis_client.client.hgetall(key)
    defaults = {
        "dark_mode": "1",
        "font_size": "normal",
        "language": "fr",
        "notification_sound": "1",
    }
    defaults.update(data)
    return {"settings": defaults}


@app.post("/api/settings")
async def update_settings(request: Request):
    """Update user settings."""
    tid = getattr(request.state, "tenant_id", 1)
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service indisponible"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})
    if not isinstance(body, dict):
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})
    allowed = {"dark_mode", "font_size", "language", "notification_sound"}
    updates = {k: str(v) for k, v in body.items() if k in allowed}
    if updates:
        key = _redis_client._key(tid, "settings")
        _redis_client.client.hset(key, mapping=updates)
    return {"success": True}


# =========================================================================
# CONTACTS ENDPOINTS
# =========================================================================

_MAX_CONTACTS_PROPRIO = 30
_MAX_CONTACTS_DEFAULT = 5

def _get_max_contacts(tid: int) -> int:
    return _MAX_CONTACTS_PROPRIO if tid == _PROPRIO_TENANT_ID else _MAX_CONTACTS_DEFAULT

@app.get("/api/contacts")
async def list_contacts(request: Request):
    """Liste les contacts de confiance"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    contacts = mgr.list_trusted_contacts()
    return {
        "contacts": [
            {
                "phone": c.phone,
                "name": c.name,
                "relation": c.relation,
                "email": c.email or "",
                "address": c.address or "",
                "preferred_channel": c.preferred_channel.value if hasattr(c.preferred_channel, "value") else str(c.preferred_channel),
                "emergency_only": c.emergency_only,
                "verified_at": c.verified_at.isoformat() if c.verified_at else None,
            }
            for c in contacts
        ],
        "count": len(contacts),
        "max": _get_max_contacts(tid),
    }


@app.post("/api/contacts")
async def add_contact(req: ContactRequest, request: Request):
    """Ajoute un contact de confiance (max 5)"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        channel = Channel(req.preferred_channel) if req.preferred_channel in [c.value for c in Channel] else Channel.SMS
        contact = mgr.add_trusted_contact(
            phone=req.phone,
            name=req.name,
            relation=req.relation,
            preferred_channel=channel,
            emergency_only=req.emergency_only,
            email=req.email or "",
            address=req.address or "",
        )
        _gamify(tid, "add_contact")
        return {"success": True, "contact": {
            "phone": contact.phone,
            "name": contact.name,
            "relation": contact.relation,
            "email": contact.email or "",
            "address": contact.address or "",
        }}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.delete("/api/contacts/{phone}")
async def delete_contact(phone: str, request: Request):
    """Supprime un contact de confiance"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    removed = mgr.remove_trusted_contact(phone)
    if not removed:
        return JSONResponse(status_code=404, content={"error": "Contact introuvable"})
    return {"success": True}


# =========================================================================
# FAMILY PACK ENDPOINTS
# =========================================================================

from core.memory.schemas import (
    FamilyMember, FamilyGroup, FamilyRole, FamilyMemberType,
    EscalationRule, EscalationStage, FamilyMessage, FamilyMessageType,
    FamilyAuditLog, DistressLevel, DistressKeywords
)
import random
import string


def _generate_otp() -> str:
    """Genere un code OTP a 6 chiffres"""
    return ''.join(random.choices(string.digits, k=6))


def _log_family_audit(action: str, actor_phone: str, actor_name: str,
                      target_phone: str = None, target_name: str = None,
                      details: dict = None, severity: str = "info",
                      tenant_id: int = None):
    """Log une action dans l'audit famille"""
    if not _redis_client:
        return
    tid = tenant_id if tenant_id is not None else TENANT_ID
    import json
    from datetime import datetime
    audit = FamilyAuditLog(
        tenant_id=tid,
        actor_phone=actor_phone,
        actor_name=actor_name,
        action=action,
        target_phone=target_phone,
        target_name=target_name,
        details=details or {},
        severity=severity,
    )
    _redis_client.add_family_audit(tid, json.dumps(audit.to_redis()))


# --- Family Group ---

@app.get("/api/family")
async def get_family(request: Request):
    """Recupere le groupe familial et ses membres"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    group_data = _redis_client.get_family_group(tid)
    if not group_data:
        return {"group": None, "members": [], "count": 0}

    group = FamilyGroup.from_redis(group_data)
    members = []
    for phone in _redis_client.get_family_members(tid):
        member_data = _redis_client.get_family_member(tid, phone)
        if member_data:
            member = FamilyMember.from_redis(member_data)
            members.append({
                "id": member.id,
                "phone": member.phone,
                "name": member.name,
                "relation": member.relation,
                "member_type": member.member_type.value,
                "role": member.role.value,
                "age": member.age,
                "is_verified": member.is_verified,
                "is_minor": member.is_minor(),
                "can_see_history": member.can_see_history,
                "can_send_messages": member.can_send_messages,
                "can_receive_alerts": member.can_receive_alerts,
                "is_active": member.is_active,
                "created_at": member.created_at.isoformat(),
            })

    return {
        "group": {
            "id": group.id,
            "name": group.name,
            "description": group.description,
            "allow_internal_messaging": group.allow_internal_messaging,
            "share_activity_summary": group.share_activity_summary,
            "auto_escalate": group.auto_escalate,
        },
        "members": members,
        "count": len(members),
    }


@app.post("/api/family")
async def create_family(request: Request):
    """Cree ou met a jour le groupe familial"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    from datetime import datetime

    # Verifier si groupe existe deja
    existing = _redis_client.get_family_group(tid)
    if existing:
        # Mise a jour
        group = FamilyGroup.from_redis(existing)
        if "name" in data:
            group.name = data["name"]
        if "description" in data:
            group.description = data["description"]
        if "allow_internal_messaging" in data:
            group.allow_internal_messaging = data["allow_internal_messaging"]
        if "share_activity_summary" in data:
            group.share_activity_summary = data["share_activity_summary"]
        if "auto_escalate" in data:
            group.auto_escalate = data["auto_escalate"]
        group.updated_at = datetime.utcnow()
    else:
        # Creation
        group = FamilyGroup(
            tenant_id=tid,
            name=data.get("name", "Ma famille"),
            description=data.get("description", ""),
        )

    _redis_client.set_family_group(tid, group.to_redis())
    _log_family_audit("family_group_created" if not existing else "family_group_updated",
                      "system", "Luna", details={"name": group.name}, tenant_id=tid)

    if not existing:
        _gamify(tid, "family_setup")

    return {"success": True, "group_id": group.id}


# --- Family Members ---

@app.get("/api/family/members")
async def list_family_members(request: Request):
    """Liste tous les membres de la famille"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    members = []
    for phone in _redis_client.get_family_members(tid):
        member_data = _redis_client.get_family_member(tid, phone)
        if member_data:
            member = FamilyMember.from_redis(member_data)
            members.append({
                "id": member.id,
                "phone": member.phone,
                "name": member.name,
                "relation": member.relation,
                "role": member.role.value,
                "member_type": member.member_type.value,
                "age": member.age,
                "is_verified": member.is_verified,
                "is_minor": member.is_minor(),
                "needs_parental_consent": member.needs_parental_consent(),
            })

    return {"members": members, "count": len(members)}


@app.post("/api/family/members")
async def add_family_member(request: Request):
    """Ajoute un membre a la famille (envoie OTP de verification)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()

    # Validation
    phone = data.get("phone", "").strip()
    name = data.get("name", "").strip()
    relation = data.get("relation", "").strip()

    if not phone or not name:
        return JSONResponse(status_code=400, content={"error": "phone et name requis"})

    # Verifier si deja membre
    if _redis_client.get_family_member(tid, phone):
        return JSONResponse(status_code=400, content={"error": "Ce membre existe deja"})

    # Valider le role
    role_str = data.get("role", "information_only")
    valid_roles = [r.value for r in FamilyRole]
    if role_str not in valid_roles:
        return JSONResponse(status_code=400, content={
            "error": f"Role invalide: {role_str}. Roles valides: {', '.join(valid_roles)}"
        })

    # Determiner le type de membre selon l'age
    age = data.get("age")
    member_type = FamilyMemberType.ADULT
    if age:
        age = int(age)
        if age < 13:
            member_type = FamilyMemberType.CHILD
        elif age < 18:
            member_type = FamilyMemberType.TEEN
        elif age >= 65:
            member_type = FamilyMemberType.SENIOR

    # Creer le membre (non verifie)
    member = FamilyMember(
        tenant_id=tid,
        phone=phone,
        name=name,
        relation=relation,
        member_type=member_type,
        role=FamilyRole(role_str),
        age=age,
        is_verified=False,
        can_see_history=data.get("can_see_history", False),
        can_send_messages=data.get("can_send_messages", True),
        can_receive_alerts=data.get("can_receive_alerts", True),
    )

    # Generer et stocker OTP
    otp = _generate_otp()
    member.verification_code = otp
    _redis_client.set_otp(phone, otp)

    # Ajouter le membre
    success = _redis_client.add_family_member(tid, phone, member.to_redis())
    if not success:
        return JSONResponse(status_code=400, content={"error": "Quota de membres atteint (max 15)"})

    # Envoyer SMS avec OTP
    sms_sent = False
    if sms_client:
        message = f"Luna Family: {name}, votre code de verification est {otp}. Valide 10 min."
        sms_sent, _ = _tracked_sms_send(phone, message, label=f"OTP famille {name}")

    _log_family_audit("member_invited", "system", "Luna",
                      target_phone=phone, target_name=name,
                      details={"role": member.role.value, "sms_sent": sms_sent},
                      tenant_id=tid)

    _gamify(tid, "family_member_added")

    return {
        "success": True,
        "member_id": member.id,
        "verification_required": True,
        "sms_sent": sms_sent,
        "message": f"Code de verification envoye a {phone}" if sms_sent else "Echec envoi SMS"
    }


@app.post("/api/family/members/{phone}/verify")
async def verify_family_member(phone: str, request: Request):
    """Verifie un membre avec son code OTP"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    otp = data.get("code", "").strip()

    if not otp:
        return JSONResponse(status_code=400, content={"error": "Code requis"})

    # Verifier OTP
    if not _redis_client.verify_otp(phone, otp):
        return JSONResponse(status_code=400, content={"error": "Code invalide ou expire"})

    # Marquer comme verifie
    from datetime import datetime
    _redis_client.update_family_member(tid, phone, {
        "is_verified": "1",
        "verified_at": datetime.utcnow().isoformat(),
        "verification_code": "",
    })

    member_data = _redis_client.get_family_member(tid, phone)
    name = member_data.get("name", "Membre") if member_data else "Membre"

    _log_family_audit("member_verified", phone, name, tenant_id=tid)

    _gamify(tid, "family_member_verified")

    return {"success": True, "message": f"{name} verifie avec succes"}


@app.patch("/api/family/members/{phone}")
async def update_family_member(phone: str, request: Request):
    """Met a jour un membre de la famille"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()

    member_data = _redis_client.get_family_member(tid, phone)
    if not member_data:
        return JSONResponse(status_code=404, content={"error": "Membre non trouve"})

    # Champs modifiables
    updates = {}
    if "name" in data:
        updates["name"] = data["name"]
    if "role" in data:
        updates["role"] = data["role"]
    if "can_see_history" in data:
        updates["can_see_history"] = "1" if data["can_see_history"] else "0"
    if "can_send_messages" in data:
        updates["can_send_messages"] = "1" if data["can_send_messages"] else "0"
    if "can_receive_alerts" in data:
        updates["can_receive_alerts"] = "1" if data["can_receive_alerts"] else "0"
    if "is_active" in data:
        updates["is_active"] = "1" if data["is_active"] else "0"

    from datetime import datetime
    updates["updated_at"] = datetime.utcnow().isoformat()

    _redis_client.update_family_member(tid, phone, updates)

    _log_family_audit("member_updated", "system", "Luna",
                      target_phone=phone, target_name=member_data.get("name"),
                      details=updates, tenant_id=tid)

    return {"success": True}


@app.delete("/api/family/members/{phone}")
async def remove_family_member(phone: str, request: Request):
    """Supprime un membre de la famille"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    member_data = _redis_client.get_family_member(tid, phone)
    if not member_data:
        return JSONResponse(status_code=404, content={"error": "Membre non trouve"})

    name = member_data.get("name", "Membre")
    _redis_client.remove_family_member(tid, phone)

    _log_family_audit("member_removed", "system", "Luna",
                      target_phone=phone, target_name=name, tenant_id=tid)

    return {"success": True}


# --- Family Messages (Internal Messaging) ---

@app.get("/api/family/messages")
async def get_family_messages(request: Request, limit: int = 50, phone: str = None):
    """Recupere les messages famille (optionnel: filtre par destinataire)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    messages = []
    for msg_id in _redis_client.get_family_messages(tid, limit=limit):
        msg_data = _redis_client.get_family_message(tid, msg_id)
        if msg_data:
            msg = FamilyMessage.from_redis(msg_data)
            # Filtrer par destinataire si specifie
            if phone and msg.to_phone and msg.to_phone != phone:
                continue
            messages.append({
                "id": msg.id,
                "from_phone": msg.from_phone,
                "from_name": msg.from_name,
                "to_phone": msg.to_phone,
                "to_name": msg.to_name,
                "content": msg.content,
                "message_type": msg.message_type.value,
                "is_read": msg.is_read,
                "read_by": msg.read_by,
                "created_at": msg.created_at.isoformat(),
                "requires_response": msg.requires_response,
                "responses": msg.responses,
            })

    return {"messages": messages, "count": len(messages)}


@app.post("/api/family/messages")
async def send_family_message(request: Request):
    """Envoie un message interne famille via Luna"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()

    from_phone = data.get("from_phone", "luna")
    from_name = data.get("from_name", "Luna")
    to_phone = data.get("to_phone")  # None = groupe
    content = data.get("content", "").strip()
    message_type = data.get("message_type", "update")

    if not content:
        return JSONResponse(status_code=400, content={"error": "Contenu requis"})

    # Creer le message
    msg = FamilyMessage(
        tenant_id=tid,
        from_phone=from_phone,
        from_name=from_name,
        to_phone=to_phone,
        to_name=data.get("to_name"),
        content=content,
        message_type=FamilyMessageType(message_type),
        requires_response=data.get("requires_response", False),
    )

    _redis_client.add_family_message(tid, msg.id, msg.to_redis())

    # Notifier les destinataires (push ou SMS selon config)
    notified = []
    if to_phone:
        # Message direct
        member_data = _redis_client.get_family_member(tid, to_phone)
        if member_data and member_data.get("can_receive_alerts") == "1":
            notified.append(to_phone)
    else:
        # Message groupe - notifier tous les membres actifs
        for phone in _redis_client.get_family_members(tid):
            if phone != from_phone:  # Ne pas notifier l'expediteur
                member_data = _redis_client.get_family_member(tid, phone)
                if member_data and member_data.get("is_active") == "1":
                    notified.append(phone)

    _log_family_audit("message_sent", from_phone, from_name,
                      target_phone=to_phone,
                      details={"type": message_type, "notified": len(notified)},
                      tenant_id=tid)

    _gamify(tid, "family_message_sent")

    return {
        "success": True,
        "message_id": msg.id,
        "notified_count": len(notified),
    }


@app.post("/api/family/messages/{msg_id}/read")
async def mark_message_read(msg_id: str, request: Request):
    """Marque un message comme lu par un membre"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    phone = data.get("phone", "") or data.get("reader_phone", "")
    phone = phone.strip() if phone else ""

    if not phone:
        return JSONResponse(status_code=400, content={"error": "phone requis"})

    msg_data = _redis_client.get_family_message(tid, msg_id)
    if not msg_data:
        return JSONResponse(status_code=404, content={"error": "Message non trouve"})

    import json
    from datetime import datetime
    read_by = json.loads(msg_data.get("read_by", "[]"))
    if phone not in read_by:
        read_by.append(phone)
        _redis_client.update_family_message(tid, msg_id, {
            "read_by": json.dumps(read_by),
            "is_read": "1" if len(read_by) > 0 else "0",
            "read_at": datetime.utcnow().isoformat(),
        })

    return {"success": True}


# --- Escalation Rules ---

@app.get("/api/family/escalation")
async def get_escalation_rules(request: Request):
    """Recupere les regles d'escalade"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    rules = []
    for rule_id in _redis_client.get_escalation_rules(tid):
        rule_data = _redis_client.get_escalation_rule(tid, rule_id)
        if rule_data:
            rule = EscalationRule.from_redis(rule_data)
            rules.append({
                "id": rule.id,
                "name": rule.name,
                "description": rule.description,
                "trigger_distress_level": rule.trigger_distress_level.value,
                "trigger_member_types": [t.value for t in rule.trigger_member_types],
                "stages": [s.to_dict() for s in rule.stages],
                "enabled": rule.enabled,
                "trigger_count": rule.trigger_count,
            })

    return {"rules": rules, "count": len(rules)}


@app.post("/api/family/escalation")
async def create_escalation_rule(request: Request):
    """Cree une regle d'escalade"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()

    # Creer les etapes
    stages = []
    for stage_data in data.get("stages", []):
        stage = EscalationStage(
            order=stage_data["order"],
            delay_minutes=stage_data.get("delay_minutes", 0),
            target_roles=[FamilyRole(r) for r in stage_data["target_roles"]],
            message_template=stage_data.get("message_template", ""),
            include_emergency_numbers=stage_data.get("include_emergency_numbers", False),
        )
        stages.append(stage)

    rule = EscalationRule(
        tenant_id=tid,
        name=data.get("name", "Regle d'escalade"),
        description=data.get("description", ""),
        trigger_distress_level=DistressLevel(data.get("trigger_distress_level", "high")),
        trigger_member_types=[FamilyMemberType(t) for t in data.get("trigger_member_types", [])],
        trigger_keywords=data.get("trigger_keywords", []),
        stages=stages,
        enabled=data.get("enabled", True),
    )

    _redis_client.add_escalation_rule(tid, rule.id, rule.to_redis())

    _log_family_audit("escalation_rule_created", "system", "Luna",
                      details={"name": rule.name, "stages": len(stages)},
                      tenant_id=tid)

    _gamify(tid, "escalation_rule_created")

    return {"success": True, "rule_id": rule.id}


@app.delete("/api/family/escalation/{rule_id}")
async def delete_escalation_rule(rule_id: str, request: Request):
    """Supprime une regle d'escalade"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    _redis_client.delete_escalation_rule(tid, rule_id)

    _log_family_audit("escalation_rule_deleted", "system", "Luna",
                      details={"rule_id": rule_id}, tenant_id=tid)

    return {"success": True}


# --- Distress Detection (for Family Pack) ---

@app.post("/api/family/detect-distress")
async def detect_distress(request: Request):
    """Analyse un texte pour detecter la detresse (ados/enfants)"""
    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    text = data.get("text", "")
    from_phone = data.get("from_phone")

    if not text:
        return JSONResponse(status_code=400, content={"error": "text requis"})

    # Detecter le niveau de detresse
    level, keywords, category = DistressKeywords.detect_level(text)

    result = {
        "level": level.value,
        "keywords_found": keywords,
        "category": category,
        "action_required": level in [DistressLevel.HIGH, DistressLevel.CRITICAL],
    }

    # Si niveau eleve, declencher l'escalade
    if level in [DistressLevel.HIGH, DistressLevel.CRITICAL] and from_phone:
        member_data = _redis_client.get_family_member(tid, from_phone) if _redis_client else None
        if member_data:
            member_name = member_data.get("name", "Membre")
            result["escalation_triggered"] = True
            result["message"] = f"Alerte detresse detectee pour {member_name}"

            # Log et notifier
            _log_family_audit("distress_detected", from_phone, member_name,
                              details={"level": level.value, "category": category},
                              severity="critical" if level == DistressLevel.CRITICAL else "alert",
                              tenant_id=tid)

    return result


# --- Family Audit Log ---

@app.get("/api/family/audit")
async def get_family_audit(request: Request, limit: int = 50):
    """Recupere le journal d'audit famille"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    import json
    entries = []
    for entry_json in _redis_client.get_family_audit(tid, limit=limit):
        try:
            entry_data = json.loads(entry_json)
            audit = FamilyAuditLog.from_redis(entry_data)
            entries.append({
                "id": audit.id,
                "actor_name": audit.actor_name,
                "action": audit.action,
                "target_name": audit.target_name,
                "details": audit.details,
                "severity": audit.severity,
                "timestamp": audit.timestamp.isoformat(),
            })
        except Exception as e:
            logger.warning(f"Audit entry parse error: {e}")

    return {"entries": entries, "count": len(entries)}


# --- Default Escalation Rule for Teens (Bullying/Distress) ---

@app.post("/api/family/setup-teen-protection")
async def setup_teen_protection(request: Request):
    """Configure automatiquement la protection ados (harcelement, detresse)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    # Creer un groupe famille si pas existant
    if not _redis_client.get_family_group(tid):
        group = FamilyGroup(tenant_id=tid, name="Ma famille")
        _redis_client.set_family_group(tid, group.to_redis())

    # Creer la regle d'escalade pour ados
    stages = [
        EscalationStage(
            order=1,
            delay_minutes=0,
            target_roles=[FamilyRole.TEEN],  # D'abord dialogue avec l'ado
            message_template="Luna a remarque que tu sembles preoccupe(e). Tu veux en parler?",
            include_emergency_numbers=False,
        ),
        EscalationStage(
            order=2,
            delay_minutes=5,
            target_roles=[FamilyRole.PRIMARY_CAREGIVER],
            message_template="[Luna Family] Votre enfant {name} a exprime des signes de detresse. Details: {summary}. Nous vous conseillons d'en discuter avec lui/elle.",
            include_emergency_numbers=False,
        ),
        EscalationStage(
            order=3,
            delay_minutes=15,
            target_roles=[FamilyRole.PRIMARY_CAREGIVER, FamilyRole.SECONDARY_CAREGIVER],
            message_template="[URGENT Luna Family] {name} necessite une attention immediate. Si vous ne pouvez pas le/la joindre, appelez le 3114 (prevention suicide) ou le 0 800 235 236 (Fil Sante Jeunes).",
            include_emergency_numbers=True,
        ),
    ]

    rule = EscalationRule(
        tenant_id=tid,
        name="Protection ados - harcelement/detresse",
        description="Detection automatique du harcelement et de la detresse chez les adolescents",
        trigger_distress_level=DistressLevel.MEDIUM,
        trigger_member_types=[FamilyMemberType.TEEN, FamilyMemberType.CHILD],
        stages=stages,
        enabled=True,
    )

    _redis_client.add_escalation_rule(tid, rule.id, rule.to_redis())

    _log_family_audit("teen_protection_enabled", "system", "Luna",
                      details={"rule_id": rule.id}, tenant_id=tid)

    return {
        "success": True,
        "rule_id": rule.id,
        "message": "Protection ados configuree avec succes. Luna va maintenant detecter les signes de harcelement et de detresse.",
    }


# --- SOS Famille ---

@app.post("/api/family/sos")
async def family_sos(request: Request):
    """
    Bouton SOS famille - Alerte tous les membres de la famille.
    N'appelle PAS les services d'urgence (interdit).
    """
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    from_phone = data.get("from_phone", "").strip()
    message = data.get("message", "").strip() or "J'ai besoin d'aide!"
    location = data.get("location")  # {lat, lng, address} optionnel
    trigger_visio = data.get("trigger_visio", False)

    if not from_phone:
        return JSONResponse(status_code=400, content={"error": "from_phone requis"})

    # Trouver le membre qui déclenche
    sender = _redis_client.get_family_member(tid, from_phone)
    sender_name = sender.get("name", "Un membre") if sender else "Un membre"

    # Récupérer tous les membres de la famille (sauf l'émetteur)
    all_members = _redis_client.get_all_family_members(tid)
    recipients = [m for m in all_members if m.get("phone") != from_phone]

    if not recipients:
        return JSONResponse(status_code=400, content={"error": "Aucun autre membre dans la famille"})

    # Construire le message SOS
    timestamp = datetime.utcnow().strftime("%H:%M")
    sos_message = f"🆘 ALERTE FAMILLE - {timestamp}\n\n"
    sos_message += f"{sender_name} a besoin d'aide!\n"
    if message != "J'ai besoin d'aide!":
        sos_message += f"Message: {message}\n"
    if location:
        if location.get("address"):
            sos_message += f"📍 {location['address']}\n"
        elif location.get("lat") and location.get("lng"):
            sos_message += f"📍 Position: {location['lat']}, {location['lng']}\n"

    # Créer un message prioritaire dans la messagerie interne
    from core.memory.schemas import FamilyMessage, FamilyMessageType
    sos_msg = FamilyMessage(
        tenant_id=tid,
        from_phone=from_phone,
        from_name=sender_name,
        to_phone="all",  # Broadcast
        content=sos_message,
        message_type=FamilyMessageType.ALERT,
        requires_response=True,
    )
    _redis_client.add_family_message(tid, sos_msg.id, sos_msg.to_redis())

    # Notifier par SMS les membres qui peuvent recevoir des alertes
    sms_sent = 0
    sms_failed = 0
    if sms_client and sms_client.is_configured:
        for member in recipients:
            if member.get("can_receive_alerts") in ["1", "true", True]:
                phone = member.get("phone")
                if phone:
                    success, _ = _tracked_sms_send(phone, sos_message, label=f"Alerte SOS")
                    if success:
                        sms_sent += 1
                    else:
                        sms_failed += 1

    # Publier événement temps réel pour l'app
    _redis_client.publish_event(tid, "sos_alert", {
        "from_phone": from_phone,
        "from_name": sender_name,
        "message": message,
        "location": location,
        "timestamp": datetime.utcnow().isoformat(),
    })

    # Log dans l'audit
    _log_family_audit("sos_triggered", from_phone, sender_name,
                      details={
                          "message": message,
                          "location": location,
                          "recipients_count": len(recipients),
                          "sms_sent": sms_sent,
                      },
                      severity="critical",
                      tenant_id=tid)

    _gamify(tid, "family_sos")

    # Optionnel: declencher une visio (Tavus puis Simli)
    visio_url = None
    visio_provider = None
    if trigger_visio:
        if tavus_client and tavus_client.is_configured:
            context = f"ALERTE SOS de {sender_name}. Message: {message}"
            success, conv_data = await tavus_client.create_conversation(
                tenant_id=tid,
                custom_greeting=f"{sender_name}, je suis la. Dis-moi ce qui se passe.",
                context=context,
            )
            if success:
                visio_url = conv_data.get("conversation_url")
                visio_provider = "tavus"
        if not visio_url:
            ok_visio, visio_payload = await _start_simli_visio(tid, sender_name)
            if ok_visio:
                visio_url = visio_payload.get("conversation_url")
                visio_provider = "simli"

    sms_expected = bool(
        sms_client and sms_client.is_configured
        and any(m.get("can_receive_alerts") in ("1", "true", True) for m in recipients)
    )
    sms_ok = (not sms_expected) or sms_sent > 0

    return {
        "success": sms_ok,
        "message_id": sos_msg.id,
        "alerted_members": len(recipients),
        "sms_sent": sms_sent,
        "sms_failed": sms_failed,
        "sms_expected": sms_expected,
        "visio_url": visio_url,
        "visio_provider": visio_provider,
        "message": (
            f"Alerte envoyee a {len(recipients)} membre(s) de la famille"
            + ("" if sms_ok else " — aucun SMS delivre")
        ),
    }


@app.get("/api/family/sos/status")
async def family_sos_status(request: Request):
    """Vérifie si le SOS est disponible et récupère les dernières alertes"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    # Récupérer les alertes récentes depuis l'audit (les entrées sont des JSON strings)
    import json
    audit_entries_raw = _redis_client.get_family_audit(tid, limit=20)
    recent_sos = []
    for entry_str in audit_entries_raw:
        try:
            entry = json.loads(entry_str)
            if entry.get("action") == "sos_triggered":
                recent_sos.append(entry)
        except (json.JSONDecodeError, TypeError):
            continue

    # Compter les membres qui peuvent recevoir des alertes
    all_members = _redis_client.get_all_family_members(tid)
    alert_receivers = len([m for m in all_members if m.get("can_receive_alerts") in ["1", "true", True]])

    return {
        "available": True,
        "sms_enabled": sms_client.is_configured if sms_client else False,
        "visio_enabled": tavus_client.is_configured if tavus_client else False,
        "members_count": len(all_members),
        "alert_receivers": alert_receivers,
        "recent_alerts": recent_sos[:5],
    }


# =========================================================================
# INSTRUCTIONS ENDPOINTS
# =========================================================================

@app.get("/api/instructions")
async def list_instructions(request: Request):
    """Liste les instructions actives"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    instructions = mgr.list_active_instructions()
    return {
        "instructions": [
            {
                "id": instr.id,
                "type": instr.type.value,
                "description": instr.description,
                "schedule": instr.schedule,
                "action": instr.action.value,
                "target": instr.target,
                "enabled": instr.enabled,
                "last_executed": instr.last_executed.isoformat() if instr.last_executed else None,
                "created_at": instr.created_at.isoformat() if instr.created_at else None,
            }
            for instr in instructions
        ],
        "count": len(instructions),
    }


@app.post("/api/instructions")
async def create_instruction(req: InstructionRequest, request: Request):
    """Cree une instruction a partir de texte naturel (francais)"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    if not _CORE_AVAILABLE:
        return JSONResponse(status_code=503, content={"error": "Parser non disponible"})

    try:
        # Parse le texte naturel en instruction structuree
        parsed = InstructionParser.parse(req.text)

        # Map parser ActionType -> schema ActionType
        # Schema has: SMS, CALL, VISIO, REMINDER, NOTE, ALERT
        action_map = {
            "reminder": SchemaActionType.REMINDER,
            "sms_contact": SchemaActionType.SMS,
            "call_contact": SchemaActionType.CALL,
            "visio_contact": SchemaActionType.VISIO,
            "wake_up": SchemaActionType.CALL,
            "check_in": SchemaActionType.REMINDER,
            "surveillance": SchemaActionType.ALERT,
            "note": SchemaActionType.NOTE,
            "daily_routine": SchemaActionType.REMINDER,
            "information": SchemaActionType.REMINDER,
            "reading": SchemaActionType.REMINDER,
            "game": SchemaActionType.REMINDER,
            "music": SchemaActionType.REMINDER,
            "gratitude": SchemaActionType.REMINDER,
        }
        action = action_map.get(parsed.action_type.value, SchemaActionType.REMINDER)

        # Map parser RecurrenceType -> schema InstructionType
        type_map = {
            "once": InstructionType.ONE_TIME,
            "daily": InstructionType.DAILY,
            "weekly": InstructionType.RECURRING,
            "monthly": InstructionType.RECURRING,
            "weekdays": InstructionType.RECURRING,
            "weekend": InstructionType.RECURRING,
        }
        instr_type = type_map.get(parsed.recurrence.value, InstructionType.ONE_TIME)

        # Genere le schedule (cron ou heure)
        schedule_str = parsed.to_cron() or ""
        if not schedule_str and parsed.scheduled_time:
            schedule_str = f"{parsed.scheduled_time.hour:02d}:{parsed.scheduled_time.minute:02d}"

        instr = mgr.add_instruction(
            description=req.text,
            action=action,
            instruction_type=instr_type,
            schedule=schedule_str,
            target=parsed.target or "self",
            message_template=parsed.message_template or "",
            priority=parsed.priority,
        )

        # Schedule dans le scheduler en memoire
        if _scheduler:
            try:
                _scheduler.schedule(
                    instruction_id=instr.id,
                    tenant_id=tid,
                    instruction=parsed,
                )
            except Exception as e:
                logger.warning(f"Scheduler schedule failed: {e}")

        # Confirmation text
        confirmation = InstructionParser.format_confirmation(parsed)

        _gamify(tid, "create_instruction")
        return {
            "success": True,
            "instruction": {
                "id": instr.id,
                "type": instr.type.value,
                "description": instr.description,
                "action": instr.action.value,
                "schedule": instr.schedule,
                "target": instr.target,
            },
            "parsed": {
                "action": parsed.action_type.value,
                "recurrence": parsed.recurrence.value,
                "time": f"{parsed.scheduled_time.hour:02d}:{parsed.scheduled_time.minute:02d}" if parsed.scheduled_time else None,
                "target": parsed.target,
                "message": parsed.message_template,
                "confidence": parsed.confidence,
                "needs_clarification": parsed.needs_clarification,
                "clarification_question": parsed.clarification_question,
            },
            "confirmation": confirmation,
        }
    except Exception as e:
        logger.error(f"Instruction creation failed: {e}")
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.delete("/api/instructions/{instr_id}")
async def delete_instruction(instr_id: str, request: Request):
    """Desactive une instruction"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    mgr.disable_instruction(instr_id)
    return {"success": True}


@app.post("/api/instructions/{instr_id}/execute")
async def execute_instruction(instr_id: str, request: Request):
    """Execute immediatement une instruction"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    instr = mgr.get_instruction(instr_id)
    if not instr:
        return JSONResponse(status_code=404, content={"error": "Instruction non trouvee"})

    # Executer reellement l'instruction via l'executor
    if _executor and _CORE_AVAILABLE:
        try:
            task = ScheduledTask(
                scheduled_at=datetime.utcnow(),
                instruction_id=instr_id,
                tenant_id=tid,
                instruction=instr,
            )
            result = await _executor.execute(task)
            mgr.mark_instruction_executed(instr_id)
            return {
                "success": result.success if hasattr(result, "success") else True,
                "message": f"Instruction '{instr.description[:50]}' executee",
                "result": str(result) if result else None,
            }
        except Exception as e:
            logger.error(f"Instruction execution error: {e}")
            return JSONResponse(status_code=500, content={"error": f"Erreur execution: {str(e)}"})
    else:
        # Fallback: marquer comme executee sans executor
        mgr.mark_instruction_executed(instr_id)
        return {"success": True, "message": f"Instruction '{instr.description[:50]}' marquee executee (executor non disponible)"}


# =========================================================================
# NOTES ENDPOINTS
# =========================================================================

@app.get("/api/notes")
async def list_notes(request: Request, limit: int = 50):
    """Liste les notes recentes"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    notes = mgr.list_notes(limit=min(limit, 200))
    return {
        "notes": [
            {
                "id": n.id,
                "content": n.content,
                "context": n.context,
                "source": n.source,
                "tags": n.tags,
                "created_at": n.created_at.isoformat() if n.created_at else None,
            }
            for n in notes
        ],
        "count": len(notes),
    }


@app.get("/api/call-reports")
async def list_call_reports(request: Request, limit: int = 30):
    """Liste les comptes-rendus d'appels et de conférences (vocaux + visio + conf)."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    notes = mgr.list_notes(limit=min(limit * 3, 200))  # large fetch pour filtrer
    _REPORT_CONTEXTS = {
        "structured_call_report", "conference", "conference_call",
        "call", "voice_call", "appel_vocal", "visio",
    }
    _REPORT_TAGS = {"rapport", "appel_vocal", "conference", "visio"}
    reports = []
    for n in notes:
        is_report = (
            n.context in _REPORT_CONTEXTS
            or bool(n.tags and _REPORT_TAGS.intersection(set(n.tags)))
        )
        if is_report:
            reports.append({
                "id": n.id,
                "content": n.content,
                "context": n.context,
                "tags": n.tags,
                "created_at": n.created_at.isoformat() if n.created_at else None,
            })
        if len(reports) >= limit:
            break
    return {"reports": reports, "count": len(reports)}


@app.post("/api/notes")
async def add_note(req: NoteRequest, request: Request):
    """Ajoute une note"""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        note = mgr.add_note(
            content=req.content,
            context=req.context,
            tags=req.tags,
        )
        _gamify(tid, "create_note")
        return {"success": True, "note": {"id": note.id, "content": note.content}}
    except Exception as e:
        return JSONResponse(status_code=400, content={"error": str(e)})


@app.delete("/api/notes/{note_id}", status_code=204)
async def delete_note(note_id: str, request: Request):
    """Supprime une note de l'utilisateur."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    found = mgr.delete_note(note_id)
    if not found:
        return JSONResponse(status_code=404, content={"error": "Note introuvable"})
    return Response(status_code=204)


# =========================================================================
# RECALL.AI — RÉUNIONS (Zoom / Meet / Teams / Webex)
# =========================================================================

class MeetingJoinRequest(BaseModel):
    meeting_url: str
    meeting_name: str = ""
    bot_name: str = "Luna (YAWatch)"

@app.post("/api/meeting/join")
async def meeting_join(req: MeetingJoinRequest, request: Request):
    """Envoie Luna dans une réunion en tant que secrétaire silencieuse."""
    if not recall_client or not recall_client.is_configured:
        return JSONResponse(status_code=503, content={"error": "Recall.ai non configuré (RECALL_API_KEY manquant)"})
    tid = getattr(request.state, "tenant_id", 1)
    platform = detect_platform(req.meeting_url)
    meeting_name = req.meeting_name or PLATFORM_NAMES.get(platform, "Réunion")
    try:
        ok, data = await recall_client.create_bot_async(
            meeting_url=req.meeting_url,
            bot_name=req.bot_name,
        )
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
    if not ok:
        return JSONResponse(status_code=502, content={"error": data.get("error", "Erreur Recall.ai"), "detail": data.get("detail")})
    bot_id = data.get("id", "")
    import datetime as _dt
    _mtg_session_set(bot_id, {
        "bot_id": bot_id,
        "meeting_url": req.meeting_url,
        "platform": platform,
        "meeting_name": meeting_name,
        "bot_name": req.bot_name,
        "tenant_id": tid,
        "status": "ready",
        "started_at": _dt.datetime.utcnow().isoformat(),
        "ended_at": None,
        "note_id": None,
    })
    _mtg_transcripts_set(bot_id, [])
    logger.info(f"Recall meeting started: {bot_id} ({platform}) tenant={tid}")
    return {"bot_id": bot_id, "platform": platform, "meeting_name": meeting_name, "status": "ready"}


@app.get("/api/meeting/active")
async def meeting_active(request: Request):
    """Liste les réunions actives et terminées (dernières 24h)."""
    tid = getattr(request.state, "tenant_id", 1)
    import datetime as _dt
    cutoff = (_dt.datetime.utcnow() - _dt.timedelta(hours=24)).isoformat()
    sessions = [
        s for s in _mtg_list_sessions(tid)
        if s.get("status") not in ("done", "fatal") or (s.get("ended_at") or "") >= cutoff
    ]
    return {"meetings": sorted(sessions, key=lambda x: x.get("started_at", ""), reverse=True)}


async def _auto_generate_report(bot_id: str, session: dict) -> None:
    """Déclenché par le poll de statut quand completed — évite d'attendre le webhook."""
    # Garde-fou : ne pas relancer si déjà en cours ou terminé
    if session.get("note_id") or session.get("_report_pending"):
        return
    session["_report_pending"] = True
    _mtg_session_set(bot_id, session)
    try:
        full_segments = await recall_client.get_transcript_async(bot_id)
        if full_segments:
            _mtg_transcripts_set(bot_id, full_segments)
            await _generate_meeting_report(bot_id, full_segments, session)
    except Exception as e:
        logger.warning(f"_auto_generate_report {bot_id}: {e}")
    finally:
        session = _mtg_session_get(bot_id) or session
        session.pop("_report_pending", None)
        _mtg_session_set(bot_id, session)


@app.get("/api/meeting/{bot_id}/status")
async def meeting_status(bot_id: str, request: Request):
    """Retourne le statut en temps réel d'un bot (polling frontend)."""
    session = _mtg_session_get(bot_id)
    if not session:
        return JSONResponse(status_code=404, content={"error": "Réunion introuvable"})
    # Rafraîchir le statut depuis l'API Recall.ai si toujours actif
    if recall_client and session.get("status") not in ("done", "completed", "fatal", "call_ended", "failed"):
        bot_data = await recall_client.get_bot_async(bot_id)
        if bot_data:
            raw = bot_data.get("status", session["status"])
            code = raw if isinstance(raw, str) else (raw or {}).get("code", session["status"])
            session["status"] = code
            _mtg_session_set(bot_id, session)
    # Si MeetingBaas a terminé et pas encore de compte rendu → déclenche Whisper+GPT
    if (session.get("status") in ("completed", "done") and
            not session.get("note_id") and recall_client):
        asyncio.create_task(_auto_generate_report(bot_id, session))
    label = STATUS_LABELS.get(session["status"], session["status"])
    return {**session, "status_label": label, "transcript_count": len(_mtg_transcripts_get(bot_id))}


@app.post("/api/meeting/{bot_id}/stop")
async def meeting_stop(bot_id: str, request: Request):
    """Demande au bot de quitter la réunion."""
    session = _mtg_session_get(bot_id)
    if not session:
        return JSONResponse(status_code=404, content={"error": "Réunion introuvable"})
    if recall_client:
        await recall_client.stop_bot_async(bot_id)
    session["status"] = "call_ended"
    _mtg_session_set(bot_id, session)
    return {"bot_id": bot_id, "status": "call_ended"}


@app.post("/api/meeting/{bot_id}/report")
async def meeting_report(bot_id: str, request: Request):
    """Génère (ou régénère) le compte rendu d'une réunion terminée.
    Fonctionne même si le webhook MeetingBaas n'est jamais arrivé :
    récupère l'audio depuis l'API, transcrit via Whisper, génère via GPT."""
    if not recall_client:
        return JSONResponse(status_code=503, content={"error": "MeetingBaas non configuré"})
    tid = getattr(request.state, "tenant_id", 1)

    # Récupère ou reconstruit la session
    session = _mtg_session_get(bot_id)
    if not session:
        # Reconstruit depuis MeetingBaas si session perdue (redéploiement)
        bot_data = await recall_client.get_bot_async(bot_id)
        if not bot_data:
            return JSONResponse(status_code=404, content={"error": "Bot introuvable sur MeetingBaas"})
        import datetime as _dt
        session = {
            "bot_id": bot_id,
            "meeting_url": bot_data.get("meeting_url", ""),
            "platform": detect_platform(bot_data.get("meeting_url", "")),
            "meeting_name": bot_data.get("meeting_name") or "Réunion",
            "bot_name": bot_data.get("bot_name", "Luna (YAWatch)"),
            "tenant_id": tid,
            "status": bot_data.get("status", "completed"),
            "started_at": bot_data.get("created_at", _dt.datetime.utcnow().isoformat()),
            "ended_at": _dt.datetime.utcnow().isoformat(),
            "note_id": None,
        }
        _mtg_session_set(bot_id, session)

    # Si déjà un compte rendu, le retourner directement
    if session.get("note_id"):
        return {"bot_id": bot_id, "note_id": session["note_id"], "status": "already_generated"}

    # Transcription Whisper
    try:
        segments = await recall_client.get_transcript_async(bot_id)
    except Exception as e:
        return JSONResponse(status_code=502, content={"error": f"Transcription échouée : {e}"})
    if not segments:
        return JSONResponse(status_code=422, content={"error": "Aucun segment audio disponible (réunion trop courte ?)"})

    _mtg_transcripts_set(bot_id, segments)
    await _generate_meeting_report(bot_id, segments, session)
    session = _mtg_session_get(bot_id)
    note_id = (session or {}).get("note_id")
    if not note_id:
        return JSONResponse(status_code=500, content={"error": "Génération du rapport échouée"})
    return {"bot_id": bot_id, "note_id": note_id, "transcript_count": len(segments), "status": "generated"}


async def _generate_meeting_report(bot_id: str, segments: list, session: dict):
    """Génère le compte rendu final via GPT et le sauvegarde en Redis."""
    if not openai_client or not segments:
        return
    transcript_text = format_transcript(segments)
    if not transcript_text.strip():
        return
    meeting_name = session.get("meeting_name", "Réunion")
    platform = PLATFORM_NAMES.get(session.get("platform", ""), "Visioconférence")
    prompt = f"""Tu es Luna, secrétaire IA de YAWatch.
Voici la transcription d'une réunion "{meeting_name}" sur {platform}.

TRANSCRIPTION :
{transcript_text[:12000]}

Génère un compte rendu structuré dans ce format exact (ne modifie pas les numéros ni le format **Titre**) :

1) **Résumé**
[2-3 phrases résumant l'essentiel]

2) **Participants**
[liste des noms détectés, un par ligne]

3) **Sujets abordés**
[liste à puces]

4) **Décisions prises**
[liste à puces, "Aucune décision formelle" si vide]

5) **Actions à faire**
[liste avec responsable et échéance si mentionnés, "Aucune" si vide]

6) **Points bloquants**
[liste à puces, "Aucun" si vide]

7) **Citations importantes**
[2-3 citations entre guillemets avec le nom, ou "Aucune" si vide]

Sois factuel et concis. Ne fabrique rien qui n'est pas dans la transcription."""
    try:
        resp = openai_client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1200,
            temperature=0.3,
        )
        content = resp.choices[0].message.content.strip()
    except Exception as e:
        logger.error(f"Meeting report GPT error: {e}")
        content = transcript_text  # fallback: transcription brute
    tid = session.get("tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if mgr:
        note = mgr.add_note(
            content=f"# Compte rendu — {meeting_name}\n\n{content}",
            source=f"Compte rendu — {meeting_name}",
            context="conference",
            tags=["conference", "rapport", "recall"],
        )
        if note:
            session["note_id"] = note.id
            _mtg_session_set(bot_id, session)
            logger.info(f"Meeting report saved: {note.id} (bot={bot_id})")


@app.post("/api/webhooks/meetingbaas")
async def meetingbaas_webhook(request: Request):
    """
    Webhook PUBLIC — MeetingBaas envoie ici les événements du bot.
    Gère : changements de statut + transcription temps réel.
    """
    try:
        payload = await request.json()
    except Exception:
        return Response(status_code=400)

    logger.debug(f"MeetingBaas webhook: {str(payload)[:200]}")

    # --- Extraction du bot_id (plusieurs formats possibles) ---
    bot_id = (
        payload.get("bot_id")
        or payload.get("id")
        or (payload.get("data") or {}).get("bot_id")
        or (payload.get("data") or {}).get("id")
        or ""
    )

    # --- Statut du bot ---
    status_code = (
        payload.get("status")
        or payload.get("event")
        or (payload.get("data") or {}).get("status")
        or ""
    )

    if bot_id and status_code:
        logger.info(f"MeetingBaas status: {bot_id} → {status_code}")
        session = _mtg_session_get(bot_id)
        if session:
            session["status"] = status_code
            _mtg_session_set(bot_id, session)

        if status_code in ("done", "completed", "call_ended", "left_call", "failed"):
            import datetime as _dt
            if session:
                session["ended_at"] = _dt.datetime.utcnow().isoformat()
                _mtg_session_set(bot_id, session)
            # Toujours tenter de récupérer la transcription (Whisper via FLAC)
            if recall_client and session:
                try:
                    full_segments = await recall_client.get_transcript_async(bot_id)
                    if full_segments:
                        _mtg_transcripts_set(bot_id, full_segments)
                except Exception as e:
                    logger.warning(f"Could not fetch full transcript: {e}")
            segments = _mtg_transcripts_get(bot_id)
            if segments and session:
                asyncio.create_task(_generate_meeting_report(bot_id, segments, session))

    # --- Transcription temps réel ---
    transcript_data = (
        payload.get("transcript")
        or (payload.get("data") or {}).get("transcript")
    )
    if transcript_data and bot_id:
        if isinstance(transcript_data, list):
            segments = transcript_data
        else:
            segments = [transcript_data] if transcript_data.get("text") or transcript_data.get("words") else []
        for seg in segments:
            if seg.get("is_final", True):
                _mtg_transcripts_append(bot_id, seg)
                logger.debug(f"MeetingBaas transcript [{bot_id}] {seg.get('speaker','?')}: {str(seg)[:60]}")

    return Response(status_code=200)


# =========================================================================
# EVENT LOG ENDPOINTS
# =========================================================================

@app.get("/api/events")
async def get_events(request: Request, limit: int = 50, offset: int = 0):
    """Retourne le journal d'evenements chronologique."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        events = mgr.get_event_log(limit=min(limit, 200), offset=offset)
        return {"events": events, "count": len(events)}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.get("/api/events/export")
async def export_events(request: Request, limit: int = 200):
    """Exporte le journal d'evenements en texte brut."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        events = mgr.get_event_log(limit=min(limit, 500))
        lines = []
        for ev in events:
            ts = ev.get("timestamp", "?")
            cat = ev.get("category", "?")
            desc = ev.get("description", "")
            reasoning = ev.get("reasoning", "")
            line = f"[{ts}] [{cat.upper()}] {desc}"
            if reasoning:
                line += f" | Raison: {reasoning}"
            lines.append(line)
        return Response(content="\n".join(lines), media_type="text/plain; charset=utf-8")
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.get("/api/events/pdf")
async def export_events_pdf(request: Request, days: int = 7):
    """Exporte le journal d'evenements en PDF pour le souscripteur."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        from fpdf import FPDF
        import io as _io
        from datetime import datetime, timedelta

        events = mgr.get_event_log(limit=500)

        # Filtrer par nombre de jours
        cutoff = (datetime.utcnow() - timedelta(days=days)).isoformat()
        events = [e for e in events if e.get("timestamp", "") >= cutoff]

        # Profil souscripteur
        profile = mgr.get_subscriber_profile()
        sub_name = "Souscripteur"
        if profile:
            sub_name = f"{profile.first_name or ''} {profile.last_name or ''}".strip() or "Souscripteur"

        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)

        # Titre
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 12, "Luna - Compte Rendu d'Activite", ln=True, align="C")
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 8, f"Souscripteur : {_clean_latin1(sub_name)}", ln=True, align="C")
        pdf.cell(0, 8, f"Periode : {days} derniers jours", ln=True, align="C")
        pdf.cell(0, 8, f"Genere le : {datetime.utcnow().strftime('%d/%m/%Y a %H:%M UTC')}", ln=True, align="C")
        pdf.ln(8)

        # Stats
        categories = {}
        for ev in events:
            cat = ev.get("category", "autre")
            categories[cat] = categories.get(cat, 0) + 1

        if categories:
            pdf.set_font("Helvetica", "B", 13)
            pdf.cell(0, 10, "Resume", ln=True)
            pdf.set_font("Helvetica", "", 10)
            for cat, count in sorted(categories.items(), key=lambda x: -x[1]):
                pdf.cell(0, 6, f"  {cat.capitalize()} : {count} evenement(s)", ln=True)
            pdf.ln(6)

        # Tableau
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 10, f"Journal detaille ({len(events)} evenements)", ln=True)
        pdf.set_font("Helvetica", "B", 9)
        pdf.cell(40, 7, "Date/Heure", border=1)
        pdf.cell(25, 7, "Type", border=1)
        pdf.cell(0, 7, "Description", border=1, ln=True)

        pdf.set_font("Helvetica", "", 8)
        for ev in events:
            ts = ev.get("timestamp", "?")[:16].replace("T", " ")
            cat = ev.get("category", "?")[:12]
            desc = _clean_latin1(ev.get("description", ""))[:90]
            pdf.cell(40, 6, ts, border=1)
            pdf.cell(25, 6, cat, border=1)
            pdf.cell(0, 6, desc, border=1, ln=True)

        pdf.ln(10)
        pdf.set_font("Helvetica", "I", 8)
        pdf.cell(0, 5, "Ce document est genere automatiquement par Luna, l'assistante IA de YAWatch.", ln=True, align="C")

        buf = _io.BytesIO()
        pdf.output(buf)
        buf.seek(0)
        filename = f"luna_rapport_{sub_name.replace(' ', '_')}_{days}j.pdf"
        return Response(
            content=buf.read(),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except ImportError:
        return JSONResponse(status_code=500, content={"error": "Module fpdf2 non installe"})
    except Exception as e:
        logger.error(f"PDF export error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


def _clean_latin1(text: str) -> str:
    """Nettoie le texte pour fpdf2 (Latin-1 compatible)."""
    return text.encode("latin-1", errors="replace").decode("latin-1")


# =========================================================================
# BUDGET GUARD — Bloque les actions si le cout API du mois depasse le plafond
# =========================================================================

async def _get_tenant_month_cost(tenant_id: int) -> float:
    """Retourne le cout API total du mois pour un tenant (en EUR) depuis Redis."""
    try:
        cortex = get_cortex() if _CORTEX_AVAILABLE else None
        if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
            usage = await cortex.cost_tracker.get_tenant_month_usage(tenant_id)
            return usage.get("total_cost", 0.0)
    except Exception:
        pass
    return 0.0

async def _check_budget_guard(tenant_id: int, plan_name: str = "essentiel") -> Optional[str]:
    """Verifie si le budget API mensuel est depasse. Retourne un message d'erreur ou None."""
    limits = _PLAN_LIMITS.get(plan_name, _PLAN_LIMITS.get("essentiel", {}))
    budget_max = limits.get("budget_api_max", 50.0)
    cost = await _get_tenant_month_cost(tenant_id)
    if cost >= budget_max:
        return f"Budget API mensuel atteint ({cost:.2f}€/{budget_max:.2f}€). Actions couteuses bloquees jusqu'au mois prochain."
    return None

# =========================================================================
# QUOTA ENDPOINT
# =========================================================================

# Limites par plan (par mois) — source unique depuis quota_guard
try:
    from core.actions.quota_guard import PLAN_SMS_LIMITS, PLAN_VOICE_LIMITS, PLAN_VISIO_LIMITS, PlanType
    _PLAN_LIMITS = {
        "fondateur": {"sms": 999999, "voice_min": 999999, "visio_min": 999999, "budget_api_max": 50.00},
        "essentiel": {"sms": PLAN_SMS_LIMITS[PlanType.ESSENTIEL], "voice_min": PLAN_VOICE_LIMITS[PlanType.ESSENTIEL], "visio_min": PLAN_VISIO_LIMITS[PlanType.ESSENTIEL], "budget_api_max": 8.15},
        "confort":   {"sms": PLAN_SMS_LIMITS[PlanType.CONFORT],   "voice_min": PLAN_VOICE_LIMITS[PlanType.CONFORT],   "visio_min": PLAN_VISIO_LIMITS[PlanType.CONFORT],   "budget_api_max": 17.40},
        "premium":   {"sms": PLAN_SMS_LIMITS[PlanType.PREMIUM],   "voice_min": PLAN_VOICE_LIMITS[PlanType.PREMIUM],   "visio_min": PLAN_VISIO_LIMITS[PlanType.PREMIUM],   "budget_api_max": 32.10},
    }
except ImportError:
    _PLAN_LIMITS = {
        "fondateur": {"sms": 999999, "voice_min": 999999, "visio_min": 999999, "budget_api_max": 50.00},
        "essentiel": {"sms": 25, "voice_min": 40, "visio_min": 12, "budget_api_max": 8.15},
        "confort":   {"sms": 50, "voice_min": 100, "visio_min": 28, "budget_api_max": 17.40},
        "premium":   {"sms": 100, "voice_min": 180, "visio_min": 55, "budget_api_max": 32.10},
    }


@app.get("/api/quota")
async def get_quota(request: Request):
    """Retourne quotas reels (stockage + usage SMS/voix/visio depuis Redis)."""
    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    try:
        # --- Quotas stockage (memoire, conversations, etc.) ---
        quota_status = mgr.get_quota_status()

        # --- Usage reel depuis Cortex Redis ---
        # Plan depuis le profil, mais verifier aussi auth record (fondateur)
        plan_name = (quota_status.get("plan") or "essentiel").lower()
        if plan_name not in _PLAN_LIMITS and _redis_client:
            auth = _redis_client.get_auth_by_tenant_id(tid)
            if auth:
                plan_name = (auth.get("plan") or "essentiel").lower()
        if tid == _PROPRIO_TENANT_ID:
            plan_name = "fondateur"
        limits = _PLAN_LIMITS.get(plan_name, _PLAN_LIMITS["essentiel"])

        real_usage = {"sms_count": 0, "voice_minutes": 0.0, "tavus_minutes": 0.0}
        cortex = get_cortex() if _CORTEX_AVAILABLE else None
        if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
            real_usage = await cortex.cost_tracker.get_tenant_month_usage(tid)

        sms_used = real_usage["sms_count"]
        sms_limit = limits["sms"]
        voice_used = real_usage.get("voice_minutes", 0)
        voice_limit = limits["voice_min"]
        visio_used = real_usage.get("tavus_minutes", 0)
        visio_limit = limits["visio_min"]

        quota_status["sms"] = {
            "used": sms_used,
            "limit": sms_limit,
            "cost_eur": real_usage.get("sms_cost", 0),
            "percentage": round((sms_used / sms_limit) * 100, 1) if sms_limit else 0,
        }
        quota_status["voice"] = {
            "used": round(voice_used, 1),
            "limit": voice_limit,
            "unit": "min",
            "cost_eur": real_usage.get("voice_cost", 0),
            "percentage": round((voice_used / voice_limit) * 100, 1) if voice_limit else 0,
        }
        quota_status["visio"] = {
            "used": round(visio_used, 1),
            "limit": visio_limit,
            "unit": "min",
            "cost_eur": real_usage.get("tavus_cost", 0),
            "percentage": round((visio_used / visio_limit) * 100, 1) if visio_limit else 0,
        }
        quota_status["openai"] = {
            "cost_eur": real_usage.get("openai_cost", 0),
            "tokens_in": real_usage.get("openai_tokens_in", 0),
            "tokens_out": real_usage.get("openai_tokens_out", 0),
        }
        quota_status["plan_limits"] = limits

        # Budget API mensuel (total reel depuis Redis)
        budget_max = limits.get("budget_api_max", 50.0)
        budget_used = real_usage.get("total_cost", 0.0)
        quota_status["budget"] = {
            "used_eur": round(budget_used, 2),
            "max_eur": budget_max,
            "percentage": round((budget_used / budget_max) * 100, 1) if budget_max else 0,
            "blocked": budget_used >= budget_max,
        }

        daily_stats = mgr.get_daily_stats()
        return {
            "quota": quota_status,
            "daily": daily_stats,
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.get("/api/weather")
async def api_weather(request: Request):
    """Retourne la meteo actuelle et previsions 3 jours (wttr.in gratuit)."""
    tid = getattr(request.state, "tenant_id", 1)
    city = request.query_params.get("city", "")
    try:
        result = await _tool_get_weather({"city": city}, tenant_id=tid)
        return result
    except Exception as e:
        logger.warning(f"Weather API error: {type(e).__name__}: {e}")
        return {"status": "error", "message": f"Erreur meteo: {type(e).__name__}"}


# =========================================================================
# CONCIERGERIE — API directe (resultats inline, pas via chat)
# =========================================================================

@app.post("/api/concierge/action")
async def api_concierge_action(request: Request):
    """Endpoint unique conciergerie: appelle la bonne fonction et retourne du JSON."""
    tid = getattr(request.state, "tenant_id", 1)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "JSON invalide"})

    action = body.get("action", "")
    params = body.get("params", {})

    DISPATCHERS = {
        "weather": lambda: _tool_get_weather(params, tenant_id=tid),
        "news": lambda: _tool_get_news(params),
        "search_places": lambda: _tool_search_places(params, tenant_id=tid),
        "search_web": lambda: _tool_search_web(params, tenant_id=tid),
        "search_flights": lambda: _tool_search_flights(params, tenant_id=tid),
        "search_hotels": lambda: _tool_search_hotels(params, tenant_id=tid),
        "book_restaurant": lambda: _tool_book_restaurant(params, tenant_id=tid),
        "get_contacts": lambda: _tool_get_contacts(tenant_id=tid),
        "get_player_stats": lambda: _tool_get_player_stats(tenant_id=tid),
        "get_active_missions": lambda: _tool_get_active_missions(tenant_id=tid),
        "get_badges": lambda: _tool_get_badges(tenant_id=tid),
        "get_friends_online": lambda: _voice_tool_get_friends_online(tid),
        "get_reminders": lambda: _sync_wrap(_tool_secretary_reminders, tid),
        "get_budget": lambda: _sync_wrap(_tool_secretary_budget, tid),
        "add_reminder": lambda: _sync_wrap(_tool_secretary_add_reminder, tid, params),
        "create_note": lambda: _tool_create_note(params, tenant_id=tid),
        "generate_document": lambda: _tool_generate_document(params, tenant_id=tid),
        "alert_contacts": lambda: _tool_alert_contacts(params, tenant_id=tid),
        "send_sms": lambda: _tool_send_sms(params, tenant_id=tid),
        "send_email": lambda: _tool_send_email(params, tenant_id=tid),
        "call_contact": lambda: _tool_call_contact(params, tenant_id=tid, session_id="concierge"),
        "book_flight": lambda: _conc_book_flight(params, tenant_id=tid),
        "book_hotel": lambda: _conc_book_hotel(params, tenant_id=tid),
    }

    handler = DISPATCHERS.get(action)
    if not handler:
        return JSONResponse(status_code=400, content={"error": f"Action inconnue: {action}"})

    try:
        import asyncio
        result = handler()
        if asyncio.iscoroutine(result) or asyncio.isfuture(result):
            result = await result
        return result
    except Exception as e:
        logger.error(f"Concierge action '{action}' error: {type(e).__name__}: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


def _sync_wrap(fn, tid, params=None):
    """Wrap synchronous secretary functions."""
    if params is not None:
        return fn(tid, params) if callable(fn) else {"status": "error"}
    return fn(tid) if callable(fn) else {"status": "error"}


async def _conc_book_flight(params: Dict, tenant_id: int = 0) -> Dict:
    """Reserve un vol via Duffel avec infos profil pre-remplies."""
    if not duffel_client or not duffel_client.is_configured:
        return {"status": "error", "message": "Reservation directe non disponible."}

    offer_id = params.get("offer_id", "")
    if not offer_id:
        return {"status": "error", "message": "ID de l'offre manquant."}

    # Recuperer le profil pour pre-remplir
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    profile = mgr.get_subscriber_profile() if mgr else None
    pf = profile.model_dump() if profile else {}

    # Construire les infos passager depuis le profil
    given_name = params.get("given_name", "") or pf.get("first_name", "")
    family_name = params.get("family_name", "") or pf.get("last_name", "")
    email = params.get("email", "") or pf.get("email", "")
    phone = params.get("phone_number", "") or pf.get("phone", "")
    born_on = params.get("born_on", "") or pf.get("birth_date", "")
    gender = params.get("gender", "") or ("m" if pf.get("gender", "").lower().startswith("m") else "f")

    if not given_name or not family_name:
        return {"status": "error", "message": "Nom et prenom requis. Complete ton profil d'abord."}
    if not email:
        return {"status": "error", "message": "Email requis pour la confirmation de reservation."}

    passengers = [{
        "given_name": given_name,
        "family_name": family_name,
        "email": email,
        "phone_number": phone or "+33600000000",
        "born_on": born_on or "1990-01-01",
        "gender": gender,
        "type": "adult",
    }]

    success, data = await duffel_client.book_flight(offer_id, passengers)
    if success:
        logger.info(f"FLIGHT_BOOKED [Duffel] tenant={tid} ref={data.get('booking_reference')}")
    return {"status": "success" if success else "error", **data}


async def _conc_book_hotel(params: Dict, tenant_id: int = 0) -> Dict:
    """Reserve un hotel via Duffel avec infos profil pre-remplies."""
    if not duffel_client or not duffel_client.is_configured:
        return {"status": "error", "message": "Reservation directe non disponible."}

    rate_id = params.get("rate_id", "")
    if not rate_id:
        return {"status": "error", "message": "ID du tarif manquant."}

    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    profile = mgr.get_subscriber_profile() if mgr else None
    pf = profile.model_dump() if profile else {}

    guest_info = {
        "given_name": params.get("given_name", "") or pf.get("first_name", ""),
        "family_name": params.get("family_name", "") or pf.get("last_name", ""),
        "email": params.get("email", "") or pf.get("email", ""),
        "phone_number": params.get("phone_number", "") or pf.get("phone", ""),
    }

    if not guest_info["given_name"] or not guest_info["family_name"]:
        return {"status": "error", "message": "Nom et prenom requis. Complete ton profil."}

    success, data = await duffel_client.book_hotel(rate_id, guest_info)
    if success:
        logger.info(f"HOTEL_BOOKED [Duffel] tenant={tid} ref={data.get('confirmation_code')}")
    return {"status": "success" if success else "error", **data}


# =========================================================================
# TAVUS WEBHOOK - Tool calling + transcription
# =========================================================================

@app.post("/api/webhook/tavus")
async def webhook_tavus(request: Request):
    """
    Webhook Tavus pour recevoir les events:
    - conversation.tool_call : Luna a appele une fonction depuis la visio
    - application.transcription_ready : transcription complete de la conversation
    """
    if not tavus_client or not tavus_client.is_configured:
        return JSONResponse(status_code=503, content={"error": "Service visio non disponible"})
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Invalid JSON"})

    # Validation: verifier que la conversation_id est connue
    conversation_id = body.get("conversation_id", "")
    if conversation_id and hasattr(tavus_client, "_active_conversations"):
        if conversation_id not in tavus_client._active_conversations:
            logger.warning(f"Tavus webhook: conversation_id inconnue {conversation_id}")

    event_type = body.get("event_type", "") or body.get("type", "")
    logger.info(f"Tavus webhook: {event_type}")

    # --- Tool Call ---
    if "tool_call" in event_type:
        return await _handle_tavus_tool_call(body)

    # --- Transcription ready ---
    if "transcription" in event_type:
        return await _handle_tavus_transcription(body)

    # --- Utterance en temps réel ---
    if "utterance" in event_type:
        return await _handle_tavus_utterance(body)

    # Acknowledge unknown events
    logger.info(f"Tavus webhook event non gere: {event_type}")
    return {"status": "ok"}


@app.post("/api/webhook/simli")
async def webhook_simli(request: Request):
    """
    Webhook Simli (fallback audio quand Tavus est down).
    Reçoit les utterances du souscripteur et applique le même filtrage intelligent.
    Body attendu: { session_id, user_message, assistant_message, tenant_id? }
    """
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"error": "Invalid JSON"})

    session_id = body.get("session_id", "")
    user_text  = body.get("user_message", "") or body.get("text", "")

    # Résolution tenant : via session_id ou champ direct
    tid = body.get("tenant_id") or TENANT_ID
    if session_id and _redis_client:
        try:
            stored = _redis_client.client.get(f"luna:simli:session:{session_id}")
            if stored:
                import json as _sj
                tid = _sj.loads(stored).get("tenant_id", tid)
        except Exception:
            pass

    if user_text:
        await _handle_simli_utterance(int(tid), user_text, session_id)

    return {"status": "ok"}


async def _handle_tavus_tool_call(body: Dict) -> Dict:
    """Route un tool call Tavus vers le bon handler backend."""
    import json as _json

    tool_name = body.get("tool_name", "") or body.get("function_name", "")
    args_raw = body.get("tool_arguments", "") or body.get("arguments", "{}")
    conversation_id = body.get("conversation_id", "")

    try:
        args = _json.loads(args_raw) if isinstance(args_raw, str) else args_raw
    except _json.JSONDecodeError:
        args = {}

    # Trouve le tenant_id depuis la conversation Tavus active
    tid = TENANT_ID
    if tavus_client and conversation_id:
        conv = tavus_client._active_conversations.get(conversation_id)
        if conv:
            tid = conv.tenant_id

    # Detecte si des invites (guests) sont presents dans cette conversation
    _has_guests = False
    if tavus_client and conversation_id:
        _conv = tavus_client._active_conversations.get(conversation_id)
        if _conv and _conv.participants:
            _has_guests = any(p.get("role") == "guest" for p in _conv.participants)

    logger.info(f"Tavus tool_call: {tool_name}({args}) [conv={conversation_id}, tenant={tid}, guests={_has_guests}]")

    # --- Protection souscripteur: bloquer les actions sensibles si des invites sont presents ---
    # Tavus ne peut pas distinguer qui parle, donc en presence d'invites,
    # on bloque les actions qui engagent le souscripteur financierement ou legalement.
    _SUBSCRIBER_ONLY_TOOLS = {
        "request_payment", "search_flights", "search_hotels", "book_restaurant",
        "send_email", "generate_document", "alert_contacts",
    }
    if _has_guests and tool_name in _SUBSCRIBER_ONLY_TOOLS:
        logger.warning(f"Tavus tool BLOCKED (guests present): {tool_name} [conv={conversation_id}]")
        result = {
            "status": "error",
            "message": (
                f"Action '{tool_name}' bloquee : des invites sont presents dans la visio. "
                f"Pour ta securite, cette action n'est disponible qu'en conversation privee avec ton souscripteur."
            ),
            "IMPORTANT": (
                "Explique au souscripteur que cette action est bloquee par securite "
                "car des invites sont dans la visio. Il doit d'abord mettre fin a l'appel "
                "avec les invites, puis relancer une visio privee pour cette demande."
            ),
        }
        return result

    result = {"status": "error", "message": "Fonction inconnue"}

    try:
        if tool_name == "send_sms":
            result = await _tool_send_sms(args, tenant_id=tid)
        elif tool_name == "create_instruction":
            result = await _tool_create_instruction(args, tenant_id=tid)
        elif tool_name == "create_note":
            result = await _tool_create_note(args, tenant_id=tid)
        elif tool_name == "get_contacts":
            result = await _tool_get_contacts(tenant_id=tid)
        elif tool_name == "generate_document":
            result = await _tool_generate_document(args, tenant_id=tid)
        elif tool_name == "alert_contacts":
            result = await _tool_alert_contacts(args, tenant_id=tid)
        elif tool_name == "report_observation":
            result = await _tool_report_observation(args, tenant_id=tid, conversation_id=conversation_id)
        elif tool_name == "send_email":
            result = await _tool_send_email(args, tenant_id=tid)
        elif tool_name == "invite_visio":
            result = await _tool_invite_visio(args, tenant_id=tid, conversation_id=conversation_id)
        elif tool_name == "call_contact":
            result = await _tool_call_contact(args, tenant_id=tid, session_id=conversation_id or "visio")
        elif tool_name == "send_conclusions":
            result = await _tool_send_conclusions(args, tenant_id=tid, conversation_id=conversation_id)
        elif tool_name == "search_web":
            result = await _tool_search_web(args, tenant_id=tid)
        elif tool_name == "search_places":
            result = await _tool_search_places(args, tenant_id=tid)
        elif tool_name == "get_page_info":
            result = await _tool_get_page_info(args, tenant_id=tid)
        elif tool_name == "request_payment":
            result = await _tool_request_payment(args, tenant_id=tid)
        elif tool_name == "get_player_stats":
            result = await _tool_get_player_stats(tenant_id=tid)
        elif tool_name == "get_active_missions":
            result = await _tool_get_active_missions(tenant_id=tid)
        elif tool_name == "get_badges":
            result = await _tool_get_badges(tenant_id=tid)
        elif tool_name == "get_weather":
            result = await _tool_get_weather(args, tenant_id=tid)
        elif tool_name == "get_news":
            result = await _tool_get_news(args)
        elif tool_name == "search_flights":
            result = await _tool_search_flights(args, tenant_id=tid)
        elif tool_name == "search_hotels":
            result = await _tool_search_hotels(args, tenant_id=tid)
        elif tool_name == "book_restaurant":
            result = await _tool_book_restaurant(args, tenant_id=tid)
        # --- Secretary tools ---
        elif tool_name == "get_documents_summary":
            result = _tool_secretary_summary(tid)
        elif tool_name == "get_budget_analysis":
            result = _tool_secretary_budget(tid)
        elif tool_name == "check_affordability":
            result = _tool_secretary_afford(tid, args)
        elif tool_name == "add_expense":
            result = _tool_secretary_add_expense(tid, args)
        elif tool_name == "get_reminders":
            result = _tool_secretary_reminders(tid)
        elif tool_name == "add_reminder":
            result = _tool_secretary_add_reminder(tid, args)
        elif tool_name == "search_documents":
            result = _tool_secretary_search(tid, args)
        elif tool_name == "list_folders":
            result = _tool_secretary_folders(tid)
        elif tool_name == "get_vision_context":
            # Retourne la derniere description vision capturee par le navigateur
            from fastapi import Request as _Req
            client_ip = (request.client.host if request.client else None) or "unknown"
            desc = _visio_scene_cache.get(client_ip, "")
            if not desc:
                # Fallback: chercher dans tous les IPs recents
                desc = next(iter(_visio_scene_cache.values()), "") if _visio_scene_cache else ""
            if desc:
                result = {"status": "ok", "message": desc}
            else:
                result = {"status": "ok", "message": "Aucune image caméra disponible pour le moment. La vision est peut-être en cours d'initialisation."}
        else:
            logger.warning(f"Tavus tool inconnu: {tool_name}")
    except Exception as e:
        logger.error(f"Tavus tool_call error ({tool_name}): {e}")
        result = {"status": "error", "message": str(e)}

    # Anti-hallucination: if action tool failed, make error explicit in result
    _TAVUS_ACTION_TOOLS = {"call_contact", "send_sms", "send_email", "alert_contacts",
                           "request_payment", "invite_visio", "generate_document"}
    if result.get("status") == "error" and tool_name in _TAVUS_ACTION_TOOLS:
        result["IMPORTANT"] = (
            f"L'action {tool_name} a ECHOUE. Tu DOIS informer le souscripteur "
            f"de l'echec. N'invente RIEN. Ne pretends PAS que l'action a reussi."
        )

    # Poste le resultat dans le fil de discussion du tenant
    if _redis_client and tid and tool_name not in ("get_contacts", "report_observation", "get_player_stats", "get_active_missions", "get_badges", "get_weather", "get_news"):
        try:
            mgr = _get_tenant_manager(tid)
            if mgr:
                status_icon = "ok" if result.get("status") == "success" else "erreur"
                action_msg = f"[Visio] {tool_name}: {result.get('message', '')}"
                mgr.add_message(
                    conv_id="visio",
                    role=MessageRole.LUNA,
                    content=action_msg,
                    channel=Channel.APP,
                )
        except Exception as e:
            logger.warning(f"Failed to post visio result to chat: {e}")

    return result


async def _tool_send_sms(args: Dict, tenant_id: int = 0) -> Dict:
    """Envoie un SMS a un contact de confiance."""
    # Service validation (redundant safety layer)
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return {"status": "error", "message": "Service non disponible"}
    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if not mgr or not sms_client.is_configured:
        return {"status": "error", "message": "Service SMS non disponible"}
    # Quota pre-check
    if _quota_guard:
        try:
            qs = _quota_guard.check(tenant_id or TENANT_ID, CoreActionType.SEND_SMS)
            if not qs.allowed:
                return {"status": "error", "message": qs.warning_message or "Quota SMS atteint pour ce mois"}
        except Exception:
            pass

    contact_name = args.get("contact_name", "")
    message = args.get("message", "")
    if not contact_name or not message:
        return {"status": "error", "message": "Nom du contact et message requis"}

    # Cherche le contact par nom avec correspondance floue
    matched_contact, all_contacts = mgr.find_contact_by_name(contact_name)
    if not matched_contact:
        available = ", ".join(c.name for c in all_contacts) if all_contacts else "aucun"
        return {"status": "error", "message": f"Contact '{contact_name}' introuvable. Contacts disponibles : {available}"}
    phone = matched_contact.phone
    matched_name = matched_contact.name

    # Recupere le prenom du souscripteur
    sub_name = _SUBSCRIBER_NAME
    if mgr and tenant_id:
        try:
            profile = mgr.get_subscriber_profile()
            if profile and profile.first_name:
                sub_name = profile.first_name
        except Exception:
            pass

    success, details = _tracked_sms_send(phone, f"[Luna pour {sub_name}] {message}", label=f"Chat SMS a {matched_name}")
    reasoning = f"Luna envoie un SMS a {matched_name} car le souscripteur l'a demande"
    if success:
        # Track cout SMS par tenant via Cortex
        try:
            cortex = get_cortex() if _CORTEX_AVAILABLE else None
            if cortex and hasattr(cortex, 'cost_tracker') and cortex.cost_tracker:
                import asyncio
                await cortex.cost_tracker.track_sms_tenant(tenant_id or 0)
        except Exception:
            pass
        try:
            mgr.add_note(
                content=f"[Action SMS] {reasoning} | Contenu: {message[:100]}",
                context="tool_call",
                tags=["sms", matched_name, "reasoning"],
            )
        except Exception:
            pass
        try:
            mgr.log_event(
                category="action",
                description=f"SMS envoye a {matched_name}: {message[:60]}",
                reasoning=reasoning,
                source="tool_call",
            )
        except Exception:
            pass
        return {"status": "success", "message": f"SMS envoye a {matched_name}", "reasoning": reasoning}
    return {"status": "error", "message": f"Echec envoi SMS: {details.get('error', 'inconnu')}"}


async def _tool_send_email(args: Dict, tenant_id: int = 0) -> Dict:
    """Envoie un email a un contact de confiance."""
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return {"status": "error", "message": "Service non disponible"}
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Service memoire non disponible"}

    # Verifie qu'au moins un service email est dispo
    has_gmail = gmail_client and gmail_client.is_configured and _redis_client and _redis_client.get_email_integration(tid)
    if not has_gmail and not email_client.is_configured:
        return {"status": "error", "message": "Aucun service email configure. Connectez Gmail ou configurez SendGrid."}

    contact_name = args.get("contact_name", "")
    subject = args.get("subject", "")
    body = args.get("body", "")
    direct_email = (args.get("email") or "").strip()
    if not contact_name or not body:
        return {"status": "error", "message": "Nom du contact et contenu requis"}

    # Recupere le prenom du souscripteur
    sub_name = _SUBSCRIBER_NAME
    if mgr and tenant_id:
        try:
            profile = mgr.get_subscriber_profile()
            if profile and profile.first_name:
                sub_name = profile.first_name
        except Exception:
            pass

    if not subject:
        subject = f"Message de {sub_name} via Luna"

    if direct_email:
        # Email fourni directement (formulaire conciergerie) — pas besoin de chercher dans les contacts
        matched_name = contact_name
        contact_email = direct_email
    else:
        # Cherche le contact par nom avec correspondance floue
        matched_contact_email, all_contacts_email = mgr.find_contact_by_name(contact_name)
        if not matched_contact_email:
            available = ", ".join(c.name for c in all_contacts_email) if all_contacts_email else "aucun"
            return {"status": "error", "message": f"Contact '{contact_name}' introuvable. Contacts disponibles : {available}"}
        matched_name = matched_contact_email.name
        contact_email = getattr(matched_contact_email, "email", None)
        if not contact_email:
            return {"status": "error", "message": f"{matched_name} n'a pas d'adresse email enregistree. Renseignez-la dans le formulaire."}

    # Utilise send_for_tenant (Gmail OAuth d'abord, puis SendGrid fallback)
    success, details = await email_client.send_for_tenant(
        tenant_id=tid,
        redis_client=_redis_client,
        gmail_client=gmail_client,
        to=contact_email,
        subject=subject,
        body_text=body,
        subscriber_name=sub_name,
    )
    reasoning = f"Luna envoie un email a {matched_name} car le souscripteur l'a demande"
    if success:
        try:
            mgr.add_note(
                content=f"[Action Email] {reasoning} | Objet: {subject[:80]} | Dest: {matched_name}",
                context="tool_call",
                tags=["email", matched_name, "reasoning"],
            )
        except Exception:
            pass
        try:
            mgr.log_event(
                category="action",
                description=f"Email envoye a {matched_name}: {subject[:60]}",
                reasoning=reasoning,
                source="tool_call",
            )
        except Exception:
            pass
        return {"status": "success", "message": f"Email envoye a {matched_name} ({contact_email})", "reasoning": reasoning}
    return {"status": "error", "message": f"Echec envoi email: {details.get('error', 'inconnu')}"}


async def _tool_invite_visio(args: Dict, tenant_id: int = 0, conversation_id: str = "") -> Dict:
    """Invite un contact en visioconference.

    Si le souscripteur a deja une visio active (conversation_id ou tenant actif),
    partage le MEME lien Daily.js pour que tout le monde soit dans la meme room.
    Sinon, cree une nouvelle conversation.
    """
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return {"status": "error", "message": "Service non disponible"}
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Memoire non disponible"}
    if not tavus_client or not tavus_client.is_configured:
        return {"status": "error", "message": "Service visio non disponible"}
    if not sms_client or not sms_client.is_configured:
        return {"status": "error", "message": "Service SMS non disponible"}

    contact_name = args.get("contact_name", "")
    if not contact_name:
        return {"status": "error", "message": "Nom du contact requis"}

    # Cherche le contact avec correspondance floue
    matched_cv, all_cv = mgr.find_contact_by_name(contact_name)
    if not matched_cv:
        available = ", ".join(c.name for c in all_cv) if all_cv else "aucun"
        return {"status": "error", "message": f"Contact '{contact_name}' introuvable. Contacts disponibles : {available}"}
    phone = matched_cv.phone
    matched_name = matched_cv.name

    # Recupere le prenom du souscripteur
    sub_name = _SUBSCRIBER_NAME
    if mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile and profile.first_name:
                sub_name = profile.first_name
        except Exception:
            pass

    # --- Cherche une conversation active existante ---
    visio_url = None
    existing_conv_id = None

    # 1) conversation_id fourni (appel depuis le webhook Tavus = souscripteur deja en visio)
    if conversation_id and tavus_client:
        url = tavus_client.get_conversation_url(conversation_id)
        if url:
            visio_url = url
            existing_conv_id = conversation_id
            logger.info(f"Invite visio: reutilise conversation existante {conversation_id}")

    # 2) Sinon cherche la conversation active du tenant
    if not visio_url and tavus_client:
        active_conv = tavus_client.get_active_conversation(tid)
        if active_conv:
            visio_url = active_conv.conversation_url
            existing_conv_id = active_conv.conversation_id
            logger.info(f"Invite visio: reutilise conversation active du tenant {tid}: {existing_conv_id}")

    # 3) Si aucune conversation active, cree une nouvelle
    if not visio_url:
        visio_max = int(os.getenv("VISIO_MAX_DURATION", "60")) * 60  # minutes -> secondes
        context = build_tavus_context(
            subscriber_name=sub_name,
            memory_manager=tavus_client.memory,
            guest_names=[matched_name],
        )
        replica_id = _get_luna_replica()
        success_tavus, data = await tavus_client.create_conversation(
            tenant_id=tid,
            custom_greeting=f"Salut {sub_name} ! {matched_name} va bientot nous rejoindre.",
            context=context,
            max_duration=visio_max,
            callback_url=TAVUS_CALLBACK_URL if TAVUS_CALLBACK_URL else None,
            replica_id=replica_id,
        )
        if not success_tavus:
            return {"status": "error", "message": f"Impossible de creer la visio: {data.get('error', 'inconnu')}"}
        visio_url = data["conversation_url"]
        existing_conv_id = data["conversation_id"]
        logger.info(f"Invite visio: nouvelle conversation creee {existing_conv_id}")

    # Track le participant invite dans la conversation Tavus
    if existing_conv_id and tavus_client:
        conv = tavus_client._active_conversations.get(existing_conv_id)
        if conv:
            conv.participants.append({
                "name": matched_name,
                "phone": phone,
                "role": "guest",
                "invited_at": datetime.utcnow().isoformat(),
            })

    # Envoie le lien par SMS directement (pas de OUI/NON — le contact clique s'il veut)
    invite_msg = (
        f"[Luna] {sub_name} t'invite en visioconference ! "
        f"Clique ici pour rejoindre : {visio_url}"
    )
    from integrations.twilio.sms_client import TwilioSMSClient
    normalized_phone = TwilioSMSClient.normalize_phone(phone)
    success_sms, sms_details = _tracked_sms_send(normalized_phone, invite_msg, label=f"Invitation visio a {matched_name}")

    if success_sms:
        # Track cout SMS par tenant via Cortex
        try:
            cortex = get_cortex() if _CORTEX_AVAILABLE else None
            if cortex and hasattr(cortex, 'cost_tracker') and cortex.cost_tracker:
                await cortex.cost_tracker.track_sms_tenant(tenant_id or 0)
        except Exception:
            pass

    if not success_sms:
        # La visio est creee mais le SMS a echoue — donne quand meme le lien
        logger.error(f"SMS invite failed but visio created: {visio_url}")
        return {
            "status": "partial",
            "message": f"Visio creee mais echec envoi SMS a {matched_name}. Voici le lien : {visio_url}",
            "visio_url": visio_url,
        }

    logger.info(f"Visio invite sent: {matched_name} ({normalized_phone}) -> {visio_url}")

    try:
        mgr.log_event(
            category="action",
            description=f"Invitation visio envoyee a {matched_name} avec lien {visio_url}",
            reasoning=f"Luna invite {matched_name} en visio a la demande de {sub_name}",
            source="tool_call",
        )
    except Exception:
        pass

    return {
        "status": "success",
        "message": f"Lien visio envoye a {matched_name} par SMS ! Vous serez dans la meme conversation video.",
        "visio_url": visio_url,
    }


async def _tool_send_conclusions(args: Dict, tenant_id: int = 0, conversation_id: str = "") -> Dict:
    """Rédige un compte-rendu professionnel et l'envoie aux participants de la visio."""
    import json as _jsc
    from datetime import datetime as _dt

    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Service mémoire non disponible"}

    subject = args.get("subject", "Compte-rendu de réunion")
    conclusions = args.get("conclusions", "")
    recipients_filter = [r.lower() for r in (args.get("recipients") or [])]
    if not conclusions:
        return {"status": "error", "message": "Contenu des conclusions requis"}

    sub_name = _tenant_subscriber_first_name(tid)
    profile = None
    try:
        profile = mgr.get_subscriber_profile()
        if profile and profile.first_name:
            sub_name = profile.first_name
    except Exception:
        pass

    now_str = _dt.now(ZoneInfo("Europe/Paris")).strftime("%d/%m/%Y à %H:%M")

    # Récupère les participants depuis Redis
    participants = []
    if _redis_client and conversation_id:
        try:
            raw = _redis_client.client.get(f"luna:conv:{conversation_id}:participants")
            if raw:
                participants = _jsc.loads(raw.decode() if isinstance(raw, bytes) else raw)
        except Exception:
            pass

    if recipients_filter and participants:
        participants = [p for p in participants
                       if any(f in p.get("name", "").lower() for f in recipients_filter)]

    # Document formaté
    doc_text = (
        f"COMPTE-RENDU DE RÉUNION\n"
        f"Rédigé par Luna IA — {now_str}\n"
        f"Animateur : {sub_name}\n"
        f"Objet : {subject}\n\n"
        f"{'─' * 40}\n\n"
        f"{conclusions}\n\n"
        f"{'─' * 40}\n"
        f"Document généré automatiquement par Luna IA (YAWatch)"
    )

    # Sauvegarder en mémoire
    try:
        mgr.add_note(
            content=f"[CR visio] {subject} ({now_str})\n{conclusions[:500]}",
            context="visio_conclusions",
            tags=["visio", "compte-rendu"],
        )
    except Exception:
        pass

    # Générer DOCX
    download_url = ""
    if _doc_generator:
        try:
            fname = _doc_generator.generate_letter(
                doc_type="compte_rendu",
                subject=subject,
                body_text=conclusions,
                profile=profile.model_dump() if profile else {},
            )
            download_url = f"/api/documents/download/{fname}"
        except Exception as _de:
            logger.warning(f"send_conclusions docx: {_de}")

    # Envoyer aux participants
    sent_to, errors = [], []
    base_url = os.getenv("BASE_URL", "")
    for p in participants:
        pname = p.get("name", "l'invité")
        phone = p.get("phone", "")
        sms_preview = conclusions[:200].replace("\n", " ")
        sms_msg = (
            f"[Luna] CR '{subject}' ({now_str}) :\n{sms_preview}…"
            + (f"\nDoc : {base_url}{download_url}" if download_url else "")
        )
        if phone:
            try:
                from integrations.twilio.sms_client import TwilioSMSClient
                norm = TwilioSMSClient.normalize_phone(phone)
                ok_sms, _ = _tracked_sms_send(norm, sms_msg[:320], label=f"CR visio {pname}")
                if ok_sms:
                    sent_to.append(pname)
                else:
                    errors.append(f"{pname}: SMS échoué")
            except Exception as _se:
                errors.append(f"{pname}: {str(_se)[:60]}")

    if not participants:
        return {
            "status": "partial",
            "message": (
                "Compte-rendu rédigé et sauvegardé en mémoire"
                + (f". Disponible : {download_url}" if download_url else "")
                + ". Aucun participant enregistré — partage le lien manuellement si besoin."
            ),
            "download_url": download_url,
        }

    logger.info(f"send_conclusions sent to {sent_to}")
    return {
        "status": "success",
        "message": (
            f"Compte-rendu envoyé à : {', '.join(sent_to) if sent_to else 'personne'}."
            + (f" Erreurs : {'; '.join(errors)}" if errors else "")
            + (f" Document : {download_url}" if download_url else "")
        ),
        "sent_to": sent_to,
        "download_url": download_url,
    }


async def _tool_call_contact(args: Dict, tenant_id: int = 0, session_id: str = "default") -> Dict:
    """Appelle un contact de confiance en audio via Twilio. Luna parle au contact et transmet un message."""
    if _license_heartbeat and (_license_heartbeat.is_blocked() or _license_heartbeat.is_degraded()):
        return {"status": "error", "message": "Service non disponible"}
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Memoire non disponible"}
    if not voice_client or not voice_client.is_configured:
        return {"status": "error", "message": "Service d'appels vocaux non configure"}
    if not VOICE_CALLBACK_URL:
        return {"status": "error", "message": "URL de callback vocal non configuree"}

    contact_name = args.get("contact_name", "")
    message = args.get("message", "")
    direct_phone = args.get("phone_number", "")
    if not contact_name:
        return {"status": "error", "message": "Nom du contact requis"}

    # Numeros d'urgence interdits (Luna ne doit pas appeler le 15, 17, 18, 112)
    _EMERGENCY_NUMBERS = {"15", "17", "18", "112", "114", "115", "119", "3114", "3977"}
    if direct_phone and direct_phone.strip().replace(" ", "") in _EMERGENCY_NUMBERS:
        return {"status": "error", "message": f"Luna ne peut pas appeler les numeros d'urgence. Suggere au souscripteur de composer le {direct_phone} lui-meme."}

    phone = None
    matched_name = contact_name
    is_admin_call = False

    if direct_phone:
        # Appel administration/service : numero fourni directement par le souscripteur
        phone = direct_phone
        matched_name = contact_name or "Administration"
        is_admin_call = True
    else:
        # Cherche dans les contacts de confiance avec correspondance floue
        matched_c, all_c = mgr.find_contact_by_name(contact_name)
        if matched_c:
            phone = matched_c.phone
            matched_name = matched_c.name
        else:
            available = ", ".join(c.name for c in all_c) if all_c else "aucun"
            return {"status": "error", "message": f"Contact '{contact_name}' introuvable. Contacts disponibles : {available}. Pour appeler un service, fournis le numero directement."}

    if not phone:
        return {"status": "error", "message": f"Contact '{contact_name}' non trouve parmi les contacts de confiance."}

    # Recupere le prenom du souscripteur
    sub_name = _SUBSCRIBER_NAME
    if mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile and profile.first_name:
                sub_name = profile.first_name
        except Exception:
            pass

    # Normaliser le numero
    from integrations.twilio.sms_client import TwilioSMSClient
    normalized_phone = TwilioSMSClient.normalize_phone(phone)

    # Mention legale rapport PDF (l'interlocuteur doit etre informe)
    _report_notice = (
        "IMPORTANT : En debut d'appel, apres ta salutation, informe l'interlocuteur que "
        "cet appel fait l'objet d'un compte-rendu qui sera transmis a {sub_name}. "
        "Dis quelque chose comme : 'Je te previens que {sub_name} recevra un petit "
        "resume de notre echange, c'est pour qu'il soit au courant de tout.' "
        "Si l'interlocuteur refuse, note-le et termine l'appel poliment."
    ).format(sub_name=sub_name)

    # Mission pour Luna pendant l'appel
    if is_admin_call:
        mission = f"Tu appelles {matched_name} pour {sub_name}. "
        if message:
            mission += f"Voici la demande de {sub_name} : {message}"
        else:
            mission += f"{sub_name} souhaite obtenir des informations."
        mission += f"\n{_report_notice}"
        greeting = f"Bonjour ! Je suis Luna, l'assistante de {sub_name}. "
        if message:
            greeting += f"J'appelle de la part de {sub_name}. {message}"
        else:
            greeting += f"J'appelle de la part de {sub_name} pour obtenir des renseignements."
    else:
        mission = f"Tu appelles {matched_name} de la part de {sub_name}. "
        if message:
            mission += f"Voici ce que {sub_name} veut que tu transmettes : {message}"
        else:
            mission += f"{sub_name} voulait prendre des nouvelles."
        mission += f"\n{_report_notice}"
        greeting = f"Bonjour {matched_name} ! C'est Luna, l'assistante de {sub_name}. "
        if message:
            greeting += f"{sub_name} m'a demande de t'appeler pour te dire : {message}"
        else:
            greeting += f"{sub_name} m'a demande de t'appeler pour prendre de tes nouvelles."

    # Lancer l'appel via Twilio
    success, data = await voice_client.initiate_call_async(normalized_phone)
    if not success:
        return {"status": "error", "message": f"Impossible d'appeler {matched_name}: {data.get('error', 'erreur inconnue')}"}

    call_sid = data.get("call_sid", "")
    # Stocker les parametres pour le bridge (lock pour thread-safety)
    import time as _ts
    _req_duration = args.get("max_duration", 180)
    if isinstance(_req_duration, str):
        try: _req_duration = int(_req_duration)
        except: _req_duration = 180
    _req_duration = max(60, min(600, _req_duration))  # clamp 1-10 min

    async with _voice_call_params_lock:
        _voice_call_params[call_sid] = {
            "mission": mission,
            "max_duration": _req_duration,
            "greeting": greeting,
            "contact_name": matched_name,
            "tenant_id": tid,
            "subscriber_name": sub_name,
            "contact_phone": normalized_phone,
            "message": message,
            "session_id": session_id,
            "_ts": _ts.time(),
        }

    logger.info(f"Voice call initiated to {matched_name} ({normalized_phone}) call_sid={call_sid}")

    # Track active call so chat knows an outgoing call is in progress
    _active_voice_calls[tid] = {
        "contact": matched_name,
        "started": time.time(),
        "call_sid": call_sid,
        "phone": normalized_phone,
        "session_id": session_id,
    }

    try:
        mgr.log_event(
            category="action",
            description=f"Appel audio lance vers {matched_name} ({normalized_phone})",
            reasoning=f"Luna appelle {matched_name} a la demande de {sub_name}: {message[:80]}",
            source="tool_call",
        )
    except Exception:
        pass

    return {
        "status": "success",
        "message": f"J'appelle {matched_name} maintenant ! L'appel est en cours. Tu recevras un compte-rendu par email quand l'appel sera termine.",
        "call_sid": call_sid,
    }


async def _tool_create_instruction(args: Dict, tenant_id: int = 0) -> Dict:
    """Cree une instruction depuis la visio."""
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr or not _CORE_AVAILABLE:
        return {"status": "error", "message": "Memoire non disponible"}

    text = args.get("text", "")
    if not text:
        return {"status": "error", "message": "Texte de l'instruction requis"}

    parsed = InstructionParser.parse(text)

    action_map = {
        "reminder": SchemaActionType.REMINDER,
        "sms_contact": SchemaActionType.SMS,
        "call_contact": SchemaActionType.CALL,
        "visio_contact": SchemaActionType.VISIO,
        "wake_up": SchemaActionType.CALL,
        "check_in": SchemaActionType.REMINDER,
        "surveillance": SchemaActionType.ALERT,
        "note": SchemaActionType.NOTE,
        "daily_routine": SchemaActionType.REMINDER,
        "information": SchemaActionType.REMINDER,
        "reading": SchemaActionType.REMINDER,
        "game": SchemaActionType.REMINDER,
        "music": SchemaActionType.REMINDER,
        "gratitude": SchemaActionType.REMINDER,
    }
    action = action_map.get(parsed.action_type.value, SchemaActionType.REMINDER)

    type_map = {
        "once": InstructionType.ONE_TIME,
        "daily": InstructionType.DAILY,
        "weekly": InstructionType.RECURRING,
        "monthly": InstructionType.RECURRING,
        "weekdays": InstructionType.RECURRING,
        "weekend": InstructionType.RECURRING,
    }
    instr_type = type_map.get(parsed.recurrence.value, InstructionType.ONE_TIME)

    schedule_str = parsed.to_cron() or ""
    if not schedule_str and parsed.scheduled_time:
        schedule_str = f"{parsed.scheduled_time.hour:02d}:{parsed.scheduled_time.minute:02d}"

    # Metadata pour les appels planifiés (durée max)
    _instr_meta = None
    _max_dur_min = args.get("max_duration_minutes")
    if _max_dur_min is not None:
        try:
            _max_dur_min = max(1, min(10, int(_max_dur_min)))
        except (ValueError, TypeError):
            _max_dur_min = 3
        _instr_meta = {"max_duration": _max_dur_min * 60}
    elif parsed.action_type.value in ("call_contact", "wake_up"):
        # Défaut 3 min pour les appels planifiés
        _instr_meta = {"max_duration": 180}

    instr = mgr.add_instruction(
        description=text,
        action=action,
        instruction_type=instr_type,
        schedule=schedule_str,
        target=parsed.target or "self",
        message_template=parsed.message_template or "",
        priority=parsed.priority,
        metadata=_instr_meta,
    )

    if _scheduler:
        try:
            _scheduler.schedule(instruction_id=instr.id, tenant_id=tid, instruction=parsed)
        except Exception:
            pass

    confirmation = InstructionParser.format_confirmation(parsed)
    if _max_dur_min:
        confirmation += f" (duree max: {_max_dur_min} min)"
    return {"status": "success", "message": confirmation, "instruction_id": instr.id}


async def _tool_join_conference(args: Dict, tenant_id: int = 0) -> Dict:
    """Lance un appel vers un bridge de conférence avec DTMF PIN + Media Stream."""
    phone_number = args.get("phone_number", "").strip()
    pin = args.get("pin", "").strip()
    conference_name = args.get("context", "Conference").strip()
    max_dur_min = int(args.get("max_duration_minutes", 60))

    if not phone_number:
        return {"status": "error", "message": "Numero de conference manquant."}
    if not voice_client or not voice_client.is_configured:
        return {"status": "error", "message": "Service d'appel non configure."}
    if not VOICE_CALLBACK_URL:
        return {"status": "error", "message": "VOICE_CALLBACK_URL manquant."}

    from urllib.parse import urlencode
    import time as _time

    qs = urlencode({"pin": pin, "conf_name": conference_name, "max_min": max_dur_min})
    twiml_url = f"{VOICE_CALLBACK_URL.rstrip('/')}/api/voice-call/conference-twiml?{qs}"

    success, data = await voice_client.make_call_to_async(phone_number, twiml_url)
    if not success:
        return {"status": "error", "message": f"Impossible de rejoindre la conference : {data.get('error')}"}

    call_sid = data.get("call_sid", "")
    async with _voice_call_params_lock:
        _voice_call_params[call_sid] = {
            "source": "conference",
            "conference_context": conference_name,
            "max_duration_minutes": max_dur_min,
            "tenant_id": tenant_id,
            "_ts": _time.time(),
        }

    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if mgr:
        mgr.add_note(
            content=f"[Conference lancee] Luna rejoint '{conference_name}' ({phone_number})",
            context="conference",
            tags=["conference", "rapport"],
        )

    return {
        "status": "ok",
        "call_sid": call_sid,
        "message": f"Luna rejoint la conference '{conference_name}'. Les notes seront disponibles dans l'onglet Rapports.",
    }


async def _tool_create_note(args: Dict, tenant_id: int = 0) -> Dict:
    """Prend une note depuis la visio."""
    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Memoire non disponible"}

    content = args.get("content", "")
    if not content:
        return {"status": "error", "message": "Contenu de la note requis"}

    reasoning = f"Luna prend une note a la demande du souscripteur"
    try:
        mgr.add_note(
            content=f"{content}\n[Raison: {reasoning}]",
            context="tool_call",
            tags=["note", "reasoning"],
        )
    except Exception as e:
        logger.warning(f"Note creation failed: {type(e).__name__}: {e}")
        if "quota" in str(e).lower() or "Quota" in type(e).__name__:
            return {"status": "error", "message": "Tu as atteint le nombre maximum de notes pour ton abonnement. Supprime des anciennes notes pour en creer de nouvelles."}
        return {"status": "error", "message": f"Impossible de sauvegarder la note: {e}"}
    return {"status": "success", "message": f"Note enregistree: {content[:50]}", "reasoning": reasoning}


async def _tool_get_contacts(tenant_id: int = 0) -> Dict:
    """Liste les contacts de confiance."""
    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Memoire non disponible"}

    contacts = mgr.list_trusted_contacts()
    if not contacts:
        return {"status": "success", "message": "Aucun contact de confiance enregistre.", "contacts": []}

    contact_list = [{"name": c.name, "relation": c.relation} for c in contacts]
    names = ", ".join([f"{c.name} ({c.relation})" for c in contacts])
    return {"status": "success", "message": f"Contacts de confiance: {names}", "contacts": contact_list}


async def _tool_generate_document(args: Dict, tenant_id: int = 0) -> Dict:
    """Genere un document depuis la visio."""
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not _doc_generator or not mgr:
        return {"status": "error", "message": "Generateur de documents non disponible"}

    doc_type = args.get("doc_type", "courrier_admin")
    subject = args.get("subject", "")
    details = args.get("details", "")

    if not subject:
        return {"status": "error", "message": "Sujet du document requis"}

    # Genere le contenu avec GPT
    profile = mgr.get_subscriber_profile()
    profile_dict = profile.model_dump() if profile else {}

    prompt = f"Redige un {doc_type} professionnel en francais.\nObjet: {subject}\n"
    if details:
        prompt += f"Details: {details}\n"
    if profile_dict.get("first_name"):
        prompt += f"Expediteur: {profile_dict.get('first_name', '')} {profile_dict.get('last_name', '')}\n"
    prompt += "Redige uniquement le corps du courrier, sans en-tete ni signature (ils seront ajoutes automatiquement)."

    try:
        gpt_resp = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1000,
            temperature=0.7,
            timeout=30,
        )
        await _track_openai_cost(gpt_resp)
        body_text = gpt_resp.choices[0].message.content
    except Exception as e:
        logger.error(f"Document generation LLM error: {e}")
        return {"status": "error", "message": "Erreur lors de la generation du document. Reessaie."}

    if doc_type == "fiche_sante" and profile_dict:
        filename = _doc_generator.generate_health_sheet(profile_dict)
    else:
        filename = _doc_generator.generate_letter(
            doc_type=doc_type,
            subject=subject,
            body_text=body_text,
            profile=profile_dict,
        )

    url = f"/api/documents/download/{filename}"

    try:
        mgr.add_note(
            content=f"Document genere: {subject} ({doc_type}) → {filename}",
            context="document",
            tags=["document", doc_type],
        )
    except Exception:
        pass

    # Envoie automatiquement le contenu par email au souscripteur si email dispo
    email_sent = False
    if profile_dict.get("email") and body_text:
        try:
            sub_email = profile_dict["email"]
            sub_name = profile_dict.get("first_name", "")
            email_subject = f"Document Luna : {subject}"
            email_body = f"Bonjour {sub_name},\n\nVoici le document que tu m'as demande :\n\n---\n{body_text}\n---\n\nBien a toi,\nLuna"
            ok, _ = await email_client.send_for_tenant(
                tenant_id=tid,
                redis_client=_redis_client,
                gmail_client=gmail_client,
                to=sub_email,
                subject=email_subject,
                body_text=email_body,
                subscriber_name=sub_name,
            )
            if ok:
                email_sent = True
                logger.info(f"Document envoye par email a {sub_email}")
        except Exception as e:
            logger.warning(f"Failed to email document: {e}")

    msg = f"Document genere : {subject}."
    if email_sent:
        msg += f" Je te l'ai aussi envoye par email a {profile_dict['email']}."
    msg += f" Tu peux le retrouver dans l'onglet Documents."

    return {"status": "success", "message": msg, "url": url, "filename": filename}


async def _tool_alert_contacts(args: Dict, tenant_id: int = 0) -> Dict:
    """Alerte tous les contacts de confiance avec position GPS si disponible."""
    mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
    if not mgr or not sms_client.is_configured:
        return {"status": "error", "message": "Service SMS non disponible"}

    reason = args.get("reason", "situation preoccupante")
    contacts = mgr.list_trusted_contacts()

    if not contacts:
        return {"status": "error", "message": "Aucun contact de confiance enregistre"}

    profile = mgr.get_subscriber_profile()
    name = profile.first_name if profile else "votre proche"

    # Récupérer la dernière position GPS connue (TTL 1h)
    location_line = ""
    try:
        import json as _json
        geo_raw = _redis_client.client.get(f"luna:{tenant_id}:geolocation") if _redis_client else None
        if geo_raw:
            geo = _json.loads(geo_raw)
            address = geo.get("address") or geo.get("city", "")
            lat = geo.get("latitude")
            lng = geo.get("longitude")
            if address:
                location_line = f"\n📍 Dernière position connue : {address}"
            elif lat and lng:
                location_line = f"\n📍 Dernière position connue : https://maps.google.com/?q={lat},{lng}"
    except Exception:
        pass

    # Heure locale Paris
    try:
        import pytz as _pytz
        _paris = _pytz.timezone("Europe/Paris")
        _now_paris = datetime.utcnow().replace(tzinfo=_pytz.utc).astimezone(_paris)
        time_line = f"\n🕐 Heure : {_now_paris.strftime('%H:%M')} (heure de Paris)"
    except Exception:
        time_line = ""

    sent = 0
    for c in contacts:
        # Ajoute l'adresse du contact si disponible
        addr_line = f"\n🏠 Votre adresse : {c.address}" if getattr(c, "address", None) else ""
        msg = (
            f"[ALERTE Luna] {name} a besoin d'aide. Raison : {reason}."
            f"{location_line}"
            f"{time_line}"
            f"{addr_line}"
            f"\nUrgence médicale : 15 (SAMU) • Pompiers : 18 • Urgences : 112"
            f"\nMerci de vérifier qu'il va bien."
        )
        success, _ = _tracked_sms_send(c.phone, msg, label="Alerte contacts urgence")
        if success:
            sent += 1
            # Track cout SMS par tenant via Cortex
            try:
                cortex = get_cortex() if _CORTEX_AVAILABLE else None
                if cortex and hasattr(cortex, 'cost_tracker') and cortex.cost_tracker:
                    await cortex.cost_tracker.track_sms_tenant(tenant_id or 0)
            except Exception:
                pass

    reasoning = f"Luna alerte les contacts car: {reason}"
    try:
        mgr.add_note(
            content=f"[Action ALERTE] {reasoning} | {sent} contact(s) alertes",
            context="alerte_urgence",
            tags=["urgence", "alerte", "reasoning"],
        )
    except Exception:
        pass
    try:
        mgr.log_event(
            category="safety",
            description=f"ALERTE envoyee a {sent} contact(s): {reason}",
            reasoning=reasoning,
            source="tool_call",
        )
    except Exception:
        pass
    return {"status": "success", "message": f"Alerte envoyee a {sent} contact(s) de confiance", "reasoning": reasoning}


async def _tool_report_observation(args: Dict, tenant_id: int = 0, conversation_id: str = "") -> Dict:
    """Log une observation visuelle Raven. Si severity==concern, déclenche la chaîne de vérification."""
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "error", "message": "Memoire non disponible"}

    observation = args.get("observation", "")
    severity = args.get("severity", "info")

    reasoning = f"Luna note une observation visuelle Raven pendant la visio: {severity}"
    if observation:
        mgr.add_note(
            content=f"[Observation visio] {observation}",
            context="visio_perception",
            tags=["perception", "visio", "raven", severity, "reasoning"],
        )
        try:
            mgr.log_perception_event({
                "type": "visio_observation",
                "severity": severity,
                "description": observation,
                "reasoning": reasoning,
                "source": "tavus_raven",
                "timestamp": datetime.utcnow().isoformat(),
            })
        except Exception:
            pass

        # CHAÎNE INCIDENT — severity concern uniquement
        if severity == "concern" and _redis_client:
            checkin_data = {
                "observation": observation,
                "severity": severity,
                "timestamp": datetime.utcnow().isoformat(),
                "conversation_id": conversation_id,
                "tenant_id": tid,
                "status": "pending",
            }
            redis_key = f"luna:{tid}:checkin_pending:{conversation_id or 'default'}"
            _redis_client.client.setex(redis_key, 45, json.dumps(checkin_data))
            asyncio.create_task(_schedule_auto_alert(tid, conversation_id or "default", observation))

            return {
                "status": "success",
                "URGENT_INSTRUCTION": (
                    f"ALERTE VISUELLE DETECTEE : {observation}. "
                    f"Demande IMMEDIATEMENT au souscripteur : 'Tu vas bien ? Reponds-moi.' "
                    f"Si pas de reponse positive dans 30 secondes, utilise la fonction alert_contacts "
                    f"avec la raison : '{observation}'"
                ),
                "reasoning": reasoning,
                "checkin_initiated": True,
            }

    return {"status": "success", "message": "Observation notee", "reasoning": reasoning}


async def _schedule_auto_alert(tenant_id: int, conversation_id: str, reason: str) -> None:
    """Attend 35s, déclenche l'alerte si le checkin n'a pas été confirmé."""
    await asyncio.sleep(35)
    if not _redis_client:
        return
    redis_key = f"luna:{tenant_id}:checkin_pending:{conversation_id}"
    checkin_raw = _redis_client.client.get(redis_key)
    if not checkin_raw:
        return
    try:
        checkin = json.loads(checkin_raw)
    except Exception:
        return
    if checkin.get("status") != "pending":
        return
    checkin["status"] = "alerted"
    checkin["alerted_at"] = datetime.utcnow().isoformat()
    _redis_client.client.setex(redis_key, 300, json.dumps(checkin))
    try:
        mgr = _get_tenant_manager(tenant_id) if tenant_id else _memory_manager
        result = await _tool_alert_contacts({"reason": reason}, tenant_id)
        logger.warning(f"ALERTE AUTO declenche apres non-reponse [tenant={tenant_id}]: {result}")
        if mgr:
            mgr.log_event(
                category="safety",
                description="Alerte automatique declenchee apres non-reponse Raven (35s)",
                reasoning=f"Observation: {reason}",
                source="checkin_auto",
            )
    except Exception as e:
        logger.exception(f"Echec alerte auto [tenant={tenant_id}]: {e}")


# --- GAMIFICATION TOOLS (chat + voix) ---

async def _tool_get_player_stats(tenant_id: int = 0) -> Dict:
    """Retourne les stats du joueur dans le Monde IA Watch."""
    tid = tenant_id or TENANT_ID
    if not _GAMIFICATION_AVAILABLE or not _redis_client:
        return {"status": "error", "message": "Monde IA Watch non disponible"}
    try:
        from core.gamification.redis_ops import GamificationRedisOps
        from core.gamification.engine import get_level_for_xp
        gops = GamificationRedisOps(_redis_client)
        player = gops.get_player(tid)
        if not player:
            return {"status": "info", "message": "Aucune progression trouvee. Le souscripteur n'a pas encore commence."}
        xp = int(player.get("xp", 0))
        lvl = get_level_for_xp(xp)
        stab = gops.get_stability(tid)
        return {
            "status": "success",
            "niveau": lvl["level"],
            "titre": lvl["title"],
            "xp": xp,
            "xp_prochain_niveau": lvl["xp_next_level"],
            "progression": f"{lvl['progress_percent']}%",
            "etoiles": int(player.get("stars", 0)),
            "serie_jours": int(player.get("streak_days", 0)),
            "meilleure_serie": int(player.get("streak_best", 0)),
            "jours_actifs": int(player.get("days_active", 0)),
            "messages_total": int(player.get("total_messages", 0)),
            "appels_total": int(player.get("total_calls", 0)),
            "stabilite": int(stab.get("score", 70)) if stab else 70,
            "tendance_stabilite": stab.get("trend", "stable") if stab else "stable",
        }
    except Exception as e:
        logger.warning(f"get_player_stats error: {e}")
        return {"status": "error", "message": "Erreur lors de la recuperation des stats"}


async def _tool_get_active_missions(tenant_id: int = 0) -> Dict:
    """Retourne les missions actives du joueur."""
    tid = tenant_id or TENANT_ID
    if not _GAMIFICATION_AVAILABLE or not _redis_client:
        return {"status": "error", "message": "Monde IA Watch non disponible"}
    try:
        from core.gamification.redis_ops import GamificationRedisOps
        gops = GamificationRedisOps(_redis_client)
        mission_ids = gops.get_active_mission_ids(tid)
        if not mission_ids:
            return {"status": "info", "missions": [], "message": "Aucune mission active pour le moment."}
        missions = []
        for mid in mission_ids:
            m = gops.get_mission(tid, mid)
            if m:
                prog = int(m.get("progress", 0)) - int(m.get("start_progress", 0))
                tgt = int(m.get("target", 1)) - int(m.get("start_progress", 0))
                missions.append({
                    "titre": m.get("title", "Mission"),
                    "description": m.get("description", ""),
                    "progression": f"{prog}/{tgt}",
                    "pourcentage": f"{min(round(prog / max(tgt, 1) * 100), 100)}%",
                    "recompense_xp": int(m.get("xp_reward", 0)),
                    "statut": m.get("status", "active"),
                })
        return {"status": "success", "missions": missions}
    except Exception as e:
        logger.warning(f"get_active_missions error: {e}")
        return {"status": "error", "message": "Erreur lors de la recuperation des missions"}


async def _tool_get_badges(tenant_id: int = 0) -> Dict:
    """Retourne les badges gagnes par le joueur."""
    tid = tenant_id or TENANT_ID
    if not _GAMIFICATION_AVAILABLE or not _redis_client:
        return {"status": "error", "message": "Monde IA Watch non disponible"}
    try:
        from core.gamification.redis_ops import GamificationRedisOps
        from core.gamification.constants import ALL_CLIENT_BADGES
        gops = GamificationRedisOps(_redis_client)
        badge_ids = gops.get_badges(tid)
        if not badge_ids:
            return {"status": "info", "badges": [], "message": "Aucun badge obtenu pour le moment. Encourage le souscripteur a explorer !"}
        badges = []
        for bid in badge_ids:
            bdef = ALL_CLIENT_BADGES.get(bid)
            if bdef:
                detail = gops.get_badge_detail(tid, bid)
                badges.append({
                    "nom": bdef["name"],
                    "description": bdef["description"],
                    "categorie": bdef["category"],
                    "rarete": bdef["rarity"],
                    "obtenu_le": detail.get("earned_at", "") if detail else "",
                })
        return {"status": "success", "badges": badges, "total": len(badges)}
    except Exception as e:
        logger.warning(f"get_badges error: {e}")
        return {"status": "error", "message": "Erreur lors de la recuperation des badges"}


# --- CONCIERGERIE: Meteo + Actualites + Recherche web + Paiement ---

async def _tool_get_weather(args: Dict, tenant_id: int = 0) -> Dict:
    """Meteo via wttr.in + fallback Open-Meteo (gratuits, pas de cle API)."""
    import httpx
    city = args.get("city", "")
    if isinstance(city, dict):
        city = ""

    # Si pas de ville, essayer la geolocalisation du souscripteur
    if not city and _redis_client and tenant_id:
        try:
            import json as _jgeo
            geo_raw = _redis_client.client.get(f"luna:{tenant_id}:geolocation")
            if geo_raw:
                geo_str = geo_raw.decode() if isinstance(geo_raw, bytes) else geo_raw
                geo = _jgeo.loads(geo_str)
                city = geo.get("city", "") or ""
        except Exception:
            pass
    if not city:
        city = "Paris"

    # Essayer wttr.in d'abord
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            resp = await client.get(
                f"https://wttr.in/{city}?format=j1&lang=fr",
                headers={"User-Agent": "Luna/2.2 (concierge assistant)"}
            )
            if resp.status_code == 200:
                data = resp.json()
                current = data.get("current_condition", [{}])[0]
                forecasts = data.get("weather", [])
                area = data.get("nearest_area", [{}])[0]
                area_name = area.get("areaName", [{}])[0].get("value", city)
                country = area.get("country", [{}])[0].get("value", "")
                weather_now = {
                    "ville": area_name, "pays": country,
                    "temperature": f"{current.get('temp_C', '?')}°C",
                    "ressenti": f"{current.get('FeelsLikeC', '?')}°C",
                    "description": current.get("lang_fr", [{}])[0].get("value", current.get("weatherDesc", [{}])[0].get("value", "")),
                    "humidite": f"{current.get('humidity', '?')}%",
                    "vent": f"{current.get('windspeedKmph', '?')} km/h",
                    "visibilite": f"{current.get('visibility', '?')} km",
                }
                previsions = []
                for day in forecasts[:3]:
                    hourly = day.get("hourly", [])
                    idx = min(4, len(hourly) - 1) if hourly else 0
                    desc = hourly[idx].get("lang_fr", [{}])[0].get("value", "") if hourly else ""
                    previsions.append({
                        "date": day.get("date", ""),
                        "max": f"{day.get('maxtempC', '?')}°C",
                        "min": f"{day.get('mintempC', '?')}°C",
                        "description": desc,
                    })
                return {"status": "success", "actuel": weather_now, "previsions": previsions}
    except Exception as e:
        logger.info(f"wttr.in failed ({type(e).__name__}), trying Open-Meteo fallback")

    # Fallback: Open-Meteo (gratuit, fiable, pas de cle API)
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            # Geocoding pour obtenir lat/lon
            geo_resp = await client.get(
                f"https://geocoding-api.open-meteo.com/v1/search?name={city}&count=1&language=fr"
            )
            if geo_resp.status_code != 200:
                return {"status": "error", "message": f"Ville introuvable: {city}"}
            geo_data = geo_resp.json()
            results = geo_data.get("results", [])
            if not results:
                return {"status": "error", "message": f"Ville introuvable: {city}"}
            lat = results[0]["latitude"]
            lon = results[0]["longitude"]
            geo_city = results[0].get("name", city)
            geo_country = results[0].get("country", "")

            # Meteo actuelle + previsions
            meteo_resp = await client.get(
                f"https://api.open-meteo.com/v1/forecast?"
                f"latitude={lat}&longitude={lon}"
                f"&current=temperature_2m,relative_humidity_2m,apparent_temperature,weather_code,wind_speed_10m"
                f"&daily=weather_code,temperature_2m_max,temperature_2m_min"
                f"&timezone=auto&forecast_days=3"
            )
            if meteo_resp.status_code != 200:
                return {"status": "error", "message": "Erreur API meteo"}
            m = meteo_resp.json()

        cur = m.get("current", {})
        # WMO weather codes -> description francaise
        wmo_code = cur.get("weather_code", 0)
        wmo_desc = _wmo_to_french(wmo_code)

        weather_now = {
            "ville": geo_city, "pays": geo_country,
            "temperature": f"{round(cur.get('temperature_2m', 0))}°C",
            "ressenti": f"{round(cur.get('apparent_temperature', 0))}°C",
            "description": wmo_desc,
            "humidite": f"{cur.get('relative_humidity_2m', '?')}%",
            "vent": f"{round(cur.get('wind_speed_10m', 0))} km/h",
        }

        previsions = []
        daily = m.get("daily", {})
        dates = daily.get("time", [])
        maxs = daily.get("temperature_2m_max", [])
        mins = daily.get("temperature_2m_min", [])
        codes = daily.get("weather_code", [])
        for i in range(min(3, len(dates))):
            previsions.append({
                "date": dates[i] if i < len(dates) else "",
                "max": f"{round(maxs[i])}°C" if i < len(maxs) else "?°C",
                "min": f"{round(mins[i])}°C" if i < len(mins) else "?°C",
                "description": _wmo_to_french(codes[i]) if i < len(codes) else "",
            })

        return {"status": "success", "actuel": weather_now, "previsions": previsions}
    except Exception as e:
        logger.warning(f"Weather fallback error: {type(e).__name__}: {e}")
        return {"status": "error", "message": f"Erreur meteo: {type(e).__name__}"}


def _wmo_to_french(code: int) -> str:
    """Convertit un WMO weather code en description francaise."""
    wmo = {
        0: "Ensoleille", 1: "Principalement degage", 2: "Partiellement nuageux",
        3: "Couvert", 45: "Brouillard", 48: "Brouillard givrant",
        51: "Bruine legere", 53: "Bruine moderee", 55: "Bruine dense",
        56: "Bruine verglacante", 57: "Bruine verglacante forte",
        61: "Pluie legere", 63: "Pluie moderee", 65: "Pluie forte",
        66: "Pluie verglacante", 67: "Pluie verglacante forte",
        71: "Neige legere", 73: "Neige moderee", 75: "Neige forte",
        77: "Grains de neige", 80: "Averses legeres", 81: "Averses moderees",
        82: "Averses violentes", 85: "Averses de neige", 86: "Fortes averses de neige",
        95: "Orage", 96: "Orage avec grele legere", 99: "Orage avec forte grele",
    }
    return wmo.get(code, f"Code {code}")


async def _tool_get_news(args: Dict) -> Dict:
    """Actualites via flux RSS francais (gratuit, pas de cle API)."""
    import httpx
    import xml.etree.ElementTree as ET

    category = args.get("category", "general")
    count = min(max(args.get("count", 5), 3), 10)

    # Sources RSS par categorie
    RSS_FEEDS = {
        "general": [
            ("France Info", "https://www.francetvinfo.fr/titres.rss"),
            ("Le Monde", "https://www.lemonde.fr/rss/une.xml"),
        ],
        "france": [
            ("France Info France", "https://www.francetvinfo.fr/france.rss"),
        ],
        "monde": [
            ("France Info Monde", "https://www.francetvinfo.fr/monde.rss"),
        ],
        "economie": [
            ("France Info Eco", "https://www.francetvinfo.fr/economie.rss"),
        ],
        "sport": [
            ("France Info Sport", "https://www.francetvinfo.fr/sports.rss"),
        ],
        "tech": [
            ("France Info Tech", "https://www.francetvinfo.fr/internet.rss"),
        ],
        "sante": [
            ("France Info Sante", "https://www.francetvinfo.fr/sante.rss"),
        ],
    }

    feeds = RSS_FEEDS.get(category, RSS_FEEDS["general"])
    articles = []

    try:
        async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
            for source_name, url in feeds:
                try:
                    resp = await client.get(url, headers={"User-Agent": "Luna/2.2"})
                    if resp.status_code != 200:
                        continue
                    root = ET.fromstring(resp.text)
                    # RSS 2.0 format
                    for item in root.findall(".//item")[:count]:
                        title = item.findtext("title", "")
                        desc = item.findtext("description", "")
                        link = item.findtext("link", "")
                        pub_date = item.findtext("pubDate", "")
                        if title:
                            # Nettoyer le HTML dans la description
                            import re
                            desc_clean = re.sub(r'<[^>]+>', '', desc)[:200]
                            articles.append({
                                "source": source_name,
                                "titre": title.strip(),
                                "resume": desc_clean.strip(),
                                "lien": link.strip(),
                                "date": pub_date,
                            })
                except Exception:
                    continue

        if not articles:
            return {"status": "error", "message": "Aucune actualite disponible pour le moment."}

        # Trier par date et limiter
        articles = articles[:count]

        return {
            "status": "success",
            "categorie": category,
            "articles": articles,
            "total": len(articles),
        }
    except Exception as e:
        logger.warning(f"News tool error: {e}")
        return {"status": "error", "message": f"Erreur actualites: {e}"}




_SERPER_API_KEY = os.getenv("SERPER_API_KEY", "")


async def _tool_search_web(args: Dict, tenant_id: int = 0) -> Dict:
    """Recherche sur Internet via Serper API (Google Search)."""
    if not _SERPER_API_KEY:
        return {"status": "error", "message": "Service de recherche non configure"}
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager

    query = args.get("query", "").strip()
    location = args.get("location", "").strip()
    if not query:
        return {"status": "error", "message": "Requete de recherche requise"}

    # Enrichir la recherche avec la localisation du souscripteur si disponible
    if not location and mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile:
                city = getattr(profile, "city", "")
                if city:
                    location = city
            # Aussi check la geoloc recente
            geo = _redis_client.client.get(f"luna:{tid}:geolocation") if _redis_client else None
            if geo:
                import json as _json
                geo_data = _json.loads(geo)
                if geo_data.get("city"):
                    location = geo_data["city"]
        except Exception:
            pass

    # Enrichir la requete avec la localisation
    search_query = query
    if location and location.lower() not in query.lower():
        search_query = f"{query} {location}"

    import httpx
    try:
        async with httpx.AsyncClient() as http:
            resp = await http.post(
                "https://google.serper.dev/search",
                headers={
                    "X-API-KEY": _SERPER_API_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "q": search_query,
                    "gl": "fr",
                    "hl": "fr",
                    "num": 5,
                },
                timeout=10,
            )
            resp.raise_for_status()
            data = resp.json()
    except Exception as e:
        logger.error(f"Serper search error: {e}")
        return {"status": "error", "message": "Erreur lors de la recherche. Reessaie."}

    # Formatter les resultats pour Luna
    results = []

    # Knowledge Graph (answer box)
    kg = data.get("knowledgeGraph", {})
    if kg:
        kg_info = kg.get("title", "")
        if kg.get("description"):
            kg_info += f" — {kg['description']}"
        if kg.get("attributes"):
            for k, v in list(kg["attributes"].items())[:5]:
                kg_info += f"\n  {k}: {v}"
        if kg_info:
            results.append(f"[Info directe] {kg_info}")

    # Answer Box
    answer = data.get("answerBox", {})
    if answer:
        ans_text = answer.get("answer") or answer.get("snippet") or answer.get("title", "")
        if ans_text:
            results.append(f"[Reponse] {ans_text}")

    # Organic results
    organic = data.get("organic", [])
    for item in organic[:5]:
        title = item.get("title", "")
        snippet = item.get("snippet", "")
        link = item.get("link", "")
        phone = ""
        # Extraire un numero de telephone s'il y en a un dans le snippet
        import re
        phone_match = re.search(r'(\+?\d[\d\s\-.]{8,15})', snippet)
        if phone_match:
            phone = f" | Tel: {phone_match.group(1).strip()}"
        link_info = f" | {link}" if link else ""
        results.append(f"{title}: {snippet}{phone}{link_info}")

    # Places (local business results)
    places = data.get("places", [])
    for place in places[:3]:
        name = place.get("title", "")
        addr = place.get("address", "")
        rating = place.get("rating", "")
        phone_p = place.get("phoneNumber", "")
        info = f"[Lieu] {name}"
        if addr:
            info += f" — {addr}"
            import urllib.parse
            maps_link = "https://www.google.com/maps/dir/?api=1&destination=" + urllib.parse.quote(f"{name} {addr}")
            info += f" | Itineraire: {maps_link}"
        if rating:
            info += f" | Note: {rating}/5"
        if phone_p:
            info += f" | Tel: {phone_p}"
        results.append(info)

    if not results:
        return {"status": "success", "message": f"Aucun resultat pour '{query}'.", "results": []}

    formatted = "\n".join(results[:8])

    # Log la recherche
    if mgr:
        try:
            mgr.add_note(
                content=f"[Recherche web] {query}\n{formatted[:300]}",
                context="search_web",
                tags=["recherche", "conciergerie"],
            )
        except Exception:
            pass

    return {
        "status": "success",
        "message": f"Resultats pour '{query}':\n{formatted}",
        "results": results[:8],
        "query": search_query,
    }


async def _tool_search_places(args: Dict, tenant_id: int = 0) -> Dict:
    """Recherche de lieux/commerces locaux via Serper Places API."""
    if not _SERPER_API_KEY:
        return {"status": "error", "message": "Service de recherche non configure"}
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager

    query = args.get("query", "").strip()
    location = args.get("location", "").strip()
    category = args.get("category", "").strip()  # restaurant, hotel, taxi, pharmacie...

    if not query:
        return {"status": "error", "message": "Requete de recherche requise"}

    # Enrichir avec la localisation du souscripteur
    if not location and mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile:
                city = getattr(profile, "city", "")
                if city:
                    location = city
            geo = _redis_client.client.get(f"luna:{tid}:geolocation") if _redis_client else None
            if geo:
                import json as _json_geo
                geo_data = _json_geo.loads(geo)
                if geo_data.get("city"):
                    location = geo_data["city"]
        except Exception:
            pass

    search_query = query
    if location and location.lower() not in query.lower():
        search_query = f"{query} {location}"

    import httpx
    try:
        async with httpx.AsyncClient() as http:
            resp = await http.post(
                "https://google.serper.dev/places",
                headers={
                    "X-API-KEY": _SERPER_API_KEY,
                    "Content-Type": "application/json",
                },
                json={
                    "q": search_query,
                    "gl": "fr",
                    "hl": "fr",
                },
                timeout=10,
            )
            resp.raise_for_status()
            data = resp.json()
    except Exception as e:
        logger.error(f"Serper places error: {e}")
        return {"status": "error", "message": "Erreur lors de la recherche de lieux."}

    places = data.get("places", [])
    if not places:
        return {"status": "success", "message": f"Aucun lieu trouve pour '{query}'.", "places": []}

    # Recuperer la geolocation du souscripteur pour les liens transport
    _user_lat, _user_lng = None, None
    if _redis_client and tid:
        try:
            import json as _json_geo2
            _geo_raw = _redis_client.client.get(f"luna:{tid}:geolocation")
            if _geo_raw:
                _geo_d = _json_geo2.loads(_geo_raw)
                _user_lat = _geo_d.get("latitude")
                _user_lng = _geo_d.get("longitude")
        except Exception:
            pass

    results = []
    for p in places[:6]:
        place_info = {
            "name": p.get("title", ""),
            "address": p.get("address", ""),
            "phone": p.get("phoneNumber", ""),
            "rating": p.get("rating", ""),
            "reviews": p.get("ratingCount", 0),
            "type": p.get("type", ""),
            "hours": p.get("openingHours", ""),
            "website": p.get("website", ""),
            "price_level": p.get("priceLevel", ""),
            "latitude": p.get("latitude"),
            "longitude": p.get("longitude"),
        }

        # Calculer la distance a vol d'oiseau si on a les deux positions
        _dist_km = None
        if _user_lat and _user_lng and place_info["latitude"] and place_info["longitude"]:
            import math
            _lat1, _lon1 = math.radians(float(_user_lat)), math.radians(float(_user_lng))
            _lat2, _lon2 = math.radians(float(place_info["latitude"])), math.radians(float(place_info["longitude"]))
            _dlat, _dlon = _lat2 - _lat1, _lon2 - _lon1
            _a = math.sin(_dlat / 2) ** 2 + math.cos(_lat1) * math.cos(_lat2) * math.sin(_dlon / 2) ** 2
            _dist_km = round(6371 * 2 * math.asin(math.sqrt(_a)), 1)
            place_info["distance_km"] = _dist_km

        # Liens Google Maps pour chaque mode de transport
        _dest = place_info["address"] or place_info["name"]
        _origin = f"{_user_lat},{_user_lng}" if _user_lat else ""
        if _origin:
            place_info["directions"] = {
                "driving": f"https://www.google.com/maps/dir/?api=1&origin={_origin}&destination={_dest}&travelmode=driving",
                "transit": f"https://www.google.com/maps/dir/?api=1&origin={_origin}&destination={_dest}&travelmode=transit",
                "walking": f"https://www.google.com/maps/dir/?api=1&origin={_origin}&destination={_dest}&travelmode=walking",
                "bicycling": f"https://www.google.com/maps/dir/?api=1&origin={_origin}&destination={_dest}&travelmode=bicycling",
            }

        # Texte lisible pour Luna (voix)
        desc = f"{place_info['name']}"
        if place_info["rating"]:
            desc += f", note {place_info['rating']} sur 5"
            if place_info["reviews"]:
                desc += f" ({place_info['reviews']} avis)"
        if _dist_km is not None:
            desc += f", a {_dist_km} km"
        if place_info["price_level"]:
            desc += f", {place_info['price_level']}"
        if place_info["address"]:
            desc += f", {place_info['address']}"
        if place_info["phone"]:
            desc += f", tel: {place_info['phone']}"
        if place_info["hours"]:
            desc += f", horaires: {place_info['hours']}"
        place_info["description"] = desc
        results.append(place_info)

    # Formater pour la reponse vocale
    voice_summary = f"J'ai trouve {len(results)} resultats pour '{query}':\n"
    for i, r in enumerate(results, 1):
        voice_summary += f"\n{i}. {r['description']}"

    if mgr:
        try:
            mgr.add_note(
                content=f"[Recherche lieux] {query} — {len(results)} resultats\n{voice_summary[:400]}",
                context="search_places",
                tags=["recherche", "conciergerie", category or "lieu"],
            )
        except Exception:
            pass

    return {
        "status": "success",
        "message": voice_summary,
        "places": results,
        "count": len(results),
    }


async def _tool_get_page_info(args: Dict, tenant_id: int = 0) -> Dict:
    """Extrait les informations cles d'une page web (menu, prix, horaires, disponibilites)."""
    url = args.get("url", "").strip()
    focus = args.get("focus", "").strip()  # ce que Luna cherche: "menu", "prix", "horaires"...

    if not url:
        return {"status": "error", "message": "URL requise"}

    import httpx
    try:
        async with httpx.AsyncClient(follow_redirects=True) as http:
            resp = await http.get(
                url,
                headers={
                    "User-Agent": "Mozilla/5.0 (compatible; Luna/2.2; concierge assistant)",
                    "Accept": "text/html,application/xhtml+xml",
                    "Accept-Language": "fr-FR,fr;q=0.9",
                },
                timeout=12,
            )
            resp.raise_for_status()
            html = resp.text
    except Exception as e:
        logger.error(f"Page fetch error: {e}")
        return {"status": "error", "message": f"Impossible d'acceder a la page: {str(e)[:80]}"}

    # Extraire le texte utile du HTML (sans scripts/styles)
    import re
    # Supprimer scripts, styles, nav, footer
    html_clean = re.sub(r'<(script|style|nav|footer|header|noscript)[^>]*>.*?</\1>', '', html, flags=re.DOTALL | re.IGNORECASE)
    # Supprimer les balises HTML
    text = re.sub(r'<[^>]+>', ' ', html_clean)
    # Supprimer les espaces multiples
    text = re.sub(r'\s+', ' ', text).strip()

    # Tronquer a un max raisonnable pour le contexte
    max_chars = 3000
    if len(text) > max_chars:
        text = text[:max_chars] + "..."

    if not text or len(text) < 20:
        return {"status": "error", "message": "Page vide ou inaccessible."}

    # Extraire les numeros de telephone
    phones = re.findall(r'(?:\+33|0)\s*[1-9](?:[\s.\-]?\d{2}){4}', text)
    phones = list(set(phones))[:3]

    # Extraire les prix (EUR)
    prices = re.findall(r'(\d{1,4}(?:[.,]\d{2})?\s*(?:€|EUR|euros?))', text, re.IGNORECASE)
    prices = list(set(prices))[:10]

    # Extraire les emails
    emails = re.findall(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    emails = list(set(emails))[:3]

    # Extraire les horaires
    hours = re.findall(r'(?:lundi|mardi|mercredi|jeudi|vendredi|samedi|dimanche|lun|mar|mer|jeu|ven|sam|dim)[^.;\n]{5,80}', text, re.IGNORECASE)
    hours = list(set(hours))[:5]

    # Extraire les adresses (patterns francais)
    addresses = re.findall(r'\d{1,4}\s+(?:rue|avenue|boulevard|place|chemin|impasse|allee|passage|cours)\s+[A-Za-zÀ-ÿ\s\-]{3,50}', text, re.IGNORECASE)
    addresses = list(set(addresses))[:3]

    result = {
        "status": "success",
        "content": text[:2000],
        "phones": phones,
        "prices": prices,
        "emails": emails,
        "hours": hours,
        "addresses": addresses,
        "url": url,
    }

    # Si un focus est specifie et que OpenAI est dispo, resumer le contenu
    if focus and openai_client:
        try:
            _summary_resp = openai_client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {"role": "system", "content": "Tu es un assistant qui extrait des informations precises d'un texte de page web. Reponds en francais, de facon concise et structuree."},
                    {"role": "user", "content": f"Voici le contenu de la page {url}.\nJe cherche : {focus}\n\nTexte:\n{text[:2500]}\n\nExtrais uniquement les informations pertinentes pour '{focus}'. Sois precis et concis."},
                ],
                max_tokens=500,
                temperature=0.1,
            )
            _summary = _summary_resp.choices[0].message.content.strip()
            result["summary"] = _summary
            result["message"] = f"Informations sur '{focus}' extraites de {url}."
        except Exception as e:
            logger.warning(f"Page summary error: {e}")
            result["message"] = f"Informations extraites de {url} (focus: {focus}). {len(phones)} telephone(s), {len(prices)} prix trouves."
    elif focus:
        result["focus"] = focus
        result["message"] = f"Informations extraites de {url} (focus: {focus}). {len(phones)} telephone(s), {len(prices)} prix trouves."
    else:
        result["message"] = f"Contenu extrait de {url}. {len(phones)} telephone(s), {len(prices)} prix trouves."

    return result


_CONCIERGE_COMMISSION_RATE = float(os.getenv("CONCIERGE_COMMISSION_RATE", "0.10"))  # 10% par defaut


async def _tool_request_payment(args: Dict, tenant_id: int = 0) -> Dict:
    """Cree un PaymentIntent Stripe pour un achat de conciergerie avec commission fondateur."""
    tid = tenant_id or TENANT_ID
    mgr = _get_tenant_manager(tid) if tid else _memory_manager

    base_amount_cents = args.get("amount_cents", 0)
    description = args.get("description", "")
    merchant_name = args.get("merchant_name", "Conciergerie Luna")

    if not base_amount_cents or base_amount_cents <= 0:
        return {"status": "error", "message": "Montant invalide"}
    if not description:
        return {"status": "error", "message": "Description de l'achat requise"}

    # Calcul commission fondateur
    commission_cents = int(round(base_amount_cents * _CONCIERGE_COMMISSION_RATE))
    total_amount_cents = base_amount_cents + commission_cents

    # Verifier le plafond du souscripteur (sur le total TTC)
    max_budget = 0
    sub_name = _SUBSCRIBER_NAME
    if mgr:
        try:
            profile = mgr.get_subscriber_profile()
            if profile:
                try:
                    max_budget = int(float(getattr(profile, "max_budget", 0) or 0))
                except (ValueError, TypeError):
                    max_budget = 0
                if profile.first_name:
                    sub_name = profile.first_name
        except Exception:
            pass

    total_eur = total_amount_cents / 100
    base_eur = base_amount_cents / 100
    commission_eur = commission_cents / 100

    # Suivi cumule des depenses du mois (Redis)
    already_spent_cents = 0
    spending_key = ""
    if _redis_client and tid:
        import datetime as _dt
        _month_key = _dt.datetime.now().strftime("%Y-%m")
        spending_key = f"{_redis_client.prefix}:{tid}:concierge:spending:{_month_key}"
        try:
            already_spent_cents = int(_redis_client.client.get(spending_key) or 0)
        except Exception:
            already_spent_cents = 0

    already_spent_eur = already_spent_cents / 100
    projected_total_eur = already_spent_eur + total_eur

    # BLOCAGE STRICT si le budget mensuel est depasse (pas de "demande-lui")
    if max_budget > 0 and total_eur > max_budget:
        return {
            "status": "error",
            "message": (
                f"REFUSE : ce paiement de {total_eur:.2f} EUR depasse le plafond "
                f"de {max_budget} EUR par transaction fixe par {sub_name}. "
                f"Je ne peux pas proceder."
            )
        }
    if max_budget > 0 and projected_total_eur > max_budget:
        remaining = max(0, max_budget - already_spent_eur)
        return {
            "status": "error",
            "message": (
                f"REFUSE : {sub_name} a deja depense {already_spent_eur:.2f} EUR ce mois-ci "
                f"(plafond: {max_budget} EUR). Ce paiement de {total_eur:.2f} EUR "
                f"porterait le total a {projected_total_eur:.2f} EUR. "
                f"Budget restant: {remaining:.2f} EUR."
            )
        }

    # Recuperer le stripe_customer_id du souscripteur (par tenant_id, pas par PROPRIO_EMAIL)
    import stripe as _stripe
    _stripe_secret = os.getenv("STRIPE_SECRET_KEY", "") or os.getenv("STRIPE_API_KEY", "")
    if not _stripe_secret:
        return {"status": "error", "message": "Service de paiement non configure"}
    _stripe.api_key = _stripe_secret

    # Chercher le customer Stripe via Redis (scan_iter, pas KEYS *)
    stripe_customer_id = None
    if _redis_client:
        try:
            auth = _redis_client.get_auth_by_tenant_id(tid)
            if auth:
                stripe_customer_id = auth.get("stripe_customer_id")
        except Exception as e:
            logger.warning(f"Stripe customer lookup error: {e}")

    if not stripe_customer_id:
        return {
            "status": "error",
            "message": f"Aucune methode de paiement enregistree pour {sub_name}. Il doit d'abord ajouter une carte dans l'onglet Parametres."
        }

    try:
        # Recuperer la methode de paiement par defaut
        customer = _stripe.Customer.retrieve(stripe_customer_id)
        default_pm = customer.get("invoice_settings", {}).get("default_payment_method")
        if not default_pm:
            pms = _stripe.PaymentMethod.list(customer=stripe_customer_id, type="card", limit=1)
            if pms.data:
                default_pm = pms.data[0].id
            else:
                return {
                    "status": "error",
                    "message": f"Aucune carte bancaire enregistree pour {sub_name}. Il doit d'abord ajouter une carte."
                }

        # Creer le PaymentIntent avec commission dans les metadata
        intent = _stripe.PaymentIntent.create(
            amount=total_amount_cents,
            currency="eur",
            customer=stripe_customer_id,
            payment_method=default_pm,
            description=f"[Luna Conciergerie] {description}",
            metadata={
                "tenant_id": str(tid),
                "merchant": merchant_name,
                "type": "concierge",
                "base_amount_cents": str(base_amount_cents),
                "commission_cents": str(commission_cents),
                "commission_rate": str(_CONCIERGE_COMMISSION_RATE),
            },
            # Ne PAS confirmer automatiquement — le souscripteur doit valider
            confirm=False,
        )

        # Incrementer le compteur de depenses mensuelles dans Redis
        if _redis_client and spending_key:
            try:
                _redis_client.client.incrby(spending_key, total_amount_cents)
                # Expire apres 35 jours (nettoyage auto)
                _redis_client.client.expire(spending_key, 35 * 86400)
            except Exception as _e:
                logger.warning(f"Failed to track spending: {_e}")

        # Log l'intent
        if mgr:
            try:
                _budget_info = ""
                if max_budget > 0:
                    _new_spent = (already_spent_cents + total_amount_cents) / 100
                    _remaining = max(0, max_budget - _new_spent)
                    _budget_info = f"\nBudget: {_new_spent:.2f}/{max_budget} EUR utilises, reste {_remaining:.2f} EUR"
                mgr.add_note(
                    content=(
                        f"[Paiement conciergerie] {total_eur:.2f} EUR TTC — {description}\n"
                        f"Commercant: {merchant_name} | Service: {base_eur:.2f} EUR + {commission_eur:.2f} EUR frais\n"
                        f"Statut: en attente de confirmation | Ref: {intent.id}"
                        f"{_budget_info}"
                    ),
                    context="payment",
                    tags=["paiement", "conciergerie", merchant_name],
                )
                mgr.log_event(
                    category="action",
                    description=f"Paiement conciergerie cree: {total_eur:.2f} EUR pour {merchant_name} (commission: {commission_eur:.2f} EUR)",
                    reasoning=f"Luna prepare un paiement de {total_eur:.2f} EUR: {description}",
                    source="tool_call",
                )
            except Exception:
                pass

        # Ajouter le budget restant dans la reponse pour que Luna informe le souscripteur
        budget_info = {}
        if max_budget > 0:
            _new_total = (already_spent_cents + total_amount_cents) / 100
            budget_info = {
                "budget_max": max_budget,
                "spent_this_month": _new_total,
                "remaining": max(0, max_budget - _new_total),
            }

        return {
            "status": "success",
            "message": (
                f"Paiement de {total_eur:.2f} EUR prepare pour {merchant_name} "
                f"(dont {commission_eur:.2f} EUR de frais de service). "
                f"{sub_name} doit confirmer le paiement dans l'application pour finaliser. "
                f"Reference: {intent.id[-8:]}"
            ),
            "payment_intent_id": intent.id,
            "amount_eur": total_eur,
            "base_amount_eur": base_eur,
            "commission_eur": commission_eur,
            **({"budget": budget_info} if budget_info else {}),
        }

    except _stripe.error.StripeError as e:
        logger.error(f"Stripe payment error: {e}")
        return {"status": "error", "message": f"Erreur de paiement: {str(e)[:100]}"}
    except Exception as e:
        logger.error(f"Payment error: {e}")
        return {"status": "error", "message": "Erreur lors de la creation du paiement. Reessaie."}


# =============================================================================
# TOOLS RESERVATION (Vols, Hotels, Restaurants)
# =============================================================================

async def _tool_search_flights(args: Dict, tenant_id: int = 0) -> Dict:
    """Recherche de vols via Duffel (prioritaire) → Amadeus → fallback web."""
    origin = args.get("origin", "")
    destination = args.get("destination", "")
    departure_date = args.get("departure_date", "")
    return_date = args.get("return_date", "")
    passengers = int(args.get("passengers", 1))
    travel_class = args.get("travel_class", "ECONOMY")

    if not origin or not destination or not departure_date:
        return {"status": "error", "message": "Il me faut la ville de depart, la destination et la date."}

    # 1) Duffel (reservation directe possible)
    if duffel_client and duffel_client.is_configured:
        cabin = {"ECONOMY": "economy", "PREMIUM_ECONOMY": "premium_economy",
                 "BUSINESS": "business", "FIRST": "first"}.get(travel_class, "economy")
        success, data = await duffel_client.search_flights(
            origin=origin, destination=destination,
            departure_date=departure_date, return_date=return_date or None,
            passengers=passengers, cabin_class=cabin,
        )
        if success and data.get("flights"):
            logger.info(f"FLIGHT_SEARCH [Duffel] {origin}→{destination} {departure_date}: {data['count']} resultats")
            return {"status": "success", **data}

    # 2) Amadeus
    if amadeus_client and amadeus_client.is_configured:
        success, data = await amadeus_client.search_flights(
            origin=origin, destination=destination,
            departure_date=departure_date, return_date=return_date or None,
            adults=passengers, travel_class=travel_class,
        )
        if success and data.get("flights"):
            logger.info(f"FLIGHT_SEARCH [Amadeus] {origin}→{destination} {departure_date}: {data['count']} resultats")
            return {"status": "success", **data}

    # 3) Fallback: recherche web + liens comparateurs
    import urllib.parse
    web_result = await _tool_search_web(
        {"query": f"vol {origin} {destination} {departure_date} prix horaires billet"},
        tenant_id=tenant_id,
    )
    _skyscanner_q = urllib.parse.quote(f"{origin} {destination}")
    booking_links = [
        f"Skyscanner: https://www.skyscanner.fr/transport/vols/{origin.lower()}/{destination.lower()}/{departure_date.replace('-', '')}/",
        f"Google Flights: https://www.google.com/travel/flights?q=vol+{_skyscanner_q}+{departure_date}",
    ]
    if web_result.get("status") == "success":
        web_result["liens_reservation"] = booking_links
        web_result["conseil"] = "Voici les resultats. Pour reserver, le souscripteur peut cliquer sur les liens ci-dessus."
    return web_result


async def _tool_search_hotels(args: Dict, tenant_id: int = 0) -> Dict:
    """Recherche d'hotels via Duffel (prioritaire) → Amadeus → fallback web."""
    city = args.get("city", "")
    check_in = args.get("check_in", "")
    check_out = args.get("check_out", "")
    guests = int(args.get("guests", 1))

    if not city or not check_in or not check_out:
        return {"status": "error", "message": "Il me faut la ville, la date d'arrivee et la date de depart."}

    # 1) Duffel Stays (reservation directe possible)
    if duffel_client and duffel_client.is_configured:
        success, data = await duffel_client.search_hotels(
            city=city, check_in=check_in, check_out=check_out, guests=guests,
        )
        if success and data.get("hotels"):
            logger.info(f"HOTEL_SEARCH [Duffel] {city} {check_in}-{check_out}: {data['count']} resultats")
            return {"status": "success", **data}

    # 2) Amadeus
    if amadeus_client and amadeus_client.is_configured:
        stars = args.get("stars")
        price_range = args.get("price_range")
        success, data = await amadeus_client.search_hotels(
            city=city, check_in=check_in, check_out=check_out,
            adults=guests, stars=stars, price_range=price_range,
        )
        if success and data.get("hotels"):
            logger.info(f"HOTEL_SEARCH [Amadeus] {city} {check_in}-{check_out}: {data['count']} resultats")
            return {"status": "success", **data}

    # 3) Fallback: recherche web + liens comparateurs
    import urllib.parse
    _city_q = urllib.parse.quote(city)
    web_result = await _tool_search_web(
        {"query": f"hotel {city} {check_in} {check_out} prix chambre disponibilite"},
        tenant_id=tenant_id,
    )
    booking_links = [
        f"Booking.com: https://www.booking.com/searchresults.fr.html?ss={_city_q}&checkin={check_in}&checkout={check_out}",
        f"Hotels.com: https://fr.hotels.com/search.do?q-destination={_city_q}&q-check-in={check_in}&q-check-out={check_out}",
    ]
    if web_result.get("status") == "success":
        web_result["liens_reservation"] = booking_links
        web_result["conseil"] = "Voici les hotels trouves. Pour reserver, le souscripteur peut utiliser les liens ci-dessus."
    return web_result


async def _tool_book_restaurant(args: Dict, tenant_id: int = 0) -> Dict:
    """Recherche et reservation de restaurant via TheFork ou fallback search_places."""
    city = args.get("city", "")
    date = args.get("date", "")
    time_str = args.get("time", "20:00")
    party_size = args.get("party_size", 2)
    cuisine = args.get("cuisine")

    if not city or not date:
        return {"status": "error", "message": "Il me faut la ville et la date pour chercher un restaurant."}

    # Tente TheFork d'abord
    if thefork_client and thefork_client.is_configured:
        success, data = await thefork_client.search_restaurants(
            city=city,
            date=date,
            time=time_str,
            party_size=party_size,
            cuisine=cuisine,
        )
        if success:
            logger.info(f"RESTAURANT_SEARCH {city} {date} {time_str}: {data.get('count', 0)} resultats (TheFork)")
            return {"status": "success", "source": "thefork", **data}

    # Fallback: search_places
    query = f"restaurant {cuisine + ' ' if cuisine else ''}{city}"
    result = await _tool_search_places(
        {"query": query, "location": city, "category": "restaurant"},
        tenant_id=tenant_id,
    )
    if result.get("status") == "success":
        result["booking_method"] = "call_contact"
        result["message"] = (
            f"Voici les restaurants trouves. Pour reserver, je peux appeler le restaurant "
            f"de ton choix avec call_contact pour reserver une table pour {party_size} "
            f"personne{'s' if party_size > 1 else ''} le {date} a {time_str}."
        )
    return result


# =========================================================================
# SECRETARY TOOLS — Documents, Budget, Rappels
# =========================================================================

def _get_secretary_ops(tid: int):
    """Retourne SecretaryRedisOps pour un tenant."""
    if not _redis_client:
        return None
    try:
        from core.secretary.redis_ops import SecretaryRedisOps
        return SecretaryRedisOps(_redis_client, tid)
    except ImportError:
        return None


def _tool_secretary_summary(tid: int) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    summary = sops.get_documents_summary()
    summary["status"] = "success"
    return summary


def _tool_secretary_budget(tid: int) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    from core.secretary.scanner import generate_budget_suggestions
    analysis = sops.get_budget_analysis()
    suggestions = generate_budget_suggestions(analysis)
    return {"status": "success", "analysis": analysis, "suggestions": suggestions}


def _tool_secretary_afford(tid: int, args: Dict) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    try:
        amount = float(args.get("amount", 0))
    except (ValueError, TypeError):
        return {"status": "error", "message": "Montant invalide"}
    label = args.get("label", "cette depense")

    analysis = sops.get_budget_analysis()
    reste = analysis.get("reste_a_vivre", 0)
    prevu_fin = analysis.get("solde_prevu_fin_mois", 0)

    can_afford = (reste - amount) > 0
    safe = (prevu_fin - amount) > 50

    if can_afford and safe:
        verdict = "oui"
        message = f"Oui, tu peux te permettre {label} ({amount:.0f}€). Il te restera environ {prevu_fin - amount:.0f}€ en fin de mois."
    elif can_afford and not safe:
        verdict = "risque"
        message = f"C'est possible mais risque. Apres {label} ({amount:.0f}€), ta fin de mois sera serree ({prevu_fin - amount:.0f}€)."
    else:
        deficit = amount - reste
        verdict = "non"
        message = f"Ce n'est pas raisonnable. {label} ({amount:.0f}€) depasse ton budget de {deficit:.0f}€."

    return {"status": "success", "verdict": verdict, "message": message, "reste_apres": round(reste - amount, 2)}


def _tool_secretary_add_expense(tid: int, args: Dict) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    try:
        montant = float(args.get("montant", 0))
    except (ValueError, TypeError):
        return {"status": "error", "message": "Montant invalide"}

    import html as _html
    entry = {
        "montant": montant,
        "direction": args.get("direction", "depense"),
        "categorie": _html.escape(args.get("categorie", "autre")),
        "label": _html.escape(args.get("label", "")[:100]),
    }
    entry_id = sops.add_budget_entry(entry)
    direction_label = "revenu" if entry["direction"] == "revenu" else "depense"
    return {"status": "success", "message": f"{direction_label.capitalize()} de {montant:.0f}€ enregistre(e) ({entry['label']})."}


def _tool_secretary_reminders(tid: int) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    upcoming = sops.get_upcoming_reminders()
    overdue = sops.get_overdue_reminders()
    return {
        "status": "success",
        "upcoming": upcoming,
        "overdue": overdue,
        "message": f"{len(overdue)} en retard, {len(upcoming)} a venir.",
    }


def _tool_secretary_add_reminder(tid: int, args: Dict) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    import html as _html
    title = _html.escape(args.get("title", "").strip()[:100])
    if not title:
        return {"status": "error", "message": "Titre requis"}
    reminder = {
        "title": title,
        "description": _html.escape(args.get("description", "")[:300]),
        "due_date": args.get("due_date", ""),
        "due_time": args.get("due_time", ""),
        "type": "luna",
    }
    sops.add_reminder(reminder)
    due_str = ""
    if reminder["due_date"]:
        due_str = f" pour le {reminder['due_date']}"
        if reminder["due_time"]:
            due_str += f" à {reminder['due_time']}"
    return {"status": "success", "message": f"Rappel cree : {title}{due_str}"}


def _tool_secretary_search(tid: int, args: Dict) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    query = args.get("query", "")
    if not query:
        return {"status": "error", "message": "Mot-cle requis"}
    docs = sops.search_documents(query)
    if not docs:
        return {"status": "success", "message": f"Aucun document trouve pour '{query}'.", "documents": []}
    return {"status": "success", "documents": docs, "message": f"{len(docs)} document(s) trouve(s) pour '{query}'."}


def _tool_secretary_folders(tid: int) -> Dict:
    sops = _get_secretary_ops(tid)
    if not sops:
        return {"status": "error", "message": "Module secretaire non disponible"}
    folders = sops.get_folders()
    if not folders:
        return {"status": "success", "message": "Aucun dossier pour l'instant. Scanne un document pour commencer.", "folders": []}
    lines = []
    for f in folders:
        indent = "  " if f.get("parent") else ""
        lines.append(f"{indent}{f['path']} ({f['count']} doc{'s' if f['count'] > 1 else ''})")
    return {"status": "success", "folders": folders, "message": "Tes dossiers :\n" + "\n".join(lines)}


# =========================================================================
# VOICE TOOLS — Social & Perception (Mode Jarvis)
# =========================================================================

async def _voice_tool_get_friends_online(tid: int) -> Dict:
    """Liste les amis Luna et indique lesquels sont en ligne."""
    if not _redis_client:
        return {"status": "error", "message": "Service social non disponible"}
    try:
        from core.social.redis_ops import SocialRedisOps
        sops = SocialRedisOps(_redis_client)
        friend_tids = sops.get_friends(tid)
        if not friend_tids:
            return {
                "status": "success",
                "friends": [],
                "count": 0,
                "message": "Tu n'as pas encore d'amis dans le Monde Luna. Partage ton code ami pour en ajouter !",
            }
        online_set = sops.get_online_tids()
        friends = []
        for f_tid in friend_tids:
            profile = sops.get_social_profile(f_tid)
            if profile:
                friends.append({
                    "nickname": profile.get("nickname", f"User{f_tid}"),
                    "is_online": f_tid in online_set,
                    "level": profile.get("level", "1"),
                })
        friends.sort(key=lambda f: (not f["is_online"], f["nickname"]))
        online_count = sum(1 for f in friends if f["is_online"])
        return {
            "status": "success",
            "friends": friends,
            "count": len(friends),
            "online_count": online_count,
            "message": f"{online_count} ami(s) en ligne sur {len(friends)} au total.",
        }
    except Exception as e:
        logger.error(f"Voice tool get_friends_online error: {e}")
        return {"status": "error", "message": "Impossible de charger les amis"}


async def _voice_tool_send_dm(tid: int, args: Dict) -> Dict:
    """Envoie un DM a un ami Luna par son pseudo."""
    friend_name = args.get("friend_name", "").strip()
    message = args.get("message", "").strip()
    if not friend_name or not message:
        return {"status": "error", "message": "Il me faut le nom de l'ami et le message."}
    if len(message) > 500:
        return {"status": "error", "message": "Message trop long (max 500 caracteres)."}
    if not _redis_client:
        return {"status": "error", "message": "Service social non disponible"}
    try:
        from core.social.redis_ops import SocialRedisOps
        sops = SocialRedisOps(_redis_client)
        friend_tids = sops.get_friends(tid)
        # Find matching friend by nickname
        target_tid = None
        target_nickname = None
        for f_tid in friend_tids:
            profile = sops.get_social_profile(f_tid)
            if profile:
                nickname = profile.get("nickname", "")
                if nickname.lower() == friend_name.lower() or friend_name.lower() in nickname.lower():
                    target_tid = f_tid
                    target_nickname = nickname
                    break
        if not target_tid:
            return {
                "status": "error",
                "message": f"Aucun ami trouve avec le nom '{friend_name}'. Verifie le pseudo exact.",
            }
        # Create or get DM room
        room_id = sops.create_dm_room(tid, target_tid)
        if not room_id:
            return {"status": "error", "message": "Impossible de creer la conversation."}
        sops.add_dm_message(room_id, tid, message)
        logger.info(f"Voice DM sent: {tid} -> {target_tid} ({target_nickname})")
        return {
            "status": "success",
            "message": f"Message envoye a {target_nickname} dans le Monde Luna.",
            "recipient": target_nickname,
        }
    except Exception as e:
        logger.error(f"Voice tool send_dm error: {e}")
        return {"status": "error", "message": "Erreur lors de l'envoi du message"}


async def _voice_tool_get_friend_code(tid: int) -> Dict:
    """Retourne le code ami du souscripteur."""
    if not _redis_client:
        return {"status": "error", "message": "Service social non disponible"}
    try:
        from core.social.redis_ops import SocialRedisOps
        sops = SocialRedisOps(_redis_client)
        code = sops.get_friend_code(tid)
        return {
            "status": "success",
            "friend_code": code,
            "message": f"Ton code ami est {code}. Epelle-le lettre par lettre pour que la personne puisse l'entrer.",
        }
    except Exception as e:
        logger.error(f"Voice tool get_friend_code error: {e}")
        return {"status": "error", "message": "Impossible de recuperer le code ami"}


async def _voice_tool_look_around(tid: int) -> Dict:
    """Regarde ce que la camera du souscripteur voit actuellement."""
    if not _memory_manager:
        return {"status": "error", "message": "Perception non disponible."}
    try:
        if not _memory_manager.is_perception_enabled():
            return {
                "status": "error",
                "message": "La camera n'est pas activee. Demande au souscripteur d'activer la camera dans les reglages.",
            }
        scene = _memory_manager.get_perception_state()
        if not scene:
            return {
                "status": "success",
                "message": "La camera est activee mais je n'ai pas encore d'image recente. Attends quelques secondes.",
            }
        result = {
            "status": "success",
            "scene_description": scene.scene_description,
            "persons_present": scene.persons_present,
            "primary_posture": scene.primary_posture,
            "objects_visible": scene.objects_visible,
        }
        if scene.abnormalities:
            result["observations"] = [
                {"severity": a["severity"], "description": a["description"]}
                for a in scene.abnormalities
            ]
        return result
    except Exception as e:
        logger.error(f"Voice tool look_around error: {e}")
        return {"status": "error", "message": "Impossible d'analyser la scene."}


async def _handle_tavus_transcription(body: Dict) -> Dict:
    """Sauvegarde la transcription d'un appel Tavus."""
    conversation_id = body.get("conversation_id", "")
    transcript = body.get("transcript", "") or body.get("data", {}).get("transcript", "")

    # Résoudre le tenant depuis la conversation active
    tid = TENANT_ID
    if tavus_client and conversation_id:
        conv = tavus_client._active_conversations.get(conversation_id)
        if conv:
            tid = conv.tenant_id
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"status": "ok"}

    if transcript:
        try:
            mgr.add_note(
                content=f"Transcription visio: {transcript[:2000]}",
                context="visio_transcription",
                tags=["visio", "transcription", conversation_id],
            )
            logger.info(f"Tavus transcription saved for {conversation_id}")
        except Exception as e:
            logger.warning(f"Failed to save transcription: {e}")

    return {"status": "ok"}


# ── Notes intelligentes : filtres + déduplication ──────────────────────────

# Patterns à ignorer (salutations, remplissages, météo banale)
_NOTE_SKIP_RE = re.compile(
    r"^\s*(?:"
    r"bonjour|bonsoir|salut|coucou|hello|hi\b|au revoir|bonne journee|bonne soiree|bonne nuit|"
    r"a bientot|a tout a l'heure|a tout heure|merci|de rien|s'il vous plait|sil vous plait|"
    r"euh+|bah+|ben+|hm+|ah+|oh+|ouais|oui|non|ok|d'accord|d'ac|"
    r"ca va|ça va|je vais bien|tout va bien|nickel|impeccable|parfait|"
    r"il fait beau|beau temps|soleil|nuageux|il pleut|meteorologie|la meteo|"
    r"c'est bien|tres bien|super|excellent|formidable"
    r")[\s!?.]*$",
    re.IGNORECASE | re.UNICODE,
)

# Catégories + mots-clés pour classification heuristique
_NOTE_CATEGORIES = {
    "RDV":   ["rendez-vous", "rdv", "docteur", "dentiste", "medecin", "hopital", "clinique",
               "appointment", "demain", "lundi", "mardi", "mercredi", "jeudi", "vendredi",
               "samedi", "dimanche", "semaine prochaine", "consultation", "examen"],
    "MED":   ["medicament", "médicament", "comprime", "comprimé", "pilule", "cachet",
               "ordonnance", "pharmacie", "dose", "prendre", "oublier", "oublie",
               "tension", "glycemie", "glycémie", "insuline", "traitement"],
    "SYM":   ["douleur", "mal", "souffre", "fatigue", "fatigué", "essoufle", "essoufflé",
               "vertige", "tete", "tête", "ventre", "dos", "jambe", "bras", "chute",
               "tombe", "tombé", "blessure", "fievre", "fièvre", "nausee", "nausée"],
    "INST":  ["rappelle", "n'oublie pas", "pense a", "pense à", "dis-moi", "previens",
               "préviens", "alerte", "contacte", "appelle", "envoie", "fais", "il faut que"],
    "FAM":   ["fils", "fille", "mari", "femme", "enfant", "petit-fils", "petite-fille",
               "famille", "frere", "sœur", "soeur", "neveu", "niece", "parent", "proche"],
    "ALERT": ["urgence", "aide", "au secours", "danger", "chute", "inconscient", "douleur forte",
               "ambulance", "pompier", "appel urgence"],
}

_NOTE_MIN_LENGTH = 25  # caractères minimum pour déclencher une note
_NOTE_SIM_THRESHOLD = 0.70  # similarité au-delà de laquelle c'est un doublon

def _note_similarity(a: str, b: str) -> float:
    """Similarité simple mot-à-mot (Jaccard sur tokens)."""
    wa = set(re.findall(r"\w+", a.lower()))
    wb = set(re.findall(r"\w+", b.lower()))
    if not wa or not wb:
        return 0.0
    return len(wa & wb) / max(len(wa), len(wb))

def _classify_note(text: str) -> str | None:
    """
    Classe le texte dans une catégorie si score suffisant.
    Retourne la catégorie (ex: 'RDV') ou None si le texte n'est pas digne d'une note.
    """
    text_l = text.lower()
    scores: dict[str, int] = {}
    for cat, keywords in _NOTE_CATEGORIES.items():
        scores[cat] = sum(1 for kw in keywords if kw in text_l)

    best_cat, best_score = max(scores.items(), key=lambda x: x[1])
    if best_score >= 1:
        return best_cat

    # Fallback : verbe d'action + phrase assez longue = INFO mémorable
    action_verbs = ["prendre", "aller", "faire", "acheter", "appeler", "venir", "passer", "voir"]
    if any(v in text_l for v in action_verbs) and len(text.strip()) >= _NOTE_MIN_LENGTH:
        return "INFO"

    return None

def _should_skip_note(text: str, mgr) -> tuple[bool, str]:
    """
    Retourne (doit_ignorer, raison).
    Vérifie : longueur min, pattern skip, classification, déduplication.
    """
    t = text.strip()

    if len(t) < _NOTE_MIN_LENGTH:
        return True, "trop_court"

    if _NOTE_SKIP_RE.match(t):
        return True, "salutation_ou_remplissage"

    category = _classify_note(t)
    if category is None:
        return True, "non_classifiable"

    # Déduplication : comparer avec les 5 dernières notes (2 min)
    if mgr:
        try:
            recent = mgr.get_recent_notes(limit=5, minutes=2)
            for note_data in recent:
                prev_content = note_data.get("content", "")
                if _note_similarity(t, prev_content) >= _NOTE_SIM_THRESHOLD:
                    return True, "doublon"
        except Exception:
            pass

    return False, category


async def _handle_tavus_utterance(body: Dict) -> Dict:
    """Utterances Tavus : filtrage intelligent + note auto + checkin."""
    data = body.get("data", {})
    speaker = data.get("speaker", "")
    text = data.get("text", "")
    conversation_id = body.get("conversation_id", "")
    if not text:
        return {"status": "ok"}

    # Résoudre le tenant
    tid = TENANT_ID
    if tavus_client and conversation_id:
        conv = tavus_client._active_conversations.get(conversation_id)
        if conv:
            tid = conv.tenant_id
    mgr = _get_tenant_manager(tid) if tid else _memory_manager

    # Note auto filtrée (uniquement ce que dit le souscripteur)
    if speaker == "user" and mgr and mgr.is_auto_note_enabled():
        skip, reason = _should_skip_note(text, mgr)
        if not skip:
            try:
                mgr.add_note(
                    content=f"[{reason}] {text[:400]}",
                    context="visio_utterance",
                    tags=["visio", "tavus", reason, conversation_id],
                )
                logger.debug(f"Note visio [{reason}] tenant={tid}: {text[:60]}")
            except Exception:
                pass
        else:
            logger.debug(f"Note ignorée ({reason}) tenant={tid}: {text[:60]}")

    # Vérification checkin si le souscripteur parle
    if speaker == "user":
        await _process_checkin_response(tid, conversation_id, text)

    return {"status": "ok"}


async def _handle_simli_utterance(tid: int, user_text: str, session_id: str = "") -> None:
    """Même logique de filtrage que Tavus, pour les appels audio Simli."""
    if not user_text or not user_text.strip():
        return
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr or not mgr.is_auto_note_enabled():
        return
    skip, reason = _should_skip_note(user_text, mgr)
    if not skip:
        try:
            mgr.add_note(
                content=f"[{reason}] {user_text[:400]}",
                context="audio_utterance",
                tags=["audio", "simli", reason, session_id],
            )
            logger.debug(f"Note audio Simli [{reason}] tenant={tid}: {user_text[:60]}")
        except Exception:
            pass


async def _process_checkin_response(tenant_id: int, conversation_id: str, text: str) -> None:
    """Évalue la réponse du souscripteur à un checkin de sécurité Raven."""
    if not _redis_client:
        return
    redis_key = f"luna:{tenant_id}:checkin_pending:{conversation_id or 'default'}"
    checkin_raw = _redis_client.client.get(redis_key)
    if not checkin_raw:
        return
    try:
        checkin = json.loads(checkin_raw)
    except Exception:
        return
    if checkin.get("status") != "pending":
        return

    text_lower = text.lower().strip()
    POSITIVE = {"oui", "ça va", "ca va", "tout va bien", "bien", "ok", "super",
                "merci", "tranquille", "pas de souci", "ras", "nickel",
                "je vais bien", "pas de problème", "np"}
    NEGATIVE = {"non", "pas bien", "mal", "aide", "au secours", "urgence",
                "chute", "tombe", "douleur", "souffre", "peux pas",
                "impossible", "bloqué", "panique"}

    is_positive = any(kw in text_lower for kw in POSITIVE)
    is_negative = any(kw in text_lower for kw in NEGATIVE)

    if is_positive and not is_negative:
        checkin["status"] = "confirmed"
        checkin["response"] = text[:200]
        checkin["confirmed_at"] = datetime.utcnow().isoformat()
        _redis_client.client.setex(redis_key, 300, json.dumps(checkin))
        logger.info(f"Checkin CONFIRME tenant={tenant_id} conv={conversation_id}")

    elif is_negative:
        checkin["status"] = "alerted"
        checkin["response"] = text[:200]
        checkin["alerted_at"] = datetime.utcnow().isoformat()
        _redis_client.client.setex(redis_key, 300, json.dumps(checkin))
        reason = checkin.get("observation", "detresse confirmee par le souscripteur")
        try:
            await _tool_alert_contacts({"reason": reason}, tenant_id)
            logger.warning(f"ALERTE IMMEDIATE (reponse negative) tenant={tenant_id}")
        except Exception as e:
            logger.exception(f"Echec alerte immediate checkin tenant={tenant_id}: {e}")


# =========================================================================
# DOCUMENTS ENDPOINTS
# =========================================================================

class DocumentRequest(BaseModel):
    doc_type: str = Field(..., max_length=50)
    subject: str = Field(default="", max_length=500)
    details: str = Field(default="", max_length=2000)
    recipient: str = Field(default="", max_length=200)


@app.post("/api/documents/generate")
async def generate_document(req: DocumentRequest, request: Request):
    """Genere un document DOCX via GPT + python-docx"""
    if not _doc_generator:
        return JSONResponse(status_code=503, content={"error": "Generateur non disponible"})

    tid = getattr(request.state, "tenant_id", 1)
    mgr = _get_tenant_manager(tid)

    profile_dict = {}
    if mgr:
        profile = mgr.get_subscriber_profile()
        if profile:
            profile_dict = profile.model_dump()

    # Special case: fiche sante (from profile directly)
    if req.doc_type == "fiche_sante" and profile_dict:
        filename = _doc_generator.generate_health_sheet(profile_dict)
        return {
            "success": True,
            "filename": filename,
            "download_url": f"/api/documents/download/{filename}",
            "type": "fiche_sante",
        }

    # Special case: export notes
    if req.doc_type == "export_notes" and mgr:
        notes = mgr.list_notes(limit=200)
        note_dicts = [{"content": n.content, "context": n.context, "created_at": n.created_at.isoformat() if n.created_at else "", "tags": n.tags} for n in notes]
        filename = _doc_generator.generate_notes_export(note_dicts)
        return {
            "success": True,
            "filename": filename,
            "download_url": f"/api/documents/download/{filename}",
            "type": "export_notes",
        }

    # General case: GPT generates the body
    prompt = f"Redige un {req.doc_type} professionnel en francais.\nObjet: {req.subject}\n"
    if req.details:
        prompt += f"Details: {req.details}\n"
    if req.recipient:
        prompt += f"Destinataire: {req.recipient}\n"
    if profile_dict.get("first_name"):
        prompt += f"Expediteur: {profile_dict.get('first_name', '')} {profile_dict.get('last_name', '')}, {profile_dict.get('address', '')}, {profile_dict.get('city', '')}\n"
    prompt += "Redige uniquement le corps du courrier, sans en-tete ni signature (ils seront ajoutes automatiquement). Ton professionnel et respectueux."

    try:
        gpt_resp = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=1500,
            temperature=0.7,
            timeout=30,
        )
        await _track_openai_cost(gpt_resp)
        body_text = gpt_resp.choices[0].message.content
    except Exception as e:
        logger.error(f"Document generation error: {e}")
        return JSONResponse(status_code=500, content={"error": "Erreur lors de la generation du document."})

    filename = _doc_generator.generate_letter(
        doc_type=req.doc_type,
        subject=req.subject,
        body_text=body_text,
        recipient=req.recipient,
        profile=profile_dict,
    )

    return {
        "success": True,
        "filename": filename,
        "download_url": f"/api/documents/download/{filename}",
        "type": req.doc_type,
        "preview": body_text[:300],
    }


@app.get("/api/documents")
async def list_documents(request: Request):
    """Liste les documents generes"""
    if not _doc_generator:
        return JSONResponse(status_code=503, content={"error": "Generateur non disponible"})
    tid = getattr(request.state, "tenant_id", 1)
    docs = _doc_generator.list_documents()
    for d in docs:
        d["download_url"] = f"/api/documents/download/{d['filename']}"
    return {"documents": docs, "count": len(docs)}


@app.get("/api/documents/download/{filename}")
async def serve_document(filename: str, request: Request):
    """Sert un document genere au telechargement (auth requise, tenant verifie)"""
    tid = getattr(request.state, "tenant_id", 1)
    # Securite: whitelist extensions + anti path traversal
    _ALLOWED_EXTENSIONS = (".pdf", ".docx")
    if not filename.endswith(_ALLOWED_EXTENSIONS):
        return JSONResponse(status_code=400, content={"error": "Type de fichier non autorise"})
    if ".." in filename or "/" in filename or "\\" in filename or "\x00" in filename:
        return JSONResponse(status_code=400, content={"error": "Nom de fichier invalide"})
    base_dir = os.path.normpath(os.path.join(os.path.dirname(__file__), "static", "documents", str(tid)))
    filepath = os.path.normpath(os.path.join(base_dir, filename))
    # Verifier que le fichier est bien dans le repertoire du tenant
    if not filepath.startswith(base_dir + os.sep) and filepath != base_dir:
        return JSONResponse(status_code=403, content={"error": "Acces refuse"})
    if not os.path.exists(filepath):
        return JSONResponse(status_code=404, content={"error": "Document non trouve"})
    media_type = "application/pdf" if filename.endswith(".pdf") else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return FileResponse(
        filepath,
        media_type=media_type,
        filename=filename,
    )


# =========================================================================
# PERCEPTION - Aide contextuelle visuelle
# =========================================================================

# Anti-spam: max 1 alerte perception par heure par tenant
_perception_alert_cooldown: Dict[int, float] = {}

def _perception_alert_contacts(tenant_id: int, abnormality: dict):
    """
    Alerte les contacts de confiance par SMS en cas de situation preoccupante.
    Genere un PDF de compte-rendu d'incident (RGPD-compliant).
    NE se substitue PAS aux services d'urgence (interdit).
    """
    import time as _time
    now = _time.time()
    last_alert = _perception_alert_cooldown.get(tenant_id, 0)
    if now - last_alert < 3600:
        logger.info(f"Perception alert cooldown active for tenant {tenant_id}")
        return
    _perception_alert_cooldown[tenant_id] = now

    if not _memory_manager or not _twilio_client:
        return

    # Recuperer le profil du souscripteur
    subscriber_name = "le souscripteur"
    if _redis_client:
        profile = _redis_client.get_profile(tenant_id)
        if profile:
            subscriber_name = profile.get("first_name", subscriber_name)

    # Recuperer les contacts de confiance
    contacts = _memory_manager.list_trusted_contacts()
    if not contacts:
        logger.warning(f"Perception alert: no trusted contacts for tenant {tenant_id}")
        return

    # Description de la situation (sans donnees sensibles)
    situation = abnormality.get("description", "Situation preoccupante detectee")
    abn_type = abnormality.get("type", "inconnu")

    # SMS aux contacts de confiance
    sms_text = (
        f"[Luna - Alerte] Situation preoccupante detectee chez {subscriber_name}. "
        f"Type: {abn_type}. {situation}. "
        f"Merci de prendre des nouvelles. "
        f"En cas d'urgence, appelez le 15 (SAMU) ou le 112."
    )

    sent_count = 0
    for c in contacts[:3]:  # Max 3 contacts alertes
        phone = c.get("phone", "")
        if not phone:
            continue
        try:
            _twilio_client.send_sms(
                to=phone,
                body=sms_text,
            )
            sent_count += 1
            logger.info(f"Perception alert SMS sent to {phone[:6]}*** for tenant {tenant_id}")
        except Exception as e:
            logger.error(f"Perception alert SMS failed to {phone[:6]}***: {e}")

    # Generer un PDF de compte-rendu d'incident
    if _doc_generator and sent_count > 0:
        try:
            _generate_incident_report(tenant_id, subscriber_name, abnormality, sent_count)
        except Exception as e:
            logger.error(f"Incident report generation failed: {e}")

    logger.info(f"Perception alert: {sent_count} SMS sent for tenant {tenant_id}, type={abn_type}")


def _generate_incident_report(tenant_id: int, subscriber_name: str, abnormality: dict, sms_count: int):
    """Genere un PDF RGPD-compliant de compte-rendu d'incident perception."""
    from fpdf import FPDF
    from datetime import datetime

    class IncidentPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(124, 58, 237)
            self.cell(0, 8, "Luna - YAWatch", align="L")
            self.set_font("Helvetica", "", 9)
            self.set_text_color(136, 136, 136)
            self.cell(0, 8, datetime.now().strftime("%d/%m/%Y %H:%M"), align="R", new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(124, 58, 237)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)
        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(170, 170, 170)
            self.cell(0, 10, f"Rapport d'incident Luna | Page {self.page_no()}/{{nb}}", align="C")

    pdf = IncidentPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    # Titre
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(244, 67, 54)
    pdf.cell(0, 12, "Compte-rendu d'incident", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Infos
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(51, 51, 51)
    now = datetime.now()
    info = [
        ("Date et heure", now.strftime("%d/%m/%Y a %H:%M:%S")),
        ("Souscripteur", subscriber_name),
        ("Type d'incident", abnormality.get("type", "Non specifie")),
        ("Niveau", abnormality.get("severity", "concern")),
        ("Contacts alertes", f"{sms_count} personne(s) de confiance"),
    ]
    for label, value in info:
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(50, 7, f"{label} :", align="R")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 7, f"  {value}", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(6)

    # Description
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(26, 26, 62)
    pdf.cell(0, 8, "Description de la situation", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(51, 51, 51)
    desc = abnormality.get("description", "Situation preoccupante detectee par la perception contextuelle Luna.")
    pdf.multi_cell(0, 6, desc, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # Actions
    pdf.set_font("Helvetica", "B", 11)
    pdf.set_text_color(26, 26, 62)
    pdf.cell(0, 8, "Actions effectuees", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(51, 51, 51)
    actions = [
        f"SMS d'alerte envoye a {sms_count} contact(s) de confiance",
        "Suggestion des numeros d'urgence (15, 18, 112) dans le SMS",
        "Incident enregistre dans les notes Luna du souscripteur",
    ]
    for a in actions:
        pdf.cell(6, 6, "-")
        pdf.cell(0, 6, f" {a}", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(8)

    # Mention legale RGPD
    pdf.set_draw_color(200, 200, 200)
    pdf.line(10, pdf.get_y(), 200, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("Helvetica", "I", 8)
    pdf.set_text_color(136, 136, 136)
    legal = (
        "Ce rapport a ete genere automatiquement par Luna (YAWatch). "
        "Luna est une aide contextuelle et NE se substitue PAS aux services d'urgence (SAMU 15, Pompiers 18, Urgences 112). "
        "Les contacts de confiance ont ete alertes pour prendre des nouvelles. "
        "Aucune donnee medicale n'est collectee ni transmise. "
        "Ce document est confidentiel, conforme au RGPD, et destine uniquement au souscripteur."
    )
    pdf.multi_cell(0, 4, legal, new_x="LMARGIN", new_y="NEXT")

    # Sauvegarder
    import uuid
    ts = now.strftime("%Y%m%d_%H%M%S")
    sid = uuid.uuid4().hex[:6]
    filename = f"incident_{ts}_{sid}.pdf"
    filepath = os.path.join(STATIC_DIR, "documents", filename)
    os.makedirs(os.path.dirname(filepath), exist_ok=True)
    pdf.output(filepath)
    logger.info(f"Incident report generated: {filepath}")

    # Ajouter dans les notes pour le souscripteur
    if _memory_manager:
        _memory_manager.add_note(
            content=f"[Alerte Perception] Rapport d'incident genere: {filename}. "
                    f"Type: {abnormality.get('type', '?')}. {sms_count} contact(s) alerte(s).",
            context="incident",
            tags=["perception", "incident", "alerte"],
        )


@app.post("/api/perception/start")
async def start_perception():
    """Active la perception (camera navigateur, analyse OpenAI Vision)."""
    if not _memory_manager:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})

    # Init le detector si pas encore fait
    if _perception_detector and not _perception_detector._initialized:
        _perception_detector.initialize()

    if not _perception_detector or not _perception_detector._initialized:
        return JSONResponse(status_code=503, content={
            "error": "Module perception non disponible"
        })

    _memory_manager.set_perception_enabled(True)
    _perception_detector.set_remote_camera_active(True)
    logger.info("Perception activated (browser camera mode)")
    return {"success": True, "message": "Perception activee - ouvre ta camera"}


@app.post("/api/perception/stop")
async def stop_perception():
    """Desactive la perception camera."""
    if not _memory_manager:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})

    _memory_manager.set_perception_enabled(False)
    if _perception_detector:
        _perception_detector.set_remote_camera_active(False)
    if _perception_analyzer:
        _perception_analyzer.reset()

    logger.info("Perception deactivated by user")
    return {"success": True, "message": "Perception desactivee"}


_perception_last_frame: dict = {}  # tid -> timestamp

@app.post("/api/perception/frame")
async def receive_perception_frame(request: Request):
    """
    Recoit une frame JPEG base64 depuis le navigateur et l'analyse.
    Le navigateur capture via getUserMedia et envoie toutes les ~10 secondes.
    """
    if not _perception_detector or not _perception_analyzer or not _memory_manager:
        return JSONResponse(status_code=503, content={"error": "Perception non disponible"})

    if not _memory_manager.is_perception_enabled():
        return JSONResponse(status_code=400, content={"error": "Perception non activee"})

    # Rate limit: max 1 frame / 5 secondes par tenant (protege la facture OpenAI Vision)
    tid = getattr(request.state, "tenant_id", 1)
    now = time.time()
    last = _perception_last_frame.get(tid, 0)
    if now - last < 5.0:
        return JSONResponse(status_code=429, content={"error": "Trop rapide. 1 frame / 5 secondes max."})
    _perception_last_frame[tid] = now

    try:
        body = await request.json()
        image_b64 = body.get("image")
        if not image_b64:
            return JSONResponse(status_code=400, content={"error": "Image manquante"})

        # Limiter la taille (max ~500KB base64 = ~375KB image)
        if len(image_b64) > 700_000:
            return JSONResponse(status_code=400, content={"error": "Image trop grande (max 500KB)"})

    except Exception:
        return JSONResponse(status_code=400, content={"error": "Body JSON invalide"})

    # Analyse via OpenAI Vision (bloquant -> executor)
    loop = asyncio.get_event_loop()
    frame_analysis = await loop.run_in_executor(
        None, _perception_detector.analyze_frame_b64, image_b64
    )

    if not frame_analysis:
        return {"success": False, "error": "Analyse echouee"}

    # Analyse temporelle (postures, anomalies)
    scene_state = _perception_analyzer.analyze(frame_analysis)
    _memory_manager.update_perception_state(scene_state)

    # Log les anomalies significatives
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    for abn in scene_state.abnormalities:
        if abn["severity"] in ("attention", "concern"):
            _memory_manager.log_perception_event(abn)
            reasoning = f"Detection automatique par la perception camera: {abn['type']}"
            try:
                _memory_manager.add_note(
                    content=f"[Perception] {abn['description']}\n[Raison: {reasoning}]",
                    context="perception",
                    tags=["perception", abn["type"], abn["severity"], "reasoning"],
                )
                _memory_manager.log_event(
                    category="perception",
                    description=abn["description"],
                    reasoning=reasoning,
                    source="perception_frame",
                    details={"severity": abn["severity"], "type": abn["type"]},
                )
            except Exception:
                pass

            # ALERTE CONTACTS DE CONFIANCE en cas de severite "concern"
            if abn["severity"] == "concern":
                _perception_alert_contacts(tid, abn)

    return {
        "success": True,
        "scene": scene_state.to_dict(),
        "inference_ms": round(frame_analysis.inference_time_ms),
    }


@app.get("/api/perception/status")
async def perception_status():
    """Statut de la perception."""
    enabled = _memory_manager.is_perception_enabled() if _memory_manager else False
    camera_ok = _perception_detector.is_camera_available() if _perception_detector else False

    state = None
    if enabled and _memory_manager:
        scene = _memory_manager.get_perception_state()
        if scene:
            state = scene.to_dict()

    return {
        "available": _perception_detector is not None and (
            _perception_detector._initialized if _perception_detector else False
        ),
        "enabled": enabled,
        "camera_connected": camera_ok,
        "current_scene": state,
    }


@app.get("/api/perception/scene")
async def get_current_scene():
    """Scene actuelle detectee."""
    if not _memory_manager:
        return JSONResponse(status_code=503, content={"error": "Memoire non disponible"})
    scene = _memory_manager.get_perception_state()
    if not scene:
        return {"scene": None, "message": "Aucune scene disponible"}
    return {"scene": scene.to_dict()}


# =========================================================================
# INSTRUCTION ENGINE - Background loop
# =========================================================================

async def _load_instructions_to_scheduler():
    """Charge les instructions actives depuis Redis dans le scheduler au demarrage"""
    if not _memory_manager or not _scheduler:
        return
    try:
        instructions = _memory_manager.list_active_instructions()
        loaded = 0
        for instr in instructions:
            try:
                # Re-parse l'instruction pour obtenir un ParsedInstruction
                parsed = InstructionParser.parse(instr.description)
                _scheduler.schedule(
                    instruction_id=instr.id,
                    tenant_id=TENANT_ID,
                    instruction=parsed,
                )
                loaded += 1
            except Exception as e:
                logger.warning(f"Could not schedule instruction {instr.id}: {e}")
        logger.info(f"Loaded {loaded}/{len(instructions)} instructions into scheduler")
    except Exception as e:
        logger.warning(f"Failed to load instructions: {e}")


_last_heartbeat_time = 0  # timestamp du dernier heartbeat

async def _instruction_loop():
    """Boucle de fond : verifie les taches dues toutes les 30s et les execute"""
    global _last_heartbeat_time
    logger.info("Instruction loop started (30s interval)")
    while True:
        try:
            await asyncio.sleep(30)

            # License heartbeat toutes les 6 heures
            if _license_heartbeat:
                now = time.time()
                if now - _last_heartbeat_time > 21600:  # 6h = 21600s
                    try:
                        result = await _license_heartbeat.check()
                        _last_heartbeat_time = now
                        action = result.get("action", "unknown")
                        logger.info(f"License heartbeat: action={action}")
                        if _license_heartbeat.is_blocked():
                            logger.critical("LICENSE BLOQUEE par le backend")
                        elif _license_heartbeat.is_degraded():
                            logger.warning("LICENSE DEGRADEE - chat seul")
                    except Exception as e:
                        logger.error(f"License heartbeat error: {e}")

            # Notification engine check (self rate-limited to every 5 min)
            if _notification_engine:
                try:
                    from core.gamification.redis_ops import GamificationRedisOps
                    await _notification_engine.check_all_tenants(GamificationRedisOps)
                except Exception as e:
                    logger.warning(f"Notification engine error: {e}")

            # Secretary reminders check — envoie des notifications pour les rappels echus
            if _redis_client:
                try:
                    from core.secretary.redis_ops import SecretaryRedisOps
                    _sec_ops = SecretaryRedisOps(_redis_client, int(TENANT_ID))
                    _overdue = _sec_ops.get_overdue_reminders()
                    for _rem in _overdue:
                        if _rem.get("notified"):
                            continue
                        import json as _json_rem
                        _notif = {
                            "id": _rem.get("id", ""),
                            "type": "secretary_reminder",
                            "title": "Rappel secretaire",
                            "body": _rem.get("title", "Rappel") + (" — " + _rem.get("description", "") if _rem.get("description") else ""),
                            "ts": datetime.utcnow().isoformat(),
                        }
                        _redis_client.client.rpush(
                            f"luna:{TENANT_ID}:notifications:pending",
                            _json_rem.dumps(_notif)
                        )
                        _redis_client.client.expire(f"luna:{TENANT_ID}:notifications:pending", 86400)
                        # Marquer comme notifie
                        _rem_key = f"luna:{TENANT_ID}:secretary:reminder:{_rem['id']}"
                        _redis_client.client.hset(_rem_key, "notified", "1")
                except Exception as e:
                    logger.debug(f"Secretary reminders check: {e}")

            if not _scheduler or not _executor:
                continue

            due_tasks = _scheduler.pop_due_tasks()
            if not due_tasks:
                continue

            logger.info(f"Instruction loop: {len(due_tasks)} task(s) due")
            for task in due_tasks:
                try:
                    # Get subscriber phone for SMS delivery
                    _exec_phone = ""
                    _exec_mgr = _get_tenant_manager(task.tenant_id) if task.tenant_id else _memory_manager
                    if _exec_mgr:
                        try:
                            _exec_profile = _exec_mgr.get_subscriber_profile()
                            if _exec_profile:
                                _exec_phone = getattr(_exec_profile, "phone", "") or ""
                        except Exception:
                            pass
                    if not _exec_phone:
                        _exec_phone = os.getenv("ADMIN_PHONE", "")

                    # Load instruction metadata (max_duration etc.) from Redis
                    _exec_ctx = {
                        "tenant_id": task.tenant_id,
                        "subscriber_phone": _exec_phone,
                    }
                    if _exec_mgr:
                        try:
                            _instr_obj = _exec_mgr.get_instruction(task.instruction_id)
                            if _instr_obj and _instr_obj.metadata:
                                _exec_ctx.update(_instr_obj.metadata)
                        except Exception:
                            pass

                    result = await _executor.execute(
                        task,
                        context=_exec_ctx,
                    )
                    logger.info(
                        f"Instruction {task.instruction_id} executed: "
                        f"{result.status.value} - {result.message}"
                    )

                    # Deliver in-app notification for followup actions
                    if result.requires_followup and _redis_client:
                        try:
                            import uuid as _uuid
                            _notif_key = f"luna:{task.tenant_id}:notifications:pending"
                            import json as _json
                            _notif = {
                                "id": str(_uuid.uuid4()),
                                "type": "reminder",
                                "title": "Rappel Luna",
                                "body": result.message[:200],
                                "ts": datetime.utcnow().isoformat(),
                            }
                            _redis_client.client.rpush(_notif_key, _json.dumps(_notif))
                            _redis_client.client.expire(_notif_key, 86400)
                        except Exception as e:
                            logger.debug(f"Notification queuing: {e}")

                    # Marque l'instruction comme executee dans Redis
                    if _memory_manager:
                        try:
                            _memory_manager.mark_instruction_executed(task.instruction_id)
                        except Exception:
                            pass

                    # Enregistre le compte-rendu comme note + event log
                    if _memory_manager:
                        try:
                            _memory_manager.add_note(
                                content=f"[Auto] {result.message}",
                                context="instruction_execution",
                                tags=["auto", result.status.value, task.instruction.action_type.value],
                            )
                            _memory_manager.log_event(
                                category="instruction",
                                description=f"Instruction executee: {result.message}",
                                reasoning=f"Declenchee par le scheduler ({task.instruction.original_text[:60]})",
                                source="instruction_loop",
                            )
                        except Exception:
                            pass

                    # Replanifie si recurrent (le scheduler gere ca)
                    if result.status.value in ("success", "partial"):
                        _scheduler.complete_task(task, result=result.message)
                    else:
                        _scheduler.fail_task(task, error=result.error or "unknown")

                except Exception as e:
                    logger.error(f"Error executing task {task.instruction_id}: {e}")
                    _scheduler.fail_task(task, error=str(e))

        except asyncio.CancelledError:
            logger.info("Instruction loop stopped")
            break
        except Exception as e:
            logger.error(f"Instruction loop error: {e}")
            # Backoff progressif en cas d'erreurs repetees
            await asyncio.sleep(30)


# =========================================================================
# MAIN
# =========================================================================

# =========================================================================
# TEST / SIMULATOR ENDPOINTS
# =========================================================================

@app.get("/api/test/scenarios")
async def list_scenarios():
    """Liste les scenarios de test disponibles."""
    try:
        from core.testing.simulator import ALL_SCENARIOS
        return {
            "scenarios": [
                {"name": s.name, "description": s.description, "steps": len(s.steps)}
                for s in ALL_SCENARIOS.values()
            ]
        }
    except ImportError:
        return JSONResponse(status_code=503, content={"error": "Module testing non disponible"})


@app.post("/api/test/scenario")
async def run_scenario(req: Request):
    """Execute un scenario de test. Protege par cle admin."""
    body = await req.json()
    scenario_name = body.get("scenario", "")
    admin_key = body.get("admin_key", "")

    # Protection par cle admin (utilise ADMIN_NUMBER comme cle simple)
    if admin_key != ADMIN_NUMBER:
        return JSONResponse(status_code=403, content={"error": "Cle admin requise"})

    try:
        from core.testing.simulator import ScenarioSimulator, ALL_SCENARIOS

        if scenario_name not in ALL_SCENARIOS:
            return JSONResponse(
                status_code=400,
                content={"error": f"Scenario inconnu: {scenario_name}", "available": list(ALL_SCENARIOS.keys())}
            )

        global _test_mode
        _test_mode = True

        async def test_chat_fn(message: str) -> str:
            """Appelle le chat en mode test."""
            try:
                test_session = f"test_{datetime.now().strftime('%H%M%S')}"
                if test_session not in conversations:
                    conversations[test_session] = [
                        {"role": "system", "content": _system_prompt_now()}
                    ]
                messages = conversations[test_session]
                messages.append({"role": "user", "content": message})
                response = openai_client.chat.completions.create(
                    model=OPENAI_MODEL,
                    messages=messages,
                    max_tokens=500,
                    temperature=0.8,
                    timeout=30,
                )
                await _track_openai_cost(response)
                luna_msg = response.choices[0].message.content
                messages.append({"role": "assistant", "content": luna_msg})
                return luna_msg
            except Exception as e:
                return f"[ERREUR TEST] {e}"

        simulator = ScenarioSimulator(chat_fn=test_chat_fn)
        scenario = ALL_SCENARIOS[scenario_name]
        result = await simulator.run_scenario(scenario)

        _test_mode = False

        return result.to_dict()

    except ImportError:
        return JSONResponse(status_code=503, content={"error": "Module testing non disponible"})
    except Exception as e:
        _test_mode = False
        return JSONResponse(status_code=500, content={"error": str(e)})


# =============================================================================
# UNIFIED API - Memoire temps reel multi-canal
# =============================================================================

def _get_or_create_session(tid: int = None) -> Optional[Dict]:
    """Recupere ou cree une session unifiee"""
    if not _redis_client:
        return None
    tid = tid if tid is not None else TENANT_ID
    session = _redis_client.get_session(tid)
    if not session:
        new_session = UnifiedSession(tenant_id=tid)
        _redis_client.set_session(tid, new_session.to_redis())
        return new_session.to_redis()
    return session


def _update_session_activity(channel: str = None, topic: str = None, mood: str = None,
                             tid: int = None):
    """Met a jour l'activite de la session"""
    if not _redis_client:
        return
    tid = tid if tid is not None else TENANT_ID
    updates = {"last_activity": datetime.utcnow().isoformat(), "status": "active"}
    if channel:
        updates["active_channel"] = channel
    if topic:
        updates["current_topic"] = topic
    if mood:
        updates["current_mood"] = mood
    _redis_client.update_session(tid, updates)


def _add_unified_message(channel: str, role: str, content: str,
                         audio_duration: float = None, mood: str = None,
                         intent: str = None, tool_calls: list = None,
                         tid: int = None):
    """Ajoute un message a l'historique unifie"""
    if not _redis_client:
        return None
    tid = tid if tid is not None else TENANT_ID
    msg = UnifiedMessage(
        tenant_id=tid,
        channel=Channel(channel),
        role=MessageRole(role),
        content=content,
        audio_duration_sec=audio_duration,
        detected_mood=MoodIndicator(mood) if mood else None,
        detected_intent=intent,
        tool_calls=tool_calls,
    )
    _redis_client.add_unified_message(tid, msg.id, msg.to_redis())
    # Publier evenement temps reel
    _redis_client.publish_event(tid, "new_message", {
        "id": msg.id, "channel": channel, "role": role, "content": content[:100]
    })
    return msg.id


@app.get("/api/unified/session")
async def get_unified_session(request: Request):
    """Recupere l'etat de la session active"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    session = _get_or_create_session(tid)
    context = _redis_client.get_context(tid) or {}

    return {
        "session": session,
        "context": context,
        "channels": {
            "app": True,
            "voice": True,
            "sms": sms_client.is_configured if sms_client else False,
            "visio": tavus_client.is_configured if tavus_client else False,
        }
    }


@app.post("/api/unified/send")
async def unified_send(request: Request):
    """Envoie un message unifie (detecte le canal automatiquement ou specifie)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    content = data.get("content", "").strip()
    channel = data.get("channel", "app")
    role = data.get("role", "subscriber")

    if not content:
        return JSONResponse(status_code=400, content={"error": "content requis"})

    # Valider le canal
    valid_channels = ["app", "voice", "sms", "call", "visio"]
    if channel not in valid_channels:
        return JSONResponse(status_code=400, content={
            "error": f"Canal invalide. Valides: {', '.join(valid_channels)}"
        })

    # Mettre a jour la session
    _update_session_activity(channel=channel, tid=tid)

    # Ajouter le message a l'historique unifie
    msg_id = _add_unified_message(channel, role, content, tid=tid)

    # Si c'est un message du subscriber, generer une reponse Luna
    response_content = None
    if role == "subscriber":
        # Recuperer le contexte recent
        recent_messages = _redis_client.get_recent_context_messages(tid, count=15)

        # Construire l'historique pour OpenAI
        messages = [{"role": "system", "content": _system_prompt_now()}]
        for msg in reversed(recent_messages):
            msg_role = "assistant" if msg.get("role") == "luna" else "user"
            messages.append({"role": msg_role, "content": msg.get("content", "")})

        # Appel OpenAI
        try:
            response = openai_client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=messages,
                max_tokens=500,
            )
            await _track_openai_cost(response, tid)
            response_content = response.choices[0].message.content

            # Ajouter la reponse Luna
            luna_msg_id = _add_unified_message(channel, "luna", response_content, tid=tid)

            # Mettre a jour le contexte
            _redis_client.update_context(tid, {
                "last_topic": content[:50],
                "last_response": response_content[:100],
                "last_channel": channel,
                "updated_at": datetime.utcnow().isoformat(),
            })

        except Exception as e:
            logger.error(f"Unified send error: {e}")
            response_content = "Je suis desolee, j'ai eu un petit souci. Pouvez-vous repeter?"

    return {
        "success": True,
        "message_id": msg_id,
        "channel": channel,
        "response": response_content,
    }


@app.get("/api/unified/history")
async def get_unified_history(request: Request, limit: int = 50, channel: str = None):
    """Recupere l'historique unifie de tous les canaux"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    messages = _redis_client.get_unified_messages(tid, limit=limit, channel=channel)

    return {
        "messages": messages,
        "count": len(messages),
        "filter_channel": channel,
    }


@app.get("/api/unified/context")
async def get_realtime_context(request: Request):
    """Recupere le contexte temps reel (pour affichage app)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    session = _redis_client.get_session(tid)
    context = _redis_client.get_context(tid)
    last_handoff = _redis_client.get_last_handoff(tid)

    # Dernier message par canal
    last_by_channel = {}
    for ch in ["app", "voice", "sms", "visio"]:
        msgs = _redis_client.get_unified_messages(tid, limit=1, channel=ch)
        if msgs:
            last_by_channel[ch] = msgs[0]

    return {
        "session": session,
        "context": context,
        "last_handoff": last_handoff,
        "last_by_channel": last_by_channel,
    }


@app.post("/api/unified/handoff")
async def channel_handoff(request: Request):
    """Change de canal avec transfert de contexte"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    from_channel = data.get("from_channel", "app")
    to_channel = data.get("to_channel")
    reason = data.get("reason", "user_request")

    if not to_channel:
        return JSONResponse(status_code=400, content={"error": "to_channel requis"})

    # Recuperer le contexte actuel
    context = _redis_client.get_context(tid) or {}
    recent_messages = _redis_client.get_recent_context_messages(tid, count=5)

    # Creer un resume du contexte
    context_summary = f"Sujet: {context.get('last_topic', 'aucun')}. "
    if recent_messages:
        last_msg = recent_messages[0]
        context_summary += f"Dernier message ({last_msg.get('channel', 'app')}): {last_msg.get('content', '')[:50]}..."

    # Enregistrer le handoff
    handoff = ChannelHandoff(
        tenant_id=tid,
        from_channel=Channel(from_channel),
        to_channel=Channel(to_channel),
        reason=reason,
        context_summary=context_summary,
    )
    _redis_client.add_handoff(tid, handoff.id, handoff.to_redis())

    # Mettre a jour la session
    _redis_client.update_session(tid, {
        "active_channel": to_channel,
        "last_activity": datetime.utcnow().isoformat(),
    })

    # Publier l'evenement
    _redis_client.publish_event(tid, "channel_handoff", {
        "from": from_channel, "to": to_channel, "reason": reason
    })

    return {
        "success": True,
        "handoff_id": handoff.id,
        "context_transferred": context_summary,
    }


# =============================================================================
# VOICE MODE - WebSocket temps reel (sans video)
# =============================================================================

# Pipeline voix legacy (/api/voice/stream) supprime — utiliser /api/voice-call (Twilio Realtime)


# =============================================================================
# SALONS FAMILLE (Rooms) — Chat, Cinema, Karaoke, Jeux
# =============================================================================

from core.rooms.models import (
    ROOM_TYPES, GAME_TYPES, QUIZ_SETS,
    generate_member_token, verify_member_token,
    new_room_data,
)
from core.rooms.redis_ops import RoomRedisOps
from core.rooms.manager import room_manager

_room_ops: Optional[RoomRedisOps] = None


def _get_room_ops() -> Optional[RoomRedisOps]:
    global _room_ops
    if _room_ops is None and _redis_client:
        _room_ops = RoomRedisOps(_redis_client)
    return _room_ops


def _get_member_name(tid: int, phone: str) -> str:
    """Lookup member name from family data or profile (if email)."""
    if not _redis_client:
        return phone
    # If phone looks like an email, lookup profile instead
    if "@" in phone:
        prof = _redis_client.get_profile(tid)
        if prof and prof.get("first_name"):
            return prof["first_name"]
        return phone.split("@")[0]
    m = _redis_client.get_family_member(tid, phone)
    if m:
        return m.get("name", phone)
    return phone


@app.get("/salon")
async def salon_page():
    """Page des salons famille."""
    salon_path = os.path.join(STATIC_DIR, "salon.html")
    if os.path.exists(salon_path):
        return FileResponse(salon_path)
    return JSONResponse(status_code=404, content={"error": "Salon non disponible"})


@app.get("/api/rooms")
async def list_rooms(request: Request):
    """Lister les salons actifs."""
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    rooms = rops.list_rooms(tid)
    result = []
    for r in rooms:
        result.append({
            "id": r.get("id"),
            "name": r.get("name"),
            "type": r.get("type"),
            "host_name": r.get("host_name"),
            "status": r.get("status", "open"),
            "participants": rops.count_participants(tid, r["id"]),
            "youtube_url": r.get("youtube_url", ""),
            "game_type": r.get("game_type", ""),
            "created_at": r.get("created_at"),
        })
    return result


@app.post("/api/rooms")
async def create_room(request: Request):
    """Créer un salon (souscripteur uniquement)."""
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    data = await request.json()
    name = data.get("name", "Salon Famille").strip()[:50]
    room_type = data.get("type", "chat")
    if room_type not in ROOM_TYPES:
        return JSONResponse(status_code=400, content={"error": f"Type invalide. Valides: {ROOM_TYPES}"})

    # Get subscriber info
    host_phone = data.get("phone", "subscriber")
    host_name = data.get("host_name", "Hôte")

    room = new_room_data(name, room_type, host_phone, host_name, tid)
    room_id = rops.create_room(tid, room)
    _gamify(tid, "room_created")

    # Auto-join host
    rops.join_room(tid, room_id, host_phone)

    # Generate invite links for family members
    invite_links = {}
    if _redis_client:
        members = _redis_client.get_all_family_members(tid)
        secret = _JWT_SECRET
        for m in members:
            phone = m.get("phone", "")
            if phone and m.get("is_active") in ("1", "True", "true"):
                token = generate_member_token(phone, tid, secret)
                invite_links[m.get("name", phone)] = {
                    "phone": phone,
                    "token": token,
                    "url": f"/salon?room={room_id}&phone={phone}&token={token}",
                }

    return {
        "success": True,
        "room_id": room_id,
        "name": name,
        "type": room_type,
        "invite_links": invite_links,
    }


@app.get("/api/rooms/{room_id}")
async def get_room(room_id: str, request: Request):
    """Détails d'un salon."""
    import json as _json
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    room = rops.get_room(tid, room_id)
    if not room:
        return JSONResponse(status_code=404, content={"error": "Salon introuvable"})

    participants = rops.get_participants(tid, room_id)
    members_info = []
    for p in participants:
        name = _get_member_name(tid, p)
        members_info.append({"phone": p, "name": name})

    messages = rops.get_messages(tid, room_id, limit=50)
    parsed_msgs = []
    for m in reversed(messages):
        try:
            parsed_msgs.append(_json.loads(m))
        except (_json.JSONDecodeError, TypeError):
            pass

    # Parse game_state JSON string to object
    result = dict(room)
    if result.get("game_state"):
        try:
            result["game_state"] = _json.loads(result["game_state"])
        except (_json.JSONDecodeError, TypeError):
            result["game_state"] = {}
    # Convert playback_time to number
    try:
        result["playback_time"] = float(result.get("playback_time", 0))
    except (ValueError, TypeError):
        result["playback_time"] = 0

    return {
        **result,
        "participants": members_info,
        "participant_count": len(participants),
        "messages": parsed_msgs,
    }


@app.post("/api/rooms/{room_id}/join")
async def join_room(room_id: str, request: Request):
    """Rejoindre un salon."""
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    data = await request.json()
    phone = data.get("phone", "")

    room = rops.get_room(tid, room_id)
    if not room:
        return JSONResponse(status_code=404, content={"error": "Salon introuvable"})

    if not rops.join_room(tid, room_id, phone):
        return JSONResponse(status_code=400, content={"error": "Salon plein"})

    return {"success": True, "room_id": room_id}


@app.delete("/api/rooms/{room_id}")
async def delete_room(room_id: str, request: Request):
    """Fermer un salon (host uniquement)."""
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    room = rops.get_room(tid, room_id)
    if not room:
        return JSONResponse(status_code=404, content={"error": "Salon introuvable"})

    # Verifier que le tenant est bien le proprietaire du salon
    room_tenant = int(room.get("tenant_id", 0))
    if room_tenant != tid:
        return JSONResponse(status_code=403, content={"error": "Seul le createur du salon peut le fermer"})

    rops.delete_room(tid, room_id)

    # Notify all connected
    await room_manager.broadcast(room_id, {
        "type": "system", "content": "Le salon a été fermé par l'hôte.",
    })

    return {"success": True}


@app.post("/api/rooms/{room_id}/invite")
async def invite_to_room(room_id: str, request: Request):
    """Invite un membre famille dans le salon par SMS."""
    rops = _get_room_ops()
    if not rops:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    room = rops.get_room(tid, room_id)
    if not room:
        return JSONResponse(status_code=404, content={"error": "Salon introuvable"})

    data = await request.json()
    phone = data.get("phone", "")
    member_name = data.get("name", "")
    if not phone:
        return JSONResponse(status_code=400, content={"error": "Numero requis"})

    if not sms_client or not sms_client.is_configured:
        return JSONResponse(status_code=503, content={"error": "Service SMS non disponible"})

    # Generate invite link
    secret = _JWT_SECRET
    token = generate_member_token(phone, tid, secret)
    base_url = os.getenv("LUNA_BASE_URL", "https://luna-beta-674304336025.europe-west1.run.app")
    invite_url = f"{base_url}/salon?room={room_id}&phone={phone}&token={token}"

    room_name = room.get("name", "Salon")
    host_name = room.get("host_name", "ton proche")
    sms_msg = f"[Luna] {host_name} t'invite dans le salon \"{room_name}\" ! Rejoins ici : {invite_url}"

    from integrations.twilio.sms_client import TwilioSMSClient
    normalized = TwilioSMSClient.normalize_phone(phone)
    success, details = _tracked_sms_send(normalized, sms_msg, label=f"Invitation salon a {member_name or phone}")

    if success:
        return {"success": True, "message": f"Invitation envoyee a {member_name or phone}"}
    else:
        logger.error(f"Room invite SMS failed: {details}")
        return JSONResponse(status_code=500, content={"error": "Echec de l'envoi du SMS"})


@app.get("/api/rooms/member-token")
async def get_member_token(request: Request):
    """Génère un token d'accès pour un membre famille (souscripteur only)."""
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    phone = request.query_params.get("phone", "")
    if not phone:
        return JSONResponse(status_code=400, content={"error": "phone requis"})
    secret = _JWT_SECRET
    token = generate_member_token(phone, tid, secret)
    return {"phone": phone, "token": token}


@app.websocket("/api/rooms/{room_id}/ws")
async def room_websocket(websocket: WebSocket, room_id: str):
    """WebSocket temps réel pour un salon famille."""
    import json as _json
    await websocket.accept()

    rops = _get_room_ops()
    if not rops:
        await websocket.close(code=1011, reason="Redis non disponible")
        return

    tid = getattr(websocket.state, "tenant_id", TENANT_ID)

    # Auth: get phone + token from query params
    phone = websocket.query_params.get("phone", "")
    token = websocket.query_params.get("token", "")

    # Detect token type: JWT (contains dots) vs HMAC member token
    jwt_payload = None
    if token and "." in token:
        # Try JWT client token first
        jwt_payload = _decode_client_token(token)
        if jwt_payload:
            tid = jwt_payload.get("tenant_id", tid)
            # Use phone from query param, or email from JWT
            if not phone:
                phone = jwt_payload.get("email", "subscriber")
        else:
            await websocket.close(code=4001, reason="Token JWT invalide")
            return
    elif token:
        # HMAC member token
        if not phone:
            phone = "subscriber"
        secret = _JWT_SECRET
        if not verify_member_token(phone, tid, token, secret):
            await websocket.close(code=4001, reason="Token membre invalide")
            return
    else:
        # No token at all
        if not phone:
            phone = "subscriber"

    # Verify room exists
    room = rops.get_room(tid, room_id)
    if not room:
        await websocket.close(code=4004, reason="Salon introuvable")
        return

    # Get member name
    if jwt_payload:
        # For JWT-authenticated users, get name from profile
        _name = ""
        if _redis_client:
            _prof = _redis_client.get_profile(tid)
            if _prof:
                _name = _prof.get("first_name", "")
        name = _name or jwt_payload.get("email", "").split("@")[0] or "Hôte"
    else:
        name = _get_member_name(tid, phone) if phone != "subscriber" else "Hôte"

    # Join
    rops.join_room(tid, room_id, phone)
    await room_manager.connect(room_id, phone, websocket)

    # Announce join
    count = rops.count_participants(tid, room_id)
    await room_manager.broadcast(room_id, {
        "type": "join", "name": name, "phone": phone, "count": count,
    }, exclude_phone=phone)
    await room_manager.broadcast(room_id, {
        "type": "system", "content": f"{name} a rejoint le salon",
    }, exclude_phone=phone)

    try:
        while True:
            raw = await websocket.receive_text()
            try:
                data = _json.loads(raw)
            except (_json.JSONDecodeError, TypeError):
                continue

            # Refresh room data for each message
            current_room = rops.get_room(tid, room_id)
            if not current_room:
                await websocket.close(code=4004, reason="Salon fermé")
                break

            events = await room_manager.handle_message(
                room_id, phone, name, data, rops, tid, current_room
            )
            for ev in events:
                _gamify(tid, ev)
    except WebSocketDisconnect:
        pass
    except Exception as e:
        logger.warning(f"Room WS error: {e}")
    finally:
        rops.leave_room(tid, room_id, phone)
        await room_manager.disconnect(room_id, phone)
        count = rops.count_participants(tid, room_id)
        await room_manager.broadcast(room_id, {
            "type": "leave", "name": name, "phone": phone, "count": count,
        })
        await room_manager.broadcast(room_id, {
            "type": "system", "content": f"{name} a quitté le salon",
        })


# =============================================================================
# SYNC - Notifications des autres canaux
# =============================================================================

@app.post("/api/sync/tavus")
async def sync_from_tavus(request: Request):
    """Tavus notifie la fin d'un appel visio"""
    if not tavus_client or not tavus_client.is_configured:
        return JSONResponse(status_code=503, content={"error": "Service visio non disponible"})
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    data = await request.json()
    # Le tenant_id peut venir du JWT (si appel client) ou du body (si callback serveur)
    tid = getattr(request.state, "tenant_id", None) or data.get("tenant_id", TENANT_ID)
    event = data.get("event", "")
    conversation_id = data.get("conversation_id", "")
    transcript = data.get("transcript", [])

    if event == "conversation.ended":
        # Ajouter les messages du transcript a l'historique unifie
        for entry in transcript:
            role = "luna" if entry.get("speaker") == "replica" else "subscriber"
            content = entry.get("text", "")
            if content:
                _add_unified_message("visio", role, content)

        # Mettre a jour la session
        _redis_client.update_session(tid, {
            "is_video_active": "0",
            "channel_session_id": "",
            "last_activity": datetime.utcnow().isoformat(),
        })

        # Track cout visio Tavus par tenant
        duration_sec = data.get("duration", 0)
        if not duration_sec and transcript:
            # Estimer la duree par le nombre de messages (~10s par message)
            duration_sec = len(transcript) * 10
        duration_min = max(1, duration_sec / 60)  # minimum 1 minute facturee
        try:
            cortex = get_cortex() if _CORTEX_AVAILABLE else None
            if cortex and hasattr(cortex, 'cost_tracker') and cortex.cost_tracker:
                await cortex.cost_tracker.track_tavus_tenant(tid, duration_min)
        except Exception:
            pass

        # Generate LLM summary of visio call
        _visio_mgr = _get_tenant_manager(tid) if tid else _memory_manager
        if transcript and _visio_mgr:
            try:
                _visio_summary = await _generate_visio_summary(transcript, conversation_id, duration_min)
                _visio_mgr.add_note(
                    content=_visio_summary,
                    context="visio_summary",
                    tags=["visio", "rapport", "resume", conversation_id],
                )
                logger.info(f"Visio summary saved for {conversation_id}")
            except Exception as e:
                logger.warning(f"Failed to generate visio summary: {e}")

            # Generate PDF report
            if _doc_generator:
                try:
                    _pdf_transcript = []
                    for entry in transcript:
                        role = "luna" if entry.get("speaker") == "replica" else "user"
                        _pdf_transcript.append({"role": role, "text": entry.get("text", "")})
                    _report_fn = _doc_generator.generate_call_report(
                        call_type="visio",
                        subscriber_name=_SUBSCRIBER_NAME,
                        contact_name="Luna (visio)",
                        contact_phone="",
                        duration_minutes=duration_min,
                        transcript=_pdf_transcript,
                        call_sid=conversation_id,
                    )
                    _report_url = f"/api/documents/download/{_report_fn}"
                    _visio_mgr.add_note(
                        content=f"[Rapport visio] Duree: {duration_min:.1f} min — PDF: {_report_url}",
                        context="visio_report",
                        tags=["rapport", "visio", conversation_id],
                    )
                    logger.info(f"Visio PDF report: {_report_fn}")
                except Exception as e:
                    logger.warning(f"Failed to generate visio PDF: {e}")

        logger.info(f"Tavus conversation ended: {conversation_id}, {len(transcript)} messages synced")
        _gamify(tid, "visio_session")

    return {"success": True}


@app.post("/api/sync/twilio")
async def sync_from_twilio(request: Request):
    """Twilio notifie la fin d'un appel"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    data = await request.json()
    tid = getattr(request.state, "tenant_id", None) or data.get("tenant_id", TENANT_ID)
    call_sid = data.get("CallSid", "")
    call_status = data.get("CallStatus", "")
    from_number = data.get("From", "")
    transcript = data.get("transcript", "")

    if call_status == "completed":
        if transcript:
            _add_unified_message("call", "subscriber", f"[Appel Twilio] {transcript}")

        _redis_client.update_session(tid, {
            "active_channel": "app",
            "last_activity": datetime.utcnow().isoformat(),
        })

        logger.info(f"Twilio call ended: {call_sid}")

    return {"success": True}


# =============================================================================
# THEMES & PERSONNALISATION
# =============================================================================

# Import des presets de themes
try:
    from core.memory.schemas import (
        THEME_PRESETS, ThemePreset, UserThemePreferences,
        ThemeStyle, LunaAvatarStyle, FontSize, VoiceStyle
    )
    _THEMES_AVAILABLE = True
except ImportError:
    _THEMES_AVAILABLE = False
    THEME_PRESETS = {}


@app.get("/api/themes")
async def list_themes(age: int = None, gender: str = None):
    """Liste tous les themes disponibles, filtres optionnellement par age/genre"""
    themes = []
    for preset_id, preset in THEME_PRESETS.items():
        # Filtrage par age
        if age is not None:
            if preset.target_age_min and age < preset.target_age_min:
                continue
            if preset.target_age_max and age > preset.target_age_max:
                continue
        # Filtrage par genre
        if gender and preset.target_gender and preset.target_gender != gender:
            continue

        themes.append({
            "id": preset.id,
            "name": preset.name,
            "description": preset.description,
            "style": preset.style.value,
            "preview": {
                "primary_color": preset.primary_color,
                "secondary_color": preset.secondary_color,
                "accent_color": preset.accent_color,
                "background_type": preset.background_type,
                "background_value": preset.background_value,
                "dark_mode": preset.dark_mode,
            },
            "target": {
                "age_min": preset.target_age_min,
                "age_max": preset.target_age_max,
                "gender": preset.target_gender,
            }
        })

    return {
        "themes": themes,
        "count": len(themes),
        "filters_applied": {"age": age, "gender": gender},
    }


@app.get("/api/themes/suggest")
async def suggest_theme(age: int = None, gender: str = None):
    """Suggere le meilleur theme selon le profil"""
    best_match = None
    best_score = -1

    for preset_id, preset in THEME_PRESETS.items():
        score = 0

        # Score par age
        if age is not None:
            if preset.target_age_min and preset.target_age_max:
                if preset.target_age_min <= age <= preset.target_age_max:
                    score += 10
            elif preset.target_age_min and age >= preset.target_age_min:
                score += 5
            elif not preset.target_age_min and not preset.target_age_max:
                score += 2  # Theme universel

        # Score par genre
        if gender:
            if preset.target_gender == gender:
                score += 8
            elif not preset.target_gender:
                score += 3  # Theme unisexe

        if score > best_score:
            best_score = score
            best_match = preset

    if not best_match:
        best_match = THEME_PRESETS.get("zen")  # Default

    return {
        "suggested_theme": best_match.id if best_match else "zen",
        "theme_name": best_match.name if best_match else "Zen",
        "match_score": best_score,
        "profile": {"age": age, "gender": gender},
    }


@app.get("/api/themes/options")
async def get_theme_options():
    """Liste toutes les options de personnalisation disponibles"""
    return {
        "styles": [{"id": s.value, "name": s.value.title()} for s in ThemeStyle],
        "avatar_styles": [{"id": a.value, "name": a.value.replace("_", " ").title()} for a in LunaAvatarStyle],
        "voice_styles": [{"id": v.value, "name": v.value.title()} for v in VoiceStyle],
        "font_sizes": [{"id": f.value, "name": f.value.title()} for f in FontSize],
        "background_types": ["solid", "gradient", "image", "pattern"],
        "suggested_colors": {
            "primary": ["#6366F1", "#EC4899", "#0EA5E9", "#10B981", "#F59E0B", "#8B5CF6"],
            "secondary": ["#A5B4FC", "#F9A8D4", "#7DD3FC", "#6EE7B7", "#FCD34D", "#C4B5FD"],
            "accent": ["#F472B6", "#22D3EE", "#FACC15", "#A78BFA", "#FB923C", "#34D399"],
        }
    }


@app.get("/api/themes/{theme_id}")
async def get_theme_details(theme_id: str):
    """Recupere les details complets d'un theme"""
    preset = THEME_PRESETS.get(theme_id)
    if not preset:
        return JSONResponse(status_code=404, content={"error": "Theme non trouve"})

    return {
        "id": preset.id,
        "name": preset.name,
        "description": preset.description,
        "style": preset.style.value,
        "colors": {
            "primary": preset.primary_color,
            "secondary": preset.secondary_color,
            "accent": preset.accent_color,
        },
        "background": {
            "type": preset.background_type,
            "value": preset.background_value,
        },
        "dark_mode": preset.dark_mode,
        "avatar": {
            "style": preset.avatar_style.value,
        },
        "voice": {
            "style": preset.voice_style.value,
        },
        "typography": {
            "font_family": preset.font_family,
            "font_size": preset.font_size.value,
        },
        "assets": {
            "icon_pack": preset.icon_pack,
            "sounds_pack": preset.sounds_pack,
        },
        "target": {
            "age_min": preset.target_age_min,
            "age_max": preset.target_age_max,
            "gender": preset.target_gender,
        }
    }


@app.get("/api/profile/theme")
async def get_user_theme(request: Request, phone: str = None):
    """Recupere le theme actuel d'un utilisateur"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    user_phone = phone or ADMIN_NUMBER
    theme_data = _redis_client.get_user_theme(tid, user_phone)

    if not theme_data:
        # Retourner le theme par defaut
        default_preset = THEME_PRESETS.get("zen")
        return {
            "user_phone": user_phone,
            "preset_id": "zen",
            "preset_name": "Zen",
            "is_custom": False,
            "theme": {
                "primary_color": default_preset.primary_color,
                "secondary_color": default_preset.secondary_color,
                "accent_color": default_preset.accent_color,
                "background_type": default_preset.background_type,
                "background_value": default_preset.background_value,
                "dark_mode": default_preset.dark_mode,
                "avatar_style": default_preset.avatar_style.value,
                "voice_style": default_preset.voice_style.value,
                "font_size": default_preset.font_size.value,
            }
        }

    # Construire le theme effectif (preset + overrides)
    preset_id = theme_data.get("preset_id")
    preset = THEME_PRESETS.get(preset_id) if preset_id else None

    effective_theme = {}
    if preset:
        effective_theme = {
            "primary_color": preset.primary_color,
            "secondary_color": preset.secondary_color,
            "accent_color": preset.accent_color,
            "background_type": preset.background_type,
            "background_value": preset.background_value,
            "dark_mode": preset.dark_mode,
            "avatar_style": preset.avatar_style.value,
            "voice_style": preset.voice_style.value,
            "font_size": preset.font_size.value,
        }

    # Appliquer les overrides custom
    if theme_data.get("custom_primary_color"):
        effective_theme["primary_color"] = theme_data["custom_primary_color"]
    if theme_data.get("custom_secondary_color"):
        effective_theme["secondary_color"] = theme_data["custom_secondary_color"]
    if theme_data.get("custom_accent_color"):
        effective_theme["accent_color"] = theme_data["custom_accent_color"]
    if theme_data.get("custom_background_type"):
        effective_theme["background_type"] = theme_data["custom_background_type"]
    if theme_data.get("custom_background_value"):
        effective_theme["background_value"] = theme_data["custom_background_value"]
    if theme_data.get("custom_avatar_style"):
        effective_theme["avatar_style"] = theme_data["custom_avatar_style"]
    if theme_data.get("custom_voice_style"):
        effective_theme["voice_style"] = theme_data["custom_voice_style"]
    if theme_data.get("custom_font_size"):
        effective_theme["font_size"] = theme_data["custom_font_size"]
    if theme_data.get("dark_mode"):
        effective_theme["dark_mode"] = theme_data["dark_mode"] == "1"

    has_custom = any(theme_data.get(f"custom_{k}") for k in
                     ["primary_color", "secondary_color", "accent_color",
                      "background_type", "background_value", "avatar_style",
                      "voice_style", "font_size"])

    return {
        "user_phone": user_phone,
        "preset_id": preset_id,
        "preset_name": preset.name if preset else None,
        "is_custom": has_custom,
        "theme": effective_theme,
        "family_photos": theme_data.get("family_photos", "[]"),
    }


@app.post("/api/profile/theme")
async def set_user_theme(request: Request):
    """Definit le theme d'un utilisateur"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    user_phone = data.get("phone", ADMIN_NUMBER)
    preset_id = data.get("preset_id")

    # Valider le preset si fourni
    if preset_id and preset_id not in THEME_PRESETS:
        return JSONResponse(status_code=400, content={
            "error": f"Theme inconnu: {preset_id}",
            "available_themes": list(THEME_PRESETS.keys())
        })

    # Construire les preferences
    prefs = UserThemePreferences(
        user_phone=user_phone,
        tenant_id=tid,
        preset_id=preset_id,
        custom_primary_color=data.get("primary_color"),
        custom_secondary_color=data.get("secondary_color"),
        custom_accent_color=data.get("accent_color"),
        custom_background_type=data.get("background_type"),
        custom_background_value=data.get("background_value"),
        custom_avatar_style=LunaAvatarStyle(data["avatar_style"]) if data.get("avatar_style") else None,
        custom_voice_style=VoiceStyle(data["voice_style"]) if data.get("voice_style") else None,
        custom_font_size=FontSize(data["font_size"]) if data.get("font_size") else None,
        dark_mode=data.get("dark_mode"),
        family_photos=data.get("family_photos", []),
    )

    _redis_client.set_user_theme(tid, user_phone, prefs.to_redis())

    # Publier evenement pour mise a jour temps reel
    _redis_client.publish_event(tid, "theme_changed", {
        "user_phone": user_phone,
        "preset_id": preset_id,
    })

    return {
        "success": True,
        "user_phone": user_phone,
        "preset_id": preset_id,
        "message": f"Theme {'personnalise' if not preset_id else THEME_PRESETS[preset_id].name} applique",
    }


@app.delete("/api/profile/theme")
async def reset_user_theme(request: Request, phone: str = None):
    """Remet le theme par defaut"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    user_phone = phone or ADMIN_NUMBER
    _redis_client.delete_user_theme(tid, user_phone)

    return {
        "success": True,
        "message": "Theme reinitialise (Zen par defaut)",
    }


# =============================================================================
# ASSISTANT PRO - Analyse, Generation, Iteration
# =============================================================================

# Import des modeles assistant
try:
    from core.memory.schemas import (
        AssistantTask, DocumentType, DocumentStatus, EmailDraft, ADMIN_TEMPLATES
    )
    _ASSISTANT_AVAILABLE = True
except ImportError:
    _ASSISTANT_AVAILABLE = False


@app.post("/api/assistant/analyze")
async def analyze_content(request: Request):
    """Analyse un document ou texte et retourne des insights"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    content = data.get("content", "").strip()
    analysis_type = data.get("type", "general")  # general, contract, email, letter, cv

    if not content:
        return JSONResponse(status_code=400, content={"error": "content requis"})

    # Construire le prompt d'analyse
    prompts = {
        "general": "Analyse ce texte et fournis: 1) Resume en 2-3 phrases, 2) Points cles, 3) Suggestions d'amelioration.",
        "contract": "Analyse ce contrat et fournis: 1) Resume des obligations, 2) Points d'attention/risques, 3) Clauses importantes, 4) Recommandations.",
        "email": "Analyse cet email et fournis: 1) Ton utilise, 2) Points cles, 3) Suggestions pour ameliorer la clarte et l'impact.",
        "letter": "Analyse ce courrier et fournis: 1) Objectif du courrier, 2) Ton et formalite, 3) Points forts, 4) Ameliorations possibles.",
        "cv": "Analyse ce CV et fournis: 1) Points forts, 2) Points faibles, 3) Competences mises en avant, 4) Suggestions d'amelioration.",
    }

    prompt = prompts.get(analysis_type, prompts["general"])
    full_prompt = f"{prompt}\n\nTexte a analyser:\n\n{content}"

    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Tu es un assistant d'analyse professionnelle. Reponds en francais de maniere structuree."},
                {"role": "user", "content": full_prompt}
            ],
            max_tokens=1500,
            temperature=0.3,
        )
        await _track_openai_cost(response, tid)
        analysis = response.choices[0].message.content

        # Sauvegarder la tache
        task = AssistantTask(
            tenant_id=tid,
            task_type="analyze",
            input_text=content[:5000],
            analysis_result={"type": analysis_type, "analysis": analysis},
            status=DocumentStatus.APPROVED,
        )
        _redis_client.add_assistant_task(tid, task.id, task.to_redis())

        return {
            "success": True,
            "task_id": task.id,
            "analysis_type": analysis_type,
            "analysis": analysis,
        }

    except Exception as e:
        logger.error(f"Analysis error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/assistant/generate")
async def assistant_generate_document(request: Request):
    """Genere un document professionnel"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    doc_type = data.get("type", "email")
    subject = data.get("subject", "")
    recipient = data.get("recipient", "")
    context = data.get("context", "")
    tone = data.get("tone", "professional")
    additional = data.get("additional_instructions", "")

    # Recuperer le profil pour personnalisation
    mgr = _get_tenant_manager(tid)
    profile = {}
    if mgr:
        p = mgr.get_subscriber_profile()
        if p:
            profile = p.model_dump()

    sender_name = f"{profile.get('first_name', '')} {profile.get('last_name', '')}".strip() or "L'expediteur"

    # Construire le prompt selon le type
    tone_map = {
        "professional": "professionnel et courtois",
        "formal": "tres formel et respectueux",
        "friendly": "amical mais professionnel",
        "casual": "decontracte",
    }
    tone_desc = tone_map.get(tone, "professionnel")

    if doc_type == "email":
        prompt = f"""Redige un email {tone_desc} en francais.
Sujet: {subject}
Destinataire: {recipient}
Contexte: {context}
{f'Instructions supplementaires: {additional}' if additional else ''}

Fournis:
1. L'objet de l'email (court)
2. Le corps de l'email (avec salutation et signature de {sender_name})"""

    elif doc_type == "cover_letter":
        prompt = f"""Redige une lettre de motivation {tone_desc} en francais.
Poste vise: {subject}
Entreprise: {recipient}
Contexte du candidat: {context}
{f'Instructions: {additional}' if additional else ''}

La lettre doit etre percutante et personnalisee."""

    elif doc_type == "cv":
        prompt = f"""Aide a rediger/ameliorer un CV en francais.
Poste cible: {subject}
Informations fournies: {context}
{f'Instructions: {additional}' if additional else ''}

Fournis un CV structure avec: Coordonnees, Profil, Experiences, Formation, Competences."""

    elif doc_type in ADMIN_TEMPLATES:
        template = ADMIN_TEMPLATES[doc_type]
        prompt = f"""Redige un(e) {template['name']} {tone_desc} en francais.
Destinataire: {template['recipient']}
Objet: {subject}
Details: {context}
{f'Instructions: {additional}' if additional else ''}

Utilise un format officiel et respectueux."""

    else:
        prompt = f"""Redige un document {tone_desc} en francais.
Type: {doc_type}
Sujet: {subject}
Destinataire: {recipient}
Contexte: {context}
{f'Instructions: {additional}' if additional else ''}"""

    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Tu es un assistant de redaction professionnelle. Tu rediges des documents clairs, bien structures et adaptes au contexte."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.7,
        )
        await _track_openai_cost(response, tid)
        generated_text = response.choices[0].message.content

        # Sauvegarder la tache
        task = AssistantTask(
            tenant_id=tid,
            task_type="generate",
            document_type=DocumentType(doc_type) if doc_type in [e.value for e in DocumentType] else None,
            subject=subject,
            recipient=recipient,
            tone=tone,
            additional_instructions=additional,
            output_text=generated_text,
            status=DocumentStatus.DRAFT,
        )
        _redis_client.add_assistant_task(tid, task.id, task.to_redis())

        return {
            "success": True,
            "task_id": task.id,
            "document_type": doc_type,
            "version": 1,
            "status": "draft",
            "content": generated_text,
            "message": "Document genere. Utilisez /api/assistant/improve pour l'ameliorer.",
        }

    except Exception as e:
        logger.error(f"Generate error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/assistant/improve")
async def improve_document(request: Request):
    """Ameliore un document existant avec feedback"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    task_id = data.get("task_id")
    feedback = data.get("feedback", "").strip()
    content = data.get("content")  # Optionnel: contenu modifie manuellement

    if not task_id:
        return JSONResponse(status_code=400, content={"error": "task_id requis"})

    if not feedback and not content:
        return JSONResponse(status_code=400, content={"error": "feedback ou content requis"})

    # Recuperer la tache originale
    original = _redis_client.get_assistant_task(tid, task_id)
    if not original:
        return JSONResponse(status_code=404, content={"error": "Tache non trouvee"})

    original_content = content or original.get("output_text", "")
    original_version = int(original.get("version", 1))

    prompt = f"""Voici un document a ameliorer:

---
{original_content}
---

Feedback de l'utilisateur: {feedback}

Ameliore le document en tenant compte du feedback. Garde le meme format et le meme ton general, sauf si le feedback demande un changement de ton."""

    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": "Tu es un assistant de redaction. Tu ameliores les documents selon le feedback sans changer leur essence."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=2000,
            temperature=0.5,
        )
        await _track_openai_cost(response, tid)
        improved_text = response.choices[0].message.content

        # Creer une nouvelle version
        new_task = AssistantTask(
            tenant_id=tid,
            task_type="improve",
            document_type=DocumentType(original["document_type"]) if original.get("document_type") else None,
            subject=original.get("subject"),
            recipient=original.get("recipient"),
            tone=original.get("tone", "professional"),
            input_text=original_content,
            output_text=improved_text,
            feedback=feedback,
            version=original_version + 1,
            parent_task_id=task_id,
            status=DocumentStatus.DRAFT,
        )
        _redis_client.add_assistant_task(tid, new_task.id, new_task.to_redis())

        return {
            "success": True,
            "task_id": new_task.id,
            "parent_task_id": task_id,
            "version": new_task.version,
            "status": "draft",
            "content": improved_text,
            "feedback_applied": feedback,
            "message": f"Version {new_task.version} generee.",
        }

    except Exception as e:
        logger.error(f"Improve error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.get("/api/assistant/task/{task_id}")
async def get_task(task_id: str, request: Request):
    """Recupere une tache et son contenu"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    task = _redis_client.get_assistant_task(tid, task_id)
    if not task:
        return JSONResponse(status_code=404, content={"error": "Tache non trouvee"})

    return {
        "task": task,
        "content": task.get("output_text") or task.get("input_text"),
    }


@app.get("/api/assistant/task/{task_id}/versions")
async def get_task_versions(task_id: str, request: Request):
    """Recupere toutes les versions d'un document"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    versions = _redis_client.get_task_versions(tid, task_id)

    return {
        "original_task_id": task_id,
        "versions": versions,
        "count": len(versions),
    }


@app.post("/api/assistant/task/{task_id}/approve")
async def approve_task(task_id: str, request: Request):
    """Approuve un document (pret a envoyer/utiliser)"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    task = _redis_client.get_assistant_task(tid, task_id)
    if not task:
        return JSONResponse(status_code=404, content={"error": "Tache non trouvee"})

    _redis_client.update_assistant_task(tid, task_id, {
        "status": DocumentStatus.APPROVED.value,
        "updated_at": datetime.utcnow().isoformat(),
    })

    return {
        "success": True,
        "task_id": task_id,
        "status": "approved",
        "message": "Document approuve et pret a l'emploi.",
    }


@app.get("/api/assistant/tasks")
async def list_tasks(request: Request, limit: int = 20, doc_type: str = None):
    """Liste les dernieres taches"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)

    tasks = _redis_client.get_assistant_tasks(tid, limit=limit)

    if doc_type:
        tasks = [t for t in tasks if t.get("document_type") == doc_type]

    return {
        "tasks": tasks,
        "count": len(tasks),
    }


@app.get("/api/assistant/templates")
async def list_templates():
    """Liste les templates administratifs disponibles"""
    return {
        "templates": ADMIN_TEMPLATES,
        "document_types": [{"id": t.value, "name": t.value.replace("_", " ").title()} for t in DocumentType],
    }


# --- Email sending (placeholder - needs SMTP config) ---

@app.post("/api/email/draft")
async def create_email_draft(request: Request):
    """Cree un brouillon email a partir d'une tache"""
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    tid = getattr(request.state, "tenant_id", 1)
    data = await request.json()
    task_id = data.get("task_id")
    to = data.get("to", [])
    cc = data.get("cc", [])
    subject = data.get("subject", "")

    if not task_id and not data.get("body"):
        return JSONResponse(status_code=400, content={"error": "task_id ou body requis"})

    body = data.get("body", "")
    if task_id:
        task = _redis_client.get_assistant_task(tid, task_id)
        if task:
            body = task.get("output_text", "")
            subject = subject or task.get("subject", "")

    draft = EmailDraft(
        tenant_id=tid,
        to=to if isinstance(to, list) else [to],
        cc=cc if isinstance(cc, list) else [cc] if cc else [],
        subject=subject,
        body_html=body.replace("\n", "<br>"),
        body_text=body,
        task_id=task_id,
    )
    _redis_client.save_email_draft(tid, draft.id, draft.to_redis())

    return {
        "success": True,
        "draft_id": draft.id,
        "subject": subject,
        "to": draft.to,
        "message": "Brouillon cree. Configurez SMTP pour envoyer.",
    }


@app.post("/api/email/send")
async def send_email_endpoint(request: Request):
    """Envoie un email via Gmail OAuth (prioritaire) ou SendGrid (fallback)"""
    tid = getattr(request.state, "tenant_id", 1)

    data = await request.json()
    to = data.get("to", "")
    subject = data.get("subject", "")
    body = data.get("body", "")
    draft_id = data.get("draft_id")

    # Si draft_id fourni, charge le brouillon
    if draft_id and _redis_client:
        draft_data = _redis_client.get_email_draft(tid, draft_id)
        if draft_data:
            to = to or (draft_data.get("to", "") if isinstance(draft_data.get("to"), str) else ",".join(draft_data.get("to", [])))
            subject = subject or draft_data.get("subject", "")
            body = body or draft_data.get("body_text", "")

    if not to or not body:
        return JSONResponse(status_code=400, content={"error": "Destinataire (to) et contenu (body) requis"})

    success, details = await email_client.send_for_tenant(
        tenant_id=tid,
        redis_client=_redis_client,
        gmail_client=gmail_client,
        to=to,
        subject=subject or "Message via Luna",
        body_text=body,
        subscriber_name=_SUBSCRIBER_NAME,
    )

    if success:
        return {"success": True, "to": to, "subject": subject, "message": "Email envoye", "via": details.get("from", "sendgrid")}
    return JSONResponse(status_code=502, content={"error": "Echec envoi", "details": details})


# =========================================================================
# GMAIL OAUTH2 ENDPOINTS (per-tenant email)
# =========================================================================

@app.get("/api/email/oauth/start")
async def gmail_oauth_start(request: Request):
    """Demarre le flow OAuth2 Gmail pour un tenant. Admin only."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Admin requis"})
    if not gmail_client.is_configured:
        return JSONResponse(status_code=501, content={
            "error": "Gmail OAuth non configure",
            "message": "Ajoutez GOOGLE_OAUTH_CLIENT_ID et GOOGLE_OAUTH_CLIENT_SECRET aux env vars.",
        })

    tenant_id = request.query_params.get("tenant_id")
    if not tenant_id:
        return JSONResponse(status_code=400, content={"error": "tenant_id requis en query param"})

    auth_url = gmail_client.get_auth_url(int(tenant_id))
    # Redirige directement vers Google
    from starlette.responses import RedirectResponse
    return RedirectResponse(url=auth_url)


@app.get("/api/email/oauth/callback")
async def gmail_oauth_callback(request: Request):
    """Callback OAuth2 Google. Stocke les tokens dans Redis."""
    code = request.query_params.get("code", "")
    state = request.query_params.get("state", "")
    error = request.query_params.get("error", "")

    if error:
        return JSONResponse(content={
            "success": False,
            "error": f"Google OAuth refuse: {error}",
        })

    if not code or not state:
        return JSONResponse(status_code=400, content={"error": "code et state requis"})

    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    result = await gmail_client.handle_callback(code, state, _redis_client)

    from html import escape as _html_escape
    if result.get("success"):
        # Affiche une page de succes simple
        email = _html_escape(str(result.get("email", "")))
        tid = _html_escape(str(result.get("tenant_id", "?")))
        html = f"""<!DOCTYPE html><html><head><title>Luna - Gmail connecte</title>
        <style>body{{font-family:sans-serif;display:flex;justify-content:center;align-items:center;height:100vh;background:#0a0a1a;color:#e0e0e0;}}
        .card{{background:#1e1e3f;padding:40px;border-radius:16px;text-align:center;}}
        .ok{{color:#4ade80;font-size:2em;}}h2{{color:#a78bfa;}}</style></head>
        <body><div class="card"><div class="ok">&#10003;</div>
        <h2>Gmail connecte !</h2><p>Compte: <strong>{email}</strong></p>
        <p>Tenant ID: {tid}</p>
        <p>Luna peut maintenant envoyer des emails depuis ce compte.</p>
        <p><a href="/admin" style="color:#a78bfa;">Retour au dashboard</a></p>
        </div></body></html>"""
        return HTMLResponse(content=html)
    else:
        error_msg = _html_escape(str(result.get("error", "Erreur inconnue")))
        html = f"""<!DOCTYPE html><html><head><title>Luna - Erreur Gmail</title>
        <style>body{{font-family:sans-serif;display:flex;justify-content:center;align-items:center;height:100vh;background:#0a0a1a;color:#e0e0e0;}}
        .card{{background:#1e1e3f;padding:40px;border-radius:16px;text-align:center;}}
        .err{{color:#f87171;font-size:2em;}}h2{{color:#f87171;}}</style></head>
        <body><div class="card"><div class="err">&#10007;</div>
        <h2>Erreur connexion Gmail</h2><p>{error_msg}</p>
        <p><a href="/admin" style="color:#a78bfa;">Retour au dashboard</a></p>
        </div></body></html>"""
        return HTMLResponse(content=html)


@app.get("/api/admin/clients/{tenant_id}/email")
async def admin_get_email_integration(tenant_id: int, request: Request):
    """Statut de l'integration email d'un tenant."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Admin requis"})
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    integration = _redis_client.get_email_integration(tenant_id)
    if not integration:
        return {"connected": False, "service": None}

    return {
        "connected": True,
        "service": integration.get("service", "unknown"),
        "email": integration.get("email", ""),
        "connected_at": integration.get("connected_at", ""),
    }


@app.delete("/api/admin/clients/{tenant_id}/email")
async def admin_delete_email_integration(tenant_id: int, request: Request):
    """Deconnecte Gmail d'un tenant."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Admin requis"})
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    _redis_client.delete_email_integration(tenant_id)
    logger.info(f"Email integration deleted: tenant={tenant_id}")
    return {"success": True, "message": f"Integration email supprimee pour tenant {tenant_id}"}


# =========================================================================
# ADMIN DASHBOARD ENDPOINTS
# =========================================================================

_ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD", "")
_server_start_time = time.time()


def _create_admin_token() -> str:
    """Cree un JWT admin valide 24h."""
    import jwt as pyjwt
    payload = {
        "role": "admin",
        "iat": int(time.time()),
        "exp": int(time.time()) + 86400,
    }
    return pyjwt.encode(payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)


def _verify_admin(request: Request) -> bool:
    """Verifie le token admin dans le header Authorization."""
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        return False
    token = auth[7:]
    try:
        import jwt as pyjwt
        payload = pyjwt.decode(token, _JWT_SECRET, algorithms=[_JWT_ALGORITHM])
        return payload.get("role") == "admin"
    except Exception:
        return False


# --- Certificat d'autonomie ---
_certificate_timestamp: float = 0  # timestamp de generation du certificat

def _generate_autonomy_certificate(pv_data: dict, reset_code: str) -> str:
    """Genere un certificat d'autonomie DOCX apres signature PV."""
    global _certificate_timestamp
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx non installe")
        return None

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Titre
    title = doc.add_heading("CERTIFICAT D'AUTONOMIE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("YAWatch Luna")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph("")

    # Identite exploitant
    doc.add_heading("Identite Exploitant", level=2)
    doc.add_paragraph(f"Nom / Raison sociale : {pv_data.get('exploitant_name', 'N/A')}")
    doc.add_paragraph(f"SIRET : {pv_data.get('exploitant_siret', 'N/A')}")

    try:
        from datetime import timezone as _tz
        sig_date = pv_data.get("date_signature", "")
        doc.add_paragraph(f"Date de signature : {sig_date}")
    except Exception:
        doc.add_paragraph(f"Date de signature : {pv_data.get('date_signature', 'N/A')}")

    doc.add_paragraph(f"Mode Luna : {pv_data.get('luna_mode', 'N/A')}")
    doc.add_paragraph(f"Version PV : {pv_data.get('version', 'N/A')}")

    # Resultats des phases
    doc.add_heading("Resultats de Recette", level=2)
    for phase_key, phase_label in [("phase_a", "Phase A - Technique"), ("phase_b", "Phase B - Legal"), ("phase_c", "Phase C - Operationnel")]:
        phase = pv_data.get(phase_key, {})
        passed = phase.get("all_passed", False)
        doc.add_paragraph(f"{phase_label} : {'[OK]' if passed else '[ECHEC]'}")

    # Hash PV
    doc.add_heading("Signature Numerique", level=2)
    doc.add_paragraph(f"SHA-256 : {pv_data.get('signature_hash', 'N/A')}")

    # RESET CODE
    doc.add_heading("CODE DE REINITIALISATION", level=2)
    warning = doc.add_paragraph("CONSERVEZ CE CODE EN LIEU SUR — Il ne sera plus jamais affiche.")
    warning.runs[0].bold = True
    warning.runs[0].font.color.rgb = RGBColor(204, 0, 0)

    code_para = doc.add_paragraph()
    code_run = code_para.add_run(reset_code)
    code_run.bold = True
    code_run.font.size = Pt(16)
    code_run.font.name = "Courier New"
    code_run.font.color.rgb = RGBColor(204, 0, 0)
    code_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    doc.add_paragraph(
        "Ce code permet de reinitialiser l'instance Luna en cas de besoin. "
        "Sans ce code, la reinitialisation est impossible. "
        "Commande : python tools/factory_reset.py --code <VOTRE_CODE>"
    )

    # Autonomie operationnelle
    doc.add_heading("Autonomie Operationnelle", level=2)
    doc.add_paragraph(
        "Ce certificat atteste que l'instance Luna est desormais autonome. "
        "Le fondateur n'a plus acces technique a cette instance. "
        "L'exploitant gere seul ses comptes API, son serveur et ses clients."
    )

    # Mise a jour
    doc.add_heading("Procedure de Mise a Jour", level=2)
    doc.add_paragraph("docker compose pull && docker compose up -d")
    doc.add_paragraph("Les mises a jour preservent toutes les donnees et la configuration.")

    # Mentions legales
    doc.add_heading("Mentions Legales", level=2)
    doc.add_paragraph(
        "Luna est un compagnon de lien social, PAS un dispositif medical. "
        "La detection de detresse est best-effort uniquement. "
        "L'exploitant doit recommander une teleassistance certifiee en complement."
    )

    # Sauvegarde
    cert_dir = Path(os.path.dirname(__file__)) / "data" / "certificates"
    cert_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"certificat_autonomie_{timestamp}.docx"
    cert_path = cert_dir / filename
    doc.save(str(cert_path))
    _certificate_timestamp = time.time()
    logger.info(f"Certificat d'autonomie genere: {cert_path}")
    return str(cert_path)


@app.get("/api/admin/certificate")
async def admin_certificate(request: Request):
    """Telecharge le certificat d'autonomie DOCX."""
    # Grace period: 10 minutes apres signature, accessible sans auth
    grace_period = _certificate_timestamp > 0 and (time.time() - _certificate_timestamp) < 600
    if not grace_period and not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    cert_dir = Path(os.path.dirname(__file__)) / "data" / "certificates"
    if not cert_dir.exists():
        return JSONResponse(status_code=404, content={"error": "Aucun certificat genere"})

    # Trouver le plus recent
    certs = sorted(cert_dir.glob("certificat_autonomie_*.docx"), reverse=True)
    if not certs:
        return JSONResponse(status_code=404, content={"error": "Aucun certificat genere"})

    return FileResponse(
        path=str(certs[0]),
        filename=certs[0].name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@app.post("/api/admin/login")
async def admin_login(request: Request):
    """Login admin avec mot de passe."""
    data = await request.json()
    password = data.get("password", "")

    if not _ADMIN_PASSWORD:
        return JSONResponse(status_code=503, content={
            "error": "ADMIN_PASSWORD non configure dans .env",
        })

    import hmac
    if not hmac.compare_digest(password, _ADMIN_PASSWORD):
        return JSONResponse(status_code=401, content={"error": "Mot de passe incorrect"})

    token = _create_admin_token()
    return {"token": token, "expires_in": 86400}


@app.get("/api/admin/dashboard")
async def admin_dashboard(request: Request):
    """Vue d'ensemble du dashboard admin."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    uptime = time.time() - _server_start_time
    hours = int(uptime // 3600)
    minutes = int((uptime % 3600) // 60)

    result = {
        "server": {
            "status": "running",
            "uptime_seconds": int(uptime),
            "uptime_human": f"{hours}h {minutes}m",
            "luna_mode": LUNA_MODE,
            "pv_signed": PV_SIGNED,
        },
        "services": {
            "openai": bool(OPENAI_API_KEY),
            "twilio": bool(sms_client and hasattr(sms_client, 'is_configured') and sms_client.is_configured),
            "tavus": bool(tavus_client and tavus_client.is_configured) if tavus_client else False,
            "redis": bool(_redis_client and _redis_client.ping()) if _redis_client else False,
        },
        "clients_count": 0,
        "alerts_today": 0,
    }

    if _memory_manager:
        try:
            result["alerts_today"] = len(_memory_manager.get_alerts(limit=100, category="safety"))
        except Exception:
            pass
    if _redis_client:
        try:
            result["clients_count"] = len(_redis_client.get_all_tenant_ids())
        except Exception:
            pass

    return result


@app.get("/api/admin/clients")
async def admin_clients(request: Request):
    """Liste des clients (tenants) — lecture directe Redis, sans MemoryManager."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    clients = []
    if _redis_client:
        try:
            tenant_ids = _redis_client.get_all_tenant_ids()
            for tid in tenant_ids:
                profile = _redis_client.get_profile(tid) or {}
                fn = profile.get("first_name", "")
                ln = profile.get("last_name", "")
                name = f"{fn} {ln}".strip() or f"Tenant {tid}"
                plan = profile.get("plan", "essentiel")
                # Lecture auth pour l'email
                auth = _redis_client.get_auth_by_tenant_id(tid)
                email = auth.get("email", "") if auth else ""
                clients.append({
                    "tenant_id": tid,
                    "name": name,
                    "email": email,
                    "plan": plan,
                })
        except Exception as e:
            logger.error(f"Admin clients error: {e}")

    return {"clients": clients, "total": len(clients)}


@app.get("/api/admin/clients/{tenant_id}")
async def admin_client_detail(tenant_id: int, request: Request):
    """Detail d'un client."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    if not _redis_client or not _CORE_AVAILABLE:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    try:
        mm = MemoryManager(tenant_id=tenant_id, redis_client=_redis_client)
        profile = mm.get_subscriber_profile()
        quota = mm.get_quota_status()
        contacts = mm.list_trusted_contacts()
        stats = mm.get_daily_stats_range(7)
        return {
            "profile": profile.dict() if profile else None,
            "quota": quota,
            "contacts": [c.dict() for c in contacts] if contacts else [],
            "daily_stats": stats,
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/admin/clients")
async def admin_create_client(request: Request):
    """Admin: cree un client manuellement."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    data = await request.json()
    email = data.get("email", "").strip().lower()
    password = data.get("password", "")
    first_name = data.get("first_name", "")
    last_name = data.get("last_name", "")
    plan = data.get("plan", "essentiel").lower()

    if not email or "@" not in email:
        return JSONResponse(status_code=400, content={"error": "Email invalide"})
    if len(password) < 6:
        return JSONResponse(status_code=400, content={"error": "Mot de passe trop court (min 6)"})
    if plan not in ("essentiel", "confort", "premium"):
        return JSONResponse(status_code=400, content={"error": "Plan invalide"})

    if _redis_client.get_auth_by_email(email):
        return JSONResponse(status_code=409, content={"error": "Email deja utilise"})

    tenant_id = _redis_client.get_next_tenant_id()
    password_hash = _hash_password(password)

    created = _redis_client.create_auth_record(email, password_hash, tenant_id, plan)
    if not created:
        return JSONResponse(status_code=409, content={"error": "Email deja utilise"})

    # Create profile
    if _CORE_AVAILABLE:
        profile = SubscriberProfile(
            tenant_id=tenant_id,
            first_name=first_name or email.split("@")[0],
            last_name=last_name,
            email=email,
        )
        mgr = _get_tenant_manager(tenant_id)
        mgr.set_subscriber_profile(profile)

    logger.info(f"ADMIN_CREATE_CLIENT tenant_id={tenant_id} email={email} plan={plan}")
    _gamify("admin", "new_client", is_admin=True)

    return {"success": True, "tenant_id": tenant_id, "email": email, "plan": plan}


@app.patch("/api/admin/clients/{tenant_id}")
async def admin_update_client(tenant_id: int, request: Request):
    """Admin: modifie le plan ou le statut d'un client."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    auth = _redis_client.get_auth_by_tenant_id(tenant_id)
    if not auth:
        return JSONResponse(status_code=404, content={"error": "Client introuvable"})

    data = await request.json()
    updates = {}

    if "plan" in data:
        plan = data["plan"].lower()
        if plan not in ("essentiel", "confort", "premium", "fondateur"):
            return JSONResponse(status_code=400, content={"error": "Plan invalide"})
        updates["plan"] = plan
        # Evict from tenant manager cache
        if tenant_id in _tenant_managers:
            del _tenant_managers[tenant_id]

    if "active" in data:
        updates["active"] = bool(data["active"])

    if not updates:
        return JSONResponse(status_code=400, content={"error": "Aucune modification"})

    email = auth.get("email", "")
    _redis_client.update_auth_record(email, updates)
    logger.info(f"ADMIN_UPDATE_CLIENT tenant_id={tenant_id} updates={updates}")
    return {"success": True, "tenant_id": tenant_id, "updates": updates}


@app.delete("/api/admin/clients/{tenant_id}")
async def admin_delete_client(tenant_id: int, request: Request):
    """Admin: supprime un client (purge complete de toutes ses donnees Redis)."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    if tenant_id == _PROPRIO_TENANT_ID:
        return JSONResponse(status_code=403, content={"error": "Impossible de supprimer le compte fondateur"})

    # Supprimer l'auth record
    auth = _redis_client.get_auth_by_tenant_id(tenant_id)
    if auth:
        email = auth.get("email", "")
        auth_key = f"{_redis_client.prefix}:auth:{email}"
        _redis_client.client.delete(auth_key)

    # Purger toutes les cles du tenant
    keys_deleted = _redis_client.purge_tenant(tenant_id)
    logger.info(f"ADMIN_DELETE_CLIENT tenant_id={tenant_id} keys_deleted={keys_deleted}")
    return {"success": True, "tenant_id": tenant_id, "keys_deleted": keys_deleted}


@app.post("/api/admin/reset-password/{tenant_id}")
async def admin_reset_password(tenant_id: int, request: Request):
    """Admin: reset le mot de passe d'un client."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Service temporairement indisponible"})

    body = await request.json()
    new_password = body.get("password", "").strip()
    if len(new_password) < 6:
        return JSONResponse(status_code=400, content={"error": "Mot de passe trop court (min 6 caracteres)"})

    auth = _redis_client.get_auth_by_tenant_id(tenant_id)
    if not auth:
        return JSONResponse(status_code=404, content={"error": f"Tenant {tenant_id} introuvable"})

    email = auth.get("email", "")
    new_hash = _hash_password(new_password)
    _redis_client.update_auth_record(email, {"password_hash": new_hash})
    logger.info(f"ADMIN_RESET_PASSWORD tenant_id={tenant_id} email={email}")
    return {"success": True, "tenant_id": tenant_id, "email": email}


@app.get("/api/admin/quotas")
async def admin_quotas(request: Request):
    """Quotas de tous les clients — lecture legere + usage reel Cortex."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    quotas = []
    if _redis_client:
        try:
            # Usage reel Cortex (un seul appel pour tous les tenants)
            all_usage = {}
            cortex = get_cortex() if _CORTEX_AVAILABLE else None
            if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
                all_usage = await cortex.cost_tracker.get_month_costs_per_tenant()

            for tid in _redis_client.get_all_tenant_ids():
                profile = _redis_client.get_profile(tid) or {}
                name = profile.get("first_name", f"Tenant {tid}")
                plan = profile.get("plan", "essentiel")
                limits = _PLAN_LIMITS.get(plan, _PLAN_LIMITS["essentiel"])
                usage = all_usage.get(str(tid), {"sms_count": 0, "sms_cost": 0, "voice_minutes": 0, "voice_cost": 0, "tavus_minutes": 0, "tavus_cost": 0})
                sms_cost = round(float(usage.get("sms_cost", 0)), 2)
                voice_cost = round(float(usage.get("voice_cost", 0)), 2)
                visio_cost = round(float(usage.get("tavus_cost", 0)), 2)
                total_cost = round(sms_cost + voice_cost + visio_cost, 2)
                quotas.append({
                    "tenant_id": tid,
                    "name": name,
                    "plan": plan,
                    "sms": {"used": usage.get("sms_count", 0), "limit": limits["sms"], "cost_eur": sms_cost},
                    "voice": {"used": round(usage.get("voice_minutes", 0), 1), "limit": limits["voice_min"], "cost_eur": voice_cost},
                    "visio": {"used": round(usage.get("tavus_minutes", 0), 1), "limit": limits["visio_min"], "cost_eur": visio_cost},
                    "total_cost_eur": total_cost,
                    "budget_max_eur": limits.get("budget_api_max", 0),
                })
        except Exception as e:
            logger.error(f"Admin quotas error: {e}")

    return {"quotas": quotas}


@app.get("/api/admin/costs")
async def admin_costs(request: Request, month: str = None):
    """
    Couts API reels par tenant pour un mois donne.
    Utilise les donnees tracees par CortexCostTracker dans Redis.
    ?month=2026-03 (defaut: mois courant)
    """
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    from datetime import date as _date
    target_date = None
    if month:
        try:
            parts = month.split("-")
            target_date = _date(int(parts[0]), int(parts[1]), 1)
        except (ValueError, IndexError):
            return JSONResponse(status_code=400, content={"error": "Format mois invalide (YYYY-MM)"})

    result = {"month": month or _date.today().strftime("%Y-%m"), "tenants": {}, "totals": {}}

    cortex = get_cortex() if _CORTEX_AVAILABLE else None
    if cortex and hasattr(cortex, "cost_tracker") and cortex.cost_tracker:
        try:
            all_costs = await cortex.cost_tracker.get_month_costs_per_tenant(target_date)
            total_sms_cost = 0
            total_voice_cost = 0
            total_visio_cost = 0
            total_sms_count = 0

            for tid_str, costs in all_costs.items():
                tid = int(tid_str) if tid_str.isdigit() else tid_str
                name = f"Tenant {tid}"
                plan = "essentiel"
                if _redis_client:
                    profile = _redis_client.get_profile(int(tid_str)) if tid_str.isdigit() else {}
                    if profile:
                        name = profile.get("first_name", name)
                        plan = profile.get("plan", plan)

                sc = round(float(costs.get("sms_cost", 0)), 2)
                vc = round(float(costs.get("voice_cost", 0)), 2)
                tc = round(float(costs.get("tavus_cost", 0)), 2)
                total = round(sc + vc + tc, 2)

                result["tenants"][tid_str] = {
                    "name": name,
                    "plan": plan,
                    "sms": {"count": int(costs.get("sms_count", 0)), "cost_eur": sc},
                    "voice": {"minutes": round(float(costs.get("voice_minutes", 0)), 1), "cost_eur": vc},
                    "visio": {"minutes": round(float(costs.get("tavus_minutes", 0)), 1), "cost_eur": tc},
                    "total_cost_eur": total,
                    "budget_max_eur": _PLAN_LIMITS.get(plan, {}).get("budget_api_max", 0),
                }

                total_sms_cost += sc
                total_voice_cost += vc
                total_visio_cost += tc
                total_sms_count += int(costs.get("sms_count", 0))

            result["totals"] = {
                "sms_cost_eur": round(total_sms_cost, 2),
                "voice_cost_eur": round(total_voice_cost, 2),
                "visio_cost_eur": round(total_visio_cost, 2),
                "total_cost_eur": round(total_sms_cost + total_voice_cost + total_visio_cost, 2),
                "sms_count": total_sms_count,
            }
        except Exception as e:
            logger.error(f"Admin costs error: {e}")
            result["error"] = str(e)

    # Tarifs de reference utilises
    result["rates"] = {
        "sms_eur": 0.07,
        "voice_eur_per_min": 0.02,
        "visio_eur_per_min": 0.05,
        "note": "Tarifs estimes Twilio FR + Tavus. Verifier avec factures reelles.",
    }

    return result


@app.get("/api/admin/alerts")
async def admin_alerts(request: Request, limit: int = 50, category: str = None):
    """Evenements et alertes."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    if not _memory_manager:
        return {"alerts": [], "counts": {}}

    alerts = _memory_manager.get_alerts(limit=limit, category=category)

    # Compter par categorie
    counts = {}
    for a in alerts:
        cat = a.get("category", "unknown")
        counts[cat] = counts.get(cat, 0) + 1

    return {"alerts": alerts, "counts": counts, "total": len(alerts)}


@app.get("/api/admin/checkin-pending")
async def admin_checkin_pending(request: Request):
    """Checkins Raven en attente de confirmation (admin)."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})
    if not _redis_client:
        return {"pending_count": 0, "pending": []}
    pattern = f"luna:*:checkin_pending:*"
    pending = []
    cursor = 0
    while True:
        cursor, keys = _redis_client.client.scan(cursor, match=pattern, count=100)
        for key in keys:
            raw = _redis_client.client.get(key)
            if not raw:
                continue
            try:
                data = json.loads(raw)
                if data.get("status") == "pending":
                    created = datetime.fromisoformat(data["timestamp"])
                    elapsed = (datetime.utcnow() - created).total_seconds()
                    pending.append({
                        "conversation_id": data.get("conversation_id"),
                        "observation": data.get("observation"),
                        "severity": data.get("severity"),
                        "tenant_id": data.get("tenant_id"),
                        "elapsed_seconds": int(elapsed),
                        "remaining_seconds": max(0, int(45 - elapsed)),
                        "timestamp": data.get("timestamp"),
                    })
            except Exception:
                pass
        if cursor == 0:
            break
    return {"pending_count": len(pending), "pending": pending}


@app.get("/api/admin/subscriber-location")
async def admin_subscriber_location(request: Request):
    """Dernière position GPS connue du souscripteur (admin)."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})
    if not _redis_client:
        return {"available": False, "message": "Redis non disponible"}
    raw = _redis_client.client.get(f"luna:{TENANT_ID}:geolocation")
    if not raw:
        return {"available": False, "message": "Position non disponible"}
    try:
        data = json.loads(raw)
        ts_str = data.get("timestamp", "")
        try:
            ts = datetime.fromisoformat(ts_str)
            age = int((datetime.utcnow() - ts).total_seconds())
        except Exception:
            age = -1
        return {
            "available": True,
            "latitude": data.get("latitude"),
            "longitude": data.get("longitude"),
            "accuracy": data.get("accuracy"),
            "city": data.get("city", ""),
            "address": data.get("address", ""),
            "timestamp": ts_str,
            "age_seconds": age,
            "is_fresh": age >= 0 and age < 300,
        }
    except Exception as e:
        return {"available": False, "message": f"Erreur lecture: {e}"}


@app.get("/api/settings/auto-note")
async def get_auto_note(request: Request):
    """Retourne l'état de la prise de note automatique en visio."""
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return {"enabled": True}
    return {"enabled": mgr.is_auto_note_enabled()}


@app.post("/api/settings/auto-note")
async def set_auto_note(request: Request):
    """Active ou désactive la prise de note automatique en visio."""
    tid = getattr(request.state, "tenant_id", TENANT_ID)
    mgr = _get_tenant_manager(tid) if tid else _memory_manager
    if not mgr:
        return JSONResponse(status_code=503, content={"error": "Service non disponible"})
    try:
        body = await request.json()
    except Exception:
        body = {}
    enabled = bool(body.get("enabled", True))
    mgr.set_auto_note_enabled(enabled)
    return {
        "status": "success",
        "enabled": enabled,
        "message": f"Prise de notes automatique {'activee' if enabled else 'desactivee'}",
    }


@app.get("/api/admin/health")
async def admin_health(request: Request):
    """Sante detaillee du serveur."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    uptime = time.time() - _server_start_time
    result = {
        "uptime_seconds": int(uptime),
        "luna_mode": LUNA_MODE,
        "pv_signed": PV_SIGNED,
        "services": {
            "openai": {"status": "ok" if OPENAI_API_KEY else "missing", "model": OPENAI_MODEL},
            "twilio": {"status": "ok" if (sms_client and hasattr(sms_client, 'is_configured') and sms_client.is_configured) else "not_configured"},
            "tavus": {"status": "ok" if (tavus_client and tavus_client.is_configured) else "not_configured"} if LUNA_MODE == "full" else {"status": "skipped"},
            "redis": {"status": "offline"},
        },
        "core_modules": {
            "memory": bool(_memory_manager),
            "safety": bool(_safety_guardian),
            "quota": bool(_quota_guard),
            "scheduler": bool(_scheduler),
            "executor": bool(_executor),
        },
    }

    if _redis_client:
        try:
            result["services"]["redis"] = _redis_client.get_info()
        except Exception:
            pass

    return result


@app.get("/api/admin/revenue")
async def admin_revenue(request: Request):
    """Estimation CA — lecture directe Redis."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})

    plan_prices = {"essentiel": 79, "confort": 149, "premium": 249}
    by_plan = {p: {"count": 0, "revenue": 0} for p in plan_prices}
    total = 0

    if _redis_client:
        try:
            for tid in _redis_client.get_all_tenant_ids():
                profile = _redis_client.get_profile(tid) or {}
                plan = profile.get("plan", "essentiel")
                if plan in by_plan:
                    by_plan[plan]["count"] += 1
                    by_plan[plan]["revenue"] += plan_prices.get(plan, 0)
                    total += plan_prices.get(plan, 0)
        except Exception as e:
            logger.error(f"Admin revenue error: {e}")

    return {"total_revenue": total, "by_plan": by_plan, "currency": "EUR"}


# =========================================================================
# CLIENT DEBUG LOGS — Remote APK monitoring
# =========================================================================
_DEBUG_LOG_KEY = "luna:debug:client_logs"
_DEBUG_LOG_MAX = 500  # max entries kept


@app.post("/api/debug/log")
async def client_debug_log(request: Request):
    """Receive debug logs from client (APK/browser). No auth required for reliability."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse(status_code=400, content={"ok": False})
    if not isinstance(body, dict):
        return JSONResponse(status_code=400, content={"ok": False})

    entry = {
        "ts": datetime.now(ZoneInfo("Europe/Paris")).strftime("%Y-%m-%d %H:%M:%S"),
        "level": body.get("level", "info"),
        "tag": body.get("tag", ""),
        "msg": str(body.get("msg", ""))[:500],
        "data": str(body.get("data", ""))[:1000],
        "ua": str(request.headers.get("user-agent", ""))[:200],
        "tid": getattr(request.state, "tenant_id", None),
    }
    if _redis_client:
        try:
            _redis_client.client.lpush(_DEBUG_LOG_KEY, json.dumps(entry, ensure_ascii=False))
            _redis_client.client.ltrim(_DEBUG_LOG_KEY, 0, _DEBUG_LOG_MAX - 1)
            _redis_client.client.expire(_DEBUG_LOG_KEY, 7 * 86400)  # 7 days
        except Exception:
            pass
    else:
        logger.info(f"[CLIENT-DEBUG] {entry['level']} [{entry['tag']}] {entry['msg']} | {entry['data']}")
    return {"ok": True}


@app.get("/api/admin/debug-logs")
async def admin_debug_logs(request: Request):
    """View client debug logs. Admin only."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})
    limit = int(request.query_params.get("limit", "100"))
    level_filter = request.query_params.get("level", "")  # error, warn, info, chat
    tag_filter = request.query_params.get("tag", "")

    logs = []
    if _redis_client:
        try:
            raw = _redis_client.client.lrange(_DEBUG_LOG_KEY, 0, min(limit, _DEBUG_LOG_MAX) - 1)
            for r in raw:
                try:
                    entry = json.loads(r)
                    if level_filter and entry.get("level") != level_filter:
                        continue
                    if tag_filter and tag_filter not in entry.get("tag", ""):
                        continue
                    logs.append(entry)
                except Exception:
                    pass
        except Exception as e:
            logger.error(f"Debug logs read error: {e}")
    return {"logs": logs, "count": len(logs)}


@app.delete("/api/admin/debug-logs")
async def admin_clear_debug_logs(request: Request):
    """Clear debug logs. Admin only."""
    if not _verify_admin(request):
        return JSONResponse(status_code=401, content={"error": "Non autorise"})
    if _redis_client:
        _redis_client.client.delete(_DEBUG_LOG_KEY)
    return {"ok": True}


# =========================================================================
# THEOCRATIE — Espace spirituel (pionnier permanent, tenant 1 uniquement)
# =========================================================================

_THEO_PIONEER_GOAL = 50  # heures/mois
_THEO_PIONEER_WEEKLY = 13  # ~50/4 heures/semaine


def _theo_month_key(tenant_id: int, month: str = None) -> str:
    if not month:
        month = datetime.now().strftime("%Y-%m")
    return f"luna:{tenant_id}:theo:hours:{month}"


def _theo_entries_key(tenant_id: int, month: str = None) -> str:
    if not month:
        month = datetime.now().strftime("%Y-%m")
    return f"luna:{tenant_id}:theo:entries:{month}"


@app.get("/api/theo/hours")
async def theo_hours(request: Request, month: str = None):
    """Heures de predication du mois (pionnier permanent)."""
    payload = _decode_client_token(_extract_bearer(request))
    tid = payload.get("tenant_id") if payload else None
    if tid != _PROPRIO_TENANT_ID:
        return JSONResponse(status_code=403, content={"error": "Acces reserve"})
    if not _redis_client:
        return {"hours": 0, "goal": _THEO_PIONEER_GOAL, "entries": []}

    if not month:
        month = datetime.now().strftime("%Y-%m")

    total = 0.0
    try:
        val = _redis_client.client.get(_theo_month_key(tid, month))
        if val:
            total = float(val)
    except Exception:
        pass

    entries = []
    try:
        raw = _redis_client.client.lrange(_theo_entries_key(tid, month), 0, -1)
        for r in raw:
            try:
                entries.append(json.loads(r))
            except Exception:
                pass
    except Exception:
        pass

    now = datetime.now()
    day_of_month = now.day
    days_in_month = 30
    try:
        import calendar
        days_in_month = calendar.monthrange(now.year, now.month)[1]
    except Exception:
        pass
    days_left = max(1, days_in_month - day_of_month)
    remaining = max(0, _THEO_PIONEER_GOAL - total)
    daily_needed = round(remaining / days_left, 1) if remaining > 0 else 0

    # Semaine courante
    week_start = now - timedelta(days=now.weekday())
    week_hours = 0.0
    for e in entries:
        try:
            edate = datetime.strptime(e.get("date", ""), "%Y-%m-%d")
            if edate >= week_start:
                week_hours += float(e.get("hours", 0))
        except Exception:
            pass

    status = "on_track"
    if total >= _THEO_PIONEER_GOAL:
        status = "completed"
    elif day_of_month > 20 and total < _THEO_PIONEER_GOAL * 0.5:
        status = "behind"
    elif day_of_month > 10 and total < _THEO_PIONEER_GOAL * 0.25:
        status = "behind"

    advice = None
    if status == "behind":
        alternatives = []
        if remaining > 10:
            alternatives.append("Ecrire des lettres de temoignage (compte dans les heures)")
            alternatives.append("Temoignage par telephone le soir")
            alternatives.append("Temoignage informel au travail, transports")
            alternatives.append("Temoignage public avec presentoir/chariot")
        advice = {
            "message": f"Il te reste {remaining:.0f}h a faire en {days_left} jours. "
                       f"Objectif: {daily_needed}h/jour.",
            "alternatives": alternatives,
        }

    return {
        "hours": round(total, 1),
        "goal": _THEO_PIONEER_GOAL,
        "percentage": round((total / _THEO_PIONEER_GOAL) * 100, 1),
        "remaining": round(remaining, 1),
        "daily_needed": daily_needed,
        "week_hours": round(week_hours, 1),
        "week_goal": _THEO_PIONEER_WEEKLY,
        "status": status,
        "advice": advice,
        "entries": entries[-15:],  # 15 derniers
        "month": month,
    }


@app.post("/api/theo/hours")
async def theo_add_hours(request: Request):
    """Ajoute des heures de predication."""
    payload = _decode_client_token(_extract_bearer(request))
    tid = payload.get("tenant_id") if payload else None
    if tid != _PROPRIO_TENANT_ID:
        return JSONResponse(status_code=403, content={"error": "Acces reserve"})
    if not _redis_client:
        return JSONResponse(status_code=503, content={"error": "Redis non disponible"})

    body = await request.json()
    hours = float(body.get("hours", 0))
    activity = body.get("activity", "predication")
    date_str = body.get("date", datetime.now().strftime("%Y-%m-%d"))
    note = body.get("note", "")

    if hours <= 0 or hours > 24:
        return JSONResponse(status_code=400, content={"error": "Heures invalides (0-24)"})

    month = date_str[:7]  # YYYY-MM
    entry = {
        "date": date_str,
        "hours": hours,
        "activity": activity,
        "note": note,
        "ts": time.time(),
    }

    try:
        _redis_client.client.incrbyfloat(_theo_month_key(tid, month), hours)
        _redis_client.client.lpush(_theo_entries_key(tid, month), json.dumps(entry))
        _redis_client.client.expire(_theo_month_key(tid, month), 365 * 86400)
        _redis_client.client.expire(_theo_entries_key(tid, month), 365 * 86400)
    except Exception as e:
        logger.error(f"Theo hours add error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})

    return {"ok": True, "added": hours, "activity": activity}


@app.post("/api/theo/prepare")
async def theo_prepare_meeting(request: Request):
    """Luna prepare la reunion. Scrape jw.org pour le contenu reel puis OpenAI pour preparer."""
    payload = _decode_client_token(_extract_bearer(request))
    tid = payload.get("tenant_id") if payload else None
    if tid != _PROPRIO_TENANT_ID:
        return JSONResponse(status_code=403, content={"error": "Acces reserve"})
    if not openai_client:
        return JSONResponse(status_code=503, content={"error": "Service IA non disponible"})

    body = await request.json()
    week = body.get("week", "")
    request_type = body.get("type", "midweek")  # midweek, watchtower, bible_reading
    question = body.get("question", "")

    # --- Etape 1: Scraper le contenu reel de jw.org ---
    jw_content = ""
    jw_urls = {
        "midweek": "https://www.jw.org/fr/biblioth%C3%A8que/programme-des-r%C3%A9unions/",
        "watchtower": "https://www.jw.org/fr/biblioth%C3%A8que/la-tour-de-garde-%C3%A9tude/",
        "bible_reading": "https://www.jw.org/fr/biblioth%C3%A8que/programme-de-lecture-de-la-bible/",
    }
    try:
        import httpx
        target_url = jw_urls.get(request_type, jw_urls["midweek"])
        async with httpx.AsyncClient(follow_redirects=True) as http:
            resp = await http.get(
                target_url,
                headers={
                    "User-Agent": "Mozilla/5.0 (Linux; Android 12) AppleWebKit/537.36 Chrome/120.0.0.0 Mobile Safari/537.36",
                    "Accept": "text/html,application/xhtml+xml",
                    "Accept-Language": "fr-FR,fr;q=0.9",
                },
                timeout=15,
            )
            if resp.status_code == 200:
                import re as _re
                _html = resp.text
                _html = _re.sub(r'<(script|style|nav|footer|header|noscript)[^>]*>.*?</\1>', '', _html, flags=_re.DOTALL | _re.IGNORECASE)
                _text = _re.sub(r'<[^>]+>', ' ', _html)
                _text = _re.sub(r'\s+', ' ', _text).strip()
                jw_content = _text[:4000]
                logger.info(f"Theo: scraped jw.org ({len(jw_content)} chars) for {request_type}")
    except Exception as e:
        logger.warning(f"Theo: jw.org scrape failed: {e}")

    # --- Etape 2: Construire le prompt avec le contenu reel ---
    now = datetime.now(ZoneInfo("Europe/Paris"))
    week_str = week or now.strftime("%d %B %Y")

    base_context = (
        "Tu es Luna, assistante spirituelle d'un Temoin de Jehovah pionnier permanent. "
        f"Date: {week_str}. "
        "Tu dois preparer du contenu CONCRET et UTILE pour la participation aux reunions. "
        "Formate avec des titres markdown clairs. Cite les versets bibliques entre parentheses."
    )

    if jw_content:
        base_context += f"\n\nVoici le contenu que j'ai recupere depuis jw.org:\n---\n{jw_content}\n---\nUtilise ce contenu pour preparer des reponses precises."

    prompt_map = {
        "midweek": (
            base_context + "\n\n"
            "Prepare la reunion Vie chretienne et ministere de cette semaine:\n"
            "## JOYAUX DE LA PAROLE DE DIEU\n"
            "- Theme principal et points cles\n"
            "- 2-3 reponses preparees pour participer\n\n"
            "## PERLES SPIRITUELLES\n"
            "- Reponses aux questions\n\n"
            "## APPLIQUE-TOI AU MINISTERE\n"
            "- Aide pour les presentations et demonstrations\n\n"
            "## VIE CHRETIENNE\n"
            "- Points cles + reponses\n\n"
            "## ETUDE BIBLIQUE DE CONGREGATION\n"
            "- Resume des paragraphes + reponses preparees\n"
        ),
        "watchtower": (
            base_context + "\n\n"
            "Prepare l'etude de La Tour de Garde de cette semaine:\n"
            "## THEME ET TEXTE CLE\n"
            "## RESUME DE L'ARTICLE\n"
            "## REPONSES PREPAREES (par paragraphe cle)\n"
            "## POINTS D'APPLICATION PRATIQUE\n"
            "## VERSETS CLES A RETENIR\n"
            "Prepare des reponses courtes et naturelles pour lever la main."
        ),
        "bible_reading": (
            base_context + "\n\n"
            "Prepare la lecture biblique de cette semaine:\n"
            "## CHAPITRES ASSIGNES\n"
            "## CONTEXTE HISTORIQUE\n"
            "## PERLES SPIRITUELLES\n"
            "## LECONS PRATIQUES POUR AUJOURD'HUI\n"
            "## VERSETS MARQUANTS\n"
            "Extrais les points les plus interessants pour la meditation personnelle."
        ),
    }

    system_prompt = prompt_map.get(request_type, prompt_map["midweek"])

    try:
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question or f"Prepare ma reunion pour la semaine du {week_str}"},
            ],
            max_tokens=2500,
            temperature=0.5,
        )
        text = response.choices[0].message.content
        await _track_openai_cost(response, tid)
        return {"response": text, "type": request_type}
    except Exception as e:
        logger.error(f"Theo prepare error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.post("/api/theo/letter")
async def theo_generate_letter(request: Request):
    """Genere une lettre de predication personnalisee."""
    payload = _decode_client_token(_extract_bearer(request))
    tid = payload.get("tenant_id") if payload else None
    if tid != _PROPRIO_TENANT_ID:
        return JSONResponse(status_code=403, content={"error": "Acces reserve"})

    body = await request.json()
    recipient_name = body.get("name", "")
    address = body.get("address", "")
    topic = body.get("topic", "esperance biblique")
    tone = body.get("tone", "chaleureux")

    system_prompt = (
        "Tu es Luna, assistante d'un Temoin de Jehovah pionnier permanent. "
        "Genere une lettre de temoignage a envoyer par courrier postal. "
        "REGLES STRICTES:\n"
        "- Lettre COURTE (150-200 mots max) — les lettres courtes sont lues\n"
        "- Ton chaleureux et respectueux, PAS commercial\n"
        "- Commence par une question engageante sur un sujet d'actualite ou biblique\n"
        "- Mentionne 1-2 versets bibliques pertinents\n"
        "- Propose de decouvrir la reponse de la Bible\n"
        "- Invite a visiter jw.org ou a accepter un cours biblique gratuit\n"
        "- NE PAS copier mot a mot un modele — chaque lettre doit etre unique\n"
        "- Inclus le nom du destinataire si fourni\n"
        "- Termine par une formule polie et ton prenom (Ludovic)\n"
        "- Format: texte pret a imprimer, avec mise en page lettre classique\n"
    )

    user_msg = f"Destinataire: {recipient_name or 'Cher voisin'}\n"
    if address:
        user_msg += f"Adresse: {address}\n"
    user_msg += f"Theme souhaite: {topic}\nTon: {tone}"

    try:
        if not openai_client:
            return JSONResponse(status_code=503, content={"error": "Service IA non disponible"})
        response = openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg},
            ],
            max_tokens=800,
            temperature=0.8,
        )
        text = response.choices[0].message.content
        await _track_openai_cost(response, tid)
        return {"letter": text, "recipient": recipient_name, "topic": topic}
    except Exception as e:
        logger.error(f"Theo letter error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})


if __name__ == "__main__":
    import uvicorn
    import subprocess as _sp

    _is_cloudrun = os.getenv("ENVIRONMENT", "").lower() == "cloudrun"
    port = int(os.getenv("PORT", os.getenv("LUNA_PORT", "8080" if _is_cloudrun else "8888")))

    ssl_kwargs = {}
    if not _is_cloudrun:
        # Local/Docker: SSL auto-signe
        ssl_dir = os.path.dirname(__file__)
        cert_file = os.getenv("SSL_CERTFILE", os.path.join(ssl_dir, "cert.pem"))
        key_file = os.getenv("SSL_KEYFILE", os.path.join(ssl_dir, "key.pem"))

        if not os.path.exists(cert_file) or not os.path.exists(key_file):
            logger.info("Certificats SSL absents - generation auto-signee...")
            try:
                _sp.run([
                    "openssl", "req", "-x509", "-newkey", "rsa:4096",
                    "-keyout", key_file, "-out", cert_file,
                    "-days", "365", "-nodes",
                    "-subj", "/CN=localhost/O=YAWatch-Luna",
                ], check=True, capture_output=True)
                logger.info("Certificats SSL generes (auto-signes, 1 an)")
            except Exception as e:
                logger.error(f"Impossible de generer les certificats SSL: {e}")

        ssl_kwargs = {"ssl_keyfile": key_file, "ssl_certfile": cert_file}
    else:
        logger.info("Cloud Run detecte — SSL desactive (gere par Google)")

    logger.info(f"Demarrage Luna Web - YAWatch-Luna (Souscripteur: {_SUBSCRIBER_NAME})")
    logger.info(f"Mode: {LUNA_MODE.upper()}" + (" (chat + voix + SMS)" if LUNA_MODE == "lite" else " (chat + voix + SMS + visio)"))
    logger.info(f"Environnement: {'CLOUD RUN' if _is_cloudrun else 'LOCAL/DOCKER'}")
    logger.info(f"PV Recette: {'SIGNE' if PV_SIGNED else 'NON SIGNE - MODE SETUP'}")
    logger.info(f"OpenAI: {'OK' if OPENAI_API_KEY else 'MANQUANT'}")
    _twilio_ok = sms_client and hasattr(sms_client, 'is_configured') and sms_client.is_configured
    logger.info(f"Twilio: {'OK' if _twilio_ok else 'NON CONFIGURE'}")
    logger.info(f"Tavus: {'OK' if (tavus_client and tavus_client.is_configured) else 'NON CONFIGURE'}" + (" (mode lite - skip)" if LUNA_MODE == "lite" else ""))
    logger.info(f"Simli: {'OK' if (_SIMLI_AVAILABLE and os.getenv('SIMLI_API_KEY')) else 'NON CONFIGURE'}")
    logger.info(f"Redis/Memory: {'OK' if _memory_manager else 'OFFLINE'}")
    logger.info(f"Safety Guardian: {'OK' if _safety_guardian else 'OFFLINE'}")
    logger.info(f"Quota Guard: {'OK' if _quota_guard else 'OFFLINE'}")
    logger.info(f"Scheduler: {'OK' if _scheduler else 'OFFLINE'}")
    logger.info(f"Executor: {'OK' if _executor else 'OFFLINE'}")
    if _pv_locked:
        logger.info(f"Setup AI: {'OK' if SETUP_OPENAI_API_KEY else 'MANQUANT (SETUP_OPENAI_API_KEY)'}")

    uvicorn.run(
        app,
        host="0.0.0.0",
        port=port,
        **ssl_kwargs,
    )
