"""
Luna Document Generator - Generation de documents DOCX et PDF

Genere des courriers, resumes, fiches et exports au format Word (.docx).
Genere des rapports d'appel au format PDF (fpdf2).
Utilise python-docx pour la mise en page et GPT pour le contenu.
"""
import os
import uuid
import logging
from datetime import datetime
from typing import Optional, Dict, Any, List

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

# Types de documents supportes
DOC_TYPES = {
    "courrier_admin": "Courrier administratif",
    "courrier_resiliation": "Lettre de resiliation",
    "resume_hebdo": "Resume hebdomadaire",
    "liste_courses": "Liste de courses",
    "fiche_sante": "Fiche sante",
    "compte_rendu": "Compte-rendu",
    "export_notes": "Export des notes",
}


class DocumentGenerator:
    """Genere des documents DOCX pour le souscripteur."""

    def __init__(self, output_dir: str, tenant_id: int):
        self.tenant_id = tenant_id
        self.output_dir = os.path.join(output_dir, str(tenant_id))
        os.makedirs(self.output_dir, exist_ok=True)

    def _new_doc(self) -> Document:
        """Cree un nouveau document avec style Luna."""
        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        style.paragraph_format.space_after = Pt(6)
        return doc

    def _add_header(self, doc: Document, title: str, subtitle: str = ""):
        """Ajoute un en-tete Luna au document."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("Luna - YAWatch")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)

        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_p.add_run(datetime.now().strftime("%d/%m/%Y"))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_paragraph()  # espace

        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x3E)

        if subtitle:
            p = doc.add_paragraph(subtitle)
            p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            p.runs[0].font.size = Pt(10)

        doc.add_paragraph()  # espace

    def _add_sender_block(self, doc: Document, profile: Dict[str, Any]):
        """Ajoute le bloc expediteur depuis le profil."""
        name = f"{profile.get('first_name', '')} {profile.get('last_name', '')}".strip()
        address = profile.get("address", "")
        city = profile.get("city", "")
        phone = profile.get("phone", "")
        email = profile.get("email", "")

        lines = [l for l in [name, address, city, phone, email] if l]
        if lines:
            p = doc.add_paragraph("\n".join(lines))
            for run in p.runs:
                run.font.size = Pt(10)
            doc.add_paragraph()

    def _save(self, doc: Document, doc_type: str) -> str:
        """Sauvegarde le document et retourne le nom de fichier."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        short_id = uuid.uuid4().hex[:6]
        filename = f"{doc_type}_{timestamp}_{short_id}.docx"
        filepath = os.path.join(self.output_dir, filename)
        doc.save(filepath)
        logger.info(f"Document genere: {filepath}")
        return filename

    def generate_letter(
        self,
        doc_type: str,
        subject: str,
        body_text: str,
        recipient: str = "",
        profile: Optional[Dict[str, Any]] = None,
    ) -> str:
        """
        Genere un courrier (administratif, resiliation, etc.)

        Args:
            doc_type: Type de document (courrier_admin, courrier_resiliation)
            subject: Objet du courrier
            body_text: Corps du courrier (genere par GPT)
            recipient: Destinataire
            profile: Profil du souscripteur

        Returns:
            Nom du fichier genere
        """
        doc = self._new_doc()

        title = DOC_TYPES.get(doc_type, "Courrier")
        self._add_header(doc, title, f"Objet : {subject}")

        if profile:
            self._add_sender_block(doc, profile)

        if recipient:
            p = doc.add_paragraph(f"A l'attention de : {recipient}")
            p.runs[0].font.bold = True
            doc.add_paragraph()

        # Corps du courrier
        for paragraph in body_text.split("\n\n"):
            paragraph = paragraph.strip()
            if paragraph:
                doc.add_paragraph(paragraph)

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        name = ""
        if profile:
            name = f"{profile.get('first_name', '')} {profile.get('last_name', '')}".strip()
        p = doc.add_paragraph(name or "Le souscripteur")
        p.runs[0].font.bold = True

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("Document genere par Luna - YAWatch")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        footer.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        return self._save(doc, doc_type)

    def generate_summary(
        self,
        notes: List[Dict[str, Any]],
        period: str = "semaine",
        profile: Optional[Dict[str, Any]] = None,
        summary_text: str = "",
    ) -> str:
        """
        Genere un resume (hebdo, mensuel).

        Args:
            notes: Liste de notes a inclure
            period: Periode du resume
            profile: Profil du souscripteur
            summary_text: Resume genere par GPT

        Returns:
            Nom du fichier genere
        """
        doc = self._new_doc()

        name = ""
        if profile:
            name = profile.get("first_name", "")
        self._add_header(doc, f"Resume de la {period}", f"Pour {name}" if name else "")

        if summary_text:
            doc.add_paragraph(summary_text)
            doc.add_paragraph()

        if notes:
            doc.add_heading("Details", level=2)
            for note in notes:
                content = note.get("content", "")
                created = note.get("created_at", "")
                date_str = ""
                if created:
                    try:
                        dt = datetime.fromisoformat(str(created))
                        date_str = dt.strftime("%d/%m %Hh%M")
                    except (ValueError, TypeError):
                        date_str = str(created)[:16]

                p = doc.add_paragraph()
                if date_str:
                    run = p.add_run(f"[{date_str}] ")
                    run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
                    run.font.size = Pt(9)
                p.add_run(content)

        return self._save(doc, "resume_hebdo")

    def generate_health_sheet(
        self,
        profile: Dict[str, Any],
    ) -> str:
        """
        Genere une fiche sante avec les infos du profil.

        Returns:
            Nom du fichier genere
        """
        doc = self._new_doc()

        name = f"{profile.get('first_name', '')} {profile.get('last_name', '')}".strip()
        self._add_header(doc, "Fiche Sante", name)

        # Medecin
        doc.add_heading("Medecin traitant", level=2)
        doctor = profile.get("doctor_name", "")
        doctor_phone = profile.get("doctor_phone", "")
        if doctor:
            doc.add_paragraph(f"Dr {doctor}" + (f" - Tel: {doctor_phone}" if doctor_phone else ""))
        pharmacy = profile.get("pharmacy", "")
        if pharmacy:
            doc.add_paragraph(f"Pharmacie : {pharmacy}")

        # Allergies
        allergies = profile.get("allergies", "")
        if allergies:
            doc.add_heading("Allergies", level=2)
            p = doc.add_paragraph(allergies)
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

        # Traitements
        treatments = profile.get("treatments", "")
        if treatments:
            doc.add_heading("Traitements en cours", level=2)
            for line in treatments.split(","):
                line = line.strip()
                if line:
                    doc.add_paragraph(line, style="List Bullet")

        # Pathologies
        conditions = profile.get("conditions", "")
        if conditions:
            doc.add_heading("Pathologies declarees", level=2)
            doc.add_paragraph(conditions)

        # Mutuelle
        mutual = profile.get("mutual_name", "")
        mutual_num = profile.get("mutual_number", "")
        carte_vitale = profile.get("carte_vitale", "")
        if mutual or carte_vitale:
            doc.add_heading("Couverture", level=2)
            if mutual:
                doc.add_paragraph(f"Mutuelle : {mutual}" + (f" (N. {mutual_num})" if mutual_num else ""))
            if carte_vitale:
                doc.add_paragraph(f"Carte Vitale : {carte_vitale}")

        # Contact urgence medical
        medical_contact = profile.get("medical_contact_person", "")
        if medical_contact:
            doc.add_heading("Personne de confiance medicale", level=2)
            doc.add_paragraph(medical_contact)

        return self._save(doc, "fiche_sante")

    def generate_notes_export(
        self,
        notes: List[Dict[str, Any]],
    ) -> str:
        """
        Exporte les notes en document DOCX.

        Returns:
            Nom du fichier genere
        """
        doc = self._new_doc()
        self._add_header(doc, "Export des notes", f"{len(notes)} note(s)")

        for note in notes:
            content = note.get("content", "")
            context = note.get("context", "")
            created = note.get("created_at", "")
            tags = note.get("tags", [])

            date_str = ""
            if created:
                try:
                    dt = datetime.fromisoformat(str(created))
                    date_str = dt.strftime("%d/%m/%Y %H:%M")
                except (ValueError, TypeError):
                    date_str = str(created)[:16]

            p = doc.add_paragraph()
            if date_str:
                run = p.add_run(f"{date_str} ")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            if context:
                run = p.add_run(f"[{context}] ")
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
            p.add_run(content)

            if tags:
                tag_p = doc.add_paragraph(f"Tags: {', '.join(tags)}")
                tag_p.runs[0].font.size = Pt(8)
                tag_p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        return self._save(doc, "export_notes")

    def generate_contact_list(
        self,
        contacts: List[Dict[str, Any]],
    ) -> str:
        """
        Genere une liste des contacts de confiance.

        Returns:
            Nom du fichier genere
        """
        doc = self._new_doc()
        self._add_header(doc, "Contacts de confiance", f"{len(contacts)} contact(s)")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Nom"
        hdr[1].text = "Relation"
        hdr[2].text = "Telephone"
        hdr[3].text = "Canal"

        for c in contacts:
            row = table.add_row().cells
            row[0].text = c.get("name", "")
            row[1].text = c.get("relation", "")
            row[2].text = c.get("phone", "")
            row[3].text = c.get("preferred_channel", "sms")

        return self._save(doc, "contacts")

    def generate_call_report(
        self,
        call_type: str,
        subscriber_name: str,
        contact_name: str,
        contact_phone: str,
        duration_minutes: float,
        transcript: List[Dict[str, str]],
        actions: List[str] = None,
        call_sid: str = "",
        message_original: str = "",
    ) -> str:
        """
        Genere un rapport PDF d'appel vocal/visio.

        Args:
            call_type: "vocal" ou "visio"
            subscriber_name: Prenom du souscripteur
            contact_name: Nom du contact appele
            contact_phone: Numero du contact (masque partiellement)
            duration_minutes: Duree de l'appel en minutes
            transcript: Liste de {"role": "user|luna", "text": "..."}
            actions: Actions executees pendant l'appel (SMS envoyes, notes, etc.)
            call_sid: Identifiant Twilio de l'appel
            message_original: Message initial demande par le souscripteur

        Returns:
            Nom du fichier PDF genere
        """
        from fpdf import FPDF

        def _sanitize_text(text: str) -> str:
            """Remplace les caracteres Unicode non supportes par latin-1."""
            replacements = {
                "\u2018": "'", "\u2019": "'",  # smart quotes
                "\u201c": '"', "\u201d": '"',
                "\u2013": "-", "\u2014": "-",  # dashes
                "\u2026": "...",  # ellipsis
                "\u2022": "-",  # bullet
                "\u00a0": " ",  # nbsp
            }
            for old, new in replacements.items():
                text = text.replace(old, new)
            # Fallback: remplacer tout caractere non-latin1
            try:
                text.encode("latin-1")
            except UnicodeEncodeError:
                text = text.encode("latin-1", errors="replace").decode("latin-1")
            return text

        class CallReportPDF(FPDF):
            """PDF personnalise avec header/footer Luna."""

            def header(self):
                self.set_font("Helvetica", "B", 10)
                self.set_text_color(124, 58, 237)  # Violet Luna
                self.cell(0, 8, "Luna - YAWatch", align="L")
                self.set_font("Helvetica", "", 9)
                self.set_text_color(136, 136, 136)
                self.cell(0, 8, datetime.now().strftime("%d/%m/%Y %H:%M"), align="R", new_x="LMARGIN", new_y="NEXT")
                self.set_draw_color(124, 58, 237)
                self.line(10, self.get_y(), 200, self.get_y())
                self.ln(4)

            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 8)
                self.set_text_color(170, 170, 170)
                ref = f"Ref: {call_sid[:12]}..." if call_sid and len(call_sid) > 12 else (call_sid or "N/A")
                self.cell(0, 10, f"Rapport Luna | {ref} | Page {self.page_no()}/{{nb}}", align="C")

        pdf = CallReportPDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        # --- Titre ---
        pdf.set_font("Helvetica", "B", 18)
        pdf.set_text_color(26, 26, 62)
        title = f"Rapport d'appel {call_type}"
        pdf.cell(0, 12, title, new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

        # --- Informations de l'appel ---
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 62)
        pdf.cell(0, 8, "Informations", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Helvetica", "", 10)
        pdf.set_text_color(51, 51, 51)

        # Masquer partiellement le numero
        masked_phone = contact_phone
        if contact_phone and len(contact_phone) > 6:
            masked_phone = contact_phone[:4] + "****" + contact_phone[-2:]

        info_lines = [
            ("Souscripteur", subscriber_name),
            ("Contact appele", f"{contact_name} ({masked_phone})"),
            ("Type", f"Appel {call_type}"),
            ("Date", datetime.now().strftime("%d/%m/%Y a %H:%M")),
            ("Duree", f"{duration_minutes:.1f} minute(s)"),
        ]
        if message_original:
            info_lines.append(("Demande initiale", message_original[:200]))

        for label, value in info_lines:
            pdf.set_font("Helvetica", "B", 10)
            pdf.cell(45, 7, f"{label} :", align="R")
            pdf.set_font("Helvetica", "", 10)
            pdf.cell(0, 7, f"  {_sanitize_text(value)}", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(6)

        # --- Transcription ---
        pdf.set_font("Helvetica", "B", 11)
        pdf.set_text_color(26, 26, 62)
        pdf.cell(0, 8, "Transcription de l'appel", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        if transcript:
            for entry in transcript:
                role = entry.get("role", "")
                text = entry.get("text", "").strip()
                if not text:
                    continue

                if role == "luna":
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(124, 58, 237)
                    speaker = "Luna"
                else:
                    pdf.set_font("Helvetica", "B", 9)
                    pdf.set_text_color(51, 51, 51)
                    speaker = contact_name or "Interlocuteur"

                pdf.cell(30, 6, f"{_sanitize_text(speaker)} :", align="R")
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(51, 51, 51)
                # Multi-cell pour les textes longs
                pdf.multi_cell(0, 6, f"  {_sanitize_text(text)}", new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
        else:
            pdf.set_font("Helvetica", "I", 10)
            pdf.set_text_color(136, 136, 136)
            pdf.cell(0, 8, "Aucune transcription disponible.", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(4)

        # --- Actions executees ---
        if actions:
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(26, 26, 62)
            pdf.cell(0, 8, "Actions executees pendant l'appel", new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)

            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 51, 51)
            for action in actions:
                pdf.cell(6, 6, "-")
                pdf.multi_cell(0, 6, f" {_sanitize_text(action)}", new_x="LMARGIN", new_y="NEXT")

            pdf.ln(4)

        # --- Mention legale ---
        pdf.ln(6)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(136, 136, 136)
        pdf.multi_cell(0, 4,
            _sanitize_text(
                "Ce rapport a ete genere automatiquement par Luna, assistante IA YAWatch. "
                "L'interlocuteur a ete informe en debut d'appel qu'un compte-rendu serait "
                "transmis au souscripteur. Ce document est confidentiel et destine uniquement "
                f"a {subscriber_name}."
            ),
            new_x="LMARGIN", new_y="NEXT",
        )

        # Sauvegarder
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        short_id = uuid.uuid4().hex[:6]
        filename = f"rapport_appel_{timestamp}_{short_id}.pdf"
        filepath = os.path.join(self.output_dir, filename)
        pdf.output(filepath)
        logger.info(f"Call report PDF generated: {filepath}")
        return filename

    def list_documents(self) -> List[Dict[str, Any]]:
        """Liste les documents generes pour ce tenant."""
        docs = []
        if not os.path.exists(self.output_dir):
            return docs
        for f in sorted(os.listdir(self.output_dir), reverse=True):
            if f.endswith((".docx", ".pdf")):
                filepath = os.path.join(self.output_dir, f)
                stat = os.stat(filepath)
                # Parse type from filename (e.g. courrier_admin_20260131_...)
                parts = f.split("_")
                doc_type = parts[0] if parts else "unknown"
                if len(parts) > 1:
                    combined = f"{parts[0]}_{parts[1]}"
                    if combined in DOC_TYPES or combined == "rapport_appel":
                        doc_type = combined
                type_label = DOC_TYPES.get(doc_type, doc_type)
                if doc_type == "rapport_appel":
                    type_label = "Rapport d'appel"
                docs.append({
                    "filename": f,
                    "type": type_label,
                    "size_kb": round(stat.st_size / 1024, 1),
                    "created_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                })
        return docs
